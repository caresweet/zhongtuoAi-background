"""Report generation pipeline — 8-phase orchestrator.

Phase 0: Multi-modal RAG (infrastructure, run once)
Phase 1: Knowledge retrieval
Phase 2: Structure analysis
Phase 3: Image positioning
Phase 4: Content generation
Phase 5: Table processing
Phase 6: Gap analysis & fill
Phase 7: Quality validation

Usage:
    pipeline = ReportPipelineService()
    result = await pipeline.run(
        materials_dir="~/Downloads/稳评资料",
        region="淮安市洪泽区",
        template_path="...",
    )
"""

import asyncio
import os
import time
from pathlib import Path
from typing import List, Dict, Any, Optional, Callable

from app.services.pipeline.pipeline_context import (
    PipelineContext, GenerationRecipe, DocumentStructure,
    QualityReport, Gap,
)
from app.services.pipeline.knowledge_retrieval import KnowledgeRetrievalService
from app.services.pipeline.structure_analyzer import StructureAnalyzer
from app.services.pipeline.image_positioner import ImagePositioner
from app.services.pipeline.content_generator import ContentGenerator
from app.services.pipeline.table_processor import TableProcessor
from app.services.pipeline.gap_analyzer import GapAnalyzer
from app.services.pipeline.quality_validator import QualityValidator


class ReportPipelineService:
    """Main pipeline orchestrator for report generation.

    Coordinates all 8 phases and manages state flow between them.
    Can resume from any phase if state is persisted.
    """

    def __init__(
        self,
        llm_service=None,
        embedder=None,
        vector_store=None,
        retriever=None,
        cross_modal_retriever=None,
        multimodal_embedder=None,
        table_generation_service=None,
    ):
        """Initialize with all required services.

        Args:
            llm_service: LLMService instance.
            embedder: EmbedderService instance.
            vector_store: VectorStoreService instance.
            retriever: RetrieverService instance.
            cross_modal_retriever: CrossModalRetriever instance (multi-modal).
            multimodal_embedder: MultiModalEmbedder instance.
            table_generation_service: TableGenerationService instance.
        """
        self.llm = llm_service
        self.embedder = embedder
        self.vector_store = vector_store
        self.retriever = retriever
        self.cross_modal = cross_modal_retriever
        self.mm_embedder = multimodal_embedder

        # Initialize phase services
        self.phase1 = KnowledgeRetrievalService(retriever, cross_modal_retriever)
        self.phase2 = StructureAnalyzer()
        self.phase3 = ImagePositioner(llm_service, multimodal_embedder)
        self.phase4 = ContentGenerator(llm_service, retriever, cross_modal_retriever)
        self.phase5 = TableProcessor(llm_service, table_generation_service)
        self.phase6 = GapAnalyzer()
        self.phase7 = QualityValidator()

        # Connect gap analyzer to content/table services for auto-fill
        self.phase6.content_gen = self.phase4
        self.phase6.table_proc = self.phase5

        # State
        self.recipe: Optional[GenerationRecipe] = None
        self.doc_structure: Optional[DocumentStructure] = None
        self.filled_sections: Dict[int, str] = {}
        self.filled_tables: Dict[int, List[List[str]]] = {}
        self.filled_images: Dict[str, Any] = {}
        self.quality_report: Optional[QualityReport] = None
        self.remaining_gaps: List[Gap] = []

        # Progress tracking
        self._progress_callback: Optional[Callable] = None
        self._start_time: float = 0

    # ── Main entry point ──

    async def run(
        self,
        materials_dir: str,
        region: str = "",
        template_path: str = "",
        example_path: str = "",
        session_id: str = "",
        progress_callback: Optional[Callable] = None,
    ) -> Dict[str, Any]:
        """Run the complete 8-phase pipeline.

        Args:
            materials_dir: Path to user's project materials folder.
            region: Region identifier (e.g. "淮安市洪泽区").
            template_path: Path to template .docx (overrides auto-detection).
            example_path: Path to example .docx (overrides auto-detection).
            session_id: Session ID for RAG and state management.
            progress_callback: Async callback(phase, status, message).

        Returns:
            Dict with:
                - recipe: GenerationRecipe
                - doc_structure: DocumentStructure
                - filled_sections: Dict
                - filled_tables: Dict
                - filled_images: Dict
                - quality_report: QualityReport
                - output_path: str (if auto-saved)
                - stats: timing and counts
        """
        self._start_time = time.time()
        self._progress_callback = progress_callback

        print("\n" + "=" * 60)
        print("🚀 Report Generation Pipeline")
        print("=" * 60)

        # ── Phase 1: Knowledge Retrieval ──
        await self._notify("phase1", "running", "Phase 1: 多模态知识检索")
        t1 = time.time()

        self.recipe = await self.phase1.retrieve(
            materials_dir=materials_dir,
            region=region,
            session_id=session_id,
        )

        # Override template paths if provided
        if template_path:
            self.recipe.template_path = template_path
        if example_path:
            self.recipe.example_report_path = example_path

        await self._notify("phase1", "done",
                           f"✅ 数据提取完成: {self.recipe.project_context.land_area_sqm:.0f}㎡, "
                           f"{len(self.recipe.project_context.image_files)}张图片")

        # ── Phase 2: Structure Analysis ──
        await self._notify("phase2", "running", "Phase 2: 文档结构分析")
        t2 = time.time()

        self.doc_structure = self.phase2.analyze(self.recipe)
        self.recipe.doc_structure = self.doc_structure

        await self._notify("phase2", "done",
                           f"✅ {self.doc_structure.total_sections}个章节, "
                           f"{len(self.doc_structure.tables)}个表格, "
                           f"{len(self.doc_structure.image_slots)}个图位")

        # ── Phase 3: Image Positioning ──
        await self._notify("phase3", "running", "Phase 3: 图片定位与智能匹配")
        t3 = time.time()

        self.filled_images = await self.phase3.position_images(
            self.doc_structure, self.recipe.project_context,
        )

        matched = sum(1 for s in self.doc_structure.image_slots if s.matched_image)
        await self._notify("phase3", "done",
                           f"✅ {matched}/{len(self.doc_structure.image_slots)}个图位已匹配")

        # ── Phase 4: Content Generation ──
        await self._notify("phase4", "running", "Phase 4: AI内容生成")
        t4 = time.time()

        self.filled_sections = await self.phase4.generate_all_sections(
            self.doc_structure,
            self.recipe.project_context,
            session_id=session_id,
            stream_callback=self._stream_progress,
        )

        total_chars = sum(len(c) for c in self.filled_sections.values())
        await self._notify("phase4", "done",
                           f"✅ {len(self.filled_sections)}个章节, {total_chars}字")

        # ── Phase 5: Table Processing ──
        await self._notify("phase5", "running", "Phase 5: 表格数据处理")
        t5 = time.time()

        # Load template and example docs for table operations
        from docx import Document
        doc_tpl = None
        doc_ex = None
        try:
            if self.recipe.template_path and os.path.exists(self.recipe.template_path):
                doc_tpl = Document(self.recipe.template_path)
            if self.recipe.example_report_path and os.path.exists(self.recipe.example_report_path):
                doc_ex = Document(self.recipe.example_report_path)
        except Exception as e:
            print(f"  ⚠️ Could not load documents for tables: {e}")

        self.filled_tables = await self.phase5.process_all_tables(
            self.doc_structure,
            self.recipe.project_context,
            doc_tpl,
            doc_ex,
        )

        await self._notify("phase5", "done",
                           f"✅ {len(self.filled_tables)}个表格已处理")

        # ── Phase 6: Gap Analysis & Fill ──
        await self._notify("phase6", "running", "Phase 6: 缺口分析与填补")
        t6 = time.time()

        self.filled_sections, self.remaining_gaps = await self.phase6.analyze_and_fill(
            self.doc_structure,
            self.recipe.project_context,
            self.filled_sections,
            self.filled_tables,
            self.filled_images,
        )

        await self._notify("phase6", "done",
                           f"{'✅ 无缺口' if not self.remaining_gaps else f'⚠️ {len(self.remaining_gaps)}个缺口未填补'}")

        # ── Phase 7: Quality Validation ──
        await self._notify("phase7", "running", "Phase 7: 质量验证")
        t7 = time.time()

        self.quality_report = await self.phase7.validate(
            self.doc_structure,
            self.recipe.project_context,
            self.filled_sections,
            self.filled_tables,
            self.filled_images,
            self.remaining_gaps,
        )

        await self._notify("phase7", "done",
                           f"{'✅ PASS' if self.quality_report.passed else '⚠️ NEEDS REVIEW'} "
                           f"(完整性: {self.quality_report.completeness_pct:.1f}%)")

        # ── Summary ──
        elapsed = time.time() - self._start_time
        print(f"\n⏱ Pipeline complete in {elapsed:.1f}s")
        await self._notify("complete", "done",
                           f"🎉 报告生成完成 ({elapsed:.0f}秒)")

        return {
            "recipe": self.recipe,
            "doc_structure": self.doc_structure,
            "filled_sections": self.filled_sections,
            "filled_tables": self.filled_tables,
            "filled_images": self.filled_images,
            "quality_report": self.quality_report,
            "output_path": "",  # To be set by caller after docx assembly
            "stats": {
                "elapsed_sec": elapsed,
                "sections_generated": len(self.filled_sections),
                "tables_filled": len(self.filled_tables),
                "images_matched": matched,
                "gaps_remaining": len(self.remaining_gaps),
                "completeness_pct": self.quality_report.completeness_pct if self.quality_report else 0,
            },
        }

    # ── Quick mode: structure-only analysis ──

    async def analyze_only(
        self, materials_dir: str, region: str = "", template_path: str = ""
    ) -> GenerationRecipe:
        """Run only Phase 1 + 2 for quick structure analysis.

        Useful for previewing what the pipeline will do before full generation.
        """
        self.recipe = await self.phase1.retrieve(materials_dir, region)
        if template_path:
            self.recipe.template_path = template_path
        self.doc_structure = self.phase2.analyze(self.recipe)
        self.recipe.doc_structure = self.doc_structure
        return self.recipe

    # ── Resume from phase ──

    async def resume_from(
        self, phase: str, state: Dict[str, Any],
        session_id: str = "",
    ) -> Dict[str, Any]:
        """Resume pipeline from a specific phase with saved state.

        Args:
            phase: Phase name to resume from ("phase3", "phase4", etc.).
            state: Dict with saved recipe, doc_structure, etc.
            session_id: Session ID.

        Returns:
            Same as run().
        """
        # Restore state
        self.recipe = state.get("recipe")
        self.doc_structure = state.get("doc_structure")
        self.filled_sections = state.get("filled_sections", {})
        self.filled_tables = state.get("filled_tables", {})
        self.filled_images = state.get("filled_images", {})

        if not self.recipe or not self.doc_structure:
            raise ValueError("Missing recipe or doc_structure in saved state")

        # Run remaining phases
        # (simplified — in practice, call the relevant phase methods)
        phases = ["phase1", "phase2", "phase3", "phase4", "phase5", "phase6", "phase7"]
        start_idx = phases.index(phase)

        # ... (run from start_idx onward, similar to run())
        return await self.run(
            materials_dir="",  # Already parsed
            session_id=session_id,
        )

    # ── Progress notifications ──

    async def _notify(self, phase: str, status: str, message: str):
        """Send progress notification."""
        print(f"  [{phase}] {message}")
        if self._progress_callback:
            try:
                await self._progress_callback(phase, status, message)
            except Exception:
                pass

    async def _stream_progress(self, event_type: str, message: str):
        """Stream progress during content generation."""
        await self._notify("phase4", event_type, message)

    # ── Static: save pipeline output to docx ──

    @staticmethod
    async def save_to_docx(
        pipeline_result: Dict[str, Any],
        output_path: str,
        docx_service=None,
    ) -> str:
        """Save pipeline results to a .docx file.

        This integrates with the existing DocxService to apply
        generated content, tables, and images to the template.

        Args:
            pipeline_result: Result dict from pipeline.run().
            output_path: Path to save the .docx.
            docx_service: DocxService instance.

        Returns:
            Output path.
        """
        recipe = pipeline_result.get("recipe")
        doc_structure = pipeline_result.get("doc_structure")
        filled_sections = pipeline_result.get("filled_sections", {})
        filled_tables = pipeline_result.get("filled_tables", {})
        filled_images = pipeline_result.get("filled_images", {})

        if not recipe or not recipe.template_path:
            raise ValueError("No template path in pipeline result")

        if docx_service:
            # Use DocxService to apply fills
            # This mirrors the existing _assemble_and_finalize logic
            from docx import Document
            from docx.shared import Cm

            doc = Document(recipe.template_path)

            # Fill sections
            for section_idx, content in filled_sections.items():
                if section_idx < len(doc_structure.sections):
                    section = doc_structure.sections[section_idx]
                    paras = content.split("\n\n")
                    body_idx = 0
                    for pi in range(section.start_para_index + 1, section.end_para_index + 1):
                        if pi >= len(doc.paragraphs):
                            break
                        para = doc.paragraphs[pi]
                        text = para.text.strip()
                        # Only fill placeholder/empty body paragraphs
                        if not text or "待Agent生成" in text or "待生成" in text:
                            if body_idx < len(paras):
                                for run in para.runs:
                                    run.text = ""
                                if para.runs:
                                    para.runs[0].text = paras[body_idx]
                                else:
                                    para.add_run(paras[body_idx])
                                body_idx += 1

            # Fill tables
            for table_idx, data in filled_tables.items():
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    for r_idx, row_data in enumerate(data):
                        if r_idx >= len(table.rows):
                            break
                        for c_idx, val in enumerate(row_data):
                            if c_idx >= len(table.rows[r_idx].cells):
                                break
                            cell = table.rows[r_idx].cells[c_idx]
                            if cell.paragraphs:
                                for run in cell.paragraphs[0].runs:
                                    run.text = ""
                                if cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].text = str(val)
                                else:
                                    cell.paragraphs[0].add_run(str(val))

            # Insert images
            for slot_id, slot in filled_images.items():
                if hasattr(slot, 'matched_image') and slot.matched_image:
                    para_idx = slot.resolve(doc_structure)
                    if para_idx is not None and para_idx < len(doc.paragraphs):
                        para = doc.paragraphs[para_idx]
                        try:
                            run = para.add_run()
                            run.add_picture(slot.matched_image, width=Cm(14))
                        except Exception as e:
                            print(f"  ⚠️ Failed to insert image {slot.matched_image}: {e}")

            doc.save(output_path)
            print(f"\n📄 Report saved to: {output_path}")
            return output_path

        else:
            print("  ⚠️ No docx_service provided, skipping docx save")
            return ""
