#!/usr/bin/env python3
"""
stability_report_generate.py — 通用社会稳定风险评估报告生成引擎

核心逻辑:
  1. 解析参考模板docx → 提取章节/表格/法规/句式
  2. 加载稳评资料 → LLM提取项目数据
  3. 按模板结构逐章生成新报告
  4. 与模板对比 → 85%+ 达标

使用:
  python3 scripts/stability_report_generate.py \
    --template "/Users/mac/Downloads/洞庭湖路征地稳评-原报告.docx" \
    --project "洪拟征告〔2026〕7号商业服务业设施用地项目"
"""

import argparse
import asyncio
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

_SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(_SCRIPT_DIR.parent))

from docx import Document as DocxDocument
from kb_engine.llm import LLMClient
from scripts.region_profiles import (
    match_all_levels, format_all_regulations, get_government_chain,
)

# ── 路径 ──────────────────────────────────────────────────────────
MATERIAL_DIR_DEFAULT = Path.home() / "Downloads" / "稳评资料"
MATERIAL_CACHE = _SCRIPT_DIR.parent.parent / "data" / "material_text_cache.json"
SEED_DATA_DIR = _SCRIPT_DIR.parent / "seed_data"
OUTPUT_DIR = _SCRIPT_DIR / "output"

# ── 真实知识库路径 ───────────────────────────────────────────────
KB_DB_PATH = _SCRIPT_DIR.parent / "data" / "knowledge_base.db"
STABILITY_DB_PATH = _SCRIPT_DIR.parent / "data" / "stability_kb.db"
CHROMA_PATH = _SCRIPT_DIR.parent / "data" / "chroma"
UPLOADS_DIR = _SCRIPT_DIR.parent / "data" / "uploads"

# seed_data 作为兜底（知识库为空时使用）
KB_FALLBACK_FILES = [
    str(SEED_DATA_DIR / "example_report_jinhu.md"),
    str(SEED_DATA_DIR / "report_writing_guide.md"),
    str(SEED_DATA_DIR / "case_study_templates.md"),
    str(SEED_DATA_DIR / "company_info_zhongtuo.md"),
]


# ═══════════════════════════════════════════════════════════════════
# 模板解析器 — 从参考docx提取完整结构
# ═══════════════════════════════════════════════════════════════════

class TemplateParser:
    """解析参考稳评报告docx，提取章节结构、表格格式、法规列表、关键句式。"""

    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = DocxDocument(docx_path)
        self.chapters: List[Dict] = []       # [{no, title, subsections, paragraphs}]
        self.tables: List[Dict] = []         # [{chapter_no, rows, cols, headers, sample_data}]
        self.regulations: List[str] = []     # 提取的法规引用
        self.company_info: Dict = {}         # 公司信息
        self.project_info: Dict = {}         # 项目信息(从模板提取，需替换)
        self.full_text: str = ""

    def parse(self) -> Dict:
        """主解析入口。返回完整模板结构描述。"""
        self._extract_full_text()
        self._extract_chapters()
        self._extract_tables()
        self._extract_regulations()
        self._extract_company_info()
        self._extract_project_info()

        return {
            "source": self.docx_path,
            "chapters": self.chapters,
            "chapter_count": len(self.chapters),
            "tables": self.tables,
            "table_count": len(self.tables),
            "regulations": self.regulations,
            "regulation_count": len(self.regulations),
            "company_info": self.company_info,
            "template_project_info": self.project_info,
        }

    def _extract_full_text(self):
        self.full_text = "\n".join(p.text for p in self.doc.paragraphs)

    def _extract_chapters(self):
        """提取章节结构：从段落样式中识别章节标题和子标题。"""
        chapters = []
        current_chapter = None

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name if para.style else ""
            # 检测章节标题 (如 "第一章 xxx")
            ch_match = re.match(r'第([一二三四五六七八九十\d]+)章\s*(.*)', text)
            if ch_match:
                if current_chapter:
                    chapters.append(current_chapter)
                cn = ch_match.group(1)
                ch_no = self._cn_to_int(cn)
                current_chapter = {
                    "no": ch_no,
                    "title": ch_match.group(2).strip() if ch_match.group(2) else text,
                    "subsections": [],
                    "paragraph_count": 0,
                    "sample_paragraphs": [],
                }
                continue

            # 检测子标题 (如 "（一）xxx" 或 "1. xxx" 或 Heading样式)
            if current_chapter:
                is_sub = (
                    style.startswith("Heading") or
                    bool(re.match(r'[（(][一二三四五六七八九十\d]+[）)]', text)) or
                    bool(re.match(r'\d+\.\s', text))
                )
                if is_sub:
                    current_chapter["subsections"].append({
                        "text": text[:100],
                        "style": style,
                    })

                current_chapter["paragraph_count"] += 1
                if len(current_chapter["sample_paragraphs"]) < 5 and len(text) > 30:
                    # 去除纯数字/日期等非正文
                    if not re.match(r'^[\d\s./\-：:]+$', text):
                        current_chapter["sample_paragraphs"].append(text[:200])

        if current_chapter:
            chapters.append(current_chapter)

        self.chapters = chapters
        print(f"  📋 解析到 {len(chapters)} 章")

    def _extract_tables(self):
        """提取表格结构，并精确映射到所在章节。"""
        tables = []
        # 构建段落位置索引: paragraph_index → body_element_order
        body = self.doc.element.body
        para_positions = {}  # paragraph_index → body_child_index
        para_idx = 0
        for bi, child in enumerate(body):
            if child.tag.endswith('}p'):
                para_positions[para_idx] = bi
                para_idx += 1

        # 找每个章节在body中的起始位置
        chapter_body_positions = {}  # ch_no → first_body_child_index
        for ch in self.chapters:
            # 找该章节第一个段落在段落列表中的索引
            for pi, para in enumerate(self.doc.paragraphs):
                text = para.text.strip()
                if f"第{ch['no']}章" in text or (ch['no'] <= 12 and f"第{self._cn_to_str(ch['no'])}章" in text):
                    if pi in para_positions:
                        chapter_body_positions[ch['no']] = para_positions[pi]
                    break

        for i, tbl in enumerate(self.doc.tables):
            rows = len(tbl.rows)
            cols = len(tbl.columns)
            headers = []
            if rows > 0:
                for cell in tbl.rows[0].cells:
                    headers.append(cell.text.strip()[:80])
            sample = []
            for ri in range(min(3, rows)):
                row_data = [tbl.rows[ri].cells[ci].text.strip()[:60]
                           for ci in range(min(cols, 8))]
                sample.append(row_data)

            # 精确映射表格到章节：找body中表格元素位置，比对章节起始位置
            tbl_elem = tbl._element
            tbl_body_idx = None
            for bi, child in enumerate(body):
                if child is tbl_elem:
                    tbl_body_idx = bi
                    break

            # 确定章节：表格属于body位置在其之前最近的章节
            ch_no = 1
            if tbl_body_idx is not None and chapter_body_positions:
                best_ch = 1
                for cn, pos in sorted(chapter_body_positions.items()):
                    if pos <= tbl_body_idx:
                        best_ch = cn
                ch_no = best_ch

            tables.append({
                "index": i + 1,
                "rows": rows,
                "cols": cols,
                "headers": headers,
                "sample_data": sample,
                "estimated_chapter": ch_no,
            })

        self.tables = tables
        # 统计每个章节的表格数
        ch_table_counts = {}
        for t in tables:
            cn = t["estimated_chapter"]
            ch_table_counts[cn] = ch_table_counts.get(cn, 0) + 1
        for cn, count in sorted(ch_table_counts.items()):
            ch_name = next((c['title'][:20] for c in self.chapters if c['no'] == cn), '?')
            print(f"  📊 第{cn}章「{ch_name}」: {count}个表格")
        print(f"  📊 共 {len(tables)} 个表格")

    def _extract_regulations(self):
        """提取模板中引用的法规列表。"""
        regs = []
        patterns = [
            r'《([^》]+)》',
            r'(DB\d+/T\s*\d+[^，,\n]*)',
        ]
        for pattern in patterns:
            found = re.findall(pattern, self.full_text)
            regs.extend(found)

        # 去重
        seen = set()
        unique = []
        for r in regs:
            key = r[:20]
            if key not in seen:
                seen.add(key)
                unique.append(r)

        self.regulations = unique
        print(f"  📜 解析到 {len(unique)} 条法规引用")

    def _extract_company_info(self):
        """提取模板中的公司固定信息。"""
        info = {}
        patterns = {
            "company_name": r'(江苏众拓项目代理咨询有限公司)',
            "legal_rep": r'(?:法定代表人|总经理)[：:]\s*(\S{2,4})',
            "credit_code": r'(?:统一社会信用代码)[：:]\s*(\w{18})',
            "address": r'(?:住所|办公场所)[：:]\s*(\S{10,80})',
        }
        for key, pat in patterns.items():
            m = re.search(pat, self.full_text)
            if m:
                info[key] = m.group(1).strip()
        self.company_info = info

    def _extract_project_info(self):
        """提取模板中的项目特定信息（需要在新报告中替换）。"""
        info = {}
        patterns = {
            "project_name": r'(?:决策事项名称|决策名称)[：:]\s*(\S{10,80})',
            "responsibility_unit": r'(?:稳评责任单位)[：:]\s*(\S{10,60})',
            "implementation_unit": r'(?:稳评实施单位)[：:]\s*(\S{10,60})',
        }
        for key, pat in patterns.items():
            m = re.search(pat, self.full_text)
            if m:
                info[key] = m.group(1).strip()
        self.project_info = info

    _CN_MAP = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6,
               "七": 7, "八": 8, "九": 9, "十": 10, "十一": 11, "十二": 12}
    _CN_REV = {v: k for k, v in _CN_MAP.items()}

    @classmethod
    def _cn_to_int(cls, cn: str) -> int:
        if cn.isdigit():
            return int(cn)
        return cls._CN_MAP.get(cn, 0)

    @classmethod
    def _cn_to_str(cls, num: int) -> str:
        return cls._CN_REV.get(num, str(num))

    def build_structure_summary(self) -> str:
        """构建供LLM使用的模板结构摘要。"""
        lines = [f"# 参考模板结构（来源: {Path(self.docx_path).name}）\n"]
        lines.append(f"## 基本信息")
        lines.append(f"- 总章节: {len(self.chapters)}章")
        lines.append(f"- 总表格: {len(self.tables)}个")
        lines.append(f"- 总段落: {len(self.doc.paragraphs)}段\n")

        lines.append("## 章节结构")
        for ch in self.chapters:
            sub_names = [s["text"][:60] for s in ch["subsections"][:8]]
            lines.append(f"\n### 第{ch['no']}章 {ch['title']}")
            lines.append(f"- 段落数: {ch['paragraph_count']}")
            if sub_names:
                lines.append(f"- 子标题: {' | '.join(sub_names[:6])}")
            if ch["sample_paragraphs"]:
                lines.append(f"- 样例文本: \"{ch['sample_paragraphs'][0][:150]}...\"")

        lines.append(f"\n## 表格结构")
        for tbl in self.tables:
            lines.append(f"\n### 表格{tbl['index']} ({tbl['rows']}行×{tbl['cols']}列, 第{tbl['estimated_chapter']}章)")
            lines.append(f"- 表头: {' | '.join(tbl['headers'][:6])}")
            if tbl["sample_data"]:
                lines.append(f"- 样例行: {' | '.join(tbl['sample_data'][0][:4])}")

        lines.append(f"\n## 法规引用（{len(self.regulations)}条）")
        for r in self.regulations[:15]:
            lines.append(f"- {r}")

        return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════
# 系统角色定义（注入LLM的System Prompt）
# ═══════════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """# 身份
你是专业化社会稳定风险评估报告生成智能体，专职从事各类征地/工程/项目稳评报告的编制工作。
你精通国家土地管理法规、各省市地方稳评标准、公文写作规范。

# 核心能力
1. 📋 解析参考模板 → 提取章节结构、表格格式、法规引用、公文句式
2. 📁 分析稳评资料 → 从PDF/图片中提取项目数据（文号/位置/面积/地类/户数）
3. ⚖️  匹配地方规范 → 四级法规体系（国家→省→市→区县）
4. ✍️  逐章生成报告 → 遵循模板结构，用项目数据填充，知识库补充合规内容
5. 🔬 质量自检 → 与模板对比，不达标自动迭代

# 核心原则
1. 模板结构不可变 — 不新增/删除/拆分/合并模板原有章节、子标题、表格
2. 数据必有来源 — 项目专属数据来自稳评资料；通用段落复用模板/知识库
3. 法规引用完整 — 完整文号+文件名，不简写、不篡改
4. 缺失必标注 — 无数据时标注【待补充】，不自编自造
5. 实施单位固定 — 江苏众拓项目代理咨询有限公司，不可修改
6. 风险评分严格 — 按DB32/T4013-2021或目标区域等效标准执行

# 写作风格
- 全程使用征地稳评官方书面公文语体
- 杜绝口语、网络用语、主观情绪化描述
- 政策引用完整文号+文件名
- 表格使用规范Markdown格式，表头加粗
- 公司简介/评估原则/应急预案框架 → 直接复用模板原文
- 仅替换项目差异化可变参数（地名/文号/面积/户数/金额）

# 四级法规体系
根据项目所在地自动匹配适用法规:
- 国家法律: 《土地管理法》《突发事件应对法》《民法典》等
- 省级标准: DB32/T4013（江苏）/ DB43/TXXXX（湖南）等
- 市级规范: 各地市的实施细则、补偿标准通知
- 区县级: 地方十四五规划、城市总体规划、控制性详细规划

# 禁止行为
- 编造征地面积、资金、户数、问卷数据、群众诉求
- 虚构不存在的政策文件、法规文号
- 修改章节结构或风险打分体系
- 未搜索就声称"未找到相关法规" """

# ═══════════════════════════════════════════════════════════════════
# 报告生成器
# ═══════════════════════════════════════════════════════════════════

class StabilityReportGenerator:
    """通用稳评报告生成器 — 学习模板结构，填充新项目数据。"""

    def __init__(self, template_path: str, project_desc: str,
                 materials_path: str = ""):
        self.template_path = template_path
        self.project_desc = project_desc
        self.materials_path = Path(materials_path) if materials_path else MATERIAL_DIR_DEFAULT

        self.llm = LLMClient()
        self.parser = TemplateParser(template_path)
        self.template_structure: Dict = {}
        self.kb_text: str = ""
        self.kb_sources: List[str] = []
        self.source_data: Dict = {}
        self.chapters: Dict[int, str] = {}
        self.ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    # ── 主流程 ──────────────────────────────────────────────────
    async def run(self) -> Dict:
        print("=" * 70)
        print(f"  通用稳评报告生成")
        print(f"  模板: {Path(self.template_path).name}")
        print(f"  项目: {self.project_desc}")
        print(f"  {datetime.now().isoformat()}")
        print("=" * 70)

        # Step 1: 解析模板
        print(f"\n📋 Step 1: 解析参考模板...")
        self.template_structure = self.parser.parse()

        # Step 2: 加载知识库
        print(f"\n📚 Step 2: 加载知识库...")
        self._load_kb()

        # Step 3: 四级法规匹配（国家→省→市→区县）
        print(f"\n⚖️  Step 3: 四级法规匹配...")
        self.regulations = match_all_levels(self.project_desc)
        self.regulation_text = format_all_regulations(self.regulations)
        gov_chain = get_government_chain(self.regulations)
        if gov_chain:
            print(f"  政府机构链: {' → '.join(gov_chain[:4])}")
        prov = self.regulations.get("province", {})
        city = self.regulations.get("city", {})
        district = self.regulations.get("district", {})
        print(f"  法规层级: 国家 → {prov.get('name', '待识别')} → {city.get('name', '待识别')} → {district.get('name', '待识别')}")

        # Step 4: 加载资料 + 提取数据
        print(f"\n📁 Step 4: 加载稳评资料...")
        materials = self._load_materials()
        self.source_data = await self._extract_source_data(materials)

        # Step 5: 构建生成大纲
        print(f"\n📋 Step 5: 构建生成大纲...")
        structure_summary = self.parser.build_structure_summary()

        # Step 6: 逐章生成
        print(f"\n✍️ Step 6: 逐章生成（{len(self.parser.chapters)}章）...")
        for ch in self.parser.chapters:
            ch_no = ch["no"]
            if ch_no == 0:
                continue
            print(f"  第{ch_no}章「{ch['title'][:30]}」...", end=" ")
            self.chapters[ch_no] = await self._generate_chapter(ch, structure_summary)
            print(f"({len(self.chapters[ch_no])}字)")

        # Step 7: 对比
        print(f"\n🔬 Step 7: 模板对比...")
        comparison = self._compare_with_template()
        score = comparison["overall_score"]
        print(f"  相似度: {score:.1%}")

        # 迭代优化
        for iteration in range(3):
            if score >= 0.85:
                break
            print(f"\n🔄 第{iteration+1}轮迭代 (当前{score:.1%})...")
            weak = comparison.get("weak_chapters", [])
            for ch_no in weak:
                ch = next((c for c in self.parser.chapters if c["no"] == ch_no), None)
                if ch:
                    print(f"  重写第{ch_no}章...")
                    self.chapters[ch_no] = await self._generate_chapter(ch, structure_summary, rewrite=True)
            comparison = self._compare_with_template()
            score = comparison["overall_score"]

        status = "✅" if score >= 0.85 else "⚠️"
        print(f"\n  {status} 最终: {score:.1%}")

        # Step 8: 输出
        print(f"\n📄 Step 8: 输出...")
        report = self._assemble_report()
        outputs = self._save(report, comparison)

        for k, v in outputs.items():
            print(f"  📄 {k}: {v}")
        return {"chapters": self.chapters, "comparison": comparison, "outputs": outputs}

    # ── Step 2: 知识库（SQLite + ChromaDB + seed_data 三层合并）──
    def _load_kb(self):
        """从真实知识库三层加载稳评相关内容。

        Layer 1: knowledge_base.db → 文档全文 (extracted_text)
        Layer 2: stability_kb.db → 章节指引 + 法规
        Layer 3: ChromaDB → 向量化片段 (用 get()+关键词过滤, 无需embedding)
        兜底:   seed_data/*.md → 基础法规/标准
        """
        parts = []
        sources = []

        # ═══════════════════════════════════════════════════════════
        # Layer 1: knowledge_base.db — 文档全文
        # ═══════════════════════════════════════════════════════════
        if KB_DB_PATH.exists():
            try:
                import sqlite3
                conn = sqlite3.connect(str(KB_DB_PATH))

                # 稳评文档 (extracted_text 有内容)
                rows = conn.execute("""
                    SELECT title, document_type, extracted_text
                    FROM knowledge_documents
                    WHERE domain='stability'
                      AND extracted_text IS NOT NULL
                      AND length(extracted_text) > 100
                    ORDER BY length(extracted_text) DESC
                """).fetchall()

                for title, dtype, text in rows:
                    # 完整报告取更多字符，规范取适量
                    limit = 12000 if dtype == 'example_report' else 6000
                    parts.append(
                        f"【📄 {title}】\n{(text or '')[:limit]}"
                    )
                    sources.append(f"SQL:{title[:30]}")

                # 稳评报告1-7 extracted_text 为空但在 ChromaDB 中有向量
                # 标记一下，从 ChromaDB 补充
                empty_rows = conn.execute("""
                    SELECT id, title FROM knowledge_documents
                    WHERE domain='stability'
                      AND (extracted_text IS NULL OR length(extracted_text) < 100)
                      AND indexed_status='completed'
                """).fetchall()
                empty_ids = [r[0] for r in empty_rows]
                empty_names = [r[1] for r in empty_rows]

                conn.close()

                kb_chars = sum(len(p) for p in parts)
                print(f"  📚 knowledge_base.db: {len(rows)} 篇全文 ({kb_chars}字)"
                      + (f" + {len(empty_ids)} 篇仅向量" if empty_ids else ""))
            except Exception as e:
                print(f"  ⚠️ knowledge_base.db: {e}")
                empty_ids, empty_names = [], []

        # ═══════════════════════════════════════════════════════════
        # Layer 2: stability_kb.db — 章节指引 + 法规
        # ═══════════════════════════════════════════════════════════
        if STABILITY_DB_PATH.exists():
            try:
                import sqlite3
                conn = sqlite3.connect(str(STABILITY_DB_PATH))
                chapters = conn.execute(
                    "SELECT chapter_no, chapter_title, writing_guide FROM learned_chapters ORDER BY chapter_no"
                ).fetchall()
                if chapters:
                    ch_lines = []
                    for cn, ctitle, guide in chapters:
                        ch_lines.append(f"第{cn}章 {ctitle}: {guide or ''}")
                    parts.append("【📋 模板章节写作指引】\n" + "\n".join(ch_lines)[:3000])
                    sources.append("SQL:chapter_guides")

                regs = conn.execute(
                    "SELECT title, doc_no, content_text FROM regulations WHERE is_active=1"
                ).fetchall()
                if regs:
                    reg_lines = ["- " + (r[0] or '') + (f" ({r[1]})" if r[1] else '')
                                for r in regs[:20]]
                    parts.append("【📜 知识库法规列表】\n" + "\n".join(reg_lines)[:2000])
                    sources.append("SQL:regulations")
                conn.close()
                if chapters:
                    print(f"  📋 stability_kb.db: {len(chapters)} 章指引 + {len(regs)} 条法规")
            except Exception as e:
                print(f"  ⚠️ stability_kb.db: {e}")

        # ═══════════════════════════════════════════════════════════
        # Layer 3: ChromaDB — 按source_file聚合还原完整文档
        # ═══════════════════════════════════════════════════════════
        if CHROMA_PATH.exists():
            try:
                import chromadb
                from collections import defaultdict
                client = chromadb.PersistentClient(path=str(CHROMA_PATH))
                collections = client.list_collections()
                kb_cols = [c for c in collections if 'knowledge_base' in c.name.lower()]

                if kb_cols:
                    col = kb_cols[0]
                    total = col.count()

                    # 分批获取全部chunks
                    all_docs = []
                    all_metas = []
                    batch_size = 400
                    for offset in range(0, total, batch_size):
                        limit = min(batch_size, total - offset)
                        data = col.get(limit=limit, offset=offset,
                                      include=['documents', 'metadatas'])
                        all_docs.extend(data.get('documents', []))
                        all_metas.extend(data.get('metadatas', []))

                    # 按 source_file 聚合还原文档
                    BIDDING_KW = ['招标', '投标', '评标', '中标', '林地报批', '采购',
                                 '批次城市建设', '竞争性磋商', '响应文件']
                    by_source = defaultdict(list)

                    for d, m in zip(all_docs, all_metas):
                        src = (m or {}).get('source_file', 'unknown')
                        d_text = d or ''

                        # 排除招标投标
                        if any(kw in d_text for kw in BIDDING_KW) or \
                           any(kw in src for kw in BIDDING_KW):
                            continue

                        by_source[src].append(d_text)

                    # 对每个来源：合并chunks，按类型和重要性截取
                    # 优先级: 稳评报告(完整) > 地方规范(核心) > 参考书(纲要)
                    PRIORITY_ORDER = {
                        'report': 0,     # 稳评报告
                        'regulation': 1,  # 地方规范
                        'guide': 2,       # 参考书/指南
                        'other': 3,
                    }

                    def classify(src_name):
                        n = src_name.lower()
                        if any(k in n for k in ['稳评报告', '社会稳定报告', '洞庭湖',
                                                  '通港路', '淮安北', '金湖', '洪泽']):
                            return 'report'
                        if any(k in n for k in ['规范', '标准', 'db32', 'db320', 'db43']):
                            return 'regulation'
                        if any(k in n for k in ['指南', '理论', '方法', '案例']):
                            return 'guide'
                        return 'other'

                    # 稳评报告：取8K — 完整的报告结构最重要
                    # 地方规范：取5K — 核心条款
                    # 参考书：取3K — 纲要即可
                    LIMITS = {'report': 8000, 'regulation': 5000, 'guide': 3000, 'other': 2000}

                    reconstructed = []
                    total_limit = 80000  # 总量控制在80K以内

                    for src, chunks in sorted(by_source.items(),
                                             key=lambda x: -len(x[1])):
                        full_text = "\n".join(chunks)
                        cat = classify(src)
                        limit = LIMITS[cat]
                        reconstructed.append({
                            'text': f"【📄 {src} ({len(full_text):,}字)】\n{full_text[:limit]}",
                            'priority': PRIORITY_ORDER[cat],
                            'chars': limit,
                            'src': src,
                        })

                    # 按优先级排序，同优先级按大小降序
                    reconstructed.sort(key=lambda x: (x['priority'], -x['chars']))

                    # 去重：相同来源只保留一份
                    seen_names = set()
                    for p in parts:
                        # 提取已有的文档名
                        name_match = p.split('\n')[0] if p else ''
                        seen_names.add(name_match[:30])

                    added_chars = 0
                    for r in reconstructed:
                        doc_key = r['text'].split('\n')[0][:30] if r['text'] else ''
                        if doc_key in seen_names:
                            continue
                        if added_chars + r['chars'] > total_limit:
                            # 达到总量限制，后续文档只取标题
                            if r['priority'] <= 1:  # 报告和规范优先保留
                                short_text = r['text'].split('\n')[0] + '\n(已达上下文限制，仅保留标题)'
                                parts.append(short_text)
                                added_chars += 200
                            continue
                        parts.append(r['text'])
                        seen_names.add(doc_key)
                        sources.append(f"Chroma:{r['src'][:30]}")
                        added_chars += r['chars']

                    chroma_added = sum(1 for r in reconstructed
                                     if any(r['text'] in p for p in parts[-len(reconstructed):]))
                    total_chars = sum(r['chars'] for r in reconstructed)
                    print(f"  🔍 ChromaDB: {total} chunks → "
                          f"{len(by_source)} 篇文档 → "
                          f"聚合 {total_chars:,} 字符")

            except Exception as e:
                print(f"  ⚠️ ChromaDB 跳过: {str(e)[:80]}")

        # ═══════════════════════════════════════════════════════════
        # Layer 4: seed_data 兜底
        # ═══════════════════════════════════════════════════════════
        if len(parts) < 3:
            print(f"  ⚠️ 知识库内容不足，补充 seed_data")
            for path in KB_FALLBACK_FILES:
                p = Path(path)
                if p.exists():
                    parts.append(f"【兜底: {p.name}】\n{p.read_text(encoding='utf-8')[:5000]}")
                    sources.append(f"fallback:{p.name}")

        self.kb_text = "\n\n" + "="*40 + "\n\n".join(parts)
        self.kb_sources = sources
        print(f"  📦 知识库总计: {len(self.kb_text)} 字符, {len(sources)} 个来源")

        # ── 4. seed_data 兜底 ──
        if not parts:
            print(f"  ⚠️ 知识库为空，使用 seed_data 兜底文件")
            for path in KB_FALLBACK_FILES:
                p = Path(path)
                if p.exists():
                    parts.append(f"【兜底文件: {p.name}】\n{p.read_text(encoding='utf-8')[:5000]}")
                    sources.append(f"fallback:{p.name}")

        self.kb_text = "\n\n---\n\n".join(parts)
        self.kb_sources = sources
        print(f"  知识库来源: {', '.join(sources[:8])}")
        print(f"  总知识库内容: {len(self.kb_text)} 字符")

    # ── Step 3: 资料 + 数据提取 ─────────────────────────────────
    def _load_materials(self) -> Dict[str, str]:
        data = {}
        if MATERIAL_CACHE.exists():
            cache = json.loads(MATERIAL_CACHE.read_text(encoding="utf-8"))
            for key, text in cache.items():
                if text and len(text) > 10:
                    data[f"pdf_{key}"] = text[:5000]
        if self.materials_path.exists():
            for subdir in self.materials_path.iterdir():
                if subdir.is_dir() and not subdir.name.startswith("."):
                    imgs = [f.name for f in subdir.glob("*")
                           if f.suffix.lower() in {".jpg", ".jpeg", ".png"}]
                    if imgs:
                        data[f"images_{subdir.name}"] = f"{subdir.name}: {len(imgs)}张"
        print(f"  加载 {len(data)} 项资料")
        return data

    async def _extract_source_data(self, materials: Dict) -> Dict:
        """LLM从资料提取项目数据。"""
        extracted = {
            "project_name": self.project_desc,
            "implementation_unit": "江苏众拓项目代理咨询有限公司",
        }

        pdf_texts = {k: v for k, v in materials.items() if k.startswith("pdf_")}
        if pdf_texts and self.llm.available:
            combined = "\n\n".join(f"【{k}】\n{v}" for k, v in list(pdf_texts.items())[:2])
            prompt = f"""从稳评资料提取项目数据(JSON):
{combined[:6000]}

提取: project_reference(文号), project_location(位置省市区街道社区),
land_use(用途), total_area_sqm, total_area_mu, decision_body(决策单位),
responsibility_unit(责任单位), involved_villages(村组), involved_households,
involved_population, land_types_summary(地类概述), compensation_info(补偿描述)

仅JSON。无数据标"待提取"."""
            try:
                result = await asyncio.wait_for(
                    self.llm.chat([{"role": "user", "content": prompt}], max_tokens=1024, temperature=0.2),
                    timeout=30)
                result = re.sub(r"```\w*\s*", "", result.strip()).rstrip("```")
                if result.startswith("{"):
                    for k, v in json.loads(result).items():
                        if v and v != "待提取":
                            extracted[k] = str(v)
                    print(f"  LLM提取 {len(extracted)} 个字段")
            except Exception as e:
                print(f"  ⚠️ 提取失败: {e}")

        return extracted

    # ── Step 5: 逐章生成 ────────────────────────────────────────
    async def _generate_chapter(self, ch: Dict, structure_summary: str,
                                rewrite: bool = False) -> str:
        """按模板章节结构生成。注入角色+法规+数据+模板样例+知识库。"""
        sub_names = [s["text"] for s in ch["subsections"][:10]]
        samples = "\n".join(f"> {s}" for s in ch["sample_paragraphs"][:3])

        # 找出本章对应的模板表格
        ch_tables = [t for t in self.parser.tables if t["estimated_chapter"] == ch["no"]]
        table_hint = ""
        if ch_tables:
            table_hint = "\n## ⚠️ 本章必须包含以下表格！\n"
            for t in ch_tables:
                table_hint += (
                    f"- 表格{t['index']}: {t['rows']}行×{t['cols']}列, "
                    f"表头: {' | '.join(t['headers'][:5])}\n"
                )
                if t["sample_data"]:
                    table_hint += f"  样例数据: {' | '.join(t['sample_data'][0][:4])}\n"
            table_hint += "必须用Markdown表格格式生成对应内容。\n"

        # 检查法规是否需要联网搜索补充
        web_search_hint = ""
        prov = self.regulations.get("province", {})
        city = self.regulations.get("city", {})
        if ch["no"] == 3:  # 第3章是编制依据，特别需要完整法规
            if not prov:
                web_search_hint = (
                    "\n## ⚠️ 本章是编制依据，法规必须完整！\n"
                    "如知识库中缺少以下内容，请标注【待联网搜索补充】:\n"
                    "- 目标省份的征地补偿标准最新文件\n"
                    "- 目标城市的社会稳定风险评估实施细则\n"
                    "- 目标区县的国土空间规划和十四五规划\n"
                )

        rewrite_hint = ""
        if rewrite:
            rewrite_hint = (
                "\n## ⚠️ 重写轮次 — 重点改进\n"
                "1. 充实每个子节的具体内容（不能只有标题）\n"
                "2. 增加数据细节和具体数值\n"
                "3. 补全所有应包含的表格\n"
                "4. 确保法规引用包含完整文号\n"
                "5. 缺失信息统一标注【待补充】\n"
            )

        prompt = f"""# 任务：生成稳评报告 — 第{ch['no']}章: {ch['title']}

## 适用法规体系（四级：国家→省→市→区县）
{self.regulation_text[:3000]}

## 新项目数据
{json.dumps(self.source_data, ensure_ascii=False, indent=2)[:1500]}

## 模板本章参考
- 子标题（必须全部覆盖）: {sub_names if sub_names else '按模板样例自行组织'}
- 段落数参考: {ch['paragraph_count']}段
- 模板样例文本:
{samples[:1000] if samples else '(从知识库范文参考)'}
{table_hint}
{web_search_hint}
## 知识库参考
{self.kb_text[:5000]}
{rewrite_hint}
## 生成指令
你是社会稳定风险评估报告编写专家。请生成「第{ch['no']}章 {ch['title']}」完整内容：

1. **结构**: 子标题与模板保持一致，逐个覆盖
2. **数据**: 使用上述新项目数据，不编造
3. **法规**: 优先引用上述四级法规体系中的对应法规，完整文号+文件名
4. **通用段落**: 公司简介/评估原则/应急预案框架等 → 复用模板原文或知识库范文
5. **表格**: {f"必须包含 {len(ch_tables)} 个Markdown表格" if ch_tables else "根据需要添加表格"}
6. **缺失**: 无数据的字段标注【待补充】
7. **实施单位**: 江苏众拓项目代理咨询有限公司
8. **文风**: 官方书面公文语体

仅输出本章正文（Markdown格式）。"""

        try:
            result = await asyncio.wait_for(
                self.llm.chat(
                    [{"role": "user", "content": prompt}],
                    system=SYSTEM_PROMPT,
                    max_tokens=2048, temperature=0.3 if not rewrite else 0.5,
                ), timeout=60)
            return result.strip()
        except Exception as e:
            print(f"❌ {e}")
            return f"## 第{ch['no']}章 {ch['title']}\n\n(生成失败: {e})\n"

    # ── Step 6: 对比 ────────────────────────────────────────────
    def _compare_with_template(self) -> Dict:
        """与参考模板对比。"""
        t_chapters = self.parser.chapters
        # 章节覆盖
        gen_nos = set(self.chapters.keys())
        tpl_nos = {c["no"] for c in t_chapters if c["no"] > 0}
        ch_cov = len(gen_nos & tpl_nos) / max(len(tpl_nos), 1)

        # 子节覆盖
        total_subs, covered_subs = 0, 0
        for ch in t_chapters:
            ch_no = ch["no"]
            content = self.chapters.get(ch_no, "")
            for sub in ch["subsections"]:
                total_subs += 1
                key = sub["text"][:6]
                if key in content:
                    covered_subs += 1
        sub_cov = covered_subs / max(total_subs, 1)

        # 表格数
        tpl_tables = len(self.parser.tables)
        gen_tables = 0
        for content in self.chapters.values():
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                if lines[i].strip().startswith("|") and "|" in lines[i][1:]:
                    if i + 2 < len(lines) and re.match(r"\|[\s\-:]+\|", lines[i+1].strip()):
                        gen_tables += 1
                        while i < len(lines) and lines[i].strip().startswith("|"):
                            i += 1
                        continue
                i += 1
        table_diff = abs(gen_tables - tpl_tables)
        table_score = 1.0 if table_diff <= 2 else (0.8 if table_diff <= 4 else 0.4)

        # 关键要素
        all_text = "\n".join(self.chapters.values())
        key_elements = [
            r'江苏众拓项目代理咨询有限公司',
            r'稳评.*责任单位', r'土地管理法',
            r'应急.*预案', r'补偿.*标准',
            r'风险.*等级|低风险|中风险|高风险',
            r'合法性.*分析|合理性.*分析|可行性.*分析|可控性.*分析',
            r'问卷|调查', r'评估结论', r'DB32.*4013|社会稳定风险评估规范',
        ]
        key_score = sum(1 for p in key_elements if re.search(p, all_text)) / len(key_elements)

        # 文风
        style_pats = [
            r'依据|根据|按照|依照', r'予以|给予', r'不得|严禁|应当|必须',
            r'〔.*〕.*号', r'第.*条|第.*款',
            r'平方米|㎡|亩', r'经.*研究|经.*审核|经.*批准',
        ]
        style_score = min(1.0, sum(1 for p in style_pats if re.search(p, all_text))
                         / len(style_pats) * 1.2)

        overall = (ch_cov * 0.35 + sub_cov * 0.25 + table_score * 0.15
                   + key_score * 0.15 + style_score * 0.10)

        # 找薄弱章节
        weak = []
        for ch in t_chapters:
            ch_no = ch["no"]
            content = self.chapters.get(ch_no, "")
            if len(content) < 100:
                weak.append(ch_no)
            elif ch["subsections"]:
                cov = sum(1 for s in ch["subsections"] if s["text"][:6] in content)
                if cov / len(ch["subsections"]) < 0.5:
                    weak.append(ch_no)
        # 也检查表格
        for tbl in self.parser.tables:
            if tbl["estimated_chapter"] > 0:
                ch_content = self.chapters.get(tbl["estimated_chapter"], "")
                if "|" not in ch_content:
                    if tbl["estimated_chapter"] not in weak:
                        weak.append(tbl["estimated_chapter"])

        return {
            "overall_score": round(overall, 4),
            "dimension_scores": {
                "chapter_coverage": round(ch_cov, 4),
                "subsection_coverage": round(sub_cov, 4),
                "table_count": round(table_score, 4),
                "key_elements": round(key_score, 4),
                "writing_style": round(style_score, 4),
            },
            "weak_chapters": list(set(weak))[:5],
            "summary": (f"综合{overall:.1%} | 章节{ch_cov:.0%} | "
                       f"子节{sub_cov:.0%} | 表格{table_score:.0%} | "
                       f"要素{key_score:.0%} | 文风{style_score:.0%}"),
            "template_tables": tpl_tables,
            "generated_tables": gen_tables,
        }

    # ── 组装 + 输出 ─────────────────────────────────────────────
    def _assemble_report(self) -> str:
        sd = self.source_data
        name = sd.get("project_name", self.project_desc)
        lines = [
            f"# {name}",
            f"# 社会稳定风险评估报告",
            "",
            f"> 编制单位: 江苏众拓项目代理咨询有限公司",
            f"> 编制日期: {datetime.now().strftime('%Y年%m月%d日')}",
            f"> 参考模板: {Path(self.template_path).name}",
            "", "---", "",
        ]
        for ch in self.parser.chapters:
            ch_no = ch["no"]
            if ch_no == 0:
                continue
            content = self.chapters.get(ch_no, "")
            lines.append(f"# 第{ch['no']}章 {ch['title']}")
            lines.append("")
            lines.append(content)
            lines.append("")
            lines.append("---")
            lines.append("")

        lines.append("## 附录: 生成说明")
        lines.append(f"- 参考模板: {Path(self.template_path).name}")
        lines.append(f"- 项目数据来源: 稳评资料 + LLM提取")
        lines.append(f"- 实施单位: 江苏众拓项目代理咨询有限公司")
        lines.append(f"- 法规引用: 知识库 + 参考模板")
        return "\n".join(lines)

    def _save(self, report: str, comparison: Dict) -> Dict[str, str]:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        safe = re.sub(r'[\\/:*?"<>|\n\r]', '_',
                      self.source_data.get("project_name", "report"))[:40]
        outputs = {}
        md_path = OUTPUT_DIR / f"{safe}_{self.ts}.md"
        md_path.write_text(report, encoding="utf-8")
        outputs["report"] = str(md_path)

        cmp_path = OUTPUT_DIR / f"comparison_{self.ts}.json"
        cmp_path.write_text(json.dumps({
            "template": self.template_path,
            "project": self.project_desc,
            "timestamp": self.ts,
            **comparison,
        }, ensure_ascii=False, indent=2), encoding="utf-8")
        outputs["comparison"] = str(cmp_path)

        return outputs


# ═══════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════

async def main():
    p = argparse.ArgumentParser(description="通用稳评报告生成")
    p.add_argument("--template", type=str, required=True, help="参考模板docx路径")
    p.add_argument("--project", type=str, required=True, help="新项目名称")
    p.add_argument("--materials", type=str, default="", help="稳评资料目录")
    p.add_argument("--compare", type=str, default="", help="仅对比模式: 已有报告路径")
    args = p.parse_args()

    if not Path(args.template).exists():
        print(f"❌ 模板不存在: {args.template}")
        sys.exit(1)

    gen = StabilityReportGenerator(
        template_path=args.template,
        project_desc=args.project,
        materials_path=args.materials,
    )
    result = await gen.run()
    score = result["comparison"]["overall_score"]
    if score >= 0.85:
        print(f"\n✅ 相似度 {score:.1%} ≥ 85%")
    else:
        print(f"\n⚠️ 相似度 {score:.1%} < 85%")


if __name__ == "__main__":
    asyncio.run(main())
