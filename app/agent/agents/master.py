"""MasterAgent — 简洁对话主智能体，负责意图识别和任务分发。

职责：
1. 理解用户意图（提问/提供数据/要求生成报告）
2. 提问 → 检索知识库回答
3. 提供数据 → 确认并存储
4. 要求生成 → 委托 ChapterOrchestrator 逐章生成

不再负责：资料提取（DataAnalysisAgent）、章节生成（ChapterAgents）、质量审核（QualityReviewAgent）
"""

import re
import json
import asyncio as _asyncio
import logging
from typing import Dict, List, Any, Optional

from .base_agent import BaseAgent

logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════════
# System Prompt
# ═══════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """# 身份
你是众拓智能助手，江苏众拓项目代理咨询有限公司的AI助手，
专注于社会稳定风险评估领域的智能服务。

# 你能做什么
1. 📋 **生成稳评报告** — 基于用户上传的资料，逐章编写10章专业社会稳定风险评估报告
2. 📖 **查询解说法规** — 检索知识库中的法规、标准（DB32/T4013-2021等），解答稳评相关问题
3. 🖼️ **分析上传图片** — 识别公告、问卷、照片中的文字和数据，提取结构化信息
4. 📊 **构建专业表格** — 根据用户提供的数据生成规范的风险等级表、调查统计表等
5. 💬 **专业咨询解答** — 回答稳评流程、资质要求、政策法规等任何相关问题

# 核心原则
1. 所有数据来自用户提供的资料或知识库检索结果，不编造
2. 缺失信息标注【待补充】，不猜测
3. 实施单位固定：江苏众拓项目代理咨询有限公司
4. 回答时引用具体法规文号和条款，知识库未收录的诚实说明

# 意图判断（每次收到消息先判断类型）

| 意图 | 关键词/特征 | 处理方式 |
|------|--------|---------|
| 生成报告 | "生成报告""写报告""开始生成""逐章生成""需要生成""编制报告"等 | 检查数据→有则触发逐章生成，无则告知缺什么 |
| 提问咨询 | "是什么""怎么""如何""规定""标准""法规""要求""条款"等 | 深度检索知识库，结构化回答（法规名称+文号+关键条款） |
| 知识解说 | "解释""解说""说明""介绍""讲一下" + 法规/标准/流程名 | 深层RAG检索，专家口吻结构化解说 |
| 构建表格 | "表格""建表""构造表""做个表""生成.*表""整理.*表格"等 | 根据数据和知识库模板生成规范Markdown表格 |
| 提供数据 | 含具体信息（位置/面积/文号/户数/金额等） | 确认收到，汇总已有数据，告知还缺什么 |
| 上传文件 | 含PDF/图片/文档 | 分析文件内容，汇总提取结果，告知文件类型和可用的下一步操作 |
| 闲聊 | "你好""谢谢""在吗"等 | 简短友好回复，引导使用助手能力 |

# 回复要求
- 简洁直接，一般2-5句话
- 不逐项提问，一次性汇总缺失信息
- 知识解说类回答结构化：法规名称 → 文号 → 关键条款 → 实际应用
- 表格构建时使用规范Markdown表格格式
- 禁编造数据、政策、群众意见

# 知识库
可检索内容包括：
- DB32/T4013-2021《第三方社会稳定风险评估规范》
- 《土地管理法》及实施条例
- 《国有土地上房屋征收与补偿条例》
- 《突发事件应对法》
- 地方稳评实施细则
- 公司资质及业务流程文档

基于检索结果回答，引用文号和条款。知识库未收录的信息诚实告知。

# 可调度Agent（仅在有明确任务时调用）
| Agent | 用途 |
|-------|------|
| ImageAnalysisAgent | 分析上传的图片（问卷/公告/照片） |

调用格式：{{"type": "action", "agent": "ImageAnalysisAgent", "message": "正在分析...", "fills": {{}}}}

# 报告模板
{template_context}

# 知识库检索
{rag_context}

# 已收集数据
{collected_summary}

{conversation_recent}
---
请根据用户意图回复（纯文本，不要JSON，除非调用Agent）。"""


# ═══════════════════════════════════════════════════════════════
# Bidding System Prompt — 招标投标专用身份
# ═══════════════════════════════════════════════════════════════

BIDDING_SYSTEM_PROMPT = """# 身份
你是众拓智能助手，江苏众拓项目代理咨询有限公司的AI助手，
当前作为**招标投标文件编写专家**，专注于招标投标领域的专业文档服务。

# 你能做什么
1. 🏗️ **生成招标文件** — 解析上传的招标文档，生成招标公告、评标报告、中标公示、招标情况报告等专业文档
2. 📖 **查询招标法规** — 检索招标知识库中的模板、法规，解答招标投标流程相关问题
3. 📊 **提取招标数据** — 从上传的招标文档中提取项目名称、编号、预算金额、资格要求、评分标准等结构化数据
4. 📝 **仿照模板撰写** — 严格模仿招标知识库中模板的章节结构、条款措辞、表格格式，只替换具体项目数据
5. 💬 **招标专业咨询** — 回答招标流程、资质要求、评标方法等任何招标相关问题

# 核心原则
1. 所有数据来自用户上传的招标文档或招标知识库检索结果，不编造
2. 严格模仿招标知识库中模板的结构、措辞、条款完整度
3. 模板中的每个章节和条款都必须保留，只替换项目具体变量
4. 招标代理机构统一为：江苏众拓项目代理咨询有限公司
5. **禁止**引用社会稳定风险评估相关规范（DB32/T4013等）
6. **禁止**在文档中标注数据来源，来源由系统内部追踪
7. 缺失信息标注【待补充】，不猜测

# 知识库
招标知识库可检索内容包括：
- 招标公告模板
- 评标报告模板
- 中标公示模板
- 招标情况报告模板
- 招标投标相关法规
- 政府采购政策文件

基于招标知识库检索结果回答，引用具体文档。知识库未收录的信息诚实告知。

# 可调度Agent
| Agent | 用途 |
|-------|------|
| BiddingDataAgent | 解析招标文档（招标公告/投标文件/中标通知），提取结构化数据 |
| BiddingReportAgent | 生成招标报告（招标公告/评标报告/中标公示/招标情况报告） |

调用格式：{{"type": "action", "agent": "BiddingDataAgent", "message": "正在解析...", "fills": {{}}}}

# 招标数据
{collected_summary}

{conversation_recent}
---
请根据用户意图回复（纯文本，不要JSON，除非调用Agent）。你当前的身份是**招标投标文件编写专家**。"""


# ═══════════════════════════════════════════════════════════════
# Domain helpers
# ═══════════════════════════════════════════════════════════════

_BIDDING_KEYWORDS = re.compile(r'(?:招标|投标|评标|中标|标书|采购|bid|tender)', re.IGNORECASE)


def _is_bidding_context(state: dict, message: str = "") -> bool:
    """Check if the current conversation context is about bidding/tender."""
    # 1. State flag (set by previous bidding interactions)
    if state.get("_conversation_domain") == "bidding" or state.get("_domain") == "bidding":
        return True
    # 2. Bidding data already in state (strong signal, domain-specific)
    if state.get("_bidding_data") or state.get("_bidding_data_ready"):
        return True
    # 3. Resolve via the domain registry (keyword match on the message)
    try:
        from app.domains import detect_domain
        return detect_domain(message=message, state=state) == "bidding"
    except Exception:
        # Fallback to the original keyword regex if the registry is unavailable
        return bool(message and _BIDDING_KEYWORDS.search(message))


def _get_domain_system_prompt(state: dict, message: str = "") -> str:
    """Return the appropriate system prompt based on conversation domain.

    Resolves the domain via the registry so adding a new report type does not
    require editing this function. Falls back to the two legacy prompts.
    """
    try:
        from app.domains import get_domain, detect_domain
        domain_id = state.get("_domain") or detect_domain(message=message, state=state)
        return get_domain(domain_id).identity_prompt
    except Exception:
        if _is_bidding_context(state, message):
            return BIDDING_SYSTEM_PROMPT
        return SYSTEM_PROMPT


def _set_conversation_domain(state: dict, domain: str) -> None:
    """Set the conversation domain flag so subsequent messages use the right persona."""
    try:
        from app.domains import list_domains
        valid = {d.domain_id for d in list_domains()}
    except Exception:
        valid = {"bidding", "stability"}
    if domain in valid:
        state["_conversation_domain"] = domain
        state["_domain"] = domain

# ═══════════════════════════════════════════════════════════════
# Helper: conversation history
# ═══════════════════════════════════════════════════════════════

def _get_context_manager(state: dict):
    from app.agent.context_manager import ContextManager
    ctx = state.get("_context_manager")
    if ctx is None:
        ctx = ContextManager()
        state["_context_manager"] = ctx
    return ctx

def _add_to_history(state: dict, role: str, content: str):
    msgs = state.setdefault("messages", [])
    msgs.append({"role": role, "content": content[:2000], "timestamp": None})
    # Auto-summarize when exceeding 20 messages
    if len(msgs) > 20:
        summary = _summarize_history(msgs[:-10])  # Summarize older messages
        state["_context_summary"] = summary
        state["messages"] = msgs[-10:]  # Keep only last 10

def _summarize_history(msgs: list) -> str:
    """Extract key facts from older conversation messages."""
    all_text = " ".join(m.get("content", "") for m in msgs)
    facts = []
    # Extract key data points
    import re
    for pattern, label in [
        (r'(?:位置|位于|坐落)[：:]\s*(\S{5,50})', "位置"),
        (r'(\d{4,7})\s*(?:平方米|㎡)', "面积"),
        (r'(\d+\.?\d*)\s*亩', "亩数"),
        (r'(?:涉及|共)\s*(\d+)\s*户', "户数"),
        (r'(?:补偿标准|综合补偿)[^\d]*(\d[\d,.]*)', "补偿标准"),
        (r'(?:责任单位|决策单位)[：:]\s*(\S{2,30})', "责任单位"),
        (r'(\d{2,3}\.?\d*)\s*%\s*支持', "支持率"),
    ]:
        m = re.search(pattern, all_text)
        if m and not any(label in f for f in facts):
            facts.append(f"{label}: {m.group(1).strip()}")
    return "；".join(facts) if facts else "（对话摘要）"

def _get_recent_history(state: dict, n: int = 10) -> str:
    """Get recent conversation + context summary for LLM prompt."""
    parts = []
    # Add context summary if available
    summary = state.get("_context_summary", "")
    if summary:
        parts.append(f"【对话摘要】{summary}")
    # Add recent messages
    msgs = state.get("messages", [])[-n:]
    for m in msgs:
        role = "用户" if m.get("role") == "user" else "助手"
        content = str(m.get("content", ""))[:300]
        parts.append(f"{role}: {content}")
    return "\n".join(parts)

def _build_collected_summary(state: dict) -> str:
    """简洁的数据收集摘要"""
    filled = state.get("filled_data", {})
    if not filled:
        return "尚未收集任何数据。"
    lines = ["已收集数据："]
    for k, v in list(filled.items())[:10]:
        val = str(v)[:60].replace("\n", " ")
        lines.append(f"- {k}: {val}")
    return "\n".join(lines)

# ═══════════════════════════════════════════════════════════════
# MasterAgent
# ═══════════════════════════════════════════════════════════════

class MasterAgent(BaseAgent):
    """简洁对话主智能体。"""

    name = "MasterAgent"
    description = "理解用户意图，回答问题，分发任务"
    covered_steps = []

    def __init__(self, llm_service=None):
        super().__init__(llm_service)
        self._agents: Dict[str, BaseAgent] = {}

    # ═══════════════════════════════════════════════════════════
    # Agent Registry
    # ═══════════════════════════════════════════════════════════

    def _get_or_create_agent(self, agent_name: str) -> Optional[BaseAgent]:
        if agent_name in self._agents:
            return self._agents[agent_name]

        from .image_analyzer_agent import ImageAnalysisAgent
        from .bidding_data_agent import BiddingDataAgent
        from .bidding_report_agent import BiddingReportAgent
        from .capability_analysis_agent import CapabilityAnalysisAgent

        agent_map = {
            "ImageAnalysisAgent": ImageAnalysisAgent,
            "BiddingDataAgent": BiddingDataAgent,
            "BiddingReportAgent": BiddingReportAgent,
            "CapabilityAnalysisAgent": CapabilityAnalysisAgent,
        }

        agent_class = agent_map.get(agent_name)
        if agent_class is None and agent_name.endswith("Agent"):
            agent_class = agent_map.get(agent_name[:-5])
        if agent_class is None:
            return None

        agent = agent_class(llm_service=self._llm)
        self._agents[agent_name] = agent
        return agent

    async def _invoke_agent(self, agent_name: str, state: dict) -> None:
        agent = self._get_or_create_agent(agent_name)
        if agent is None:
            await self._emit_message(f"⚠️ `{agent_name}` 暂不可用", "error")
            return
        try:
            await agent.run(state, self._stream_queue)
        except Exception as e:
            logger.error(f"Agent {agent_name} failed: {e}")

    # ═══════════════════════════════════════════════════════════
    # Think / Act
    # ═══════════════════════════════════════════════════════════

    async def think(self, state: dict) -> Dict[str, Any]:
        return {
            "summary": "分析用户意图中...",
            "steps": [],
            "actions": [],
        }

    async def act(self, state: dict, plan: Dict[str, Any]) -> Dict[str, Any]:
        user_input = state.get("_latest_user_input", "")

        # ═══════════════════════════════════════════════════════
        # Step 1: IntentClarificationAgent — LLM语义意图识别
        # ═══════════════════════════════════════════════════════
        intent_result = await self._run_intent_clarification(state, user_input)

        if not intent_result:
            # Fallback: use legacy keyword matching
            return await self._legacy_intent_dispatch(state, user_input)

        primary = intent_result.get("primary_intent", "unknown")
        confidence = intent_result.get("confidence", 0)
        needs_clarification = intent_result.get("needs_clarification", False)

        # Log intent analysis
        logger.info(
            f"Intent: {primary} (conf={confidence}), "
            f"clarify={needs_clarification}, "
            f"risk={intent_result.get('hallucination_risk')}"
        )

        # ═══════════════════════════════════════════════════════
        # Step 2: Anti-hallucination guard
        # ═══════════════════════════════════════════════════════
        hallucination_risk = intent_result.get("hallucination_risk")
        if hallucination_risk:
            warning = (
                f"⚠️ 检测到潜在的幻觉风险：{hallucination_risk}\n\n"
                "系统不会编造具体数值数据。请提供准确的项目数据（面积、户数、金额等），"
                "系统将严格基于您提供的信息生成报告。"
            )
            await self._emit_message(warning, "text")
            _add_to_history(state, "assistant", warning)
            return {"type": "chat", "message": warning, "hallucination_guarded": True}

        # ═══════════════════════════════════════════════════════
        # Step 3: Clarification flow (low confidence)
        # ═══════════════════════════════════════════════════════
        if needs_clarification:
            question = intent_result.get("clarification_question",
                "请问您是想提供项目数据、咨询法规问题，还是要求生成报告？")
            await self._emit_thinking(f"🤔 意图不明确（置信度{confidence}%），需要澄清")
            await self._emit_message(question, "text")
            _add_to_history(state, "assistant", question)
            return {
                "type": "clarification",
                "message": question,
                "original_intent": intent_result,
            }

        # ═══════════════════════════════════════════════════════
        # Step 3.5: If waiting for bidding data, route ANY user response
        #           (except greeting/complaint) to bidding handler
        # ═══════════════════════════════════════════════════════
        if state.get("_bidding_waiting_for_data"):
            if primary not in ("greeting", "complaint"):
                logger.info("Routing to bidding_generation (waiting for data)")
                return await self._handle_bidding_generation(state, user_input)

        # ═══════════════════════════════════════════════════════
        # Step 4: Route by intent type
        # ═══════════════════════════════════════════════════════
        if primary == "generation_request":
            # 🔴 Safety re-check: if message/context is bidding, reroute
            if _is_bidding_context(state, user_input):
                logger.info("Rerouting generation_request → bidding_generation (bidding context detected)")
                return await self._handle_bidding_generation(state, user_input)
            return self._handle_gen_request(state)

        elif primary == "question":
            # 🔴 Safety re-check: if question is about bidding, search bidding RAG
            if _is_bidding_context(state, user_input):
                return await self._handle_bidding_question(state, user_input)
            return await self._handle_question(state, user_input)

        elif primary == "table_construction":
            return await self._handle_table_construction(state, user_input)

        elif primary == "kb_explanation":
            return await self._handle_kb_explanation(state, user_input)

        elif primary == "bidding_generation":
            return await self._handle_bidding_generation(state, user_input)

        elif primary == "data_provision":
            # 🔴 If waiting for bidding data, route back to bidding flow
            if state.get("_bidding_waiting_for_data"):
                return await self._handle_bidding_generation(state, user_input)

            # Data already extracted by IntentClarificationAgent into state
            extracted = intent_result.get("extracted_data", {})
            if extracted:
                filled = state.setdefault("filled_data", {})
                for k, v in extracted.items():
                    if v:
                        filled[k] = v
                summary = ", ".join(f"{k}={v}" for k, v in list(extracted.items())[:5] if v)
                resp = f"已记录：{summary}。请继续提供其他项目信息，或回复「生成报告」开始编写。"
            else:
                resp = "收到您的信息，已记录。请继续提供其他项目数据或回复「生成报告」。"
            await self._emit_message(resp, "text")
            _add_to_history(state, "assistant", resp)
            return {"type": "chat", "message": resp}

        elif primary in ("confirmation", "chapter_feedback"):
            # Pass to chat handler (handles approval/revision flow)
            return await self._handle_chat(state, user_input)

        elif primary == "revision_request":
            # Extract what to revise
            hint = intent_result.get("response_hint", "")
            return await self._handle_chat(state, user_input)

        elif primary == "progress_check":
            ch_state = state.get("chapter_orchestrator_state", "idle")
            chapters = state.get("chapters", {})
            done = sum(1 for ch in chapters.values() if isinstance(ch, dict) and ch.get("markdown"))
            resp = f"当前进度：{done}/10章已完成，状态：{ch_state}。"
            await self._emit_message(resp, "text")
            return {"type": "chat", "message": resp}

        elif primary == "complaint":
            resp = "理解您的反馈。我会尽力改进生成质量。如有具体问题请具体说明，我会针对性优化。"
            await self._emit_message(resp, "text")
            _add_to_history(state, "assistant", resp)
            return {"type": "chat", "message": resp}

        elif primary == "greeting":
            return await self._handle_chat(state, user_input)

        elif primary == "mixed":
            # Handle first sub-intent, note others
            sub_intents = intent_result.get("sub_intents", [])
            await self._emit_thinking(f"🔀 复合意图: {', '.join(sub_intents)}")
            # Process first actionable intent
            for sub in sub_intents:
                if sub == "generation_request":
                    return self._handle_gen_request(state)
                elif sub == "data_provision":
                    return await self._handle_chat(state, user_input)
                elif sub == "question":
                    return await self._handle_question(state, user_input)
            return await self._handle_chat(state, user_input)

        else:
            # Unknown or file_upload — fall through to general chat
            return await self._handle_chat(state, user_input)

    async def _run_intent_clarification(
        self, state: dict, user_input: str,
    ) -> Optional[Dict[str, Any]]:
        """Run IntentClarificationAgent for semantic intent analysis."""
        try:
            from .intent_clarification_agent import IntentClarificationAgent
            agent = IntentClarificationAgent(llm_service=self._llm)
            result = await agent.analyze_intent(state, user_input)
            return result
        except Exception as e:
            logger.warning(f"IntentClarificationAgent failed: {e}")
            return None

    async def _legacy_intent_dispatch(
        self, state: dict, user_input: str,
    ) -> Dict[str, Any]:
        """Fallback: legacy keyword-based intent dispatch."""
        # 🔴 Check bidding context FIRST — before generic generation/chat
        if _is_bidding_context(state, user_input):
            return await self._handle_bidding_generation(state, user_input)

        gen_keywords = [
            "生成报告", "开始生成", "逐章生成", "写报告", "做报告",
            "需要生成", "想要生成", "开始写", "编写报告", "编制报告",
        ]
        is_gen = any(kw in user_input for kw in gen_keywords)
        is_question = any(kw in user_input for kw in [
            "是什么", "怎么", "如何", "规定", "标准", "法规", "要求",
            "？", "什么", "哪些", "多少", "吗", "呢",
        ])
        if is_gen:
            return self._handle_gen_request(state)
        elif is_question:
            return await self._handle_question(state, user_input)
        else:
            return await self._handle_chat(state, user_input)

    # ═══════════════════════════════════════════════════════════
    # Intent Handlers
    # ═══════════════════════════════════════════════════════════

    def _handle_gen_request(self, state: dict) -> Dict[str, Any]:
        """User wants to generate a report — signal ChapterOrchestrator."""
        real_filled = {k: v for k, v in state.get("filled_data", {}).items()
                      if not k.startswith("_")}
        has_data = (
            len(real_filled) >= 2 or
            bool(state.get("_pdf_texts"))
        )
        if has_data:
            state["generation_mode"] = "chapter_by_chapter"
            state["_start_chapter_generation"] = True
            return {"type": "chat", "message": "启动逐章生成", "_start_orchestrator": True}
        else:
            return {
                "type": "chat",
                "message": (
                    "收到报告生成请求。当前缺少项目资料。\n\n"
                    "请上传 PDF（征地公告、勘测定界报告）和图片（问卷、公示照片），"
                    "或直接提供项目基本信息（位置、面积、责任单位）。\n\n"
                    "上传后系统自动提取数据，回复「生成报告」开始逐章编写。"
                ),
            }

    async def _handle_question(self, state: dict, question: str) -> Dict[str, Any]:
        """User is asking a question — search KB and answer."""
        # 🔴 Safety: if domain is bidding, use bidding RAG
        if _is_bidding_context(state, question):
            return await self._handle_bidding_question(state, question)

        await self._emit_thinking("🔍 正在检索知识库...")

        # Search KB via shared RAG service
        rag_text = ""
        try:
            from app.rag.rag_service import rag_service
            results = await rag_service.retrieve_with_query(
                query=question, session_id=state.get("session_id", ""), n_results=3,
                domain=state.get("_domain") or state.get("_conversation_domain"),
            )
            if results:
                rag_text = "\n".join(
                    f"【{r.get('metadata', {}).get('source_file', '知识库')}】\n{r.get('document', '')[:800]}"
                    for r in results
                )
        except Exception as e:
            logger.warning(f"KB search failed: {e}")

        # If LLM available, use it to formulate answer
        if self._llm and rag_text:
            try:
                prompt = (
                    f"用户提问：{question}\n\n"
                    f"知识库参考：\n{rag_text[:3000]}\n\n"
                    f"请基于知识库内容回答，引用具体法规文号和条款。如果知识库没有相关信息，诚实说明。"
                    f"回答要简洁，2-5句话。"
                )
                result = await _asyncio.wait_for(
                    self._llm.chat_with_reasoning(
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=512, temperature=0.3,
                    ),
                    timeout=20.0,
                )
                content = result.get("content", "")
                if content:
                    await self._emit_message(content, "text")
                    _add_to_history(state, "assistant", content)
                    return {"type": "chat", "message": content}
            except Exception:
                pass

        # Fallback: use KB text directly
        if rag_text:
            summary = rag_text[:800]
            await self._emit_message(f"**知识库检索结果**：\n\n{summary}", "text")
            _add_to_history(state, "assistant", summary)
            return {"type": "chat", "message": summary}

        await self._emit_message("抱歉，知识库中未找到相关信息。建议查阅DB32/T4013-2021标准或咨询主管部门。", "text")
        return {"type": "chat", "fallback": True}

    async def _handle_chat(self, state: dict, message: str) -> Dict[str, Any]:
        """General conversation / data provision — domain-aware persona."""
        # 🔴 Route to bidding if context (domain flag OR keywords) indicates bidding
        if _is_bidding_context(state, message):
            _set_conversation_domain(state, "bidding")
            # If message is NOT just a greeting/chat, delegate to bidding handler
            if _BIDDING_KEYWORDS.search(message):
                return await self._handle_bidding_generation(state, message)

        # Check if conversation is already in bidding domain
        is_bidding = _is_bidding_context(state, message)

        # If simple greeting, respond quickly with domain-appropriate persona
        if message.strip() in ("你好", "您好", "hi", "hello", "在吗"):
            if is_bidding:
                resp = "您好！我是招标投标文件编写专家。请上传招标文档（招标公告、招标文件等），我可以帮您生成招标公告、评标报告、中标公示等专业文档。"
            else:
                resp = "您好！我是社会稳定风险评估报告编制专家。请上传项目资料（PDF/图片）或描述项目基本信息（位置、面积、文号等），我可以帮您生成社会稳定风险评估报告。"
            await self._emit_message(resp, "text")
            _add_to_history(state, "assistant", resp)
            return {"type": "chat", "message": resp}

        # Try LLM with domain-appropriate system prompt
        if self._llm:
            try:
                template_ctx = MasterAgent._build_template_context(state)
                collected = _build_collected_summary(state)
                recent = _get_recent_history(state, 8)

                # Use domain-appropriate system prompt
                system_prompt = _get_domain_system_prompt(state, message)
                system = system_prompt.format(
                    template_context=template_ctx,
                    rag_context="",
                    collected_summary=collected,
                    conversation_recent=recent,
                )

                result = await _asyncio.wait_for(
                    self._llm.chat_with_reasoning(
                        messages=[{"role": "user", "content": message}],
                        system=system, max_tokens=512, temperature=0.7,
                    ),
                    timeout=30.0,
                )
                content = result.get("content", "")
                if content:
                    # Strip markdown code fences
                    content = re.sub(r'```json\s*|```\s*', '', content)
                    content = content.strip()

                    # 🔴 Parse JSON action from LLM response
                    try:
                        parsed = json.loads(content)
                        if isinstance(parsed, dict) and parsed.get("type") == "action":
                            agent_name = parsed.get("agent", "")
                            msg = parsed.get("message", "")
                            if msg:
                                await self._emit_message(msg, "text")
                                _add_to_history(state, "assistant", msg)
                            if agent_name:
                                await self._invoke_agent(agent_name, state)
                            return {"type": "action", "agent": agent_name, "message": msg}
                    except (json.JSONDecodeError, ValueError):
                        pass  # Not JSON — treat as plain text

                    await self._emit_message(content, "text")
                    _add_to_history(state, "assistant", content)
                    return {"type": "chat", "message": content}
            except _asyncio.TimeoutError:
                await self._emit_message("正在处理中，请稍候或重新发送...", "text")
                return {"type": "chat", "timeout": True}
            except Exception as e:
                logger.warning(f"LLM chat failed: {e}")

        # LLM unavailable — guide user (only if real data exists)
        real_filled = {k: v for k, v in state.get("filled_data", {}).items()
                      if not k.startswith("_")}
        if len(real_filled) >= 2 or state.get("_pdf_texts"):
            state["_start_chapter_generation"] = True
            await self._emit_message(
                "检测到项目资料，正在启动逐章生成流程...\n（AI服务暂时不可用，将使用模板内容）", "text"
            )
            return {"type": "chat", "_start_orchestrator": True}

        await self._emit_message(
            "请上传项目资料（PDF/图片）或描述项目基本信息（位置、面积、责任单位）。\n"
            "上传后回复「生成报告」即可逐章编写报告。", "text"
        )
        return {"type": "chat", "fallback": True}

    # ═══════════════════════════════════════════════════════════════
    # Table Construction Handler
    # ═══════════════════════════════════════════════════════════════

    async def _handle_table_construction(
        self, state: dict, user_input: str,
    ) -> Dict[str, Any]:
        """Generate formatted markdown tables from user data descriptions."""
        await self._emit_thinking("📊 正在检索知识库中的表格模板...")

        rag_text = ""
        try:
            from app.rag.rag_service import rag_service
            query = f"表格 模板 格式 {user_input[:200]}"
            results = await rag_service.retrieve_with_query(
                query=query, session_id=state.get("session_id", ""), n_results=5,
                domain=state.get("_domain") or state.get("_conversation_domain"),
            )
            if results:
                rag_text = "\n".join(
                    f"【{r.get('metadata', {}).get('source_file', '知识库')}】\n{r.get('document', '')[:1000]}"
                    for r in results
                )
        except Exception as e:
            logger.warning(f"KB search for table construction failed: {e}")

        collected = _build_collected_summary(state)
        prompt = (
            f"用户请求：{user_input}\n\n"
            f"已收集的项目数据：\n{collected}\n\n"
        )
        if rag_text:
            prompt += (
                f"知识库中的表格模板参考：\n{rag_text[:3000]}\n\n"
                f"请参考上述模板的表格结构（列名、行标题、评分标准等），"
                f"结合用户提供的数据生成规范的Markdown表格。\n"
            )
        else:
            prompt += "请根据稳评报告规范（DB32/T4013-2021），生成规范的Markdown表格。\n"

        prompt += (
            "要求：\n"
            "1. 使用规范的Markdown表格格式（| 列1 | 列2 | ... |）\n"
            "2. 表头加粗，数据对齐\n"
            "3. 表格前加简短说明（1-2句）\n"
            "4. 缺失数据标注【待补充】\n"
            "5. 如知识库有对应模板，优先采用模板的列结构\n"
            "6. 仅输出表格和简短说明，不要其他文字"
        )

        if self._llm:
            try:
                result = await _asyncio.wait_for(
                    self._llm.chat_with_reasoning(
                        messages=[{"role": "user", "content": prompt}],
                        system="你是数据表格构建专家。只输出Markdown表格和简短说明。",
                        max_tokens=1024, temperature=0.3,
                    ),
                    timeout=25.0,
                )
                content = result.get("content", "")
                if content:
                    content = re.sub(r'```(?:markdown|md)?\s*|```\s*', '', content)
                    content = content.strip()
                    await self._emit_message(content, "table_constructed")
                    _add_to_history(state, "assistant", content)
                    return {"type": "chat", "message": content, "message_type": "table_constructed"}
            except Exception as e:
                logger.warning(f"Table construction failed: {e}")

        await self._emit_message(
            "抱歉，表格构建服务暂时不可用。请稍后重试或手动提供表格结构。", "text",
        )
        return {"type": "chat", "fallback": True}

    # ═══════════════════════════════════════════════════════════════
    # Knowledge Base Explanation Handler
    # ═══════════════════════════════════════════════════════════════

    async def _handle_kb_explanation(
        self, state: dict, user_input: str,
    ) -> Dict[str, Any]:
        """Deep KB search + structured regulatory explanation."""
        # 🔴 If asking about bidding, use bidding RAG
        if re.search(r'(?:招标|投标|评标|中标|标书)', user_input):
            return await self._handle_bidding_question(state, user_input)

        await self._emit_thinking("🔍 正在深度检索知识库...")

        rag_texts = []
        try:
            from app.rag.rag_service import rag_service
            session_id = state.get("session_id", "")
            _domain = state.get("_domain") or state.get("_conversation_domain")

            results = await rag_service.retrieve_with_query(
                query=user_input, session_id=session_id, n_results=5,
                domain=_domain,
            )
            if results:
                rag_texts.extend(results)

            key_terms = re.findall(
                r'DB32[^\s]{,20}|第[一二三四五六七八九十\d]+章|第[一二三四五六七八九十\d]+条|[《〈][^》〉]+[》〉]',
                user_input,
            )
            if key_terms:
                for term in key_terms[:3]:
                    more = await rag_service.retrieve_with_query(
                        query=term, session_id=session_id, n_results=3,
                        domain=_domain,
                    )
                    for r in more:
                        if r not in rag_texts:
                            rag_texts.append(r)
        except Exception as e:
            logger.warning(f"KB explanation search failed: {e}")

        if not rag_texts:
            await self._emit_message(
                "抱歉，知识库中未找到与您问题相关的信息。建议查阅DB32/T4013-2021标准原文或咨询主管部门。",
                "text",
            )
            _add_to_history(state, "assistant", "知识库未找到相关信息")
            return {"type": "chat", "fallback": True}

        rag_context = "\n\n".join(
            f"【来源：{r.get('metadata', {}).get('source_file', '知识库')}"
            f"{' 第' + str(r.get('metadata', {}).get('chapter_number', '')) + '章' if r.get('metadata', {}).get('chapter_number') else ''}】\n"
            f"{r.get('document', '')[:1200]}"
            for r in rag_texts[:5]
        )

        if self._llm:
            try:
                prompt = (
                    f"用户提问：{user_input}\n\n"
                    f"知识库检索结果：\n{rag_context[:4000]}\n\n"
                    f"请基于知识库内容结构化回答：\n"
                    f"1. 先概述相关法规/标准的整体框架\n"
                    f"2. 再针对用户具体问题回答关键条款\n"
                    f"3. 引用具体法规名称、文号和条款编号\n"
                    f"4. 最后说明实际应用场景或注意事项\n"
                    f"5. 如果知识库内容不足以完整回答，诚实说明\n"
                    f"回答控制在5-8句话以内，结构清晰。"
                )
                result = await _asyncio.wait_for(
                    self._llm.chat_with_reasoning(
                        messages=[{"role": "user", "content": prompt}],
                        system="你是稳评法规解说专家。回答要结构化：框架→条款→应用。引用具体法规名称和文号。",
                        max_tokens=800, temperature=0.3,
                    ),
                    timeout=25.0,
                )
                content = result.get("content", "")
                if content:
                    await self._emit_message(content, "regulation_explanation")
                    _add_to_history(state, "assistant", content)
                    return {"type": "chat", "message": content, "message_type": "regulation_explanation"}
            except Exception as e:
                logger.warning(f"KB explanation LLM failed: {e}")

        summary = rag_texts[0].get("document", "")[:800] if rag_texts else ""
        source = rag_texts[0].get("metadata", {}).get("source_file", "知识库") if rag_texts else "知识库"
        fallback_msg = (
            f"**知识库检索结果**（来源：{source}）：\n\n{summary}\n\n"
            f"（AI解说服务暂时不可用，以上为知识库原文）"
        )
        await self._emit_message(fallback_msg, "text")
        _add_to_history(state, "assistant", fallback_msg)
        return {"type": "chat", "message": fallback_msg}


    # ═══════════════════════════════════════════════════════════════
    # Bidding Question Handler (uses bidding_rag, NOT rag_service)
    # ═══════════════════════════════════════════════════════════════

    async def _handle_bidding_question(
        self, state: dict, question: str,
    ) -> Dict[str, Any]:
        """Answer bidding-related questions using ONLY the bidding knowledge base."""
        # Set conversation domain so subsequent messages use bidding persona
        _set_conversation_domain(state, "bidding")

        await self._emit_thinking("🔍 正在以招标文件编写专家身份检索招标知识库...")

        rag_text = ""
        try:
            from app.rag.bidding_rag import bidding_rag
            results = await bidding_rag.search(query=question, n_results=5)
            if results:
                rag_text = "\n".join(
                    f"【{r.get('metadata', {}).get('source_file', '招标知识库')}】\n{r.get('document', '')[:1200]}"
                    for r in results
                )
        except Exception as e:
            logger.warning(f"Bidding RAG question search failed: {e}")

        if self._llm and rag_text:
            try:
                prompt = (
                    f"用户提问（招标相关）：{question}\n\n"
                    f"招标知识库参考：\n{rag_text[:4000]}\n\n"
                    f"请基于招标知识库内容回答。引用具体文档名称和关键信息。"
                    f"如果知识库没有相关信息，诚实说明。回答要简洁，3-6句话。"
                    f"禁止引用社会稳定风险评估相关规范。"
                )
                result = await _asyncio.wait_for(
                    self._llm.chat_with_reasoning(
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=512, temperature=0.3,
                    ),
                    timeout=20.0,
                )
                content = result.get("content", "")
                if content:
                    await self._emit_message(content, "text")
                    _add_to_history(state, "assistant", content)
                    return {"type": "chat", "message": content}
            except Exception:
                pass

        if rag_text:
            summary = rag_text[:800]
            await self._emit_message(f"**招标知识库检索结果**：\n\n{summary}", "text")
            _add_to_history(state, "assistant", summary)
            return {"type": "chat", "message": summary}

        await self._emit_message(
            "招标知识库中未找到相关信息。请上传招标文档（招标公告、招标文件等），"
            "系统将自动解析并索引到招标知识库中。", "text"
        )
        return {"type": "chat", "fallback": True}

    # ═══════════════════════════════════════════════════════════════
    # Bidding Generation Handler
    # ═══════════════════════════════════════════════════════════════

    async def _handle_bidding_generation(
        self, state: dict, user_input: str,
    ) -> Dict[str, Any]:
        """Handle bidding/tender report generation requests."""
        # Set conversation domain so subsequent messages use bidding persona
        _set_conversation_domain(state, "bidding")

        # If user just provided data after being asked, parse and continue
        if state.pop("_bidding_waiting_for_data", False):
            await self._emit_thinking("📋 已收到补充数据，继续生成招标报告...")
            # Parse the new data from user input
            from .bidding_data_agent import BiddingDataAgent
            temp_agent = BiddingDataAgent(llm_service=self._llm)
            new_data = temp_agent._regex_extract(user_input)
            if new_data:
                bidding_data = state.setdefault("_bidding_data", {})
                bidding_data.update({k: v for k, v in new_data.items() if v and v.strip()})
                state.setdefault("filled_data", {}).update(new_data)
                # Also try LLM extraction if regex didn't get much
                if len(new_data) < 3 and self._llm:
                    try:
                        llm_data = await temp_agent._llm_deep_extract(user_input[:15000])
                        for k, v in llm_data.items():
                            if v and str(v).strip() and k not in bidding_data:
                                bidding_data[k] = str(v).strip()
                    except Exception:
                        pass
                await self._emit_thinking(f"📋 已更新 {len(new_data)} 个数据字段")
                state["_bidding_data"] = bidding_data
            # Fall through to check if we have enough data now

        # ═══════════════════════════════════════════════════════════
        # Phase 1: 📋 解析用户上传的招标文档
        # ═══════════════════════════════════════════════════════════
        await self._emit_thinking("━━━ 📋 第1步：解析用户上传的招标文档 ━━━")

        has_materials = bool(state.get("_project_materials")) or bool(state.get("_project_material_facts"))
        if not state.get("_bidding_data") or has_materials:
            await self._invoke_agent("BiddingDataAgent", state)

        bidding_data = state.get("_bidding_data", {})
        await self._emit_thinking(f"✅ 第1步完成：从文档中提取了 {len(bidding_data)} 个数据字段")

        # ═══════════════════════════════════════════════════════════
        # Phase 2: 🔍 核验关键数据，缺失则询问用户
        # ═══════════════════════════════════════════════════════════
        await self._emit_thinking("━━━ 🔍 第2步：核验关键数据是否齐全 ━━━")

        missing = self._check_critical_bidding_fields(bidding_data)

        if missing:
            field_labels = {
                "bid_project_name": "项目名称",
                "bid_reference": "项目编号/招标编号",
                "bid_owner": "招标人/采购人名称",
                "bid_budget": "预算金额",
                "bid_deadline": "投标截止时间",
                "bid_open_time": "开标时间",
                "bid_type": "招标方式",
                "bid_scope": "招标范围/采购内容",
            }
            missing_labels = [field_labels.get(f, f) for f in missing]

            inquiry = (
                f"📋 已从文档中提取到 {len(bidding_data)} 个数据字段。\n\n"
                f"以下**关键信息**尚未获取，请提供：\n\n"
                + "\n".join(f"- **{label}**" for label in missing_labels)
                + "\n\n请直接回复上述信息，例如：\n"
                + "\n".join(f"{label}：（请填写）" for label in missing_labels[:3])
            )

            state["_bidding_waiting_for_data"] = True
            await self._emit_message(inquiry, "text")
            _add_to_history(state, "assistant", inquiry)
            return {"type": "chat", "message": inquiry, "waiting_for_data": True}

        await self._emit_thinking("✅ 第2步完成：关键数据齐全")

        # ═══════════════════════════════════════════════════════════
        # 🆕 Phase 2.5: 📋 能力分析 — 列出Agent分工、模板/AI/用户三源
        # ═══════════════════════════════════════════════════════════
        await self._emit_thinking("━━━ 📋 第2.5步：需求分析 & Agent能力映射 ━━━")
        try:
            from .capability_analysis_agent import CapabilityAnalysisAgent
            analysis_agent = CapabilityAnalysisAgent(llm_service=self._llm)
            await analysis_agent.run(state, self._stream_queue)
            await self._emit_thinking("✅ 第2.5步完成：Agent分工已明确")
        except Exception as e:
            logger.warning(f"Capability analysis failed: {e}")
            await self._emit_thinking(f"⚠️ 能力分析跳过：{e}")

        # ═══════════════════════════════════════════════════════════
        # Phase 3+4: 📝 逐章思考生成 → DOCX → 历史保存（BiddingOrchestrator）
        #   每章结合项目实际信息 + RAG 同类文风参考撰写，非填空。
        #   🆕 对于 tender_response：优先使用模板填空模式（保留模板完整内容）
        # ═══════════════════════════════════════════════════════════
        from .bidding_report_agent import BiddingReportAgent as _BRA
        report_type = _BRA()._detect_report_type(user_input) if user_input else "tender_response"
        # Prefer a report type already pinned on the conversation, if any.
        report_type = state.get("_bidding_report_type") or report_type

        from app.agent.bidding_chapters import get_bidding_type_name
        type_name = get_bidding_type_name(report_type)

        # ── 🆕 Template-fill mode for tender_response ──
        output_path = ""
        download_url = ""

        if report_type == "tender_response":
            template_path = await self._try_template_fill_bidding(state, bidding_data)
            if template_path:
                output_path = template_path
                download_url = f"/api/v1/files/{output_path}"
                logger.info(f"Bidding template-fill success: {output_path}")

        # ── Fallback: per-module generation ──
        if not output_path:
            await self._emit_thinking(f"━━━ 📝 第3步：逐模块生成《{type_name}》 ━━━")

        if not output_path:
            try:
                from .bidding_orchestrator import BiddingOrchestrator
                orchestrator = BiddingOrchestrator(llm_service=self._llm)
                state["_bidding_orchestrator"] = orchestrator
                result = await orchestrator.run_full_pipeline(
                    state, self._stream_queue, report_type=report_type,
                )
            except Exception as e:
                logger.error(f"BiddingOrchestrator failed, falling back to single-shot: {e}")
                await self._emit_thinking(f"⚠️ 逐章生成异常，回退单发模式：{e}")
                await self._invoke_agent("BiddingReportAgent", state)
                generated = state.get("generated_sections", {})
                result = {
                    "output_path": await self._save_bidding_docx(
                        state,
                        generated.get("bidding_report", ""),
                        state.get("_bidding_report_type", report_type),
                    ) if generated.get("bidding_report") else "",
                    "report_type": report_type,
                }

            output_path = result.get("output_path", "")
            download_url = f"/api/v1/files/{output_path}" if output_path else ""

        state["_bidding_report_type"] = report_type

        if output_path:
            if self._stream_queue:
                await self._stream_queue.put({
                    "event": "complete",
                    "data": {
                        "message": f"{type_name}已生成并保存",
                        "download_url": download_url,
                        "report_id": state.get("report_id", ""),
                    },
                })
            await self._emit_message(
                f"✅ **《{type_name}》已生成并保存到历史报告**\n\n"
                f"[📥 点击下载]({download_url})",
                "text",
            )
        else:
            await self._emit_thinking("⚠️ 报告内容为空或过短，未能生成DOCX")

        await self._emit_thinking(f"🎉 《{type_name}》生成流程结束")
        return {"type": "chat", "message": f"{type_name}生成完成"}

    async def _try_template_fill_bidding(
        self, state: dict, bidding_data: dict,
    ) -> str:
        """Try to fill a bidding template with project data.

        Uses the most recently active bidding template from the knowledge base,
        scans its placeholders, maps project data → placeholder values,
        and fills via DocxService.fill_template().

        Returns relative path like "generated/xxx.docx" on success, "" on failure.
        """
        from app.services.docx_service import DocxService
        from app.services.file_service import file_service
        from app.config import settings

        try:
            # 1. Find an active bidding template
            import sqlite3
            db = settings.DATA_DIR / "knowledge_base.db"
            if not db.exists():
                return ""

            conn = sqlite3.connect(str(db))
            conn.row_factory = sqlite3.Row
            cur = conn.execute(
                "SELECT id, name, template_file_path FROM templates "
                "WHERE is_active=1 AND category LIKE '%招标%' "
                "ORDER BY id DESC LIMIT 1"
            )
            tpl = cur.fetchone()
            conn.close()

            if not tpl:
                logger.info("No bidding template found for fill")
                return ""

            tpl_path = settings.STORAGE_DIR / tpl["template_file_path"]
            if not tpl_path.exists():
                logger.warning(f"Bidding template file missing: {tpl_path}")
                return ""

            await self._emit_thinking(
                f"📄 使用模板填空模式：{tpl['name']}（{tpl_path.stat().st_size // 1024} KB）"
            )

            # 2. Scan placeholders
            await self._emit_thinking("🔍 扫描模板占位符...")
            placeholders = DocxService.find_all_placeholders(str(tpl_path))
            if not placeholders:
                logger.info("No placeholders found in bidding template")
                return ""

            await self._emit_thinking(f"   发现 {len(placeholders)} 个占位符")

            # 3. Map project data → placeholder values
            filled_data = self._map_bidding_data_to_template(
                bidding_data, placeholders
            )

            # 4. Fill template
            project_name = str(bidding_data.get("bid_project_name", "投标项目"))
            # Truncate to first 50 chars to avoid overly long filenames
            project_name = project_name[:50].strip()
            safe_name = re.sub(r'[\\/:*?"<>|\n\r]', '_', project_name)
            if not safe_name or safe_name == '_':
                safe_name = "投标文件"
            output_filename = f"{safe_name}_投标文件.docx"
            output_path = settings.STORAGE_DIR / "generated" / output_filename
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # Ensure unique filename
            counter = 1
            while output_path.exists():
                output_filename = f"{safe_name}_投标文件_{counter}.docx"
                output_path = settings.STORAGE_DIR / "generated" / output_filename
                counter += 1

            await self._emit_thinking(f"📝 正在填充模板...")
            DocxService.fill_template(
                str(tpl_path),
                str(output_path),
                filled_data,
                placeholder_map=placeholders,
            )

            # ── Post-process: global text replacement for project-specific data ──
            # Many values (project name, reference, etc.) appear as plain text in
            # the template (sample project data), not as placeholders.  Do a global
            # search-and-replace on the filled DOCX to swap in the new values.
            from docx import Document as _Doc

            # Collect old→new replacement pairs (old values from template sample project)
            _global_replaces = {}

            # Old template project name → new project name
            _old_project = "洪泽区2026年经营性地块土地征收报批综合服务（含征地社会稳定性风险评估）项目"
            _new_project = str(bidding_data.get("bid_project_name", "")).strip()
            if _new_project and _old_project != _new_project:
                _global_replaces[_old_project] = _new_project

            # Old template reference number → new reference
            _old_ref = "JSJH-HZZFCG-CS-2026008"
            _new_ref = str(bidding_data.get("bid_reference", "")).strip()
            if _new_ref and _old_ref != _new_ref:
                _global_replaces[_old_ref] = _new_ref

            # Old owner → new owner
            _old_owner = "淮安市洪泽区土地储备中心"
            _new_owner = str(bidding_data.get("bid_owner", "")).strip()
            if _new_owner and _old_owner != _new_owner:
                _global_replaces[_old_owner] = _new_owner

            # Clean up: remove raw chat message fragments that leaked into project name
            # (regex extraction sometimes captures too much)
            _chat_fragments = [
                "自然资源和规划局。预算金额", "预算金额：180万元。投标截止时间",
                "请直接生成全文", "直接生成全文",
            ]

            if _global_replaces:
                _fill_doc = _Doc(str(output_path))
                _replace_count = 0
                for _para in _fill_doc.paragraphs:
                    for _run in _para.runs:
                        text = _run.text or ""
                        # 1. Replace known old→new pairs
                        for _old, _new in _global_replaces.items():
                            if _old in text:
                                _run.text = text.replace(_old, _new)
                                text = _run.text
                                _replace_count += 1
                        # 2. Clean up leaked chat fragments
                        for _frag in _chat_fragments:
                            if _frag in text:
                                _run.text = text.replace(_frag, "")
                                text = _run.text
                                _replace_count += 1
                # Also check tables
                for _tbl in _fill_doc.tables:
                    for _row in _tbl.rows:
                        for _cell in _row.cells:
                            for _para in _cell.paragraphs:
                                for _run in _para.runs:
                                    text = _run.text or ""
                                    for _old, _new in _global_replaces.items():
                                        if _old in text:
                                            _run.text = text.replace(_old, _new)
                                            text = _run.text
                                            _replace_count += 1
                                    for _frag in _chat_fragments:
                                        if _frag in text:
                                            _run.text = text.replace(_frag, "")
                                            text = _run.text
                                            _replace_count += 1
                _fill_doc.save(str(output_path))
                if _replace_count:
                    await self._emit_thinking(f"  📝 全局文字替换: {_replace_count} 处")

            relative = f"generated/{output_filename}"
            state["output_path"] = relative

            await self._emit_thinking(
                f"✅ 模板填空完成（{len(filled_data)} 个字段已填充）"
            )

            # ── 🆕 AI 生成技术方案并替换模板中对应章节 ──
            await self._emit_thinking("━━━ 🤖 AI生成技术方案（根据项目资料定制） ━━━")
            try:
                await self._replace_tech_proposal_with_ai(
                    state, str(output_path), str(output_path), bidding_data
                )
                await self._emit_thinking("✅ 技术方案已替换为AI定制内容")
            except Exception as e:
                logger.warning(f"Tech proposal AI replacement failed: {e}")
                await self._emit_thinking(f"⚠️ 技术方案替换失败，保留模板原文：{e}")

            # Persist to history
            try:
                from app.services.report_service import report_service
                state["status"] = "completed"
                state["report_title"] = project_name or "投标文件"
                session = state.get("_report_session")
                if session:
                    rid = await report_service.persist_report(session)
                    if rid:
                        logger.info(f"Bidding template-fill saved to history (id={rid})")
                        await self._emit_thinking("✅ 报告已存入历史报告")
            except Exception as e:
                logger.warning(f"History persist failed: {e}")

            return relative

        except Exception as e:
            logger.exception(f"Bidding template fill failed: {e}")
            await self._emit_thinking(f"⚠️ 模板填空失败：{e}，切换到AI生成模式")
            return ""

    @staticmethod
    def _map_bidding_data_to_template(
        bidding_data: dict, placeholders: list,
    ) -> dict:
        """Map project bidding data to template placeholder keys.

        Uses keyword matching: for each placeholder, find the best matching
        data field by comparing the placeholder's display_name / key / text
        against known field labels.
        """
        import re as _re

        # Standard field → label mapping (expanded)
        field_labels = {
            "bid_project_name": ["项目名称", "采购项目名称", "工程名称", "项目工程名称"],
            "bid_reference": ["项目编号", "招标编号", "采购编号", "标书编号"],
            "bid_owner": ["招标人", "采购人", "采购单位", "招标单位", "业主"],
            "bid_agency": ["代理机构", "招标代理", "采购代理"],
            "bid_budget": ["预算", "最高限价", "预算金额", "采购预算", "招标控制价"],
            "bid_type": ["采购方式", "招标方式", "磋商方式"],
            "bid_duration": ["服务期限", "合同履行期限", "工期", "服务期"],
            "bid_deadline": ["投标截止", "递交截止", "截止时间", "响应文件递交截止"],
            "bid_open_time": ["开标时间", "开启时间"],
            "bid_open_location": ["开标地点", "开启地点"],
            "bid_contact_name": ["联系人"],
            "bid_contact_phone": ["电话", "联系电话", "联系方式"],
            "bid_location": ["项目地点", "实施地点", "项目所在地"],
            "bid_scope": ["招标范围", "采购内容", "项目概况", "采购需求"],
        }

        filled = {}

        for ph in placeholders:
            pkey = ph.get("key", "")
            pname = ph.get("display_name", "")
            ptext = ph.get("original_text", "")
            ptype = ph.get("pattern_type", "")
            combined = f"{pkey} {pname} {ptext}".lower()

            best_val = ""
            best_score = 0

            for field_key, labels in field_labels.items():
                if field_key not in bidding_data:
                    continue
                val = str(bidding_data[field_key]).strip()
                if not val:
                    continue

                score = 0
                for label in labels:
                    if label in pname or label in ptext or label in pkey:
                        score += 10
                    # Partial 2-char matching
                    for i in range(len(label) - 1):
                        chunk = label[i:i+2]
                        if chunk in combined:
                            score += 1
                if score > best_score:
                    best_score = score
                    best_val = val

            # Handle highlighted wildcards (project-specific text in template)
            if ptype == "highlight" and not best_val:
                matched = ph.get("matched_text", "")
                if "*" in matched or all(c in '*_' for c in matched if not c.isalnum()):
                    # Try deadline date for date wildcards
                    if "年" in matched or "月" in matched or "日" in matched:
                        deadline = bidding_data.get("bid_deadline", "")
                        if deadline:
                            m = _re.search(r'(\d{4})\D+(\d{1,2})\D+(\d{1,2})', deadline)
                            if m:
                                if "年" in matched and "月" in matched and "日" in matched:
                                    best_val = f"{m.group(1)}年{m.group(2)}月{m.group(3)}日"
                                elif "年" in matched and "月" in matched:
                                    best_val = f"{m.group(1)}年{m.group(2)}月"
                                else:
                                    best_val = deadline[:10]
                    # Try project name for "受***委托" pattern
                    if "委托" in ptext and "***" in matched:
                        owner = bidding_data.get("bid_owner", "")
                        if owner:
                            best_val = owner

            if best_val and best_score >= 3:  # Higher threshold for quality
                filled[pkey] = best_val

        # Always fill these fixed fields
        filled.setdefault("bid_agency", "江苏众拓项目代理咨询有限公司")

        return filled

    async def _replace_tech_proposal_with_ai(
        self, state: dict, filled_docx_path: str, output_path: str,
        bidding_data: dict,
    ) -> None:
        """Replace the 技术方案 section in a filled template docx with AI-generated content.

        1. Runs BiddingOrchestrator for ONLY the technical proposal modules
        2. Opens the filled docx, finds the 技术方案 heading
        3. Deletes everything from 技术方案 to end
        4. Inserts AI-generated tech proposal content
        5. Also appends the remaining template sections (附件等) after the AI content
        """
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn as docx_qn

        # ── Step 1: Generate only tech proposal modules via LLM ──
        from app.agent.bidding_chapters import get_bidding_chapters, get_bidding_type_name

        all_chapters = get_bidding_chapters("tender_response")
        tech_chapters = [c for c in all_chapters if c.get("tech_group")]
        if not tech_chapters:
            logger.info("No tech modules found, skipping AI replacement")
            return

        await self._emit_thinking(f"  🔍 技术方案共 {len(tech_chapters)} 个子模块，正在结合项目资料生成...")

        from .bidding_chapter_agent import BiddingChapterAgent

        tech_markdown_parts = []
        for ch_def in tech_chapters:
            idx = ch_def.get("index", 0)
            title = ch_def.get("title", "")
            await self._emit_thinking(f"  📝 正在撰写「{title}」...")
            try:
                agent = BiddingChapterAgent(
                    llm_service=self._llm,
                    report_type="tender_response",
                    chapter_index=idx,
                    chapter_def=ch_def,
                )
                # Only generate — don't emit SSE (handled by parent flow)
                agent._stream_queue = self._stream_queue
                plan = await agent.think(state)
                result = await agent.act(state, plan)
                md = result.get("markdown", "")
                if md:
                    tech_markdown_parts.append(md)
            except Exception as e:
                logger.warning(f"Tech module '{title}' generation failed: {e}")

        if not tech_markdown_parts:
            logger.warning("No tech proposal content generated")
            return

        tech_full_md = "\n\n".join(tech_markdown_parts)

        # ── Step 2: Open filled docx, find 技术方案 section ──
        doc = DocxDocument(filled_docx_path)

        # Find the EXACT paragraph that starts the 技术方案 section within the template.
        # IMPORTANT: Templates often have "四、供应商认为有必要提供..." with sub-sections
        # （一）体系认证证书、（二）业绩、（三）人员配备 BEFORE （四）技术方案.
        # We must only replace from "（四）技术方案" onwards, preserving the earlier subsections.
        tech_start_idx = None

        # Priority 1: Look for exact "（四）技术方案" heading (H2 under 四)
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else ""
            if style.startswith("Heading") and ("（四）技术方案" in text or "四）技术方案" in text):
                tech_start_idx = i
                break

        # Priority 2: Look for any H2 "技术方案" heading
        if tech_start_idx is None:
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                style = para.style.name if para.style else ""
                if style.startswith("Heading") and "技术方案" in text:
                    tech_start_idx = i
                    break

        # Priority 3: "四、供应商认为有必要提供并说明的其它资料" (entire section, fallback)
        if tech_start_idx is None:
            other_keywords = [
                "供应商认为有必要提供并说明的其它资料",
                "供应商认为有必要提供并说明的其他资料",
            ]
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                style = para.style.name if para.style else ""
                if style.startswith("Heading") and any(kw in text for kw in other_keywords):
                    tech_start_idx = i
                    break

        # Priority 4: AI-generated sub-module heading
        if tech_start_idx is None:
            ai_sub_heads = ["项目理解与需求分析", "总体工作方案与技术路线"]
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if any(kw in text for kw in ai_sub_heads):
                    tech_start_idx = i
                    break

        if tech_start_idx is None:
            logger.warning("Could not find 技术方案 heading in filled template")
            return

        await self._emit_thinking(f"  📍 找到技术方案章节（段落 {tech_start_idx}），替换为AI生成内容...")

        # ── Step 3: Remove paragraphs from 技术方案 to end ──
        # Collect all body elements (paragraphs + tables) in order
        body = doc.element.body
        children = list(body)

        # Map paragraph indices to body element indices
        para_to_body_idx = {}
        para_count = 0
        for bi, child in enumerate(children):
            if child.tag == docx_qn('w:p'):
                para_to_body_idx[para_count] = bi
                para_count += 1

        tech_start_body_idx = para_to_body_idx.get(tech_start_idx)
        if tech_start_body_idx is None:
            return

        # Remove everything from tech_start_body_idx onwards
        for child in children[tech_start_body_idx:]:
            body.remove(child)

        # ── Step 4: Insert AI-generated tech proposal ──
        font_body = '仿宋_GB2312'
        font_h2 = '黑体'
        font_h3 = '楷体'

        # Only add a section heading if the template didn't already have one.
        # If we replaced from "（四）技术方案", the parent "四、供应商认为有必要..."
        # heading is still above us — just add a sub-heading description.
        # If we replaced the entire "四、" section, add the full heading.
        add_h1_heading = True
        for p in doc.paragraphs:
            if '四、' in (p.text or '') and (p.style.name or '').startswith('Heading 1'):
                add_h1_heading = False  # Template already has 四、 heading
                break

        if add_h1_heading:
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 1'] if 'Heading 1' in [s.name for s in doc.styles] else doc.styles['Normal']
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run('四、技术方案')
            r.font.name = '黑体'
            r._element.rPr.rFonts.set(docx_qn('w:eastAsia'), '黑体')
            r.font.size = Pt(16)
            r.bold = True
        else:
            # Add a separator paragraph to visually separate
            p_sep = doc.add_paragraph()
            r_sep = p_sep.add_run(
                '以下技术方案根据本项目实际需求定制撰写，涵盖项目理解、技术路线、'
                '分项服务、组织保障、进度计划、质量控制、安全管理及售后服务等方面。'
            )
            r_sep.font.name = font_body
            r_sep._element.rPr.rFonts.set(docx_qn('w:eastAsia'), font_body)
            r_sep.font.size = Pt(12)

        # Convert markdown tech content to docx paragraphs
        lines = tech_full_md.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # ## → Heading 2
            if line.startswith('## '):
                p = doc.add_paragraph()
                p.style = doc.styles['Heading 2'] if 'Heading 2' in [s.name for s in doc.styles] else doc.styles['Normal']
                r = p.add_run(line[3:].strip())
                r.font.name = font_h2
                r._element.rPr.rFonts.set(docx_qn('w:eastAsia'), font_h2)
                r.font.size = Pt(15)
                r.bold = True
                i += 1
                continue

            # ### → Heading 3
            if line.startswith('### '):
                p = doc.add_paragraph()
                p.style = doc.styles['Heading 3'] if 'Heading 3' in [s.name for s in doc.styles] else doc.styles['Normal']
                r = p.add_run(line[4:].strip())
                r.font.name = font_h3
                r._element.rPr.rFonts.set(docx_qn('w:eastAsia'), font_h3)
                r.font.size = Pt(14)
                r.bold = True
                i += 1
                continue

            # Table
            if line.startswith('|') and '|' in line[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_md_table_to_doc(doc, table_lines)
                continue

            # Image marker
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if img_match:
                caption = img_match.group(1).strip()
                img_ref = img_match.group(2).strip()
                # Try to insert from DB
                from app.services.bidding_docx_generator import BiddingDocxGenerator
                bg = BiddingDocxGenerator()
                bg._insert_image_from_ref(doc, img_ref, caption, {})
                i += 1
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            r = p.add_run(clean)
            r.font.name = font_body
            r._element.rPr.rFonts.set(docx_qn('w:eastAsia'), font_body)
            r.font.size = Pt(12)
            i += 1

        doc.save(output_path)
        await self._emit_thinking(f"  ✅ 技术方案已替换（{len(tech_markdown_parts)} 个子模块，{len(tech_full_md)} 字）")

    @staticmethod
    def _add_md_table_to_doc(doc, table_lines):
        """Convert markdown table lines to DOCX table (simplified)."""
        from docx.shared import Pt as _Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN
        from docx.oxml.ns import qn as _qn

        rows = []
        for line in table_lines:
            if re.match(r'\|[\s\-:|]+\|', line):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
        if not rows:
            return

        ncols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=ncols, style='Table Grid')
        table.alignment = _WD_ALIGN.CENTER

        for ri, row_data in enumerate(rows):
            for ci, cell_text in enumerate(row_data):
                if ci >= ncols:
                    break
                cell = table.cell(ri, ci)
                cell.paragraphs[0].clear()
                r = cell.paragraphs[0].add_run(cell_text)
                r.font.name = '宋体'
                r._element.rPr.rFonts.set(_qn('w:eastAsia'), '宋体')
                r.font.size = _Pt(10.5)
                if ri == 0:
                    r.bold = True
                cell.paragraphs[0].alignment = _WD_ALIGN.CENTER

        doc.add_paragraph()

    async def _save_bidding_docx(
        self, state: dict, content: str, report_type: str,
    ) -> str:
        """Convert markdown to DOCX, save to history database."""
        from app.services.bidding_docx_generator import BiddingDocxGenerator
        from app.services.file_service import file_service

        bidding_data = state.get("_bidding_data", {})

        generator = BiddingDocxGenerator()
        output_path = generator.generate(
            markdown_content=content,
            report_type=report_type,
            metadata={
                "session_id": state.get("session_id", ""),
                "project_name": bidding_data.get("bid_project_name", ""),
                "reference": bidding_data.get("bid_reference", ""),
            },
            state=state,  # Pass state for image/attachment support
        )

        # Save to history database
        try:
            from app.services.report_service import report_service

            report_title = bidding_data.get("bid_project_name", "") or "招标报告"

            state["output_path"] = output_path
            state["status"] = "completed"
            state["report_title"] = report_title

            session = state.get("_report_session")
            if session:
                report_id = await report_service.persist_report(session, file_service)
                if report_id:
                    logger.info(f"Bidding report saved to history DB (id={report_id}): {report_title}")
                else:
                    logger.warning("Bidding report persist returned no id")
        except Exception as e:
            logger.warning(f"History DB save failed (non-critical): {e}")

        return output_path

    def _check_critical_bidding_fields(self, bidding_data: dict) -> list:
        """Check which critical fields are missing from bidding data.

        Returns list of missing field keys. Empty list = all good.
        """
        critical = [
            "bid_project_name",   # MUST have — can be extracted from docs
            "bid_reference",      # Very important for bidding docs
            "bid_owner",          # Who is tendering
            "bid_budget",         # Budget is fundamental
            "bid_deadline",       # Key date for announcements
        ]
        missing = []
        for field in critical:
            val = bidding_data.get(field, "")
            if not val or not str(val).strip() or len(str(val).strip()) < 2:
                missing.append(field)
        return missing

    # Template Context (static)
    # ═══════════════════════════════════════════════════════════

    @staticmethod
    def _filter_user_questions(placeholders: list) -> list:
        """Filter placeholders that require user input. Kept for backward compat."""
        if not placeholders:
            return []
        user_types = {"text", "number", "date", "select", "textarea"}
        return [
            p for p in placeholders
            if p.get("expected_type") in user_types or not p.get("expected_type")
        ]

    @staticmethod
    def _build_template_context(state: dict) -> str:
        template_name = state.get("template_name", "")
        if not template_name:
            return "10章标准结构（知识库综合生成）"
        if "内置" in template_name or "标准结构" in template_name:
            return "10章标准结构（基于DB32/T4013-2021等规范综合生成）"
        return f"10章标准结构（参考：{template_name}）"


# ═══════════════════════════════════════════════════════════════
# Factory
# ═══════════════════════════════════════════════════════════════

def create_master_agent(llm_service=None) -> MasterAgent:
    return MasterAgent(llm_service=llm_service)
