"""Phase 2: Document structure analysis.

Extracts heading hierarchy from template and example report,
aligns them, classifies sections, and builds a DocumentStructure
that drives all subsequent phases.

Section-relative positioning is the key innovation:
- Each ImageSlot is anchored to a section, not an absolute paragraph index.
- Each TableDef knows which section it belongs to.
- Content generation is scoped to section boundaries.
"""

import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from docx import Document
from docx.oxml.ns import qn

from app.services.pipeline.pipeline_context import (
    DocumentStructure, SectionDef, TableDef, ImageSlot,
    ParagraphInfo, SectionType, GenerationRecipe,
)


class StructureAnalyzer:
    """Phase 2: Extract and align document structure from template + example."""

    # Patterns for detecting headings in plain text paragraphs (not just styled headings)
    HEADING_PATTERNS = [
        (re.compile(r'^第[一二三四五六七八九十\d]+章'), 1),          # 第X章
        (re.compile(r'^[一二三四五六七八九十]+[、．.]'), 2),          # 一、二、三、
        (re.compile(r'^[（(][一二三四五六七八九十\d]+[）)]'), 3),     # (一)(二)
        (re.compile(r'^\d+\.\d+'), 2),                                 # 1.1
        (re.compile(r'^\d+[、．.]'), 2),                                # 1.
    ]

    @staticmethod
    def _detect_heading_level(text: str) -> int:
        """Detect if text is a heading and return level (0 = not a heading).

        Filters out false positives like list items and short fragments."""
        text_stripped = text.strip()

        # Skip very short text (list items, labels)
        if len(text_stripped) < 5:
            return 0

        # Skip single-digit numbered items that are list entries, not headings
        if re.match(r'^\d{1,2}[、．.\s]', text_stripped) and len(text_stripped) < 30:
            return 0

        for pattern, level in StructureAnalyzer.HEADING_PATTERNS:
            if pattern.match(text_stripped):
                # Additional check: real headings usually have 5+ chars after the marker
                if level == 1 and len(text_stripped) >= 8:
                    return level
                elif level >= 2 and len(text_stripped) >= 6:
                    return level
        return 0

    # ── Main entry point ──

    def analyze(self, recipe: GenerationRecipe) -> DocumentStructure:
        """Run Phase 2: analyze template + example structure.

        Args:
            recipe: GenerationRecipe from Phase 1 (has template_path, example_path).

        Returns:
            DocumentStructure with sections, tables, and image slots.
        """
        template_path = recipe.template_path
        example_path = recipe.example_report_path

        if not template_path or not Path(template_path).exists():
            print("  ⚠️ Template not found, creating minimal structure")
            return DocumentStructure()

        doc_tpl = Document(template_path)
        doc_ex = Document(example_path) if example_path and Path(example_path).exists() else None

        # Step 1: Extract headings from both documents
        tpl_headings = self._extract_headings(doc_tpl)
        ex_headings = self._extract_headings(doc_ex) if doc_ex else []

        # Step 2: Align template and example headings by text
        heading_map = self._align_headings(tpl_headings, ex_headings)

        # Step 3: Build sections from template headings
        sections = self._build_sections(doc_tpl, tpl_headings, heading_map, doc_ex)

        # Step 4: Detect tables and assign to sections
        tables = self._detect_tables(doc_tpl, sections)

        # Step 5: Detect image slots (section-relative)
        image_slots = self._detect_image_slots(doc_tpl, doc_ex, sections, heading_map)

        # Step 6: Classify sections
        self._classify_sections(sections, doc_tpl)

        return DocumentStructure(
            sections=sections,
            tables=tables,
            image_slots=image_slots,
            total_paragraphs=len(doc_tpl.paragraphs),
            total_sections=len(sections),
            template_path=template_path,
            example_path=example_path,
        )

    # ── Heading extraction ──

    def _extract_headings(self, doc: Document) -> List[Tuple[int, str, int]]:
        """Extract headings from a document.

        Returns:
            List of (para_index, text, level).
        """
        headings = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Styled headings
            if style_name.startswith("Heading") or style_name.startswith("heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                headings.append((i, text, level))
                continue

            # Plain Text chapter titles
            if style_name == "Plain Text" and ("第" in text and "章" in text):
                headings.append((i, text, 1))
                continue

            # Pattern-based headings (for documents without proper styles)
            level = self._detect_heading_level(text)
            if level > 0:
                headings.append((i, text, level))

        return headings

    # ── Heading alignment ──

    def _align_headings(
        self,
        tpl_headings: List[Tuple[int, str, int]],
        ex_headings: List[Tuple[int, str, int]],
    ) -> Dict[str, Dict[str, Any]]:
        """Align template headings with example headings by text matching.

        Returns:
            {tpl_heading_text: {"tpl_index": int, "tpl_level": int,
                                "ex_index": int or None, "ex_level": int or None}}
        """
        # Build text → index maps
        ex_hmap = {h[1]: (h[0], h[2]) for h in ex_headings if h[1]}

        result = {}
        for tpl_idx, tpl_text, tpl_level in tpl_headings:
            entry = {
                "tpl_index": tpl_idx,
                "tpl_level": tpl_level,
                "ex_index": None,
                "ex_level": None,
            }
            if tpl_text in ex_hmap:
                entry["ex_index"], entry["ex_level"] = ex_hmap[tpl_text]
            result[tpl_text] = entry

        return result

    # ── Section building ──

    def _build_sections(
        self,
        doc_tpl: Document,
        tpl_headings: List[Tuple[int, str, int]],
        heading_map: Dict[str, Dict[str, Any]],
        doc_ex: Optional[Document],
    ) -> List[SectionDef]:
        """Build SectionDef objects from template headings.

        Each section spans from its heading to the next heading (exclusive).
        """
        sections = []

        for h_idx, (h_pi, h_text, h_level) in enumerate(tpl_headings):
            # Determine end paragraph index (start of next heading, or end of doc)
            if h_idx + 1 < len(tpl_headings):
                end_pi = tpl_headings[h_idx + 1][0] - 1
            else:
                end_pi = len(doc_tpl.paragraphs) - 1

            # Collect paragraph info for this section
            paragraphs = []
            for pi in range(h_pi + 1, end_pi + 1):
                if pi < len(doc_tpl.paragraphs):
                    para = doc_tpl.paragraphs[pi]
                    text = para.text.strip()
                    level = self._detect_heading_level(text)
                    paragraphs.append(ParagraphInfo(
                        para_index=pi,
                        style=para.style.name if para.style else "",
                        text_preview=text[:200],
                        heading_level=level,
                        section_index=len(sections),
                    ))

            # Detect sub-headings
            sub_headings = [
                p.text_preview[:80] for p in paragraphs
                if p.heading_level >= 2 and len(p.text_preview) > 3
            ]

            sections.append(SectionDef(
                index=len(sections),
                title=h_text,
                level=h_level,
                start_para_index=h_pi,
                end_para_index=end_pi,
                paragraphs=paragraphs,
                sub_sections=sub_headings,
            ))

        return sections

    # ── Table detection ──

    def _detect_tables(
        self, doc_tpl: Document, sections: List[SectionDef]
    ) -> List[TableDef]:
        """Detect tables and assign them to sections."""
        tables = []

        for t_idx, table in enumerate(doc_tpl.tables):
            # Get table caption (search nearby paragraphs)
            caption = ""
            section_idx = -1

            # Try to find the table's position in the document
            # Tables in python-docx don't have a paragraph index,
            # so we search for paragraphs that reference the table
            for s_idx, section in enumerate(sections):
                for para in section.paragraphs:
                    if f"表{t_idx+1}" in para.text_preview or f"表 {t_idx+1}" in para.text_preview:
                        caption = para.text_preview
                        section_idx = s_idx
                        break
                if section_idx >= 0:
                    break

            # Build preview
            preview = []
            for row in table.rows[:5]:
                row_data = [cell.text[:50] for cell in row.cells]
                preview.append(row_data)

            # Determine table type
            table_type = self._classify_table(caption, preview)

            # Determine which column to agent-fill (last column by default)
            fill_col = len(table.columns) - 1 if table_type == "copy_n_minus_1" else -1

            tables.append(TableDef(
                table_index=t_idx,
                rows=len(table.rows),
                cols=len(table.columns),
                section_index=section_idx,
                caption=caption,
                table_type=table_type,
                fill_column=fill_col,
                preview=preview,
            ))

        return tables

    def _classify_table(self, caption: str, preview: List[List[str]]) -> str:
        """Classify a table as copy_full, copy_n_minus_1, or agent_generate."""
        caption_lower = caption.lower() if caption else ""

        # Company/team tables → copy_full
        if any(kw in caption for kw in ["营业执照", "人员", "分工", "公司简介", "资质"]):
            return "copy_full"

        # Survey/statistics tables → copy_n_minus_1 (last column = analyzed data)
        if any(kw in caption for kw in ["问卷", "调查", "统计", "汇总"]):
            return "copy_n_minus_1"

        # Scoring/risk tables → copy_n_minus_1 (last column = project-specific score)
        if any(kw in caption for kw in ["评分", "风险", "等级", "量化", "措施"]):
            return "copy_n_minus_1"

        # Generic data tables → copy_n_minus_1
        if any(kw in caption for kw in ["面积", "分类", "汇总"]):
            return "copy_n_minus_1"

        # Otherwise → agent_generate
        return "agent_generate"

    # ── Image slot detection ──

    def _detect_image_slots(
        self,
        doc_tpl: Document,
        doc_ex: Optional[Document],
        sections: List[SectionDef],
        heading_map: Dict[str, Dict[str, Any]],
    ) -> List[ImageSlot]:
        """Detect image placeholders as section-relative ImageSlots.

        Instead of absolute paragraph indices, each slot is anchored to
        a section + anchor text. Resolution happens at insertion time.
        """
        slots = []
        figure_counter = 1

        # ── Strategy 1: Figure caption patterns in template ──
        for s_idx, section in enumerate(sections):
            for para in section.paragraphs:
                text = para.text_preview

                # Match figure captions: "图X-Y" or "图X-Y caption text"
                fig_match = re.match(r'(图\d+[-–—]\d+)(.*)', text)
                if fig_match:
                    fig_num = fig_match.group(1)
                    fig_rest = fig_match.group(2).strip()

                    slot_id = f"figure_{s_idx}_{para.para_index}"
                    caption_template = f"{{{fig_num}}} {{{{caption_text}}}}"

                    slots.append(ImageSlot(
                        slot_id=slot_id,
                        section_title=section.title,
                        section_index=s_idx,
                        position_type="before_caption",
                        anchor_text=fig_num,  # Anchor on the figure number
                        relative_offset=para.para_index - section.start_para_index,
                        caption_template=caption_template,
                        suggested_type=self._guess_image_type(text),
                        priority=1,
                    ))
                    continue

                # Match "照片" / "示意图" / etc. patterns
                if any(kw in text for kw in ["照片", "示意图", "平面图", "红线图", "效果图"]):
                    if len(text) < 80:
                        slot_id = f"photo_{s_idx}_{para.para_index}"
                        slots.append(ImageSlot(
                            slot_id=slot_id,
                            section_title=section.title,
                            section_index=s_idx,
                            position_type="before_caption",
                            anchor_text=text[:40],
                            relative_offset=para.para_index - section.start_para_index,
                            caption_template=f"图{{{{num}}}} {{{{project_name}}}}{text[:30]}",
                            suggested_type=self._guess_image_type(text),
                            priority=2,
                        ))

        # ── Strategy 2: Compare with example to find missing images ──
        if doc_ex:
            ex_headings = self._extract_headings(doc_ex)
            ex_hmap = {h[1]: h[0] for h in ex_headings if h[1]}

            for s_idx, section in enumerate(sections):
                if section.title not in ex_hmap:
                    continue

                ex_start = ex_hmap[section.title]
                # Find next ex heading
                ex_end = len(doc_ex.paragraphs)
                for ex_h in ex_headings:
                    if ex_h[0] > ex_start:
                        ex_end = ex_h[0]
                        break

                # Check for images in example section that are missing in template
                tpl_has_images = any(
                    doc_tpl.paragraphs[p.para_index]._p.findall('.//' + qn('w:drawing'))
                    for p in section.paragraphs
                    if p.para_index < len(doc_tpl.paragraphs)
                )

                if not tpl_has_images:
                    for ex_pi in range(ex_start + 1, ex_end):
                        if ex_pi >= len(doc_ex.paragraphs):
                            break
                        ex_para = doc_ex.paragraphs[ex_pi]
                        ex_text = ex_para.text.strip()

                        has_drawing = bool(
                            ex_para._p.findall('.//' + qn('w:drawing'))
                        )

                        if has_drawing:
                            # Found an image in example but not in template
                            fig_caption = ex_text if len(ex_text) < 80 else ""
                            slot_id = f"ex_diff_{s_idx}_{figure_counter}"
                            slots.append(ImageSlot(
                                slot_id=slot_id,
                                section_title=section.title,
                                section_index=s_idx,
                                position_type="before_caption",
                                anchor_text=section.title,
                                relative_offset=1,
                                caption_template=fig_caption or f"图{{{{num}}}}",
                                suggested_type=self._guess_image_type(
                                    section.title + " " + fig_caption
                                ),
                                priority=1,
                            ))
                            figure_counter += 1

        # ── Strategy 3: Attachment section images ──
        for s_idx, section in enumerate(sections):
            if any(kw in section.title for kw in ["附件", "附图"]):
                for para in section.paragraphs:
                    text = para.text_preview
                    # Match numbered attachment items: "1.xxx照片" etc.
                    if re.match(r'^\d+\.\s*\S', text) and len(text) < 80:
                        if any(kw in text for kw in ["照片", "图", "扫描", "问卷", "纪要"]):
                            slot_id = f"attachment_{s_idx}_{para.para_index}"
                            slots.append(ImageSlot(
                                slot_id=slot_id,
                                section_title=section.title,
                                section_index=s_idx,
                                position_type="before_caption",
                                anchor_text=text[:40],
                                relative_offset=para.para_index - section.start_para_index,
                                caption_template=text,
                                suggested_type="attachment",
                                priority=2,
                            ))

        return slots

    # ── Section classification ──

    def _classify_sections(
        self, sections: List[SectionDef], doc_tpl: Document
    ) -> None:
        """Tag each section with its type."""
        for section in sections:
            text_lower = section.title.lower() if section.title else ""

            # Preserved: company info, licenses, certificates
            if any(kw in section.title for kw in [
                "营业执照", "证书", "备案", "人员", "分工", "公司简介",
                "资质", "案卷", "目录",
            ]):
                section.section_type = SectionType.PRESERVED
                continue

            # Preserved: TOC
            if any(kw in section.title for kw in ["目录", "目  录"]):
                section.section_type = SectionType.PRESERVED
                continue

            # Check if section only has tables
            has_body_text = any(
                len(p.text_preview) > 30 and p.heading_level == 0
                for p in section.paragraphs
            )
            if not has_body_text and len(section.paragraphs) > 0:
                # Check for tables near this section
                # (simplified: if section has < 3 body paragraphs, treat as table_only)
                body_count = sum(
                    1 for p in section.paragraphs
                    if len(p.text_preview) > 30 and p.heading_level == 0
                )
                if body_count <= 2:
                    section.section_type = SectionType.TABLE_ONLY
                    continue

            # Default: agent generate
            section.section_type = SectionType.AGENT_GENERATE

    # ── Helpers ──

    @staticmethod
    def _guess_image_type(text: str) -> str:
        """Guess the type of image from caption/section text."""
        text_lower = text.lower() if text else ""
        if any(kw in text for kw in ["位置", "地图", "红线", "示意图", "平面"]):
            return "location_map"
        if any(kw in text for kw in ["公示", "公告", "张贴"]):
            return "notice"
        if any(kw in text for kw in ["座谈", "会议", "村民", "开会"]):
            return "meeting"
        if any(kw in text for kw in ["现场", "实地", "勘察", "地块"]):
            return "site_photo"
        if any(kw in text for kw in ["专家", "评审"]):
            return "expert_review"
        if any(kw in text for kw in ["问卷", "调查"]):
            return "survey"
        if any(kw in text for kw in ["附件"]):
            return "attachment"
        return "general"
