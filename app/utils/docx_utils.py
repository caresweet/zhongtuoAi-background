"""Docx utility functions — placeholder detection, section iteration, etc."""
import re
from typing import List, Dict, Any, Optional, Tuple
from docx import Document


# Regex patterns for detecting placeholders in docx text
PLACEHOLDER_PATTERNS = [
    (re.compile(r'\{\{(\w+)\}\}'), "double_brace"),            # {{project_name}}
    (re.compile(r'\{\{([^}]+)\}\}'), "double_brace_cn"),      # {{项目名称}}
    (re.compile(r'[［【]([^］】]+)[］】]'), "bracket_cn"),     # 【项目名称】
    (re.compile(r'<([^>]+)>'), "angle_bracket"),               # <project_name>
    (re.compile(r'_{3,}'), "underscore_blank"),               # ________
]


def find_placeholders_in_text(text: str) -> List[Dict[str, str]]:
    """Find all placeholder patterns in a text string."""
    found = []
    for pattern, pattern_type in PLACEHOLDER_PATTERNS:
        for match in pattern.finditer(text):
            key = match.group(1) if match.lastindex else match.group(0)
            found.append({
                "key": key.strip() if key else match.group(0),
                "matched_text": match.group(0),
                "pattern_type": pattern_type,
                "start": match.start(),
                "end": match.end(),
            })
    return found


def iter_sections(doc: Document) -> List[Dict[str, Any]]:
    """Iterate through document sections grouped by headings."""
    sections = []
    current_section = None

    for p_idx, paragraph in enumerate(doc.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""

        if style_name.startswith("Heading") or style_name.startswith("heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1

            if current_section:
                current_section["end_para_index"] = p_idx - 1
                sections.append(current_section)

            current_section = {
                "index": len(sections),
                "title": paragraph.text.strip(),
                "level": level,
                "start_para_index": p_idx,
                "end_para_index": None,
                "paragraphs": [],
            }
        elif current_section is not None:
            placeholders = find_placeholders_in_text(paragraph.text)
            current_section["paragraphs"].append({
                "para_index": p_idx,
                "style": style_name,
                "text_preview": paragraph.text[:200],
                "has_placeholder": len(placeholders) > 0,
                "placeholders": placeholders,
            })

    if current_section:
        current_section["end_para_index"] = len(doc.paragraphs) - 1
        sections.append(current_section)

    return sections


def get_heading_hierarchy(doc: Document) -> List[Dict[str, Any]]:
    """Get the heading hierarchy of a document."""
    headings = []
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name.startswith("Heading") or style_name.startswith("heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            headings.append({
                "level": level,
                "text": p.text.strip(),
                "style": style_name,
            })
    return headings


def get_paragraph_run_count(doc: Document, para_index: int) -> int:
    """Get the number of runs in a specific paragraph."""
    if para_index < len(doc.paragraphs):
        return len(doc.paragraphs[para_index].runs)
    return 0


def get_run_text(doc: Document, para_index: int, run_index: int) -> Optional[str]:
    """Get the text of a specific run. Returns None if indices are out of range."""
    try:
        return doc.paragraphs[para_index].runs[run_index].text
    except IndexError:
        return None
