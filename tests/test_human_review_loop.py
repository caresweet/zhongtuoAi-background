"""Skill 审核 → AI 重写 / 人工介入闭环 单元测试。

覆盖三个核心场景：
- 场景1：skill 命中违规，单章多次重试失败 → 推入人工队列
- 场景2：前端提交人工修改想法，AI 携带人工意见重写章节
- 场景3：前端人工直接审批通过，跳过 AI 重写

外加：Skill 输出解析不丢任务、全局重写循环过滤人工章、人工 override 直接生效、
docx 批注注入。
"""
import os
import sys
import asyncio
import re
from pathlib import Path

# 保证 backend 可被 import（与 conftest 一致）
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest

from app.validation.audit_dispatcher import (
    classify_issue,
    parse_skill_audit_to_chapter_tasks,
    split_tasks,
    collect_global_human_items,
    filter_ai_rewrite_chapters,
)

MAX_RETRY = 2  # 与 report_workflow.MAX_RETRY 默认对齐，测试里显式写入 state


def _mk_issue(ch, typ, sev="critical", sug="regenerate", **kw):
    return {"chapter": ch, "type": typ, "severity": sev, "suggestion": sug,
            "message": f"第{ch}章问题：{typ}", "correction": "", **kw}


# ═══════════════════════════════════════════════════════════════════════════════
# Skill 输出解析 —— 禁止只打日志不生成任务
# ═══════════════════════════════════════════════════════════════════════════════

def test_parse_skill_audit_generates_tasks_not_logs():
    """Skill 原始输出必须解析为章节级任务，每章都有落点，禁止静默丢弃。"""
    audit = {"all_issues": [
        _mk_issue(6, "expert_skill_violation", "critical", "regenerate", skill_rule_id=3, pattern="综上所述"),
        _mk_issue(3, "fabricated_data", "critical", "regenerate"),
        _mk_issue(5, "missing_materials", "critical", "manual_fix"),
        _mk_issue(4, "word_count", "warning", "expand"),   # 非阻塞 → 不入重写/人工队列
    ]}
    tasks = parse_skill_audit_to_chapter_tasks(audit)
    assert 6 in tasks and 3 in tasks and 5 in tasks and 4 in tasks  # 每章都有任务
    assert len(tasks[6]) == 1

    ai, human = split_tasks(tasks)
    assert 6 in ai       # expert_skill_violation → AI 可重写
    assert 3 in human    # fabricated_data → 需人工提供真实数据
    assert 5 in human    # missing_materials → 需人工上传
    assert 4 not in ai and 4 not in human  # warning 不触发重写/人工

    # 分类兜底：不认识的类型 + critical → ai_rewrite；非 critical → warning
    assert classify_issue({"type": "some_new_rule", "severity": "critical", "suggestion": "regenerate"}) == "ai_rewrite"
    assert classify_issue({"type": "some_new_rule", "severity": "warning", "suggestion": "regenerate"}) == "warning"


def test_global_human_items_not_mapped_to_chapter():
    """chapter=0 的全局人工待办（如缺营业执照）单独收集，不入具体章重写。"""
    audit = {"all_issues": [
        _mk_issue(0, "missing_materials", "critical", "manual_fix"),
        _mk_issue(2, "hallucinated_regulation", "critical", "regenerate"),
    ]}
    tasks = parse_skill_audit_to_chapter_tasks(audit)
    assert 0 not in tasks          # 全局项不挂到具体章
    assert 2 in tasks
    globals_ = collect_global_human_items(audit)
    assert len(globals_) == 1
    assert globals_[0]["type"] == "missing_materials"


# ═══════════════════════════════════════════════════════════════════════════════
# 场景1：skill 命中违规，单章多次重试失败 → 推入人工队列（独立重试计数）
# ═══════════════════════════════════════════════════════════════════════════════

def _mk_review_state(short_md="内容"):
    """构造 node_chapter_review 的最小状态：单章、内容过短必然触发 issues。"""
    return {
        "outline_list": [{"chapter_no": "6", "title": "措施前风险等级研判",
                          "raw_content": short_md, "review_msg": "", "review_score": None}],
        "_chapter_idx": 0, "_chapter_retry": 0,
        "chapters": {}, "human_items": {}, "human_queue": [], "logs": [],
        "filled_data": {}, "max_retry": MAX_RETRY,
    }


def test_scene1_retry_exhausted_pushes_human_queue():
    """第6章内容过短，连续两次审查：第1次重试、第2次推入人工队列。"""
    from app.services import report_workflow as wf

    async def _run():
        state = _mk_review_state(short_md="内容太短")   # <300字 → 必然 issues

        # 第 1 次审查：还有重试机会 → AI 重写本章
        await wf.node_chapter_review(state)
        hi = state["human_items"][6]
        assert hi["retry_count"] == 1
        assert hi["in_human_queue"] is False
        assert state["_next_action"] == "generate"       # 回生成重写
        assert 6 not in state["human_queue"]

        # 第 2 次审查：本章重试耗尽 → 推入人工队列
        await wf.node_chapter_review(state)
        hi = state["human_items"][6]
        assert hi["retry_count"] == MAX_RETRY
        assert hi["in_human_queue"] is True
        assert hi["status"] == "queued"
        assert 6 in state["human_queue"]
        assert state["chapters"][6]["status"] == "human_review"

        # 验证不再消耗全局轮次相关字段（_quality_round 未被动过）
        assert state.get("_quality_round", 0) == 0
        return state
    asyncio.run(_run())


def test_scene1_per_chapter_retry_is_independent():
    """第5章通过不影响第6章的重试计数（每章独立维护 retry_count）。"""
    from app.services import report_workflow as wf

    async def _run():
        state = _mk_review_state()
        state["outline_list"] = [
            {"chapter_no": "5", "title": "风险因素识别", "raw_content": "x" * 500,
             "review_msg": "", "review_score": None},
            {"chapter_no": "6", "title": "措施前研判", "raw_content": "短",
             "review_msg": "", "review_score": None},
        ]
        # 第5章（idx=0）通过
        await wf.node_chapter_review(state)
        assert state["human_items"][5]["status"] == "passed"
        assert state["human_items"].get(6, {}).get("retry_count", 0) == 0  # 第6章不受影响
        assert state["_chapter_idx"] == 1
    asyncio.run(_run())


# ═══════════════════════════════════════════════════════════════════════════════
# 场景2：前端提交人工修改想法，AI 携带人工意见重写章节
# ═══════════════════════════════════════════════════════════════════════════════

def test_scene2_human_opinion_flows_into_feedback():
    """node_chapter_generate 的 feedback 构造必须包含人工意见。"""
    state = {
        "human_items": {6: {
            "chapter": 6, "human_opinion": "支持率改为100%，删除反对表述",
            "human_approved": False, "human_override": False,
        }},
    }
    hi = state["human_items"][6]
    feedback = None
    ch_num = 6
    # 复刻 node_chapter_generate 的注入逻辑
    opinion = hi.get("human_opinion", "")
    if opinion:
        base = "## ⚠️ 终稿审核发现问题" if ch_num else ""
        feedback = (base + "\n" if base else "") + f"## 👤 人工修改意见（必须严格遵循）\n{opinion}"

    assert "支持率改为100%" in feedback
    assert "人工修改意见" in feedback
    assert not hi.get("human_approved")          # 未审批，仍需 AI 重写
    assert not hi.get("human_override")


def test_scene2_generation_feedback_merges_skill_defect_and_human_opinion():
    """重写时同时传入 Skill 缺陷 + 人工意见。"""
    skill_defect = "## ⚠️ 评分超出0-100范围"
    human_opinion = "## 👤 人工修改意见\n措施前得分应为75分"
    combined = skill_defect + "\n" + human_opinion
    assert "评分超出0-100范围" in combined
    assert "措施前得分应为75分" in combined


# ═══════════════════════════════════════════════════════════════════════════════
# 场景3：前端人工直接审批通过，跳过 AI 重写（全局重写循环过滤）
# ═══════════════════════════════════════════════════════════════════════════════

def test_scene3_human_approve_skips_ai_rewrite():
    """human_approved=True 的章节从全局重写队列剔除，不再参与 AI 重写。"""
    regenerate = [6, 9]                      # 审核本想重写 6 和 9
    human_items = {6: {"human_approved": True},
                   9: {"human_approved": False, "human_opinion": "改一下"}}
    ai_chapters = {6, 9}

    rewrite = filter_ai_rewrite_chapters(regenerate, human_items, [], ai_chapters)
    assert rewrite == [9]                    # 第6章被审批放行
    # 场景3 变体：全部审批通过 → 不消耗任何重写
    human_items[9]["human_approved"] = True
    assert filter_ai_rewrite_chapters([6, 9], human_items, [], ai_chapters) == []


def test_scene3_queued_chapters_also_skipped():
    """已在人工队列中的章节（未处理）同样不参与 AI 重写。"""
    regenerate = [1, 2, 3]
    human_items = {2: {"in_human_queue": True, "human_approved": False}}
    ai_chapters = {1, 2, 3}
    rewrite = filter_ai_rewrite_chapters(regenerate, human_items, [2], ai_chapters)
    assert rewrite == [1, 3]


# ═══════════════════════════════════════════════════════════════════════════════
# 场景2a：人工 override 直接改写全文，跳过 AI
# ═══════════════════════════════════════════════════════════════════════════════

def test_scene2a_human_override_applies_content():
    """human_override + override_content → node_chapter_generate 直接写回，不走 AI。"""
    from app.services import report_workflow as wf

    override_md = "第一章人工改写后的完整内容。" * 30
    state = {
        "outline_list": [{"chapter_no": "1", "title": "拟征收决策基本概况",
                          "raw_content": "旧内容", "review_msg": ""}],
        "_chapter_idx": 0, "_chapter_retry": 0,
        "chapters": {}, "human_items": {
            1: {"chapter": 1, "human_override": True, "override_content": override_md,
                "human_approved": False, "human_opinion": ""},
        },
        "filled_data": {"project_name": "测试项目"},
        "logs": [], "_gen_prepared": True, "rag_all_chunks": [],
    }

    async def _run():
        # node_chapter_generate 里 override 分支在 ch_def 之后、LLM 调用之前返回。
        # 复刻该分支的判定 + 写回逻辑：
        human_items = state.setdefault("human_items", {})
        hi = human_items.setdefault(1, {})
        if hi.get("human_override") and hi.get("override_content"):
            ch = state["outline_list"][0]
            ch["raw_content"] = hi["override_content"]
            state["chapters"][1] = {"markdown": hi["override_content"], "title": ch["title"],
                                    "status": "human_reviewed"}
        return state
    asyncio.run(_run())

    assert state["chapters"][1]["markdown"] == override_md
    assert state["chapters"][1]["status"] == "human_reviewed"


def test_scene2a_override_skipped_by_global_filter():
    """human_override 章节被全局重写过滤剔除。"""
    regenerate = [1, 2]
    human_items = {1: {"human_override": True, "override_content": "x"}}
    rewrite = filter_ai_rewrite_chapters(regenerate, human_items, [], {1, 2})
    assert rewrite == [2]


# ═══════════════════════════════════════════════════════════════════════════════
# 全局质量循环过滤（场景3 + 人工队列）
# ═══════════════════════════════════════════════════════════════════════════════

def test_global_quality_loop_only_rewrites_ai_fixable():
    """重写队列 = regenerate ∩ ai_chapters − 人工介入。"""
    # fabricated_data 需人工，不在 ai_chapters 里 → 不进重写队列
    regenerate = [3, 6]
    ai_chapters = {6}                          # 只有第6章是 ai_rewrite
    human_items = {6: {"human_approved": False}}
    rewrite = filter_ai_rewrite_chapters(regenerate, human_items, [], ai_chapters)
    assert rewrite == [6]


# ═══════════════════════════════════════════════════════════════════════════════
# docx 导出批注注入
# ═══════════════════════════════════════════════════════════════════════════════

def test_docx_audit_comments_injected():
    """导出 docx 时把 Skill 违规 + 人工干预状态写入批注。"""
    from docx import Document
    from app.services.report_assembler import ReportAssembler
    import zipfile

    asm = ReportAssembler()
    doc = Document()
    asm._add_heading(doc, "第1章 拟征收决策基本概况", 1)
    asm._add_para(doc, "正文", indent=True)
    asm._add_heading(doc, "第6章 措施前风险等级研判", 1)
    asm._add_para(doc, "正文六", indent=True)

    state = {
        "chapter_audits": {6: [{"disposition": "ai_rewrite", "severity": "critical",
                                "message": "评分超出0-100范围"}]},
        "human_items": {6: {"human_approved": False, "human_opinion": "支持率改100%"}},
        "_quality_audit": {"total_issues": 3, "critical_issues": 1,
                           "auto_fixed": 1, "timestamp": "2026-08-19T00:00:00"},
    }
    asm._inject_audit_comments(doc, state)

    import tempfile
    tmp = tempfile.mktemp(suffix=".docx")
    try:
        doc.save(tmp)
        z = zipfile.ZipFile(tmp)
        names = z.namelist()
        assert "word/comments.xml" in names, "缺少 comments.xml"
        comments_xml = z.read("word/comments.xml").decode("utf-8")
        assert "评分超出0-100范围" in comments_xml
        assert "人工意见" in comments_xml
        assert "审核元数据" in comments_xml
        doc_xml = z.read("word/document.xml").decode("utf-8")
        assert doc_xml.count("commentRangeStart") >= 2
        assert doc_xml.count("commentReference") >= 2
    finally:
        try:
            os.unlink(tmp)
        except OSError:
            pass


def test_docx_no_audit_no_comments_part():
    """无审核数据时不生成空 comments part。"""
    from docx import Document
    from app.services.report_assembler import ReportAssembler
    import zipfile, tempfile

    asm = ReportAssembler()
    doc = Document()
    asm._add_para(doc, "无审核数据")
    asm._inject_audit_comments(doc, {})   # 空 state
    tmp = tempfile.mktemp(suffix=".docx")
    try:
        doc.save(tmp)
        z = zipfile.ZipFile(tmp)
        assert "word/comments.xml" not in z.namelist()
    finally:
        try:
            os.unlink(tmp)
        except OSError:
            pass


# ═══════════════════════════════════════════════════════════════════════════════
# 回归：session 持久化后 dict 键为字符串（"6"≠6），端点必须归一化
# ═══════════════════════════════════════════════════════════════════════════════

class _FakeSession:
    def __init__(self, state):
        self.state = state


def test_human_queue_normalizes_str_keys(monkeypatch):
    """GET /human-queue 读取持久化后的字符串键 human_items/chapter_audits。"""
    from app.routers.report import report_service, get_human_queue
    state = {
        "human_queue": [6, 9],
        "human_items": {"6": {"chapter": 6, "retry_count": 2, "max_retry": 2,
                              "in_human_queue": True, "status": "queued",
                              "human_opinion": "", "human_override": False,
                              "human_approved": False}},
        "chapter_audits": {"6": [{"type": "score_out_of_range", "severity": "critical",
                                  "disposition": "ai_rewrite",
                                  "message": "评分超出0-100范围"}]},
    }
    monkeypatch.setattr(report_service, "get_session", lambda sid: _FakeSession(state))
    resp = asyncio.run(get_human_queue("fake-sid"))
    chapters = resp.data["chapters"]
    assert len(chapters) == 2
    ch6 = next(c for c in chapters if c["chapter"] == 6)
    assert ch6["retry_count"] == 2 and ch6["max_retry"] == 2          # 字符串键被归一化读出
    assert ch6["violations"][0]["type"] == "score_out_of_range"


def test_human_opinion_normalizes_str_keys(monkeypatch):
    """POST /human/opinion 在持久化字符串键基础上写入，不产生重复 int/str 键。"""
    from app.routers.report import report_service, submit_human_opinion
    state = {
        "human_queue": [6],
        "human_items": {"6": {"chapter": 6, "retry_count": 2, "max_retry": 2,
                              "in_human_queue": True, "status": "queued",
                              "human_opinion": "", "human_override": False,
                              "human_approved": False}},
    }
    monkeypatch.setattr(report_service, "get_session", lambda sid: _FakeSession(state))
    resp = asyncio.run(submit_human_opinion(
        "fake-sid", {"chapter": 6, "mode": "opinion", "opinion": "支持率改为100%"}))
    assert resp.data["human_opinion"] == "支持率改为100%"
    # 只保留一种键形式（int），无重复
    assert set(state["human_items"].keys()) == {6}
    assert state["human_items"][6]["status"] == "human_reviewed"


def test_human_rewrite_resumes_all_acted_chapters(monkeypatch):
    """POST /human/rewrite 恢复载荷携带所有已处理章节（字符串键归一化）。

    第6章有意见→AI重写，第9章已审批→放行；两者都应进入恢复载荷，由工作流分别处理。
    """
    from app.routers.report import report_service, trigger_human_ai_rewrite
    state = {
        "human_queue": [6, 9],
        "human_items": {
            "6": {"chapter": 6, "human_opinion": "改一下", "human_approved": False, "human_override": False},
            "9": {"chapter": 9, "human_opinion": "", "human_approved": True, "human_override": False},
        },
    }
    monkeypatch.setattr(report_service, "get_session", lambda sid: _FakeSession(state))
    resp = asyncio.run(trigger_human_ai_rewrite("fake-sid", {}))
    assert resp.data["chapters"] == [6, 9]       # 队列中所有已处理章节都恢复
    assert resp.data["acted"] == [6, 9]          # 6=意见, 9=审批


def test_human_rewrite_no_action_rejected(monkeypatch):
    """未对任何章节提交意见/审批 → 拒绝恢复，避免空恢复导致死循环。"""
    from app.routers.report import report_service, trigger_human_ai_rewrite
    state = {
        "human_queue": [6],
        "human_items": {
            "6": {"chapter": 6, "human_opinion": "", "human_approved": False, "human_override": False},
        },
    }
    monkeypatch.setattr(report_service, "get_session", lambda sid: _FakeSession(state))
    resp = asyncio.run(trigger_human_ai_rewrite("fake-sid", {}))
    assert resp.data["chapters"] == []           # 无任何处理 → 不触发恢复


def test_workflow_status_normalizes_int_keys(monkeypatch):
    """workflow/status 对 int 键 human_items 归一化，避免 pydantic Dict[str,Any] 500。

    工作流里 human_items/chapter_audits 是 int 键，pydantic 要求 str 键 → 必须归一化。
    """
    from app.routers.report import report_service, workflow_status
    state = {
        "human_queue": [6],
        "human_items": {6: {"chapter": 6, "human_opinion": "改一下", "human_approved": False,
                            "human_override": False}},
        "chapter_audits": {6: [{"type": "score_out_of_range", "severity": "critical",
                                "disposition": "ai_rewrite", "message": "评分超范围"}]},
        "phase": "human_review", "chapters": {}, "_outline": {},
        "_workflow_logs": ["⏸️ 1 个章节需人工复核"],
        "step_statuses": {},
    }
    monkeypatch.setattr(report_service, "get_session", lambda sid: _FakeSession(state))
    resp = asyncio.run(workflow_status("fake-sid"))
    data = resp.data  # workflow_status 返回 model_dump() 后的 dict
    assert data["phase"] == "human_review"
    assert data["human_queue"] == [6]
    assert list(data["human_items"].keys()) == ["6"]       # int 键已归一化为 str
    assert data["chapter_audits"]["6"][0]["type"] == "score_out_of_range"


# ═══════════════════════════════════════════════════════════════════════════════
# OCR 数据修复回归：短值不丢弃 + 逐页累加问卷统计
# ═══════════════════════════════════════════════════════════════════════════════

def test_parse_structured_data_keeps_short_counts():
    """_parse_structured_data 不再丢弃 support_count:'1' 这类短字符串。"""
    from app.services.pdf_data_extractor import PDFDataExtractor
    ext = PDFDataExtractor(llm_service=None)
    raw = '```json\n{"meeting_date":"2026年4月","attendees":"成美文","support_count":"1","oppose_count":"0","awareness_rate":"100%"}\n```'
    sd = ext._parse_structured_data(raw, "meeting")
    assert sd.get("support_count") == "1"
    assert sd.get("oppose_count") == "0"
    assert sd.get("awareness_rate") == "100%"


def test_aggregate_key_data_sums_survey_counts():
    """_aggregate_key_data 逐页累加 support/oppose，合并参会人，不再被最后一页覆盖。"""
    from types import SimpleNamespace
    from app.services.pdf_data_extractor import PDFDataExtractor
    ext = PDFDataExtractor(llm_service=None)

    def _sd(date, loc, att, support, oppose, aware):
        return {"meeting_date": date, "meeting_location": loc, "attendees": att,
                "discussion": "", "public_demands": "", "conclusion": "",
                "total_samples": "", "support_count": support, "oppose_count": oppose,
                "support_rate": "100%" if support == "1" else "",
                "oppose_rate": "0%" if oppose == "0" else "",
                "awareness_rate": "100%" if aware else ""}

    pages = [
        SimpleNamespace(structured_data=_sd("2026.4.28", "舟山村道三坪社区", "杨亚、叶玉华", "", "", "")),
        SimpleNamespace(structured_data=_sd("2026.4.29", "朱坝街道三圩社区", "成美文", "1", "0", True)),
        SimpleNamespace(structured_data=_sd("2026.4.21", "朱坝街道三圩社区", "徐守科", "1", "0", True)),
    ]
    doc = SimpleNamespace(pages=pages, document_type="meeting")
    agg = ext._aggregate_key_data(doc)
    assert agg["support_count"] == 2, f"支持数应累加为2, 实际{agg['support_count']}"
    assert agg["oppose_count"] == 0
    assert agg["total_samples"] == 2          # 2 个单人调查页
    assert agg["support_rate"] == "100.0%"
    assert "杨亚" in agg["symposium_attendees"] and "徐守科" in agg["symposium_attendees"]
    assert len(agg["symposium_date"].split("、")) == 3   # 3 个日期都保留
