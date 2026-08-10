"""BiddingDocxGenerator — converts bidding report markdown to formatted DOCX.

Simpler than ReportAssembler (which handles 10-chapter stability reports).
This handles bidding documents: announcements, evaluations, awards, summaries,
and tender response files (投标文件).

v2: Added full image support — detects image markers in markdown, inserts
    uploaded project images with captions, and adds company asset images
    (certificates, licenses, etc.) in the appendix.
"""

import base64 as _base64
import os, re, logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, List

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _asset_image_to_dict(row) -> Dict:
    """Convert AssetImage ORM row to a simple dict (avoids import issues)."""
    return {
        "data": _base64.b64decode(row.image_data),
        "mime_type": row.mime_type,
        "name": row.image_name,
        "category": row.category,
    }

# Chinese number mapping
_CN_NUM = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}

FONT_H1 = '黑体'
FONT_H2 = '黑体'
FONT_H3 = '楷体'
FONT_BODY = '仿宋_GB2312'
FONT_TABLE = '宋体'
FONT_CAPTION = '楷体'

REPORT_TYPE_NAMES = {
    "tender_response": "投标文件",
    "negotiation": "竞争性磋商文件",
    "announcement": "招标公告",
    "evaluation": "评标报告",
    "award": "中标结果公示",
    "summary": "招标情况报告",
}

# ── Image keyword → section mapping for intelligent placement ──
SECTION_IMAGE_MAP = {
    "营业执照": ["营业执照", "营业执照副本", "法人营业执照"],
    "资质证书": ["资质证书", "资质", "测绘资质", "资质等级"],
    "法人证明": ["法定代表人", "法人证明", "法人资格"],
    "授权委托": ["授权委托书", "授权书", "委托书"],
    "承诺函": ["承诺函", "声明函", "承诺书"],
    "财务": ["财务报告", "审计报告", "财务报表"],
    "社保纳税": ["社保", "纳税", "税收", "社保证明"],
    "人员": ["人员证书", "资格证书", "职称证书", "执业资格"],
    "业绩": ["业绩证明", "合同", "验收报告", "业绩"],
    "设备": ["设备清单", "仪器", "设备"],
    "办公场所": ["办公场所", "场所证明", "租赁合同"],
}

# ── Image caption templates per category ──
CATEGORY_CAPTIONS = {
    "营业执照": "营业执照",
    "资质证书": "资质证书",
    "法人证明": "法定代表人资格证明",
    "授权委托": "授权委托书",
    "承诺函": "各类承诺函/声明函",
    "财务": "财务状况证明",
    "社保纳税": "社保及纳税证明",
    "人员": "项目人员资质证书",
    "业绩": "类似项目业绩证明",
    "设备": "主要设备清单及证明",
    "办公场所": "办公场所证明",
    "other": "附件材料",
}


class BiddingDocxGenerator:
    """Generate formatted DOCX from bidding report markdown content."""

    def __init__(self, storage_dir: Optional[Path] = None):
        from app.config import settings
        self.storage_dir = storage_dir or settings.STORAGE_DIR
        self.images_dir = self.storage_dir / "images"
        self.company_dir = self.storage_dir / "extracted_imgs"
        self.output_dir = self.storage_dir / "generated"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ═══════════════════════════════════════════════════════════════
    # Public API
    # ═══════════════════════════════════════════════════════════════

    def generate(
        self,
        markdown_content: str,
        report_type: str = "summary",
        metadata: Optional[Dict[str, str]] = None,
        state: Optional[Dict] = None,
    ) -> str:
        """Convert markdown to DOCX and save.

        Args:
            markdown_content: Full markdown of the bidding report.
            report_type: announcement | evaluation | award | summary | tender_response.
            metadata: {project_name, reference, session_id, ...}
            state: Optional agent state dict with _uploaded_files, _bidding_data, etc.

        Returns:
            Relative path to generated DOCX (e.g. "generated/xxx.docx").
        """
        metadata = metadata or {}
        state = state or {}
        report_name = REPORT_TYPE_NAMES.get(report_type, "招标报告")
        session_id = metadata.get("session_id", "report")
        project_name = metadata.get("project_name", "")
        reference = metadata.get("reference", "")

        # ── Collect available images ──
        available_images = self._get_session_images(state)
        company_images = self._get_company_images()

        doc = Document()
        self._setup_page(doc)

        # ── Cover Page ──
        self._add_cover(doc, report_name, project_name, reference)

        # ── Body Content ──
        sections = self._parse_markdown_sections(markdown_content)
        for i, (title, content) in enumerate(sections):
            self._add_section(doc, title, content, available_images, i)
            if i < len(sections) - 1:
                doc.add_page_break()

        # ── Company Asset Images (证书/资质/证明等附件) ──
        if company_images:
            doc.add_page_break()
            self._add_company_assets_section(doc, company_images, available_images)

        # ── Save ──
        safe_name = re.sub(r'[\\/:*?"<>|]', '_', project_name or session_id)
        filename = f"{safe_name}_{report_name}.docx"
        outpath = self.output_dir / filename

        # Ensure unique filename
        counter = 1
        while outpath.exists():
            filename = f"{safe_name}_{report_name}_{counter}.docx"
            outpath = self.output_dir / filename
            counter += 1

        doc.save(str(outpath))
        logger.info(f"Bidding DOCX saved: {outpath} ({len(doc.paragraphs)} paras, {len(doc.tables)} tables)")

        return f"generated/{filename}"

    # ═══════════════════════════════════════════════════════════════
    # Page Setup
    # ═══════════════════════════════════════════════════════════════

    def _setup_page(self, doc: Document):
        for sec in doc.sections:
            sec.page_width = Cm(21)
            sec.page_height = Cm(29.7)
            sec.top_margin = Cm(2.54)
            sec.bottom_margin = Cm(2.54)
            sec.left_margin = Cm(3.18)
            sec.right_margin = Cm(3.18)

    # ═══════════════════════════════════════════════════════════════
    # Cover Page
    # ═══════════════════════════════════════════════════════════════

    def _add_cover(self, doc, report_name, project_name, reference):
        for _ in range(5):
            doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(report_name)
        r.font.name = FONT_H1
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H1)
        r.font.size = Pt(26)
        r.bold = True

        doc.add_paragraph()

        # Project info
        if project_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'项目名称：{project_name}')
            r.font.name = FONT_H3
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H3)
            r.font.size = Pt(16)

        if reference:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'项目编号：{reference}')
            r.font.name = FONT_H3
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H3)
            r.font.size = Pt(14)

        for _ in range(4):
            doc.add_paragraph()

        # Footer info
        for text in [
            '招标代理机构：江苏众拓项目代理咨询有限公司',
            f'编制日期：{datetime.now().strftime("%Y年%m月%d日")}',
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.name = FONT_BODY
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
            r.font.size = Pt(14)

        doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # Markdown Parsing → Sections
    # ═══════════════════════════════════════════════════════════════

    def _parse_markdown_sections(self, markdown: str) -> list:
        """Split markdown into (title, content) pairs by chapter headers.

        Chapters are '# 第X章 标题' (single #). In-chapter headings (##/###)
        are kept inside content and rendered as sub-headings by _add_section.
        Falls back to splitting on ## when no single-# chapter headers exist
        (backward compat with the old single-shot output).
        """
        markdown = re.sub(r'```[^`]*```', '', markdown)
        markdown = re.sub(r'```\s*', '', markdown)

        # Prefer single-# chapter headers (new per-chapter output).
        has_h1 = re.search(r'(?m)^#\s+\S', markdown) is not None
        split_pat = r'\n(?=#\s)' if has_h1 else r'\n(?=##\s)'
        head_pat = r'#\s+(.+)' if has_h1 else r'##\s+(.+)'

        sections = []
        for part in re.split(split_pat, markdown):
            part = part.strip()
            if not part:
                continue
            title_match = re.match(head_pat, part)
            if title_match:
                title = title_match.group(1).strip()
                content = part[title_match.end():].strip()
            else:
                title, content = "", part
            if content or title:
                sections.append((title, content))
        return sections

    # ═══════════════════════════════════════════════════════════════
    # Section Rendering
    # ═══════════════════════════════════════════════════════════════

    def _add_section(self, doc, title, content, available_images=None, section_idx=0):
        """Add a section with title and markdown-like content.

        Supports image markers: ![caption text](image_key) and
        ![caption text](path/to/image.jpg).
        """
        available_images = available_images or {}

        # Chapter title — use Heading 1 style so WPS recognizes the hierarchy
        if title:
            p = doc.add_paragraph()
            try:
                p.style = doc.styles['Heading 1']
            except Exception:
                pass
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(title)
            r.font.name = FONT_H1
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H1)
            r.font.size = Pt(16)
            r.bold = True

        # Parse content lines
        lines = content.split('\n')
        # Track image counter for figure numbering
        img_counter = [0]  # use list for mutable closure
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # ── Image detection: ![caption](path_or_key) ──
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if img_match:
                caption_text = img_match.group(1).strip()
                image_ref = img_match.group(2).strip()
                self._insert_image_from_ref(doc, image_ref, caption_text, available_images)
                i += 1
                continue

            # ── Image detection: 图X-X 说明 text (figure reference) ──
            fig_match = re.match(r'(图\d+[-–—]\d+)\s*(.+)', line)
            if fig_match and len(line) < 80:
                fig_label = fig_match.group(1)
                fig_desc = fig_match.group(2).strip()
                # Try to find a matching image
                self._insert_image_from_ref(doc, fig_desc, f"{fig_label} {fig_desc}", available_images)
                # If no image found, the _insert_image_from_ref outputs a placeholder text
                i += 1
                continue

            # Table detection: line starts with |
            if line.startswith('|') and '|' in line[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # In-chapter section heading: '## ' → Heading 2
            if line.startswith('## '):
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles['Heading 2']
                except Exception:
                    pass
                r = p.add_run(line[3:].strip())
                r.font.name = FONT_H2
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H2)
                r.font.size = Pt(15)
                r.bold = True
                i += 1
                # Check if next line is an image
                if i < len(lines):
                    next_line = lines[i].strip()
                    img2 = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', next_line)
                    if img2:
                        self._insert_image_from_ref(
                            doc, img2.group(2).strip(),
                            img2.group(1).strip(), available_images
                        )
                        i += 1
                continue

            # Sub-header: ### → Heading 3
            if line.startswith('### '):
                p = doc.add_paragraph()
                try:
                    p.style = doc.styles['Heading 3']
                except Exception:
                    pass
                r = p.add_run(line[4:].strip())
                r.font.name = FONT_H3
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H3)
                r.font.size = Pt(14)
                r.bold = True
                i += 1
                # Check if next line is an image
                if i < len(lines):
                    next_line = lines[i].strip()
                    img3 = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', next_line)
                    if img3:
                        self._insert_image_from_ref(
                            doc, img3.group(2).strip(),
                            img3.group(1).strip(), available_images
                        )
                        i += 1
                continue

            # Numbered list item
            if re.match(r'^\d+[\.\、\)]', line):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                r = p.add_run(line)
                r.font.name = FONT_BODY
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
                r.font.size = Pt(12)
                i += 1
                continue

            # Regular paragraph
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            # Handle inline bold markers
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = FONT_BODY
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
                r.font.size = Pt(12)

            i += 1

    # ═══════════════════════════════════════════════════════════════
    # Table Rendering
    # ═══════════════════════════════════════════════════════════════

    def _add_table(self, doc, table_lines):
        """Convert markdown table lines to DOCX table."""
        if len(table_lines) < 2:
            return

        # Parse rows from markdown table
        rows = []
        for line in table_lines:
            # Skip separator lines like |---|---|
            if re.match(r'\|[\s\-:|]+\|', line):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)

        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        num_rows = len(rows)

        table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= num_cols:
                    break
                cell = table.cell(row_idx, col_idx)
                # Clear default paragraph
                cell.paragraphs[0].clear()
                r = cell.paragraphs[0].add_run(cell_text)

                if row_idx == 0:
                    # Header row
                    r.font.name = FONT_H2
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H2)
                    r.font.size = Pt(10.5)
                    r.bold = True
                else:
                    r.font.name = FONT_TABLE
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TABLE)
                    r.font.size = Pt(10.5)

                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing after table

    # ═══════════════════════════════════════════════════════════════
    # Image Support
    # ═══════════════════════════════════════════════════════════════

    def _insert_image_from_ref(
        self, doc: Document, image_ref: str, caption: str,
        available_images: Dict[str, List[str]],
    ) -> bool:
        """Try to resolve an image reference and insert it.

        Resolution order:
        1. File on disk (uploaded images, company images)
        2. Database (asset_images table — from knowledge-base extraction)

        image_ref can be:
        - A keyword like "营业执照", "资质证书", "现场照片"
        - A filename like "company_P25_rId6.jpeg"
        - A relative path like "images/some_photo.jpg"

        Returns True if an image was inserted, False otherwise.
        """
        image_path = self._resolve_image_path(image_ref, available_images)

        if image_path:
            self._add_image(doc, image_path, caption or "")
            return True
        else:
            # ── DB fallback: search asset_images table ──
            db_image_data = self._get_image_from_db(image_ref, caption)
            if db_image_data:
                # Insert from DB binary data
                self._add_image_from_bytes(
                    doc, db_image_data["data"], db_image_data["mime_type"],
                    caption or db_image_data.get("name", image_ref),
                )
                return True

            # Output placeholder text explaining what image is needed
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(f'【图片：{caption or image_ref}】')
            r.font.name = FONT_CAPTION
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CAPTION)
            r.font.size = Pt(9)
            r.italic = True
            return False

    def _add_image(self, doc: Document, image_path: str, caption: str = ""):
        """Add a centered image with caption below it.

        Args:
            doc: python-docx Document.
            image_path: Absolute or relative path to image file.
            caption: Figure caption text (e.g. "图1-1 营业执照").
        """
        # Resolve the absolute path
        abs_path = None
        if os.path.isabs(image_path) and os.path.exists(image_path):
            abs_path = image_path
        else:
            # Try images_dir
            candidate = self.images_dir / image_path
            if candidate.exists():
                abs_path = str(candidate)
            else:
                # Try company_dir
                candidate = self.company_dir / image_path
                if candidate.exists():
                    abs_path = str(candidate)
                else:
                    # Try output_dir / storage_dir root
                    candidate = self.storage_dir / image_path
                    if candidate.exists():
                        abs_path = str(candidate)
                    else:
                        # Try as bare filename in images_dir
                        fname = Path(image_path).name
                        candidate = self.images_dir / fname
                        if candidate.exists():
                            abs_path = str(candidate)

        if not abs_path:
            logger.warning(f"Image not found: {image_path}")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'【图片待插入：{caption or Path(image_path).name}】')
            r.font.name = FONT_CAPTION
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CAPTION)
            r.font.size = Pt(9)
            r.italic = True
            return

        # Add spacing before image
        doc.add_paragraph()

        # Insert image centered
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            run = p.add_run()
            run.add_picture(abs_path, width=Inches(4.5))
        except Exception as e:
            logger.warning(f"Failed to add image {abs_path}: {e}")
            # Fallback: text placeholder
            p.clear()
            r = p.add_run(f'【图片加载失败：{Path(abs_path).name}】')
            r.font.name = FONT_CAPTION
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CAPTION)
            r.font.size = Pt(9)
            r.italic = True

        # Add caption below image
        if caption:
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pc.paragraph_format.space_before = Pt(4)
            pc.paragraph_format.space_after = Pt(14)
            rc = pc.add_run(caption)
            rc.font.name = FONT_CAPTION
            rc._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CAPTION)
            rc.font.size = Pt(9)
            rc.italic = True

    def _resolve_image_path(
        self, image_ref: str, available_images: Dict[str, List[str]]
    ) -> Optional[str]:
        """Try to find a real image file matching the given reference.

        Resolution order:
        1. Exact filename match in available_images
        2. Keyword match in available_images
        3. Direct file existence check
        """
        if not image_ref:
            return None

        # 1. Check if image_ref itself is a valid path
        if os.path.exists(image_ref):
            return image_ref

        # Check in images_dir and company_dir
        for base in [self.images_dir, self.company_dir]:
            candidate = base / image_ref
            if candidate.exists():
                return str(candidate)
            # Try just the filename
            fname = Path(image_ref).name
            candidate = base / fname
            if candidate.exists():
                return str(candidate)

        # 2. Search in available_images
        all_images = []
        for cat, imgs in available_images.items():
            all_images.extend(imgs)

        # Exact filename match
        ref_lower = image_ref.lower().replace(' ', '')
        for img in all_images:
            fname = Path(img).name.lower()
            if fname == ref_lower or Path(img).stem.lower() == ref_lower:
                # Resolve the stored path
                return self._find_image_file(img)

        # 3. Keyword-based matching
        keywords = self._extract_image_keywords(image_ref)
        for img in all_images:
            fname = Path(img).name.lower()
            folder = Path(img).parent.name.lower() if '/' in img else ''
            search_text = folder + ' ' + fname
            if any(kw.lower() in search_text for kw in keywords):
                return self._find_image_file(img)

        # 4. Try matching by SECTION_IMAGE_MAP categories
        for cat_key, cat_kws in SECTION_IMAGE_MAP.items():
            if any(kw in image_ref for kw in cat_kws):
                for img in all_images:
                    fname = Path(img).name.lower()
                    folder = Path(img).parent.name.lower() if '/' in img else ''
                    search_text = folder + ' ' + fname
                    if any(kw.lower() in search_text for kw in cat_kws):
                        return self._find_image_file(img)

        return None

    @staticmethod
    def _extract_image_keywords(image_ref: str) -> List[str]:
        """Extract meaningful keywords from an image reference string."""
        # Remove common prefixes/suffixes
        cleaned = re.sub(r'[【】\[\]（）()\s]', '', image_ref)
        # Split Chinese text into 2-char keywords
        keywords = []
        for i in range(len(cleaned) - 1):
            keywords.append(cleaned[i:i+2])
        # Also try the full cleaned string
        keywords.insert(0, cleaned)
        return keywords

    def _find_image_file(self, img_path: str) -> Optional[str]:
        """Resolve an image path from state to an actual file on disk."""
        candidates = [
            Path(img_path),
            self.images_dir / img_path,
            self.images_dir / Path(img_path).name,
            self.company_dir / img_path,
            self.company_dir / Path(img_path).name,
            self.storage_dir / img_path,
        ]
        for c in candidates:
            if c.exists():
                return str(c)
        return None

    # ═══════════════════════════════════════════════════════════════
    # Database Image Lookup (asset_images table)
    # ═══════════════════════════════════════════════════════════════

    # ── Keywords that must NOT appear in bidding document images ──
    _STABILITY_EXCLUDE = [
        "%社会稳定%", "%稳评%", "%风险评估%", "%地质灾害%", "%征地补偿%",
        "%被征地%", "%农民社会保障%", "%补偿登记%", "%征地报批%",
        "%补偿安置方案%", "%土地征收%", "%勘测定界%",
    ]

    @staticmethod
    def _get_image_from_db(image_ref: str, caption: str = "") -> Optional[Dict]:
        """Search the asset_images table for a matching image (sync, thread-safe).

        Uses keyword matching against image_name, search_keywords, and category.
        Automatically excludes stability-assessment content.
        Returns dict with keys: data (bytes), mime_type, name, or None.
        """
        import sqlite3, base64 as _b64
        from app.config import settings

        # Build search terms from both image_ref and caption
        search_terms = []
        for src in [image_ref, caption]:
            cleaned = re.sub(r'[【】\[\]（）()\s]', '', src)
            if cleaned:
                search_terms.append(cleaned)
                # Also add 2-char substrings
                for i in range(len(cleaned) - 1):
                    search_terms.append(cleaned[i:i + 2])

        search_terms = list(dict.fromkeys(search_terms))[:30]  # dedup + limit

        # Build exclusion clause
        exclude_clause = " AND ".join(
            f"(image_name NOT LIKE ? AND search_keywords NOT LIKE ?)"
            for _ in BiddingDocxGenerator._STABILITY_EXCLUDE
        )
        exclude_params = []
        for pat in BiddingDocxGenerator._STABILITY_EXCLUDE:
            exclude_params.extend([pat, pat])

        db_path = settings.DATA_DIR / "knowledge_base.db"
        if not db_path.exists():
            return None

        try:
            conn = sqlite3.connect(str(db_path))
            conn.row_factory = sqlite3.Row
            cur = conn.cursor()

            # Try exact category match first (with stability exclusion)
            for term in [image_ref, caption]:
                clean = re.sub(r'[【】\[\]（）()\s]', '', term)
                if clean:
                    cur.execute(
                        f"SELECT image_name, category, image_data, mime_type FROM asset_images "
                        f"WHERE category = ? AND is_active = 1 AND {exclude_clause} LIMIT 1",
                        (clean, *exclude_params),
                    )
                    row = cur.fetchone()
                    if row:
                        conn.close()
                        return {
                            "data": _b64.b64decode(row["image_data"]),
                            "mime_type": row["mime_type"],
                            "name": row["image_name"],
                            "category": row["category"],
                        }

            # Try LIKE search on image_name and search_keywords (with stability exclusion)
            for term in search_terms[:15]:
                if len(term) < 2:
                    continue
                like_pat = f"%{term}%"
                cur.execute(
                    f"SELECT image_name, category, image_data, mime_type FROM asset_images "
                    f"WHERE is_active = 1 AND {exclude_clause} "
                    f"AND (image_name LIKE ? OR search_keywords LIKE ? OR category LIKE ?) LIMIT 1",
                    (*exclude_params, like_pat, like_pat, like_pat),
                )
                row = cur.fetchone()
                if row:
                    conn.close()
                    return {
                        "data": _b64.b64decode(row["image_data"]),
                        "mime_type": row["mime_type"],
                        "name": row["image_name"],
                        "category": row["category"],
                    }

            conn.close()
        except Exception as e:
            logger.warning(f"DB image lookup failed: {e}")

        return None

    def _add_image_from_bytes(
        self, doc: Document, image_bytes: bytes, mime_type: str,
        caption: str = "",
    ):
        """Insert an image from binary data (from DB) with caption."""
        import tempfile, os as _os

        # Write to a temp file so python-docx can read it
        ext = mime_type.split("/")[-1] if "/" in mime_type else "png"
        if ext == "jpeg":
            ext = "jpg"
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
                tmp.write(image_bytes)
                tmp_path = tmp.name

            self._add_image(doc, tmp_path, caption)
        except Exception as e:
            logger.warning(f"Failed to insert image from DB bytes: {e}")
            # Fallback placeholder
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f'【图片加载失败：{caption}】')
            r.font.name = FONT_CAPTION
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CAPTION)
            r.font.size = Pt(9)
            r.italic = True
        finally:
            if tmp_path and _os.path.exists(tmp_path):
                try:
                    _os.unlink(tmp_path)
                except Exception:
                    pass

    # ═══════════════════════════════════════════════════════════════
    # Image Collection from State
    # ═══════════════════════════════════════════════════════════════

    def _get_session_images(self, state: Dict) -> Dict[str, List[str]]:
        """Classify uploaded images by type for document placement.

        Returns dict with categories like:
        {"business_license": [...], "qualification": [...], "personnel": [...], ...}
        """
        images = {
            "business_license": [],
            "qualification": [],
            "legal_rep": [],
            "authorization": [],
            "commitment": [],
            "financial": [],
            "social_tax": [],
            "personnel": [],
            "performance": [],
            "equipment": [],
            "office": [],
            "other": [],
        }

        uploaded = state.get("_uploaded_files", []) or []
        for fpath in uploaded:
            if not isinstance(fpath, str):
                continue
            fname = fpath.split('/')[-1] if '/' in fpath else fpath
            ext = fname.rsplit('.', 1)[-1].lower() if '.' in fname else ''
            if ext not in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp', 'tiff', 'tif'):
                continue
            self._classify_image(images, fpath, fname)

        # Also check _bidding_data for image references
        bidding_data = state.get("_bidding_data", {}) or {}
        for key in bidding_data:
            if 'image' in key.lower() or 'img' in key.lower() or 'photo' in key.lower():
                val = bidding_data[key]
                if isinstance(val, str) and val:
                    fname = val.split('/')[-1] if '/' in val else val
                    self._classify_image(images, val, fname)

        # Also scan the images directory directly
        if self.images_dir.exists():
            for f in self.images_dir.iterdir():
                if f.is_file() and f.suffix.lower() in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'):
                    fpath_str = str(f.relative_to(self.storage_dir))
                    if fpath_str not in uploaded:
                        self._classify_image(images, fpath_str, f.name)

        return images

    def _classify_image(self, images: Dict[str, List[str]], fpath: str, fname: str):
        """Classify a single image into the correct category."""
        name_lower = fname.lower()
        folder = Path(fpath).parent.name.lower() if '/' in fpath else ''

        # Check by keyword
        if any(k in name_lower for k in ['营业执照', 'license', '执照']):
            images['business_license'].append(fpath)
        elif any(k in name_lower for k in ['资质', 'qualification', 'cert', '证书', '测绘']):
            images['qualification'].append(fpath)
        elif any(k in name_lower for k in ['法人', '法定', 'legal', '代表']):
            images['legal_rep'].append(fpath)
        elif any(k in name_lower for k in ['授权', '委托', 'authorization', '授权书']):
            images['authorization'].append(fpath)
        elif any(k in name_lower for k in ['承诺', '声明', '函', 'commitment', '无重大']):
            images['commitment'].append(fpath)
        elif any(k in name_lower for k in ['财务', '审计', '报告', 'financial', '报表']):
            images['financial'].append(fpath)
        elif any(k in name_lower for k in ['社保', '纳税', '税收', 'tax', 'social']):
            images['social_tax'].append(fpath)
        elif any(k in name_lower for k in ['人员', 'personnel', '身份证', '职称', '资格证', '建造师']):
            images['personnel'].append(fpath)
        elif any(k in name_lower for k in ['业绩', '合同', '验收', 'performance']):
            images['performance'].append(fpath)
        elif any(k in name_lower for k in ['设备', '仪器', 'equipment', '车辆']):
            images['equipment'].append(fpath)
        elif any(k in name_lower for k in ['办公', '场所', '租赁', 'office']):
            images['office'].append(fpath)
        else:
            images['other'].append(fpath)

    def _get_company_images(self) -> List[tuple]:
        """Get company asset images (certificates, licenses) from company_dir.

        Returns list of (path, caption) tuples.
        """
        company_imgs = []
        if self.company_dir.exists():
            for f in sorted(self.company_dir.iterdir()):
                if f.is_file() and f.suffix.lower() in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'):
                    # Derive caption from filename
                    stem = f.stem
                    # Try to find a descriptive caption
                    caption = self._derive_company_img_caption(stem)
                    company_imgs.append((str(f.relative_to(self.storage_dir)), caption))
        return company_imgs

    @staticmethod
    def _derive_company_img_caption(filename_stem: str) -> str:
        """Derive a Chinese caption from company image filename."""
        stem_lower = filename_stem.lower()
        if '营业执照' in stem_lower or 'license' in stem_lower:
            return '营业执照'
        elif '资质' in stem_lower or 'qualification' in stem_lower:
            return '资质证书'
        elif '法人' in stem_lower or 'legal' in stem_lower:
            return '法定代表人资格证明'
        elif '授权' in stem_lower or 'authorization' in stem_lower or '委托' in stem_lower:
            return '授权委托书'
        elif '承诺' in stem_lower or 'commitment' in stem_lower:
            return '承诺函/声明函'
        elif '财务' in stem_lower or 'financial' in stem_lower:
            return '财务状况证明'
        elif '社保' in stem_lower or '纳税' in stem_lower or 'tax' in stem_lower:
            return '社保及纳税证明'
        elif '人员' in stem_lower or 'personnel' in stem_lower:
            return '项目人员资质'
        elif '业绩' in stem_lower or '合同' in stem_lower or 'performance' in stem_lower:
            return '类似业绩证明'
        elif '设备' in stem_lower or 'equipment' in stem_lower:
            return '设备证明'
        elif '办公' in stem_lower or '场所' in stem_lower or 'office' in stem_lower:
            return '办公场所证明'
        elif '备案' in stem_lower:
            return '平台备案信息'
        elif '分工' in stem_lower:
            return '项目工作组分工'
        else:
            return '附件材料'

    def _add_company_assets_section(
        self, doc: Document, company_images: List[tuple],
        available_images: Dict[str, List[str]],
    ):
        """Add a section with company asset images at the end of the document."""
        # Section heading
        p = doc.add_paragraph()
        try:
            p.style = doc.styles['Heading 1']
        except Exception:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run('附件：投标人资格证明材料')
        r.font.name = FONT_H1
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H1)
        r.font.size = Pt(16)
        r.bold = True

        # Description
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        r = p.add_run('以下为投标人（江苏众拓项目代理咨询有限公司）相关资格证书、证明材料扫描件：')
        r.font.name = FONT_BODY
        r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        r.font.size = Pt(12)

        doc.add_paragraph()

        # Insert each company image with caption
        inserted_count = 0
        for img_path, caption in company_images:
            if self._insert_image_from_ref(doc, img_path, caption, available_images):
                inserted_count += 1

        if inserted_count == 0:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run('（暂无资格证明材料图片，请上传相关资料）')
            r.font.name = FONT_CAPTION
            r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CAPTION)
            r.font.size = Pt(10)
            r.italic = True
