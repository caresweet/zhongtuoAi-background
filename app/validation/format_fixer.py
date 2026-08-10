"""Format Fixer — applies spec-based font/size/alignment corrections to generated docx.

Runs after template fill, before finalization. Reads the format specification
document and applies corrections to fix common issues:
- Headings: Times New Roman → 宋体/黑体/仿宋
- Body text: wrong size → 四号 (14pt)
- Figure/table captions: wrong font/size → 宋体小四

Target spec rules (from 社会稳定评估报告格式):
- 封面: 黑体 小二 居中
- 1级标题: 宋体 小三 加黑 居中
- 2级标题: 宋体 四号 加黑 靠左
- 3级标题: 仿宋 四号 加黑 靠左
- 正文: 仿宋 四号 首行缩进2格
- 图名: 宋体 小四 居中
- 表名: 宋体 小四 居中
- 表内容: 仿宋小四
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# Chinese font size → EMU (English Metric Units, used by python-docx)
SIZE_MAP: Dict[str, Pt] = {
    "小二": Pt(18),
    "小三": Pt(15),
    "四号": Pt(14),
    "小四": Pt(12),
    "五号": Pt(10.5),
    "小五": Pt(9),
    "小一": Pt(24),
    "一号": Pt(26),
    "二号": Pt(22),
    "三号": Pt(16),
}

# Alignment mapping
ALIGN_MAP: Dict[str, int] = {
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "靠左": WD_ALIGN_PARAGRAPH.LEFT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ── Spec Rules (hardcoded from the 格式规范 document for reliability) ──

SPEC_RULES = [
    # (target_pattern, font, size, bold, alignment, indent)
    # target_pattern is matched against paragraph text or style name

    # Cover (first 15 paragraphs) — per 文字格式: 封面: 小二，黑体，居中
    {"target": "cover_title",        "font": "黑体", "size": "小二", "bold": False, "align": "居中"},
    {"target": "cover_subtitle",     "font": "黑体", "size": "小一", "bold": False, "align": "居中"},

    # Heading levels — per 文字格式:
    #   1级标题: 宋体 小三 加黑 居中
    #   2级标题(2级表题): 宋体 四号 加黑 靠左
    #   3级标题: 仿宋 四号 加黑 靠左
    {"target": "heading_1",          "font": "宋体", "size": "小三", "bold": True,  "align": "居中"},
    {"target": "heading_2",          "font": "宋体", "size": "四号", "bold": True,  "align": "靠左"},
    {"target": "heading_3",          "font": "仿宋", "size": "四号", "bold": True,  "align": "靠左"},

    # Body — per 文字格式: 正文: 仿宋，四号，首行缩进2格
    {"target": "body_text",          "font": "仿宋", "size": "四号", "bold": False, "align": None, "indent": True},
    # 目录: 宋体，四号，加黑，居中
    {"target": "toc",                "font": "宋体", "size": "四号", "bold": True,  "align": "居中"},

    # Table headings — per 文字格式:
    #   1级表题: 宋体，小四，加黑，靠左
    #   2级表题: 宋体，5号，靠左，缩进2格
    {"target": "1级表题",            "font": "宋体", "size": "小四", "bold": True,  "align": "靠左"},
    {"target": "2级表题",            "font": "宋体", "size": "五号", "bold": False, "align": "靠左", "indent": True},

    # Figure & Table captions — per 文字格式:
    #   图名: 宋体，小四，居中
    #   表名: 宋体，小四，居中
    #   表内容: 仿宋小四（第一行加黑）
    {"target": "图名",               "font": "宋体", "size": "小四", "bold": False, "align": "居中"},
    {"target": "表名",               "font": "宋体", "size": "小四", "bold": False, "align": "居中"},
    {"target": "表内容",             "font": "仿宋", "size": "小四", "bold": False, "align": None},

    # Review form — per 文字格式:
    #   评审表封面: 黑体，小一，居中
    #   评审表内容: 黑体，小三，两端对齐
    {"target": "评审表封面",         "font": "黑体", "size": "小一", "bold": False, "align": "居中"},
    {"target": "评审表内容",         "font": "黑体", "size": "小三", "bold": False, "align": "两端对齐"},
]


class FormatFixer:
    """Applies format corrections to a generated docx based on spec rules."""

    @classmethod
    def fix(cls, docx_path: str) -> str:
        """Fix fonts, sizes, alignment in the generated docx.

        Args:
            docx_path: Path to the generated .docx file (modified in-place).

        Returns:
            The same path (for chaining).
        """
        doc = Document(docx_path)

        stats = {"fixed_headings": 0, "fixed_body": 0, "fixed_captions": 0}

        for i, p in enumerate(doc.paragraphs):
            style_name = p.style.name if p.style else ""
            text = p.text.strip()

            # ── Heading 1 ──
            if style_name.startswith("Heading 1") or style_name.startswith("heading 1"):
                if cls._apply_rule(p, "heading_1"):
                    stats["fixed_headings"] += 1

            # ── Heading 2 ──
            elif style_name.startswith("Heading 2") or style_name.startswith("heading 2"):
                if cls._apply_rule(p, "heading_2"):
                    stats["fixed_headings"] += 1

            # ── Heading 3 ──
            elif style_name.startswith("Heading 3") or style_name.startswith("heading 3"):
                if cls._apply_rule(p, "heading_3"):
                    stats["fixed_headings"] += 1

            # ── Cover (first 5 paragraphs) ──
            elif i < 5 and text:
                if cls._apply_rule(p, "cover_title"):
                    stats["fixed_headings"] += 1

            # ── Figure captions (图N-N ...) ──
            elif cls._is_figure_caption(text):
                if cls._apply_rule(p, "图名"):
                    stats["fixed_captions"] += 1

            # ── Table captions (表N ... or 表格...) ──
            elif cls._is_table_caption(text):
                if cls._apply_rule(p, "表名"):
                    stats["fixed_captions"] += 1

            # ── Body text (all non-heading, non-cover, non-caption paragraphs) ──
            elif text and i > 15:
                # Skip figure/table captions (already handled above)
                if cls._is_figure_caption(text) or cls._is_table_caption(text):
                    continue
                if cls._apply_rule(p, "body_text"):
                    stats["fixed_body"] += 1

        # ── Fix table cell text ──
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text.strip()
                        if not text:
                            continue
                        # First row = header (bold), rest = content (normal)
                        is_header = row == table.rows[0]
                        if is_header:
                            if cls._apply_rule(p, "1级表题"):
                                stats["fixed_captions"] += 1
                        else:
                            if cls._apply_rule(p, "表内容"):
                                stats["fixed_captions"] += 1

        # ── Remove blank pages caused by empty section breaks ──
        # NOTE: Temporarily disabled — blank section removal shifts paragraph
        # indices after template fill, causing content to land in wrong positions.
        # TODO: Re-enable after making it content-aware (skip filled paragraphs).
        # removed_sections = cls._remove_blank_sections(doc)
        # stats["fixed_headings"] += removed_sections * 100  # Track separately
        removed_sections = 0  # Disabled — see comment above

        # ── Apply East-Asian font fallback ──
        cls._set_east_asian_fonts(doc)

        doc.save(docx_path)

        total = sum(stats.values()) - removed_sections * 100 + removed_sections
        print(f"FormatFixer: fixed {total} paragraphs "
              f"(headings={stats['fixed_headings']}, "
              f"body={stats['fixed_body']}, "
              f"captions={stats['fixed_captions']}, "
              f"blank_sections={removed_sections})")

        return docx_path

    @classmethod
    def _apply_rule(cls, p, rule_key: str) -> bool:
        """Apply a single spec rule to a paragraph. Returns True if changes made."""
        rule = next((r for r in SPEC_RULES if r["target"] == rule_key), None)
        if not rule:
            return False

        changed = False

        for run in p.runs:
            if not run.text.strip():
                continue

            # Font
            if rule.get("font"):
                run.font.name = rule["font"]
                changed = True

            # Size
            if rule.get("size"):
                pt = SIZE_MAP.get(rule["size"])
                if pt:
                    run.font.size = pt
                    changed = True

            # Bold
            if rule.get("bold") is not None:
                run.font.bold = rule["bold"]
                changed = True

        # Alignment (paragraph-level)
        if rule.get("align"):
            align_val = ALIGN_MAP.get(rule["align"])
            if align_val is not None:
                p.alignment = align_val
                changed = True

        # Indent (first-line indent for body text)
        if rule.get("indent"):
            pf = p.paragraph_format
            if pf.first_line_indent is None:
                pf.first_line_indent = Cm(0.74)  # ~2 Chinese characters
                changed = True

        return changed

    @classmethod
    def _set_east_asian_fonts(cls, doc: Document):
        """Set East-Asian font fallback in document-level styles.

        This ensures that even if the Western font is set to something like
        Times New Roman, the CJK characters use the correct font.
        """
        for p in doc.paragraphs:
            for run in p.runs:
                if not run.text.strip():
                    continue
                font_name = run.font.name
                if not font_name:
                    continue

                # For CJK fonts, set the East-Asian font name too
                if any(cjk in font_name for cjk in ["宋体", "黑体", "仿宋", "楷体"]):
                    try:
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            from lxml import etree
                            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
                        rFonts.set(qn('w:eastAsia'), font_name)
                    except Exception:
                        pass  # Best effort

    @classmethod
    def _is_figure_caption(cls, text: str) -> bool:
        """Check if text looks like a figure caption (图1-1, 图1, Figure 1, etc.)."""
        return bool(
            re.match(r'图\s*\d+', text) or
            re.match(r'Figure\s*\d+', text, re.IGNORECASE) or
            '决策位置' in text or
            '公示内容' in text or
            '展示图' in text
        )

    @classmethod
    def _remove_blank_sections(cls, doc: Document) -> int:
        """Remove section breaks that create blank pages.

        Two fixes applied:
        1. Merge consecutive empty sections (both ending with sectPr) by removing
           the first section's sectPr — eliminates blank pages between sections.
        2. Remove leading empty paragraphs at the start of each section to
           reduce unnecessary blank space after forced page breaks.

        Returns count of blank sections merged.
        """
        from lxml import etree

        removed = 0
        body = doc.element.body
        para_elements = body.findall(qn('w:p'))
        if not para_elements:
            return 0

        # ── Fix 1: Merge consecutive empty sections ──
        # When two consecutive paragraphs BOTH end with sectPr and BOTH are
        # empty, the first section is a blank page. Remove its sectPr to merge.
        for i, p_elem in enumerate(para_elements):
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                continue
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is None:
                continue

            # Is this paragraph empty?
            texts = p_elem.findall('.//' + qn('w:t'))
            if any(t.text and t.text.strip() for t in texts if t.text):
                continue  # Has content — not a blank section

            # Check the next paragraph
            if i + 1 >= len(para_elements):
                continue
            next_p = para_elements[i + 1]
            next_pPr = next_p.find(qn('w:pPr'))
            if next_pPr is None:
                continue
            next_sectPr = next_pPr.find(qn('w:sectPr'))
            if next_sectPr is None:
                continue

            # Both this and next paragraph end sections AND are empty
            # → remove this one's sectPr to merge sections
            pPr.remove(sectPr)
            removed += 1

        # ── Fix 2: Remove leading empty paragraphs in sections ──
        # After a section break (NEW_PAGE), empty paragraphs before the first
        # content create unnecessary blank space at the top of the new page.
        in_new_section = True
        for p_elem in list(para_elements):
            texts = p_elem.findall('.//' + qn('w:t'))
            has_text = any(t.text and t.text.strip() for t in texts if t.text)

            pPr = p_elem.find(qn('w:pPr'))
            has_sect = pPr is not None and pPr.find(qn('w:sectPr')) is not None

            if has_text:
                in_new_section = False
                continue

            if has_sect:
                # Section-ending paragraph — next paragraph starts a new section
                in_new_section = True
                continue

            if in_new_section and not has_text:
                # Empty leading paragraph in a section — remove it
                body.remove(p_elem)

        return removed

    @classmethod
    def _is_table_caption(cls, text: str) -> bool:
        """Check if text looks like a table caption (表1, 表格, 汇总表, etc.)."""
        return bool(
            re.match(r'表\s*\d+', text) or
            re.match(r'表格\s*\d+', text) or
            '汇总表' in text or
            '等级表' in text
        )


# ── Convenience ──

def fix_report_format(docx_path: str) -> str:
    """Apply format corrections to a generated report.

    Args:
        docx_path: Absolute path to the .docx file.

    Returns:
        The same path.
    """
    if not Path(docx_path).exists():
        raise FileNotFoundError(f"Report not found: {docx_path}")

    return FormatFixer.fix(docx_path)
