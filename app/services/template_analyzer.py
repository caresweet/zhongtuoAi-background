"""Template Analysis Service — deep analysis of report .docx templates.

Extracts the complete structure of a social stability risk assessment report:
- Section hierarchy and cross-references
- Table definitions and data source mappings
- Preserved content (company info, certificates, seals)
- User-fillable fields (classified by category)
- Auto-generated sections (RAG-powered analysis)
- Location-dependent content (for global replace)
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

from docx import Document


# ═══════════════════════════════════════════════════════════════════════════════
# Data Classes
# ═══════════════════════════════════════════════════════════════════════════════


@dataclass
class SectionInfo:
    """Information about a report section."""
    number: str  # e.g. "1.1", "4.2"
    title: str   # e.g. "决策名称"
    heading_level: int  # 1, 2, or 3
    para_index: int     # paragraph index in document
    content: str = ""   # section body text
    tables_in_section: List[int] = field(default_factory=list)
    child_sections: List[str] = field(default_factory=list)  # sub-section numbers
    depends_on: List[str] = field(default_factory=list)  # cross-reference sections
    fill_strategy: str = "ask_user"  # ask_user | auto_generate | preserved | derived


@dataclass
class TableInfo:
    """Information about a table in the report."""
    index: int
    rows: int
    cols: int
    headers: List[str]
    section: str  # parent section number
    data_source: str  # user_input | derived_from_sections | auto_generated | preserved
    source_sections: List[str] = field(default_factory=list)  # sections that feed this table
    sample_rows: List[List[str]] = field(default_factory=list)


@dataclass
class PreservedContent:
    """Content that must be preserved exactly as-is from the template."""
    category: str  # company_info | personnel | certificate | administrative | seal | legal_text
    identifier: str  # unique identifier
    content_preview: str  # first 120 chars
    para_indices: List[int]  # paragraph indices
    reason: str = ""  # why this is preserved


@dataclass
class TemplateAnalysis:
    """Complete analysis of a report template."""
    template_path: str
    template_name: str = ""
    total_paragraphs: int = 0
    total_tables: int = 0
    sections: List[SectionInfo] = field(default_factory=list)
    tables: List[TableInfo] = field(default_factory=list)
    preserved_content: List[PreservedContent] = field(default_factory=list)
    location_names: Set[str] = field(default_factory=set)  # all location names to replace
    company_names: Set[str] = field(default_factory=set)    # all company names to preserve/replace
    section_keyword_map: Dict[str, List[str]] = field(default_factory=dict)


# ═══════════════════════════════════════════════════════════════════════════════
# Preserved Content Patterns — content that must NEVER be modified
# ═══════════════════════════════════════════════════════════════════════════════

PRESERVED_PATTERNS = [
    # Company identity — preserved for all reports using this template
    {
        "category": "company_info",
        "keywords": ["江苏众拓项目代理咨询有限公司"],
        "reason": "稳评第三方机构名称，所有报告共用同一家评估公司",
    },
    {
        "category": "personnel",
        "keywords": ["陈春"],
        "full_names": ["陈春(总经理、高级工程师、估价师、经济师)"],
        "reason": "稳评负责人（董事长），固定人员",
    },
    {
        "category": "personnel",
        "keywords": ["程诗茹"],
        "full_names": ["程诗茹(评估专业人员)"],
        "reason": "稳评联系人，固定人员",
    },
    {
        "category": "personnel",
        "keywords": ["18252573739"],
        "reason": "联系人电话号码，固定",
    },
    {
        "category": "certificate",
        "keywords": ["公司营业执照"],
        "reason": "公司营业执照图片位置，保留不修改",
    },
    {
        "category": "certificate",
        "keywords": ["稳评平台备案"],
        "reason": "稳评平台备案信息，保留不修改",
    },
    {
        "category": "certificate",
        "keywords": ["人员证书"],
        "reason": "人员资质证书，保留不修改",
    },
    {
        "category": "administrative",
        "keywords": ["项目工作组人员及分工情况"],
        "reason": "工作组人员架构，固定格式保留",
    },
    {
        "category": "administrative",
        "keywords": ["编制说明"],
        "reason": "编制说明章节，标准文本保留",
    },
    {
        "category": "administrative",
        "keywords": [
            "本表系江苏省稳评工作规范化运作统一制式",
            "事项责任单位、稳评责任单位、稳评实施单位、稳评结论备案单位各执1份",
        ],
        "reason": "标准表尾说明文字，保留不修改",
    },
    {
        "category": "administrative",
        "keywords": ["备案情况是指稳评责任单位将稳评结论"],
        "reason": "备案流程说明，标准文本保留",
    },
    {
        "category": "seal",
        "keywords": ["年月日"],
        "reason": "日期签章位置，保留空白由用户填写",
    },
    {
        "category": "seal",
        "keywords": ["（盖章）", "（公章）", "（签章）"],
        "reason": "签章位置，保留格式",
    },
    # Key paragraph ranges that are template boilerplate
    {
        "category": "legal_text",
        "keywords": ["《关于加强新形势下重大决策社会稳定风险评估机制建设实施意见》"],
        "reason": "法规依据列表，标准文本，根据地区可能需要微调",
    },
]

# ═══════════════════════════════════════════════════════════════════════════════
# Section Analysis — maps sections to their data sources and fill strategies
# ═══════════════════════════════════════════════════════════════════════════════

SECTION_STRATEGIES = {
    # Chapter 1: Basic info — mostly user input
    "1.1": {"strategy": "ask_user", "fields": ["决策名称", "报告标题", "项目名称"]},
    "1.2": {"strategy": "ask_user", "fields": ["决策单位", "决策主体"]},
    "1.3": {"strategy": "ask_user", "fields": ["拟征地位置", "地理位置"], "needs_image": True},
    "1.4": {"strategy": "ask_user", "fields": ["征收范围", "面积", "地上附着物", "地类"]},
    "1.5": {"strategy": "ask_user", "fields": ["资金筹措", "投资金额"]},
    "1.6": {"strategy": "ask_user", "fields": ["实施周期", "工期"]},
    "1.7": {"strategy": "ask_user", "fields": ["红线图", "征地红线图"], "needs_image": True},

    # Chapter 2: Assessment process — partially user, partially preserved
    "2.1": {"strategy": "ask_user", "fields": ["评估过程", "评估方法"]},
    "2.2": {"strategy": "preserved", "fields": []},  # 法规依据列表

    # Chapter 3: Survey — user provides data/images, agent analyzes
    "3.1": {"strategy": "preserved", "fields": []},  # 调查方法（标准描述）
    "3.2": {"strategy": "ask_user", "fields": ["调查范围", "调查对象"]},
    "3.3": {"strategy": "auto_generate", "fields": [], "agent": "SurveyAnalyzer",
            "depends_on": ["3.2", "survey_images"]},  # 调查内容表格 — from image analysis
    "3.4": {"strategy": "auto_generate", "fields": [], "agent": "SurveyAnalyzer",
            "depends_on": ["3.3"]},  # 调查过程
    "3.5": {"strategy": "auto_generate", "fields": [], "agent": "SurveyAnalyzer",
            "depends_on": ["3.3", "3.4"]},  # 意见建议
    "3.6": {"strategy": "ask_user", "fields": ["利益相关者", "诉求"]},

    # Chapter 4: Four-dimensional analysis — ALL auto-generated via RAG
    "4.1": {"strategy": "auto_generate", "agent": "RationalityAgent", "rag_required": True},
    "4.2": {"strategy": "auto_generate", "agent": "RationalityAgent", "rag_required": True},
    "4.3": {"strategy": "auto_generate", "agent": "RationalityAgent", "rag_required": True},
    "4.4": {"strategy": "auto_generate", "agent": "RationalityAgent", "rag_required": True},

    # Chapter 5: Risk factor identification — auto-generated from survey + analysis
    "5.1": {"strategy": "auto_generate", "agent": "RiskScorer",
            "depends_on": ["3.0", "4.0"]},
    "5.2": {"strategy": "auto_generate", "agent": "RiskScorer",
            "depends_on": ["5.1"]},

    # Chapter 6: Risk scoring (pre-measures) — auto-generated
    "6.0": {"strategy": "auto_generate", "agent": "RiskScorer",
            "depends_on": ["5.0"]},

    # Chapter 7: Mitigation measures — auto-generated + user review
    "7.1": {"strategy": "preserved", "fields": []},  # 基本原则（标准文本）
    "7.2": {"strategy": "auto_generate", "agent": "RiskScorer",
            "depends_on": ["6.0"]},

    # Chapter 8: Post-measure risk assessment — auto-generated
    "8.1": {"strategy": "preserved", "fields": []},  # 分析方法（标准描述）
    "8.2": {"strategy": "auto_generate", "agent": "RiskScorer",
            "depends_on": ["7.0"]},

    # Chapter 9: Conclusions — auto-generated + user review
    "9.1": {"strategy": "auto_generate", "agent": "RiskScorer", "depends_on": ["8.0"]},
    "9.2": {"strategy": "auto_generate", "agent": "RiskScorer", "depends_on": ["8.0"]},
    "9.3": {"strategy": "auto_generate", "agent": "RiskScorer", "depends_on": ["8.0"]},

    # Chapter 10: Emergency plan — preserved from template, can be user-customized
    "10.0": {"strategy": "preserved", "fields": []},
}

# ═══════════════════════════════════════════════════════════════════════════════
# Table Data Source Mappings
# ═══════════════════════════════════════════════════════════════════════════════

TABLE_DATA_SOURCES = {
    # Table 0 (22r x 7c): 稳评事项基本情况表 — pulls from all sections
    # This is the BIG review form table at the end of the report
    0: {
        "name": "稳评事项基本情况表（文章评审表）",
        "data_source": "derived_from_sections",
        "mappings": {
            "事项名称": "1.1",        # 决策名称
            "拟征地位置": "1.3",      # 位置
            "征收范围": "1.4",        # 面积、附着物
            "资金筹措": "1.5",        # 资金
            "决策实施周期": "1.6",    # 周期
            "稳评责任单位": "1.2",    # 责任单位
            "稳评实施单位": "1.2",    # 实施单位
            "合法性分析": "4.1",      # 合法性
            "合理性分析": "4.2",      # 合理性
            "可行性分析": "4.3",      # 可行性
            "可控性分析": "4.4",      # 可控性
            "风险等级": "6.0/8.0",   # 风险等级
            "评估结论": "9.1",        # 结论
        },
    },
    # Table 1 (5r x 4c): 调查层级/范围/对象/方法 — from user input (3.2)
    1: {
        "name": "风险调查范围及对象表",
        "data_source": "user_input",
        "section": "3.2",
    },
    # Table 2 (7r x 2c): 情形/结论 — preserved (standard risk criteria)
    2: {
        "name": "风险等级判定标准表",
        "data_source": "preserved",
    },
    # Table 3 (22r x 6c): 测评指标/权重/评分/标准/得分 — RiskScorer generates
    3: {
        "name": "措施前风险等级量化评分表",
        "data_source": "auto_generated",
        "agent": "RiskScorer",
    },
    # Table 4 (4r x 4c): 反对率/风险发生概率 scoring
    4: {
        "name": "反对率风险概率得分表",
        "data_source": "derived_from_sections",
        "source": "survey_results",
    },
    # Table 5 (22r x 6c): measure-before/after comparison
    5: {
        "name": "措施前后风险等级对比表",
        "data_source": "auto_generated",
        "agent": "RiskScorer",
    },
    # Table 6 (4r x 4c): post-measure probability scoring
    6: {
        "name": "措施后反对率风险概率得分表",
        "data_source": "derived_from_sections",
        "source": "post_measure_survey",
    },
}


# ═══════════════════════════════════════════════════════════════════════════════
# Analyzer Class
# ═══════════════════════════════════════════════════════════════════════════════

class TemplateAnalyzer:
    """Deep analysis of report .docx templates.

    Extracts section hierarchy, table metadata, preserved content,
    and cross-reference dependencies for intelligent agent guidance.
    """

    def __init__(self):
        self._cache: Dict[str, TemplateAnalysis] = {}

    def analyze(self, template_path: str) -> TemplateAnalysis:
        """Analyze a template .docx file and return structured analysis.

        Results are cached by path.
        """
        abs_path = str(Path(template_path).resolve())
        if abs_path in self._cache:
            return self._cache[abs_path]

        doc = Document(abs_path)
        analysis = TemplateAnalysis(template_path=abs_path)

        # Extract basic stats
        analysis.total_paragraphs = len(doc.paragraphs)
        analysis.total_tables = len(doc.tables)

        # Phase 1: Extract section hierarchy
        analysis.sections = self._extract_sections(doc)

        # Phase 2: Extract table info
        analysis.tables = self._extract_tables(doc, analysis.sections)

        # Phase 3: Identify preserved content
        analysis.preserved_content = self._find_preserved_content(doc)

        # Phase 4: Extract location names
        analysis.location_names = self._extract_location_names(doc)

        # Phase 5: Extract company names
        analysis.company_names = self._extract_company_names(doc)

        # Phase 6: Build section keyword map
        analysis.section_keyword_map = self._build_section_keyword_map(analysis)

        self._cache[abs_path] = analysis
        return analysis

    def _extract_sections(self, doc: Document) -> List[SectionInfo]:
        """Extract section hierarchy from document headings."""
        sections = []
        heading_stack: List[SectionInfo] = []  # parent sections

        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            if not style_name.startswith("Heading"):
                continue

            level_str = style_name.replace("Heading", "").strip()
            try:
                level = int(level_str)
            except ValueError:
                continue

            text = para.text.strip()
            if not text:
                continue

            # Parse section number from heading text
            number_match = re.match(r'(\d+(?:\.\d+)*)\s*(.*)', text)
            if number_match:
                number = number_match.group(1)
                title = number_match.group(2)
            else:
                number = text[:10]
                title = text

            section = SectionInfo(
                number=number,
                title=title,
                heading_level=level,
                para_index=i,
            )

            # Assign fill strategy
            strategy_info = self._get_section_strategy(number)
            section.fill_strategy = strategy_info.get("strategy", "ask_user")

            # Track parent-child relationships
            while heading_stack and heading_stack[-1].heading_level >= level:
                heading_stack.pop()

            if heading_stack:
                heading_stack[-1].child_sections.append(number)

            heading_stack.append(section)
            sections.append(section)

        return sections

    def _extract_tables(
        self, doc: Document, sections: List[SectionInfo]
    ) -> List[TableInfo]:
        """Extract table metadata with section association."""
        tables = []

        # Build para_index → section mapping
        section_at_para: Dict[int, SectionInfo] = {}
        for s in sections:
            section_at_para[s.para_index] = s

        for ti, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)

            # Extract headers from first row
            headers = [cell.text.strip().replace('\n', '')[:60]
                      for cell in table.rows[0].cells]

            # Determine which section this table belongs to
            # Look backwards from table's location in document
            section_num = "0"
            # Tables are associated with preceding headings
            for i in range(len(doc.paragraphs)):
                if doc.paragraphs[i].style and doc.paragraphs[i].style.name.startswith("Heading"):
                    # Find the closest heading before this table
                    pass

            # Use predefined table info if available
            table_info = TABLE_DATA_SOURCES.get(ti, {
                "name": f"表格{ti+1}",
                "data_source": "unknown",
            })

            # Extract sample rows
            sample_rows = []
            for ri in range(min(3, rows)):
                sample_rows.append([
                    cell.text.strip()[:50] for cell in table.rows[ri].cells
                ])

            tables.append(TableInfo(
                index=ti,
                rows=rows,
                cols=cols,
                headers=headers,
                section=table_info.get("section", ""),
                data_source=table_info.get("data_source", "unknown"),
                source_sections=table_info.get("mappings", {}).values() if "mappings" in table_info else [],
                sample_rows=sample_rows,
            ))

        return tables

    def _find_preserved_content(self, doc: Document) -> List[PreservedContent]:
        """Identify content that must be preserved in the generated report."""
        preserved = []

        for pattern in PRESERVED_PATTERNS:
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue

                # Check if any keyword matches
                matched = False
                for kw in pattern.get("keywords", []):
                    if kw in text:
                        matched = True
                        break

                if matched:
                    preserved.append(PreservedContent(
                        category=pattern["category"],
                        identifier=pattern.get("full_names", [pattern["keywords"][0]])[0][:80],
                        content_preview=text[:120],
                        para_indices=[i],
                        reason=pattern.get("reason", ""),
                    ))

        return preserved

    def _extract_location_names(self, doc: Document) -> Set[str]:
        """Extract all location names from the template for global replace.

        Looks for county/district/town/street level names that might change
        between different projects.
        """
        location_patterns = [
            r'金湖县', r'金北街道', r'金湖大道', r'戴楼街道',
            r'盱眙县', r'洪泽区',
        ]

        found = set()
        for para in doc.paragraphs:
            for pat in location_patterns:
                if pat in para.text:
                    found.add(pat)

        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for pat in location_patterns:
                        if pat in cell.text:
                            found.add(pat)

        return found

    def _extract_company_names(self, doc: Document) -> Set[str]:
        """Extract company/organization names for preservation."""
        company_patterns = [
            "江苏众拓项目代理咨询有限公司",
        ]
        found = set()
        for para in doc.paragraphs:
            for pat in company_patterns:
                if pat in para.text:
                    found.add(pat)
        return found

    def _get_section_strategy(self, section_num: str) -> dict:
        """Get the fill strategy for a section by number."""
        # Try exact match first
        if section_num in SECTION_STRATEGIES:
            return SECTION_STRATEGIES[section_num]

        # Try matching by major chapter number
        major = section_num.split(".")[0]
        for key, value in SECTION_STRATEGIES.items():
            if key.startswith(major + "."):
                # Return first match for this chapter as default
                return value

        return {"strategy": "ask_user"}

    def _build_section_keyword_map(
        self, analysis: TemplateAnalysis
    ) -> Dict[str, List[str]]:
        """Build a map of section numbers to keywords for content matching."""
        keyword_map = {}

        for section in analysis.sections:
            keywords = []
            title = section.title

            # Extract meaningful keywords from section title
            for word in re.split(r'[，,、\s]', title):
                word = word.strip()
                if len(word) >= 2:
                    keywords.append(word)

            # Add strategy-specific keywords
            strategy = SECTION_STRATEGIES.get(section.number, {})
            keywords.extend(strategy.get("fields", []))

            keyword_map[section.number] = list(set(keywords))

        return keyword_map

    def get_user_question_order(self, analysis: TemplateAnalysis) -> List[dict]:
        """Generate the recommended order of questions to ask the user.

        Returns a list of {section, question, field_type, needs_image} dicts
        in the optimal asking order.
        """
        questions = []

        # Priority 1: Basic project info (Chapter 1)
        basic_sections = [s for s in analysis.sections
                         if s.number.startswith("1.") and s.fill_strategy == "ask_user"]
        for s in basic_sections:
            strategy = SECTION_STRATEGIES.get(s.number, {})
            questions.append({
                "section": s.number,
                "title": s.title,
                "question": f"请提供{s.title}的信息",
                "fields": strategy.get("fields", []),
                "needs_image": strategy.get("needs_image", False),
                "priority": 1,
            })

        # Priority 2: Survey scope (3.2)
        survey_sections = [s for s in analysis.sections
                          if s.number in ("3.2", "3.6")]
        for s in survey_sections:
            questions.append({
                "section": s.number,
                "title": s.title,
                "question": f"请提供{s.title}的信息",
                "fields": SECTION_STRATEGIES.get(s.number, {}).get("fields", []),
                "needs_image": False,
                "priority": 2,
            })

        # Priority 3: Images (1.3 红线图, survey images, notice photos)
        questions.append({
            "section": "images",
            "title": "附图附件",
            "question": "请上传相关图片材料：拟征地位置图（红线图）、公示照片、问卷调查表照片",
            "fields": ["红线图", "公示照片", "问卷调查表"],
            "needs_image": True,
            "priority": 3,
            "multi_image": True,
        })

        return sorted(questions, key=lambda q: q["priority"])

    def get_preserved_summary(self, analysis: TemplateAnalysis) -> str:
        """Generate a human-readable summary of preserved content."""
        if not analysis.preserved_content:
            return ""

        by_category: Dict[str, List[PreservedContent]] = {}
        for pc in analysis.preserved_content:
            by_category.setdefault(pc.category, []).append(pc)

        lines = ["**以下内容保留模板原文，无需用户填写**："]
        category_labels = {
            "company_info": "🏢 公司信息",
            "personnel": "👤 人员信息",
            "certificate": "📜 资质证书",
            "administrative": "📋 管理文件",
            "seal": "🖊️ 签章日期",
            "legal_text": "⚖️ 法规依据",
        }

        for cat, items in by_category.items():
            label = category_labels.get(cat, cat)
            lines.append(f"\n{label}（{len(items)}项）：")
            for item in items[:3]:
                lines.append(f"  • {item.content_preview[:80]}")

        return "\n".join(lines)

    def get_table_summary(self, analysis: TemplateAnalysis) -> str:
        """Generate a summary of tables and their data sources."""
        lines = ["**报告表格及数据来源**："]

        for table in analysis.tables:
            if table.index in TABLE_DATA_SOURCES:
                info = TABLE_DATA_SOURCES[table.index]
                lines.append(
                    f"- 表格{table.index+1}「{info.get('name', '')}」"
                    f"（{table.rows}r×{table.cols}c）→ {info.get('data_source', '')}"
                )

        return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════════════
# Global Location Replace
# ═══════════════════════════════════════════════════════════════════════════════

class LocationReplacer:
    """Handles global replacement of location names in report content.

    When a user provides a new location (e.g., 盱眙县), this replaces all
    instances of the template location (e.g., 金湖县) across the report.
    """

    # Common location type patterns to replace — these are known template defaults
    LOCATION_PATTERNS = [
        (r'(金湖|洪泽|盱眙|涟水)县', '{prefix}县'),
        (r'(金北|戴楼|黎城|金南)街道', '{prefix}街道'),
        (r'(金湖|洪泽)大道', '{prefix}大道'),
    ]

    # All known county names in the template — used to detect "new" locations
    KNOWN_TEMPLATE_COUNTIES = {"金湖", "洪泽", "盱眙", "涟水"}
    KNOWN_TEMPLATE_STREETS = {"金北", "戴楼", "黎城", "金南"}
    KNOWN_TEMPLATE_ROADS = {"金湖大道", "洪泽大道"}

    @staticmethod
    def build_replace_map(
        new_location_prefix: str,  # e.g. "盱眙"
        new_street: str = "",       # e.g. "盱城"
        new_road: str = "",         # e.g. "盱眙"
        template_locations: set = None,  # known template locations to replace
    ) -> Dict[str, str]:
        """Build a comprehensive replacement map based on new location info.

        Automatically detects all known template locations and replaces them
        with the corresponding new names.

        Args:
            new_location_prefix: New county/city name (e.g. "盱眙")
            new_street: New street/subdistrict name
            new_road: New main road name
            template_locations: Set of known template location names (optional)

        Returns:
            Dict mapping old text → new text for global replacement.
        """
        replace_map = {}
        template_counties = template_locations or LocationReplacer.KNOWN_TEMPLATE_COUNTIES
        template_streets = LocationReplacer.KNOWN_TEMPLATE_STREETS

        if new_location_prefix:
            # Replace ALL known template counties with the new one
            for old_county in template_counties:
                if old_county != new_location_prefix:
                    replace_map[f"{old_county}县"] = f"{new_location_prefix}县"
                    # Also replace bare county name (without 县 suffix)
                    # But be careful — only replace when used as location context

        if new_street:
            for old_street in template_streets:
                if old_street != new_street:
                    replace_map[f"{old_street}街道"] = f"{new_street}街道"

        if new_road:
            # Derive new road name
            road_name = f"{new_road}大道"
            for old_road in LocationReplacer.KNOWN_TEMPLATE_ROADS:
                if old_road != road_name:
                    replace_map[old_road] = road_name

        return replace_map

    @staticmethod
    def apply_replacements(text: str, replace_map: Dict[str, str]) -> str:
        """Apply all replacements to a text string."""
        result = text
        for old, new in replace_map.items():
            result = result.replace(old, new)
        return result

    @classmethod
    def extract_location_from_input(cls, user_input: str) -> Optional[Dict[str, str]]:
        """Extract location from user input using simple heuristics.

        Priority: explicit suffixes (XX县/市/区) > after action verbs.
        Returns dict with 'county' key, or None.
        """
        import re
        result = {}

        # Strategy 1: "XX县", "XX市", "XX区" — exact 2 chars before suffix
        for suffix in ["县", "市", "区"]:
            m = re.search(rf'([一-鿿]{{2}}){suffix}', user_input)
            if m:
                candidate = m.group(1)
                # "XX片区", "XX园区" contain "区" but are not district names
                if suffix == "区" and ("片区" in m.group(0) or "园区" in m.group(0)
                                       or "工业园" in m.group(0) or "开发区" in m.group(0)):
                    continue
                # Exclude non-location words
                _BAD = {"某", "本", "该", "我", "你", "全", "各", "所在", "拟征", "工业"}
                if candidate not in _BAD:
                    result["county"] = candidate
                    return result

        # Strategy 2: After action verb — "生成/写/做XX（稳评/报告/...）"
        # Pattern: verb(2-4 chars) + location(2 chars) + project_keyword(2+ chars)
        m = re.search(
            r'(?:生成|写|做|编制|帮我生成|帮我写|帮我做|给我|帮我)'
            r'([一-鿿]{2})'
            r'(?:稳评|社会|征地|土地|风险|报告|片区|地区|辖区|项目|工业|园区)',
            user_input
        )
        if m:
            candidate = m.group(1)
            _VERBS = {"生成", "稳评", "社会", "征地", "风险", "报告", "土地", "项目",
                      "评估", "决策", "高铁", "帮我", "给我", "一个", "一份"}
            if candidate not in _VERBS:
                result["county"] = candidate
                return result

        # Strategy 3: "XX街道"
        m = re.search(r'([一-鿿]{2})街道', user_input)
        if m:
            result["street"] = m.group(1)

        return result if result else None


# ═══════════════════════════════════════════════════════════════════════════════
# Singleton
# ═══════════════════════════════════════════════════════════════════════════════

template_analyzer = TemplateAnalyzer()
location_replacer = LocationReplacer()
