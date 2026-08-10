"""Stability Report Template Filler — 用模板填空 + AI分析章节模式生成稳评报告。

比纯AI生成的报告质量更高，因为:
1. 保留模板的专业排版、格式、固定话术
2. 只替换项目特定数据和图片
3. AI只用于分析性章节（合法性分析、风险措施、应急预案）
"""

import os, re, asyncio, logging
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ── Placeholder → 项目数据字段映射 ──
TEXT_PLACEHOLDER_MAP = {
    "{{project_name}}": "project_name",
    "{{project_name_full}}": "project_name",
    "{{project_name_short}}": "doc_reference",
    "{{decision_name}}": "project_name",
    "{{location_prefecture}}": "land_location",
    "{{location_pref_short}}": "land_location",
    "{{report_year}}": None,  # current year
    "{{report_year_cn}}": None,  # 二〇二六年
    "{{remove_ninglian}}": None,  # remove text
    "{{remove_s350}}": None,  # remove text
    "{{implementation_bg}}": "implement_unit",
}

# ── 图片占位符 → 稳评资料图片类别映射 ──
IMAGE_PLACEHOLDER_MAP = {
    "{{img_public_notice_1}}": ("public_notices", 0),
    "{{img_public_notice_2}}": ("public_notices", 1),
    "{{img_public_notice_3}}": ("public_notices", 2),
    "{{img_public_notice_4}}": ("public_notices", 3),
    "{{img_public_notice_5}}": ("public_notices", 4),
    "{{img_site_photo_1}}": ("site_photos", 0),
    "{{img_site_photo_2}}": ("site_photos", 1),
    "{{img_site_photo_3}}": ("site_photos", 2),
    "{{img_meeting_photo_1}}": ("meeting_photos", 0),
    "{{img_meeting_photo_2}}": ("meeting_photos", 1),
    "{{img_location_map_1}}": ("location_maps", 0),
    "{{img_location_map_2}}": ("location_maps", 1),
    "{{img_location_map_3}}": ("location_maps", 2),
    "{{img_survey_scan_1}}": ("survey_questionnaires", 0),
    "{{img_expert_review_1}}": ("expert_opinions", 0),
    "{{img_expert_review_2}}": ("expert_opinions", 1),
    "{{img_expert_review_3}}": ("expert_opinions", 2),
}

# ── 分析性章节占位符 — 需要 AI 生成 ──
AI_ANALYSIS_PLACEHOLDERS = {
    "{{legality_analysis}}": (
        "合法性分析",
        "请根据以下项目信息，撰写征地项目的合法性分析（200-300字），从征收主体资格、征收目的、规划符合性、程序合规性四个方面分析：\n"
        "项目：{project_name}\n"
        "位置：{land_location}\n"
        "面积：{land_area_mu} 亩（{land_area_sqm} 平方米）\n"
        "文号：{doc_reference}\n"
        "责任单位：{decision_unit}\n"
        "征收目的：{purpose_text}\n"
        "请用正式公文风格，引用土地管理法相关条款。"
    ),
    "{{risk_mitigation}}": (
        "风险防范与化解措施",
        "请根据以下项目信息，撰写风险防范与化解措施（300-400字），覆盖补偿方案、资金分配、社保落实、信访舆情四个方向：\n"
        "项目：{project_name}\n"
        "位置：{land_location}\n"
        "涉及农户：{household_count} 户\n"
        "补偿标准：{compensation_standard}\n"
        "支持率：{support_rate}%\n"
        "请明确每项措施的责任主体和可执行方案。"
    ),
    "{{emergency_plan}}": (
        "应急预案",
        "请根据以下项目信息，撰写突发事件应急预案（400-500字），包括：\n"
        "1. 编制目的和依据\n2. 适用范围\n3. 组织领导体系\n4. 预警预防机制\n5. 现场处置方案\n6. 舆情处置\n7. 保障措施\n8. 奖惩机制\n"
        "项目：{project_name}\n"
        "位置：{land_location}\n"
        "责任单位：{decision_unit}\n"
    ),
}


class StabilityTemplateFiller:
    """用洞庭湖稳评模板填空生成报告。"""

    def __init__(self):
        from app.config import settings
        self.storage_dir = settings.STORAGE_DIR
        self.images_dir = self.storage_dir / "images"
        self.output_dir = self.storage_dir / "generated"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.template_path = self._find_template()

    def _find_template(self) -> Optional[Path]:
        """找到洞庭湖稳评模板。"""
        import sqlite3
        from app.config import settings
        db_path = settings.DATA_DIR / "knowledge_base.db"
        if not db_path.exists():
            return None
        conn = sqlite3.connect(str(db_path))
        conn.row_factory = sqlite3.Row
        cur = conn.execute(
            "SELECT template_file_path FROM templates "
            "WHERE domain='stability' AND is_active=1 "
            "AND template_file_path LIKE '%.docx' "
            "AND name LIKE '%洞庭湖%' "
            "ORDER BY id LIMIT 1"
        )
        row = cur.fetchone()
        conn.close()
        if row:
            tpl = settings.STORAGE_DIR / row["template_file_path"]
            if tpl.exists():
                return tpl
        return None

    def fill(self, project_data: dict, categorized_images: dict,
             llm_service=None) -> str:
        """填充模板生成报告。

        Args:
            project_data: 项目结构化数据
            categorized_images: 分类后的图片 {category: [path, ...]}
            llm_service: LLM服务（用于生成分析章节）

        Returns:
            相对路径如 "generated/xxx.docx"
        """
        if not self.template_path:
            raise FileNotFoundError("未找到洞庭湖稳评模板")

        doc = Document(str(self.template_path))

        # Step 1: 全文文字替换（项目名称、地点、年份等）
        self._global_text_replace(doc, project_data)

        # Step 2: 替换占位符文字
        self._fill_text_placeholders(doc, project_data)

        # Step 3: 替换图片占位符
        self._fill_image_placeholders(doc, categorized_images)

        # Step 4: 用 AI 生成分析性章节
        if llm_service:
            self._fill_ai_analysis(doc, project_data, llm_service)

        # Step 5: 清理残余占位符
        self._cleanup_remaining(doc)

        # Save
        doc_ref = project_data.get("doc_reference", "report")
        safe_name = re.sub(r'[\\/:*?"<>|]', '_', doc_ref)[:50]
        filename = f"{safe_name}_社会稳定风险评估报告.docx"
        outpath = self.output_dir / filename

        # Ensure unique name
        counter = 1
        while outpath.exists():
            filename = f"{safe_name}_社会稳定风险评估报告_{counter}.docx"
            outpath = self.output_dir / filename
            counter += 1

        doc.save(str(outpath))
        logger.info(f"Template-filled report: {outpath}")
        return f"generated/{filename}"

    # ═══════════════════════════════════════════════════════════════
    # Text Replacement
    # ═══════════════════════════════════════════════════════════════

    def _global_text_replace(self, doc: Document, data: dict):
        """全局文字替换 — 在段落和表格中替换项目特定文字。"""
        import datetime

        # Old → New mapping
        replaces = {}

        # Project name
        old_name = "金征预告〔2026〕3号（高铁枢纽北片区开发地块项目）土地征收"
        new_name = data.get("project_name", "")
        if new_name:
            replaces[old_name] = new_name
            # Also try other known old values from the template
            old_alt = "金征预告〔2026〕3号"
            if old_alt != new_name:
                replaces[old_alt] = new_name

        # Location
        old_loc = "淮安市洪泽区朱坝街道、高良涧街道"
        new_loc = data.get("land_location", "")
        if new_loc:
            replaces[old_loc] = new_loc[:len(old_loc)]

        # Year
        year = data.get("report_year") or str(datetime.datetime.now().year)
        replaces["2024"] = year

        # Company (should stay the same, but ensure)
        replaces["江苏众拓项目代理咨询有限公司"] = "江苏众拓项目代理咨询有限公司"

        # Apply to paragraphs
        for p in doc.paragraphs:
            for run in p.runs:
                text = run.text or ""
                for old, new in replaces.items():
                    if old in text and old != new:
                        run.text = text.replace(old, new)
                        text = run.text

        # Apply to tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            text = run.text or ""
                            for old, new in replaces.items():
                                if old in text and old != new:
                                    run.text = text.replace(old, new)
                                    text = run.text

    def _fill_text_placeholders(self, doc: Document, data: dict):
        """替换文字占位符。"""
        import datetime

        year = str(datetime.datetime.now().year)
        year_cn_map = {
            "2024": "二〇二四", "2025": "二〇二五", "2026": "二〇二六",
            "2027": "二〇二七", "2028": "二〇二八",
        }

        text_values = {}
        for ph, field in TEXT_PLACEHOLDER_MAP.items():
            clean_ph = ph.strip("{}")
            if field and field in data:
                text_values[ph] = str(data[field])
            elif ph == "{{report_year}}":
                text_values[ph] = year
            elif ph == "{{report_year_cn}}":
                text_values[ph] = year_cn_map.get(year, f"二〇二六")
            elif ph in ("{{remove_ninglian}}", "{{remove_s350}}"):
                text_values[ph] = ""  # Remove these markers

        # Replace in paragraphs
        for p in doc.paragraphs:
            for run in p.runs:
                text = run.text or ""
                for ph, val in text_values.items():
                    if ph in text:
                        run.text = text.replace(ph, val)
                        text = run.text

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            text = run.text or ""
                            for ph, val in text_values.items():
                                if ph in text:
                                    run.text = text.replace(ph, val)
                                    text = run.text

        # Also handle {{implementation_bg}} specially — replace with company intro
        impl_val = text_values.pop("{{implementation_bg}}", None)

    def _fill_image_placeholders(self, doc: Document, categorized_images: dict):
        """替换图片占位符为实际项目图片。"""
        for ph, (category, index) in IMAGE_PLACEHOLDER_MAP.items():
            images = categorized_images.get(category, [])
            if index >= len(images):
                continue  # Not enough images

            img_path = images[index]
            if isinstance(img_path, dict):
                img_path = img_path.get("path", "")

            # Resolve the image path
            from pathlib import Path as _Path
            abs_path = None
            if img_path.startswith("/"):
                abs_path = _Path(img_path)
            elif "/" in img_path:
                abs_path = self.storage_dir / img_path
            else:
                abs_path = self.images_dir / img_path

            if not abs_path or not abs_path.exists():
                continue

            # Find the placeholder text in the document and replace with image
            self._replace_placeholder_with_image(doc, ph, str(abs_path))

    def _replace_placeholder_with_image(self, doc: Document, placeholder: str, image_path: str):
        """在文档中找到占位符文本，替换为图片。"""
        for p in doc.paragraphs:
            for run in p.runs:
                if placeholder in (run.text or ""):
                    # Clear the placeholder text
                    run.text = run.text.replace(placeholder, "")
                    # Add image in a new run
                    try:
                        img_run = p.add_run()
                        img_run.add_picture(image_path, width=Inches(4.0))
                    except Exception as e:
                        logger.warning(f"Failed to insert image {placeholder}: {e}")
                    return

    def _fill_ai_analysis(self, doc: Document, data: dict, llm_service):
        """用 LLM 生成分析性章节，替换对应占位符。"""
        import asyncio as _asyncio

        for ph, (title, prompt_template) in AI_ANALYSIS_PLACEHOLDERS.items():
            try:
                # Build prompt with project data
                prompt = prompt_template.format(
                    project_name=data.get("project_name", ""),
                    land_location=data.get("land_location", ""),
                    land_area_mu=data.get("land_area_mu", ""),
                    land_area_sqm=data.get("land_area_sqm", ""),
                    doc_reference=data.get("doc_reference", ""),
                    decision_unit=data.get("decision_unit", ""),
                    purpose_text=data.get("purpose_text", ""),
                    household_count=data.get("household_count", ""),
                    compensation_standard=data.get("compensation_standard", ""),
                    support_rate=data.get("support_rate", ""),
                )

                # Call LLM
                result = _asyncio.run(
                    llm_service.chat_with_reasoning(
                        messages=[{"role": "user", "content": prompt}],
                        system=f"你是社会稳定风险评估专家，请撰写「{title}」章节。用正式公文风格，200-500字。",
                        max_tokens=1024,
                        temperature=0.3,
                    )
                )
                content = result.get("content", "")
                if not content:
                    continue

                # Clean up
                content = re.sub(r'```[^`]*```', '', content)
                content = content.strip()

                # Replace placeholder in document
                for p in doc.paragraphs:
                    for run in p.runs:
                        if ph in (run.text or ""):
                            run.text = run.text.replace(ph, content)
                            logger.info(f"AI filled: {title} ({len(content)} chars)")
                            return

            except Exception as e:
                logger.warning(f"AI fill failed for {title}: {e}")

    def _cleanup_remaining(self, doc: Document):
        """清理残留的占位符。"""
        for p in doc.paragraphs:
            for run in p.runs:
                text = run.text or ""
                # Replace remaining {{...}} with 【待补充】
                text = re.sub(r'\{\{[\w_]+\}\}', '【待补充】', text)
                # Remove empty image-only paragraphs
                if text != run.text:
                    run.text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            text = run.text or ""
                            text = re.sub(r'\{\{[\w_]+\}\}', '【待补充】', text)
                            if text != run.text:
                                run.text = text


# Singleton
stability_template_filler = StabilityTemplateFiller()
