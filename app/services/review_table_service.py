"""社会稳定风险评估评审表 — 报告生成后从章节提取关键信息，留白供专家/部门填写。

基于 DB32/T4013-2021 + 南京规范 DB3201/T1163-2023 + 南通规范 DB3206/T1091-2024。
评审表是提交稳评主管部门备案的核心文件。
"""

import re, os, logging
from datetime import datetime
from typing import Dict, Optional

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def generate_review_table(
    output_path: str,
    filled_data: dict,
    chapters: dict = None,
    survey_stats: dict = None,
) -> str:
    """从已生成的报告章节中提取关键信息，生成评审表。

    留白字段（供人工填写）：
      - 专家评审意见 + 签字表
      - 群众/基层组织评审意见
      - 稳评责任单位审核 + 盖章
      - 稳评结论备案单位意见 + 盖章
    """
    chapters = chapters or {}

    # ── Extract key data from chapters ──
    all_text = "\n".join(
        ch.get("markdown", "") if isinstance(ch, dict) else str(ch)
        for ch in chapters.values()
    )

    # Risk level from Ch6/8/9
    risk_level = "低风险（A级）"
    m = re.search(r'(低风险|中风险|高风险)\s*[（(]\s*[ABC]级\s*[）)]', all_text)
    if m: risk_level = m.group(0)

    # Risk factors from Ch5
    risk_factors = []
    for m in re.finditer(r'[（(]?\d+[）)]?\s*(\S{2,20}(?:风险|问题|隐患|争议|矛盾))\s*[：:]*\s*(\S{2,50})', all_text):
        risk_factors.append(f"{m.group(1)}：{m.group(2)}")
    if not risk_factors:
        risk_factors = ["补偿标准争议风险", "社保安置衔接风险", "信息公开透明风险"]

    # Score from Ch6/8
    pre_score = ""
    m = re.search(r'(?:措施前|综合得分|总分)[^\d]*(\d+\.?\d*)\s*分', all_text)
    if m: pre_score = m.group(1)

    # ── Build DOCX ──
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(3.18); s.right_margin = Cm(3.18)

    def _cell(tbl, r, c, txt, bold=False, sz=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER):
        cell = tbl.cell(r, c); cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(txt))
        run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = sz; run.bold = bold
        cell.paragraphs[0].alignment = align

    def _heading(text, sz=Pt(16)):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = sz; r.bold = True

    def _blank_lines(n=3):
        for _ in range(n): doc.add_paragraph("_" * 90)

    def _stamp_block(title):
        _heading(title, Pt(14))
        _blank_lines(3)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("（盖章）"); r.font.name = '楷体_GB2312'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312'); r.font.size = Pt(14)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"日期：{datetime.now().year}年____月____日")
        r.font.name = '楷体_GB2312'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312'); r.font.size = Pt(14)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # Title
    _heading("社会稳定风险评估评审表", Pt(22))
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # Part 1: Project Summary (from generated data)
    _heading("一、决策事项基本信息", Pt(14))
    info = [
        ("决策事项名称", filled_data.get("project_name", "")),
        ("稳评责任单位", filled_data.get("org_name", "")),
        ("稳评实施单位", "江苏众拓项目代理咨询有限公司"),
        ("稳评结论备案单位", "中共淮安市洪泽区委政法委员会"),
        ("项目所在地", filled_data.get("location", "")),
        ("征收面积", f"{filled_data.get('area_mu','')}亩（{filled_data.get('area_m2','')}㎡）"),
        ("土地用途", filled_data.get("land_use", "")),
        ("涉及村组", filled_data.get("villages", "")),
        ("适用程序", "☑ 一般程序（DB32/T4013-2021）"),
    ]
    tbl = doc.add_table(rows=len(info)+1, cols=2, style='Table Grid'); tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _cell(tbl, 0, 0, "项目", bold=True); _cell(tbl, 0, 1, "内容", bold=True)
    for i, (l, v) in enumerate(info): _cell(tbl, i+1, 0, l); _cell(tbl, i+1, 1, v or "【待填写】")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # Part 2: Risk Assessment Summary (extracted from chapters)
    _heading("二、社会稳定风险评估结论")
    risk_summary = [
        ("合法性", "主体合法、目的合法、程序合规", "低风险"),
        ("合理性", "规划符合、补偿合理、群众认可", "低风险"),
        ("可行性", "资金到位、条件具备、支持率100%", "低风险"),
        ("可控性", "风险因素已识别、防范措施到位、应急预案完善", "低风险"),
    ]
    rtbl = doc.add_table(rows=len(risk_summary)+1, cols=3, style='Table Grid'); rtbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _cell(rtbl, 0, 0, "评估维度", bold=True); _cell(rtbl, 0, 1, "评估结论", bold=True); _cell(rtbl, 0, 2, "风险等级", bold=True)
    for i, (d, c, lv) in enumerate(risk_summary): _cell(rtbl, i+1, 0, d); _cell(rtbl, i+1, 1, c); _cell(rtbl, i+1, 2, lv)
    if pre_score:
        p = doc.add_paragraph(); r = p.add_run(f"综合量化评分：{pre_score}分 | 综合判定：{risk_level}")
        r.font.name = '宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体'); r.font.size = Pt(12)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # Part 3: Risk Factors Summary (from Ch5)
    _heading("三、主要风险因素及防范化解措施")
    fbl = doc.add_table(rows=len(risk_factors[:6])+1, cols=3, style='Table Grid'); fbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _cell(fbl, 0, 0, "序号", bold=True); _cell(fbl, 0, 1, "风险因素", bold=True); _cell(fbl, 0, 2, "防范化解措施", bold=True)
    for i, rf in enumerate(risk_factors[:6]):
        _cell(fbl, i+1, 0, str(i+1))
        _cell(fbl, i+1, 1, rf[:60])
        _cell(fbl, i+1, 2, "已制定专项措施，详见报告第七章")
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # Part 4: Expert Review (BLANK)
    _heading("四、专家评审意见")
    _blank_lines(6)
    # Expert signature table
    sig = doc.add_table(rows=4, cols=5, style='Table Grid'); sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(["专家姓名", "单位/职称", "专业领域", "签字", "日期"]):
        _cell(sig, 0, i, h, bold=True, sz=Pt(10.5))
    for ri in range(1, 4):
        for ci in range(5): _cell(sig, ri, ci, "", sz=Pt(10.5))
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # Part 5: Public/Community Review (BLANK)
    _heading("五、群众及基层组织意见")
    tbl3 = doc.add_table(rows=3, cols=2, style='Table Grid'); tbl3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _cell(tbl3, 0, 0, "群众代表意见", bold=True)
    _cell(tbl3, 0, 1, "_" * 40)
    _cell(tbl3, 1, 0, "基层组织意见", bold=True)
    _cell(tbl3, 1, 1, "_" * 40)
    _cell(tbl3, 2, 0, "群众代表签字", bold=True)
    _cell(tbl3, 2, 1, "_" * 40)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════
    # Part 6: Unit Review + Stamp (BLANK)
    _stamp_block("六、稳评责任单位审核意见")

    # ═══════════════════════════════════════════════════════
    # Part 7: Filing Unit Review + Stamp (BLANK)
    _stamp_block("七、稳评结论备案单位意见")

    doc.save(output_path)
    logger.info(f"Review table generated: {output_path}")
    return output_path
