"""docx_writer.py — 将 Markdown 报告写入 Word 文档

简化版：解析 Markdown 标题/段落/表格，生成带样式的 .docx。
"""

import re
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxWriter:
    """Markdown → Word 转换器。"""

    def write(self, markdown: str, output_path: str, title: str = ""):
        doc = Document()
        self._setup_styles(doc)

        # 标题
        if title:
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        lines = markdown.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            i += 1

            # 跳过 markdown 标题行 # （第一个 # 已作为标题）
            if line.startswith("# ") and title:
                continue

            # 标题
            if line.startswith("###"):
                doc.add_heading(line.lstrip("# ").strip(), level=3)
            elif line.startswith("##"):
                doc.add_heading(line.lstrip("# ").strip(), level=2)
            elif line.startswith("#"):
                doc.add_heading(line.lstrip("# ").strip(), level=1)
            # 分隔线
            elif line.strip() in ("---", "***", "___"):
                pass  # 跳过分隔线
            # 表格
            elif line.startswith("|") and i < len(lines) and lines[i].startswith("|"):
                table_rows = [line]
                while i < len(lines) and lines[i].startswith("|"):
                    table_rows.append(lines[i])
                    i += 1
                self._add_table(doc, table_rows)
            # 引用
            elif line.startswith(">"):
                p = doc.add_paragraph(line.lstrip("> ").strip())
                p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else p.style
            # 列表
            elif re.match(r"^[-*]\s+", line):
                doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
            elif re.match(r"^\d+\.\s+", line):
                doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            # 空行
            elif not line.strip():
                pass
            # 普通段落
            else:
                doc.add_paragraph(line)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    def _setup_styles(self, doc):
        """设置默认字体样式。"""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)
        # 中文字体
        try:
            style.element.rPr.rFonts.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
            )
        except Exception:
            pass

    def _add_table(self, doc, rows: list):
        """从 Markdown 表格行创建 Word 表格。"""
        # 解析
        parsed = []
        for row in rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            parsed.append(cells)

        # 去掉分隔行（|---|---|）
        parsed = [r for r in parsed if not all(re.match(r"^[-:]+$", c) for c in r)]
        if not parsed:
            return

        max_cols = max(len(r) for r in parsed)
        table = doc.add_table(rows=len(parsed), cols=max_cols)
        table.style = "Table Grid"
        for ri, row in enumerate(parsed):
            for ci in range(max_cols):
                cell = table.cell(ri, ci)
                cell.text = row[ci] if ci < len(row) else ""
        doc.add_paragraph("")  # 表后空行
