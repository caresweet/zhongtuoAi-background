"""Knowledge-base Q&A chat endpoint with learning and context memory.

- RAG-powered Q&A with auto web-search fallback
- User correction learning: detects when user says "不对/错了/应该是..."
  and stores corrected knowledge for future queries
- Conversation memory: persistent multi-turn context with auto-summarization
"""

import asyncio, json, time, uuid, re, os, pathlib
from datetime import datetime, timezone
from typing import Optional, List, Tuple
from fastapi import APIRouter, Query, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.services.llm_service import llm_service
from app.rag.embedder import EmbedderService
from app.rag.vector_store import VectorStoreService
from app.database.knowledge_db import async_session
from sqlalchemy import text

router = APIRouter(prefix="/api/knowledge", tags=["知识问答"])


# ── Request/Response models ──────────────────────────────────────────────────

class ChatMessage(BaseModel):
    role: str = "user"  # "user" | "assistant"
    content: str

class KnowledgeChatRequest(BaseModel):
    message: str = Field(..., description="用户问题")
    history: List[ChatMessage] = Field(default_factory=list, description="对话历史")
    domain: str = Field(default="stability", description="知识库领域: stability | bidding")
    top_k: int = Field(default=5, ge=1, le=10, description="检索数量")
    session_id: str = Field(default="", description="会话ID，用于上下文记忆")


# ── System prompts ───────────────────────────────────────────────────────────

SYSTEM_PROMPT_BASE = """你是"小拓智能体"，一个专业的社会稳定风险评估（稳评）知识助手。

## 身份与能力
- 知识来源于法律法规、政策文件、技术标准、稳评报告案例
- 具备学习能力：用户纠正后记住并更新知识
- 记住对话上下文，理解代词指代

## 回答规则（重要！）
1. **优先数据**：知识库中有具体数字、金额、百分比时，必须原样列出，不要模糊概括
2. **列表呈现**：涉及步骤、标准、分类时用编号列表，不要用段落叙述
3. **引用来源**：标注文档名称或编号（如"DB32/T 4013-2021""淮政规〔2026〕1号"）
4. **具体优于笼统**：说"54000元/亩"而不是"约5万多元"；说"三步程序"而不是"有一套程序"
5. **流程必列**：用户问"流程""步骤""程序"时，必须用1.2.3.编号列出每一步
6. **上下文关联**：理解代词指代（"它""这个""上面那个"），关联前文

## 注意事项
- 不要回答与稳评、征地完全无关的问题，礼貌说明职责范围
- 不确定的信息明确说明，不要编造
- 用户纠正你时，感谢并记住正确的信息
- 回答末尾不需要"如果您还有其他问题..."之类的客套话"""

SYSTEM_PROMPT_WITH_KB = SYSTEM_PROMPT_BASE + """

【知识库参考内容】
{context}

{learned_corrections}

请根据以上知识库内容回答用户问题。必须引用具体数据，用列表呈现流程步骤。"""

SYSTEM_PROMPT_NO_KB = SYSTEM_PROMPT_BASE + """

⚠️ 知识库中未找到直接匹配的内容。请基于训练知识给出通用的政策框架，注明"仅供参考，建议核实最新政策"。
当前日期: {current_date}

{learned_corrections}"""


# ── SSE helpers ───────────────────────────────────────────────────────────────

def _sse(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


# ── Retrieval quality thresholds ──────────────────────────────────────────────

KNOWLEDGE_GOOD_THRESHOLD = 0.45
KNOWLEDGE_WEAK_THRESHOLD = 0.55


# ── Correction detection patterns ─────────────────────────────────────────────

CORRECTION_PATTERNS = [
    r'(?:你|上面|刚才|这个|那个).{0,5}(?:说错|错了|不对|不正确|有误|搞错)',
    r'(?:不对|错了|不正确|有误|搞错了)',
    r'(?:应该是|正确的是|正确的|其实是|实际上是)',
    r'(?:纠正|更正|修改|改正|更新)',
    r'(?:不是.{0,10}(?:而是|应该是))',
]

def detect_correction(message: str) -> bool:
    """Check if a user message is a correction."""
    for pat in CORRECTION_PATTERNS:
        if re.search(pat, message):
            return True
    return False


# ── Learning: store and retrieve corrections ──────────────────────────────────

async def _store_correction(
    original_query: str, original_answer: str,
    user_correction: str, domain: str
) -> dict:
    """Analyze and store a user correction as learned knowledge."""
    if not llm_service.is_available:
        return {"stored": False, "reason": "LLM not available"}

    try:
        # Use LLM to extract the corrected knowledge
        analysis_prompt = f"""分析用户纠正，提取"正确的知识"作为一条可复用的知识点。

## 原始问题
{original_query[:500]}

## AI原始回答（被纠正的）
{original_answer[:800]}

## 用户纠正
{user_correction}

## 要求
1. 提取用户纠正后的正确知识点（1-3句话，包含具体数据/标准/流程）
2. 提取3-5个关键词用于后续检索
3. 返回JSON: {{"corrected_knowledge": "...", "keywords": ["kw1", "kw2", ...]}}
"""

        result = await llm_service.chat_with_reasoning(
            messages=[{"role": "user", "content": analysis_prompt}],
            system="你是知识提取助手，只返回JSON格式。",
            max_tokens=512,
            temperature=0.1,
        )

        content = result.get("content", "").strip()
        # Extract JSON
        json_match = re.search(r'\{[\s\S]*\}', content)
        if json_match:
            data = json.loads(json_match.group())
        else:
            data = {
                "corrected_knowledge": user_correction[:500],
                "keywords": ["用户纠正"],
            }

        corrected = data.get("corrected_knowledge", user_correction[:500])
        keywords = ",".join(data.get("keywords", ["用户纠正"]))

        # Store in DB
        async with async_session() as db:
            await db.execute(text("""
                INSERT INTO learned_corrections
                (original_query, original_answer, user_correction, corrected_knowledge, topic_keywords, domain)
                VALUES (:q, :a, :c, :ck, :kw, :dom)
            """), {
                "q": original_query[:2000], "a": original_answer[:3000],
                "c": user_correction[:2000], "ck": corrected[:2000],
                "kw": keywords, "dom": domain,
            })
            await db.commit()

        return {"stored": True, "corrected_knowledge": corrected[:200], "keywords": keywords}

    except Exception as e:
        return {"stored": False, "reason": str(e)}


async def _get_learned_corrections(query: str, domain: str, top_k: int = 3) -> str:
    """Retrieve relevant learned corrections for a query."""
    try:
        async with async_session() as db:
            # Simple keyword match search
            result = await db.execute(text("""
                SELECT corrected_knowledge, topic_keywords FROM learned_corrections
                WHERE domain = :dom AND is_active = 1
                ORDER BY created_at DESC LIMIT 20
            """), {"dom": domain})
            rows = result.fetchall()

        if not rows:
            return ""

        # Score by keyword overlap
        query_lower = query.lower()
        scored = []
        for ck, kw in rows:
            kw_list = (kw or "").split(",")
            score = sum(1 for k in kw_list if k.strip().lower() in query_lower)
            if score > 0:
                scored.append((score, ck))

        scored.sort(key=lambda x: x[0], reverse=True)
        top = scored[:top_k]

        if not top:
            return ""

        lines = ["\n【已学习的纠正知识】"]
        for i, (_, ck) in enumerate(top):
            lines.append(f"{i+1}. {ck}")
        return "\n".join(lines)

    except Exception:
        return ""


# ── Conversation memory ───────────────────────────────────────────────────────

async def _save_conversation(session_id: str, role: str, content: str, msg_type: str = "chat"):
    """Persist a conversation turn."""
    try:
        async with async_session() as db:
            await db.execute(text("""
                INSERT INTO conversation_memory (session_id, role, content, message_type)
                VALUES (:sid, :role, :content, :mtype)
            """), {"sid": session_id, "role": role, "content": content[:5000], "mtype": msg_type})
            await db.commit()
    except Exception:
        pass


async def _get_conversation_summary(session_id: str, max_turns: int = 6) -> str:
    """Get recent conversation context for a session."""
    try:
        async with async_session() as db:
            result = await db.execute(text("""
                SELECT role, content FROM conversation_memory
                WHERE session_id = :sid
                ORDER BY created_at DESC LIMIT :lim
            """), {"sid": session_id, "lim": max_turns * 2})
            rows = list(result.fetchall())

        if not rows:
            return ""

        rows.reverse()  # chronological order
        lines = []
        for role, content in rows:
            label = "用户" if role == "user" else "小拓"
            # Extract key info from AI responses, keep user messages full
            if role == "assistant":
                # Keep first 2 sentences (most informative)
                sentences = content.replace('\n', ' ').split('。')
                short = '。'.join(sentences[:2]) + '。'
                short = short[:250]
            else:
                short = content[:200]
            lines.append(f"{label}: {short}")
        return "\n".join(lines)

    except Exception:
        return ""


# ── Document-code aware retrieval ─────────────────────────────────────────────

# Patterns for detecting document codes in user queries
DOC_CODE_PATTERNS = [
    r'(?:DB\d{2}\/?T\s*\d+[-–]\d+)',                        # DB32/T 4013-2021
    r'(?:DB\s*\d{2}\/?T\s*\d+)',                            # DB32/T4013
    r'(?:DB\s*\d{4}\/?T\s*\d+)',                            # DB3206/T1091
    r'[\w一-鿿]+〔\d{4}〕\d+号',                   # 淮政办发〔2012〕85号 / 发改投资〔2012〕2492号 / 苏政发〔2021〕87号
]

def _extract_doc_codes(query: str) -> list:
    """Extract document codes from query for precise matching."""
    codes = []
    for pat in DOC_CODE_PATTERNS:
        matches = re.findall(pat, query)
        codes.extend(matches)
    return list(set(codes))

async def _retrieve_by_doc_code(query: str, domain: str) -> str:
    """Direct SQL lookup for documents matching codes in the query."""
    codes = _extract_doc_codes(query)
    if not codes:
        return ""

    try:
        async with async_session() as db:
            all_texts = []
            for code in codes:
                # Search by title containing the code
                result = await db.execute(text("""
                    SELECT title, cleaned_text, raw_text FROM knowledge_documents
                    WHERE is_active = 1 AND domain = :dom
                    AND (title LIKE :code OR raw_text LIKE :code2)
                    LIMIT 3
                """), {"dom": domain, "code": f"%{code}%", "code2": f"%{code}%"})
                rows = result.fetchall()
                for title, cleaned, raw in rows:
                    doc_text = cleaned or raw or ""
                    if doc_text:
                        all_texts.append(f"【精确匹配】文档={title}\n{doc_text[:2000]}")

            if all_texts:
                return "\n\n---\n\n".join(all_texts)
    except Exception as e:
        print(f"[DocCode] lookup error: {e}")
    return ""


# ── Core RAG function ─────────────────────────────────────────────────────────

async def _retrieve_context(query: str, domain: str, top_k: int) -> Tuple[str, list, float]:
    """Retrieve relevant context from the knowledge base."""
    embedder = EmbedderService()
    vs = VectorStoreService()

    q_embeddings = await embedder.embed_texts([query])
    if not q_embeddings or not q_embeddings[0]:
        return "", [], 1.0

    col = vs.get_or_create_collection("knowledge_base")
    results = vs.query(col, q_embeddings[0], n_results=top_k)

    if not results or not results.get("documents") or not results["documents"][0]:
        return "", [], 1.0

    documents = results["documents"][0]
    metadatas = results.get("metadatas", [[]])[0] if results.get("metadatas") else []
    distances = results.get("distances", [[]])[0] if results.get("distances") else []

    chunks = []
    sources = []
    seen_titles = set()
    best_distance = 1.0

    for i, (doc, meta, dist) in enumerate(zip(documents, metadatas, distances)):
        if dist < best_distance:
            best_distance = dist
        source_file = meta.get("source_file", "") if isinstance(meta, dict) else ""
        doc_type = meta.get("document_type", "") if isinstance(meta, dict) else ""
        chunks.append(f"【参考{i+1}】来源={source_file} 类型={doc_type}\n{doc[:1500]}")
        if source_file and source_file not in seen_titles:
            seen_titles.add(source_file)
            sources.append({"title": source_file[:80], "type": doc_type, "score": round(1 - dist, 3)})

    return "\n\n---\n\n".join(chunks), sources, best_distance


def _check_knowledge_quality(best_distance: float, sources: list) -> dict:
    if not sources:
        return {"sufficient": False, "level": "none", "enable_search": True}
    if best_distance <= KNOWLEDGE_GOOD_THRESHOLD:
        return {"sufficient": True, "level": "good", "enable_search": False}
    elif best_distance <= KNOWLEDGE_WEAK_THRESHOLD:
        return {"sufficient": True, "level": "moderate", "enable_search": False}
    else:
        return {"sufficient": False, "level": "weak", "enable_search": True}


# ── API Endpoints ─────────────────────────────────────────────────────────────

@router.post("/extract-file")
async def extract_file_for_chat(file: UploadFile = File(...)):
    """提取上传文件的文本内容（用于知识问答中附加文件）。支持 PDF/DOCX/TXT/图片。"""
    import shutil, tempfile
    from app.services.file_service import file_service
    from app.services.pdf_data_extractor import PDFDataExtractor
    from app.services.llm_service import llm_service

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1])
    try:
        shutil.copyfileobj(file.file, tmp)
        tmp.close()
        path = tmp.name
        ext = os.path.splitext(file.filename)[1].lower()

        text = ""
        if ext == '.pdf':
            ext_obj = PDFDataExtractor(llm_service=llm_service)
            doc = await ext_obj.extract_pdf(path)
            text = doc.full_text or ""
        elif ext in ('.docx', '.doc'):
            text = file_service.extract_docx_text(path) or ""
        elif ext in ('.txt', '.md', '.csv'):
            text = pathlib.Path(path).read_text(encoding='utf-8', errors='ignore')
        elif ext in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'):
            # 🔴 Only OCR images likely to contain text (scanned docs, forms, certificates)
            # Skip scenery photos, notice board photos — they have no useful text
            fname = file.filename.lower()
            HAS_TEXT_KEYWORDS = ['扫描', '签字', '问卷', '调查表', '意见', '评审', '签到',
                                 '证书', '备案', '执照', '合同', '表格', '报告', '批复']
            if any(kw in fname for kw in HAS_TEXT_KEYWORDS):
                import base64
                img_b64 = base64.b64encode(pathlib.Path(path).read_bytes()).decode()
                result = await llm_service.chat_with_image(
                    text="请提取这张图片中的所有文字内容",
                    image_base64=img_b64, media_type=f"image/{ext[1:]}", max_tokens=2048,
                )
                text = result or ""
            else:
                text = f"[图片文件: {file.filename}]"
        else:
            os.unlink(path)
            return {"code": 0, "message": f"不支持的文件类型: {ext}", "data": {"text": "", "filename": file.filename}}

        os.unlink(path)
        return {"code": 0, "message": "ok", "data": {"text": text[:8000], "filename": file.filename,
                "size": len(text), "preview": text[:500]}}
    except Exception as e:
        if os.path.exists(tmp.name): os.unlink(tmp.name)
        raise HTTPException(status_code=500, detail=f"文件提取失败: {e}")


@router.post("/chat")
async def knowledge_chat(request: KnowledgeChatRequest):
    """RAG-powered Q&A with learning and context memory.

    Streams SSE events: thinking, content, done, learned, error
    """
    if not llm_service.is_available:
        return StreamingResponse(
            _error_stream("LLM 服务未配置，请检查 API Key"),
            media_type="text/event-stream",
        )

    # Generate session_id if not provided
    session_id = request.session_id or str(uuid.uuid4())

    async def event_generator():
        try:
            # ── Save user message ────────────────────────────────────────
            await _save_conversation(session_id, "user", request.message)

            # ── 1. Check if this is a correction ───────────────────────────
            is_correction = detect_correction(request.message)
            correction_result = None

            if is_correction and request.history:
                # Find the last AI response to correct
                last_ai_msg = ""
                last_user_msg = ""
                for msg in reversed(request.history):
                    if msg.role == "assistant" and not last_ai_msg:
                        last_ai_msg = msg.content
                    if msg.role == "user" and not last_user_msg:
                        last_user_msg = msg.content

                if last_ai_msg:
                    yield _sse("thinking", {
                        "status": "learning",
                        "message": "检测到纠正信息，正在分析学习...",
                    })
                    correction_result = await _store_correction(
                        original_query=last_user_msg or request.message,
                        original_answer=last_ai_msg,
                        user_correction=request.message,
                        domain=request.domain,
                    )
                    if correction_result.get("stored"):
                        yield _sse("learned", {
                            "message": "✅ 已学习并记住纠正信息",
                            "knowledge": correction_result.get("corrected_knowledge", ""),
                        })

            # ── 2. RAG retrieval ──────────────────────────────────────────
            yield _sse("thinking", {
                "status": "retrieving",
                "message": "正在检索知识库...",
            })

            context, sources, best_distance = await _retrieve_context(
                request.message, request.domain, request.top_k
            )

            # Check for document-code precise match
            doc_code_context = await _retrieve_by_doc_code(request.message, request.domain)
            if doc_code_context:
                # Merge: doc-code text + vector results for complete coverage
                context = doc_code_context + "\n\n---\n【向量检索补充内容】\n" + context if context else doc_code_context
                # Boost quality
                if best_distance > KNOWLEDGE_GOOD_THRESHOLD:
                    best_distance = KNOWLEDGE_GOOD_THRESHOLD - 0.05

            quality = _check_knowledge_quality(best_distance, sources)
            enable_search = quality["enable_search"]

            # Get learned corrections
            learned = await _get_learned_corrections(request.message, request.domain)

            # Get conversation context
            conv_context = await _get_conversation_summary(session_id)

            # Emit retrieval status
            status_msg = {
                "good": f"检索到 {len(sources)} 条知识（匹配度: 高）",
                "moderate": f"检索到 {len(sources)} 条知识（匹配度: 中）",
                "weak": f"知识库匹配度较低，将结合网络搜索回答",
                "none": "知识库未找到内容，将通过网络搜索回答",
            }.get(quality["level"], "检索中...")

            yield _sse("thinking", {
                "status": "retrieved" if quality["sufficient"] else "retrieved_weak",
                "message": status_msg,
                "sources": sources,
                "quality": quality["level"],
                "web_search": enable_search,
            })

            # ── 3. Build system prompt ────────────────────────────────────
            from datetime import date
            today = date.today().isoformat()

            if context and quality["sufficient"]:
                system = SYSTEM_PROMPT_WITH_KB.format(
                    context=context,
                    learned_corrections=learned,
                )
            elif context and not quality["sufficient"]:
                system = SYSTEM_PROMPT_NO_KB.format(
                    current_date=today,
                    learned_corrections=learned,
                )
                system += f"\n\n【知识库弱匹配内容（仅供参考）】\n{context}"
            else:
                system = SYSTEM_PROMPT_NO_KB.format(
                    current_date=today,
                    learned_corrections=learned,
                )

            # Add conversation context
            if conv_context:
                system += f"\n{conv_context}"

            # ── 4. Build messages ────────────────────────────────────────
            messages = []
            for msg in request.history[-20:]:
                messages.append({"role": msg.role, "content": msg.content})
            # Enhance with structured output hint
            enhanced = request.message
            messages.append({"role": "user", "content": enhanced})

            # ── 5. Stream LLM response ────────────────────────────────────
            full_response = ""
            async for chunk in llm_service.chat_stream(
                messages=messages,
                system=system,
                max_tokens=1536,
                temperature=0.2,
                enable_search=enable_search,
            ):
                if chunk.get("type") == "content":
                    delta = chunk.get("delta", "")
                    if delta:
                        full_response += delta
                        yield _sse("content", {"delta": delta})

            # ── 6. Save AI response ───────────────────────────────────────
            await _save_conversation(session_id, "assistant", full_response)

            # ── 7. Done ───────────────────────────────────────────────────
            yield _sse("done", {
                "message": "回答完成",
                "sources": sources,
                "quality": quality["level"],
                "web_search_used": enable_search,
                "learned": correction_result is not None,
                "session_id": session_id,
                "full_answer": full_response,
            })

        except Exception as e:
            yield _sse("error", {
                "message": f"回答生成失败: {str(e)}",
                "retryable": True,
            })

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ═══════════════════════════════════════════════════════════════
# Report Review Endpoints
# ═══════════════════════════════════════════════════════════════

class ReviewReportRequest(BaseModel):
    text: str = Field(default="", description="报告文本（直接粘贴）")
    domain: str = Field(default="stability")


@router.post("/review-report")
async def review_report(request: ReviewReportRequest):
    """审核报告：检查数据合理性 + 格式问题，返回 SSE 流式结果。"""
    text = (request.text or "").strip()
    if not text or len(text) < 100:
        return StreamingResponse(
            _sse_error("请提供至少100字的报告内容"), media_type="text/event-stream")

    async def gen():
        try:
            yield _sse("thinking", {"status":"extracting","message":"正在解析报告内容...","progress":10})
            filled = _extract_fields(text)
            yield _sse("thinking", {"status":"reviewing","message":f"提取到{len(filled)}个关键字段，正在审核...","progress":30,"filled_data":filled})

            from app.validation.content_guardrails import full_guardrail_check
            report = full_guardrail_check(markdown=text, filled_data=filled)
            yield _sse("thinking", {"status":"reviewing","message":f"发现{report['total_issues']}个问题","progress":60})

            score = _review_score(report)
            passed = score >= 60
            yield _sse("review_result", {
                "score": score, "passed": passed, "can_learn": score >= 80,
                "total_issues": report["total_issues"],
                "sections": _review_sections(report),
                "summary": _review_summary(score, passed, report),
            })
            yield _sse("done", {"message":"审核完成" if passed else "审核未通过，建议修改后重新提交","score":score,"passed":passed,"can_learn":score>=80,"report_json":json.dumps(report,ensure_ascii=False,default=str)})
            # Record feedback for learning
            try:
                from app.services.learning_service import learning_service
                await learning_service.record_feedback(
                    report_title=_extract_fields(text).get("project_name", "用户粘贴报告"),
                    domain=request.domain,
                    quality_audit={"all_issues": [{"type": k, "message": v.get("message","")} for k in ("blocking","data_validity","fabricated","numeric_ranges","regulations") for v in report.get(k,[])], "total_issues": report["total_issues"]},
                    passed=passed, full_text=text[:50000],
                )
            except Exception: pass
        except Exception as e:
            yield _sse("error", {"message":f"审核失败: {e}","retryable":True})

    return StreamingResponse(gen(), media_type="text/event-stream",
        headers={"Cache-Control":"no-cache","Connection":"keep-alive","X-Accel-Buffering":"no"})


@router.post("/review-report-file")
async def review_report_file(file: UploadFile = File(...), domain: str = "stability"):
    """上传 DOCX/PDF 文件进行全面审核：文本+图片+表格+格式。"""
    import shutil, tempfile, os as _os
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=_os.path.splitext(file.filename or ".docx")[1])
    try:
        shutil.copyfileobj(file.file, tmp); tmp.close()
        path, ext = tmp.name, _os.path.splitext(file.filename or ".docx")[1].lower()

        text = ""; images = []; tables = []; image_count = 0; table_count = 0

        if ext == '.pdf':
            from app.services.pdf_data_extractor import PDFDataExtractor
            from app.services.llm_service import llm_service as _llm
            doc = await PDFDataExtractor(llm_service=_llm).extract_pdf(path)
            text = doc.full_text or ""
            images = doc.extracted_images or []
            tables = doc.extracted_tables or []
            image_count = len(images); table_count = len(tables)
        elif ext in ('.docx','.doc'):
            # Extract text
            from app.services.file_service import file_service
            text = file_service.extract_docx_text(path) or ""
            # Extract images and tables from DOCX
            try:
                from docx import Document as DocxDocument
                from docx.opc.constants import RELATIONSHIP_TYPE as RT
                import base64
                d = DocxDocument(path)
                # Images
                for rel in d.part.rels.values():
                    if "image" in str(rel.reltype).lower():
                        try:
                            img_data = rel.target_part.blob
                            img_path = path + f"_img_{len(images)}.png"
                            with open(img_path, 'wb') as f: f.write(img_data)
                            images.append(img_path)
                        except: pass
                image_count = len(images)
                # Tables
                for t in d.tables:
                    try:
                        rows = [[c.text.strip()[:200] for c in row.cells] for row in t.rows]
                        if rows: tables.append({"headers": rows[0] if rows else [], "rows": rows[1:] if len(rows)>1 else [], "row_count": len(rows)})
                    except: pass
                table_count = len(tables)
            except Exception as e:
                logger.warning(f"DOCX image/table extraction failed (non-critical): {e}")
        else:
            _os.unlink(path)
            return StreamingResponse(_sse_error(f"不支持的文件类型: {ext}，请上传 DOCX 或 PDF"), media_type="text/event-stream")

        if not text or len(text.strip()) < 100:
            _os.unlink(path)
            return StreamingResponse(_sse_error("未能提取到足够文本（至少100字）"), media_type="text/event-stream")

        # ── Run review with image/table context ──
        async def gen():
            try:
                yield _sse("thinking", {"status":"extracting","message":f"已提取: {len(text)}字文本, {image_count}张图片, {table_count}个表格","progress":15})

                filled = _extract_fields(text)
                yield _sse("thinking", {"status":"reviewing","message":f"提取到{len(filled)}个关键字段，正在全面审核...","progress":25,"filled_data":filled})

                from app.validation.content_guardrails import full_guardrail_check
                report = full_guardrail_check(markdown=text, filled_data=filled)
                yield _sse("thinking", {"status":"reviewing","message":f"文本审核: {report['total_issues']}个问题","progress":50})

                # Image review
                img_sections = _review_images(images, image_count, text)
                yield _sse("thinking", {"status":"reviewing","message":f"图片审核: {len(images)}张","progress":65})

                # Table review
                tbl_sections = _review_tables(tables, table_count, text)
                yield _sse("thinking", {"status":"reviewing","message":f"表格审核: {table_count}个","progress":80})

                # Format review via MCP
                fmt_sections = _review_format(path)
                yield _sse("thinking", {"status":"reviewing","message":"格式审核完成","progress":90})

                # Combine
                all_sections = _review_sections(report) + img_sections + tbl_sections + fmt_sections
                total_issues = report["total_issues"] + sum(1 for s in all_sections if s.get("status")=="fail")
                score = max(0, min(100, _review_score(report) - (10 if image_count < 1 else 0) - (5 if table_count < 1 else 0)))

                yield _sse("review_result", {
                    "score":score, "passed":score>=60, "can_learn":score>=80,
                    "total_issues":total_issues,
                    "stats": {"text_chars":len(text), "images":image_count, "tables":table_count},
                    "sections": all_sections,
                    "summary": _review_summary_ext(score, score>=60, report, image_count, table_count, total_issues),
                })
                yield _sse("done", {"message":"审核完成" if score>=60 else "审核未通过","score":score,"passed":score>=60,"can_learn":score>=80})
                # Record feedback for learning
                try:
                    from app.services.learning_service import learning_service
                    await learning_service.record_feedback(
                        report_title=filled.get("project_name", file.filename or "上传文件报告"),
                        domain=domain,
                        quality_audit={"all_issues": [{"type": k, "message": v.get("message","")} for k in ("blocking","data_validity","fabricated","numeric_ranges","regulations") for v in report.get(k,[])], "total_issues": report["total_issues"]},
                        passed=score>=60, full_text=text[:50000],
                    )
                except Exception: pass
            except Exception as e:
                yield _sse("error", {"message":f"审核失败: {e}","retryable":True})

        # Keep file for format review (MCP needs it), clean up afterwards
        result = StreamingResponse(gen(), media_type="text/event-stream",
            headers={"Cache-Control":"no-cache","Connection":"keep-alive","X-Accel-Buffering":"no"})
        # Schedule cleanup
        import asyncio as _asyncio
        async def _cleanup():
            await _asyncio.sleep(5)
            if _os.path.exists(path): _os.unlink(path)
            for img in images:
                if _os.path.exists(img): _os.unlink(img)
        _asyncio.ensure_future(_cleanup())
        return result

    except Exception as e:
        if _os.path.exists(tmp.name): _os.unlink(tmp.name)
        return StreamingResponse(_sse_error(f"文件处理失败: {e}"), media_type="text/event-stream")


class LearnReportRequest(BaseModel):
    text: str = Field(..., description="报告全文")
    title: str = Field(default="")
    domain: str = Field(default="stability")
    source: str = Field(default="用户上传")


@router.post("/learn-report")
async def learn_report(request: LearnReportRequest):
    """将优秀报告存入知识库供后续学习。"""
    text = (request.text or "").strip()
    if len(text) < 500:
        return {"code":1,"message":"报告内容不足（至少500字）"}
    title = request.title or f"用户上传报告_{uuid.uuid4().hex[:8]}"
    try:
        from app.database.knowledge_db import async_session
        from sqlalchemy import text as _text
        from app.rag.rag_service import rag_service
        async with async_session() as db:
            now = datetime.now(timezone.utc).isoformat()
            await db.execute(_text("""INSERT INTO knowledge_documents (title,document_type,domain,file_path,raw_text,cleaned_text,extraction_status,indexed_status,is_active,created_at,updated_at) VALUES (:t,:dt,:d,:fp,:r,:c,'completed','pending',1,:n,:n)"""),{"t":title[:200],"dt":"report","d":request.domain,"fp":f"learned/{title[:50]}.txt","r":text[:100000],"c":text[:100000],"n":now})
            await db.commit()
            row = await db.execute(_text("SELECT last_insert_rowid()"))
            doc_id = row.scalar()
        if doc_id:
            try: await rag_service.index_document(document_id=doc_id, text=text[:50000], metadata={"title":title,"domain":request.domain,"source":request.source})
            except Exception as e: logger.warning(f"Reindex failed (non-critical): {e}")
        return {"code":0,"message":f"✅ 报告「{title[:50]}」已存入知识库（ID={doc_id}）","document_id":doc_id}
    except Exception as e:
        logger.error(f"Learn report failed: {e}")
        return {"code":1,"message":f"学习失败: {e}"}


# ── Report Review Helpers ──────────────────────────────────────────────────────

def _extract_fields(text: str) -> dict:
    fields = {}
    pats = {
        "project_name": [r'(?:项目名称|决策名称|报告标题)[：:]\s*(.+?)(?:\n|$)'],
        "org_name": [r'(?:责任单位|决策单位|征收主体)[：:]\s*(.+?)(?:\n|$)',r'(淮安|洪泽|涟水|金湖|盱眙|清江浦|淮阴)\S{0,20}(?:人民政府|区政府|街道办事处|镇政府)'],
        "location": [r'(?:位置|坐落|地址|位于)[：:]\s*(.+?)(?:\n|$)'],
        "area_mu": [r'(\d+\.?\d*)\s*亩'],
        "household_count": [r'(\d+)\s*(?:户|农户)'],
        "total_samples": [r'(?:发放|回收|调查|问卷).*?(\d+)\s*(?:份|人)'],
        "support_rate": [r'支持率[^%]*?(\d+\.?\d*)\s*%'],
    }
    for key, patterns in pats.items():
        for p in patterns:
            m = re.search(p, text)
            if m: fields[key] = m.group(1).strip(); break
    return fields


def _review_score(report: dict) -> int:
    s = 100
    s -= len(report.get("blocking",[]))*10
    s -= len(report.get("data_validity",[]))*8
    s -= len(report.get("fabricated",[]))*10
    s -= len(report.get("numeric_ranges",[]))*5
    s -= len(report.get("regulations",[]))*5
    return max(0,min(100,s))


def _review_sections(report: dict) -> list:
    cats = [("数据有效性","data_validity","检查负数、反对率、年份错误等"),
            ("数据溯源","fabricated","检查数值是否可追溯到用户资料"),
            ("数值范围","numeric_ranges","检查面积/户数/评分等是否合理"),
            ("法规引用","regulations","检查法规文号是否真实存在"),
            ("内容规范","blocking","检查占位符、口语表达、AI套词")]
    return [{"name":n,"category":"data","description":d,"status":"pass" if not report.get(k,[]) else "fail","issues":[{"severity":i.get("severity","warning"),"message":i.get("message",i.get("description",""))} for i in report.get(k,[])[:5]],"count":len(report.get(k,[]))} for n,k,d in cats]


def _review_images(images: list, count: int, text: str) -> list:
    """Review extracted images."""
    sections = []
    # Image count check
    status = "fail" if count == 0 else "pass"
    issues = []
    if count == 0:
        issues.append({"severity":"error","message":"报告中未检测到任何图片。正式报告应包含：位置示意图、现场照片、公示照片、问卷调查/座谈会照片等。"})
    elif count < 3:
        issues.append({"severity":"warning","message":f"仅检测到{count}张图片。正式稳评报告通常需5张以上图片（位置图、公示照、座谈会照、现场勘查照等）。"})
    sections.append({"name":"图片数量","category":"image","description":f"检测到{count}张图片","status":status,"issues":issues,"count":len(issues)})

    # Check for required image types in text
    img_types = {"位置图":r'(?:位置|坐落|红线|勘测).*(?:图|示意图)',"公示照":r'(?:公示|公告).*(?:照片|图)',"座谈会照":r'(?:座谈|会议).*(?:照片|图|签到)',"现场照":r'(?:现场|勘查|踏勘).*(?:照片|图)'}
    for name, pat in img_types.items():
        if not re.search(pat, text):
            sections.append({"name":f"缺少{name}","category":"image","description":f"文本中未找到{name}的图片标记","status":"warning","issues":[{"severity":"warning","message":f"建议在报告中添加{name}"}],"count":1})
    return sections


def _review_tables(tables: list, count: int, text: str) -> list:
    """Review extracted tables."""
    sections = []
    status = "pass" if count > 0 else "fail"
    issues = []
    if count == 0:
        issues.append({"severity":"error","message":"报告中未检测到表格。正式稳评报告应包含：风险因素识别表、措施前后评分表、问卷调查统计表等。"})
    else:
        # Check for garbled tables
        for i, t in enumerate(tables):
            h = t.get("headers",[])
            if not h or len(h) < 2:
                issues.append({"severity":"error","message":f"表{i+1}缺少表头，格式不完整"})
            rows = t.get("rows",[])
            if not rows:
                issues.append({"severity":"warning","message":f"表{i+1}无数据行"})
            # Check for placeholder-only tables
            all_text = " ".join(str(c) for c in h) + " " + " ".join(str(c) for r in rows for c in r if r)
            if all_text and re.search(r'【待补充】|待补充|—\s*—', all_text):
                issues.append({"severity":"error","message":f"表{i+1}包含占位符，数据未填充"})
    sections.append({"name":"表格完整性","category":"table","description":f"检测到{count}个表格","status":"fail" if issues else "pass","issues":issues,"count":len(issues)})

    # Check for expected table types in text
    expected = {"风险因素表":r'风险.*?因素|风险.*?识别',"评分表":r'措施前.*?评分|量化.*?评分',"调查统计表":r'问卷.*?统计|调查.*?结果.*?表'}
    for name, pat in expected.items():
        if not re.search(pat, text):
            sections.append({"name":f"可能缺少{name}","category":"table","description":f"文本中未检测到{name}","status":"warning","issues":[{"severity":"warning","message":f"建议添加{name}"}],"count":1})
    return sections


def _review_format(docx_path: str) -> list:
    """Review document format (fonts, headings). Simple heuristics for now."""
    sections = []
    try:
        from docx import Document as D
        d = D(docx_path)
        # Check heading count
        headings = [p for p in d.paragraphs if p.style.name.startswith('Heading')]
        sections.append({"name":"标题层级","category":"format","description":f"检测到{len(headings)}个标题","status":"pass" if len(headings)>=5 else "warning","issues":[] if len(headings)>=5 else [{"severity":"warning","message":f"仅{len(headings)}个标题，正式报告建议10章以上"}],"count":0 if len(headings)>=5 else 1})
        # Check font usage
        body_count = sum(1 for p in d.paragraphs if p.text.strip())
        sections.append({"name":"正文段落","category":"format","description":f"共{body_count}段","status":"pass" if body_count>=20 else "warning","issues":[] if body_count>=20 else [{"severity":"warning","message":f"仅{body_count}段，内容不完整"}],"count":0 if body_count>=20 else 1})
    except Exception:
        sections.append({"name":"格式检查","category":"format","description":"无法读取DOCX格式","status":"warning","issues":[{"severity":"warning","message":"格式检查失败，请用WPS打开人工检查"}],"count":1})
    return sections


def _review_summary_ext(score: int, passed: bool, report: dict, img_count: int, tbl_count: int, total_issues: int) -> str:
    emoji = "✅" if passed else "❌"
    level = "达标" if score>=80 else ("基本达标" if score>=60 else "未达标")
    cats = [("数据有效性","data_validity"),("数据溯源","fabricated"),("数值范围","numeric_ranges"),("法规引用","regulations"),("内容规范","blocking")]
    parts = [f"## 📋 报告全面审核结果\n**综合评分: {score}/100 {emoji} {level}**\n"]
    parts.append(f"**文档统计**: {report.get('chapter',0)}章 · {len(report.get('blocking',[]))+len(report.get('data_validity',[]))+len(report.get('fabricated',[]))}个文本问题 · {img_count}张图片 · {tbl_count}个表格\n")
    for n,k in cats:
        iss = report.get(k,[])
        parts.append(f"- {'✅' if not iss else '❌'} {n}: {'通过' if not iss else f'{len(iss)}个问题'}")
    parts.append(f"- {'✅' if img_count>=3 else '⚠️'} 图片: {img_count}张{'（不足）' if img_count<3 else ''}")
    parts.append(f"- {'✅' if tbl_count>=1 else '❌'} 表格: {tbl_count}个{'（缺少）' if tbl_count<1 else ''}")
    if total_issues > 0:
        parts.append(f"\n### 🔍 共发现 {total_issues} 个问题")
        for n,k in cats:
            for i in report.get(k,[])[:3]:
                parts.append(f"- **[{n}]** {i.get('message',i.get('description',''))}")
    if passed and score >= 80:
        parts.append("\n> 💡 该报告质量优秀，可存入知识库供后续学习参考。")
    return "\n".join(parts)


async def _sse_error(message: str):
    yield _sse("error", {"message": message, "retryable": False})


async def _error_stream(message: str):
    yield _sse("error", {"message": message, "retryable": False})
