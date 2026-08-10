"""material_reader.py — 用户资料读取

支持：
  PDF  → PyMuPDF 文本提取；文本过少时渲染页面为图片 → 视觉 OCR
  DOCX → python-docx
  DOC  → macOS textutil
  图片 → 视觉模型直接理解
"""

import base64
import io
import os
import subprocess
import tempfile
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument


class MaterialReader:
    """读取各类用户资料，输出文本。"""

    def __init__(self, llm=None):
        self._llm = llm  # LLMClient, 用于 OCR

    # ── 统一入口 ────────────────────────────────────────────────
    async def read(self, file_path: str, max_chars: int = 20000) -> dict:
        """返回 {path, name, text, pages, method}。"""
        p = Path(file_path)
        name = p.name
        ext = p.suffix.lower()
        result = {"path": str(p), "name": name, "text": "", "pages": 0, "method": "unknown"}

        if ext == ".pdf":
            await self._read_pdf(p, result, max_chars)
        elif ext == ".docx":
            self._read_docx(p, result, max_chars)
        elif ext == ".doc":
            self._read_doc(p, result, max_chars)
        elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".webp", ".gif"):
            await self._read_image(p, result)
        elif ext in (".txt", ".md"):
            result["text"] = p.read_text(encoding="utf-8", errors="ignore")[:max_chars]
            result["method"] = "text"
        else:
            result["text"] = f"[不支持的文件类型: {ext}]"

        return result

    async def read_many(self, paths: List[str], max_chars_each: int = 15000) -> List[dict]:
        results = []
        for fp in paths:
            try:
                r = await self.read(fp, max_chars=max_chars_each)
            except Exception as e:
                r = {"path": fp, "name": os.path.basename(fp), "text": f"[读取失败: {e}]",
                     "pages": 0, "method": "error"}
            results.append(r)
        return results

    # ── PDF ─────────────────────────────────────────────────────
    async def _read_pdf(self, p: Path, result: dict, max_chars: int):
        import fitz  # PyMuPDF
        doc = fitz.open(str(p))
        result["pages"] = len(doc)
        texts = []
        need_ocr_pages = []
        for i, page in enumerate(doc):
            t = page.get_text().strip()
            if len(t) > 30:
                texts.append(t)
            else:
                need_ocr_pages.append((i, page))
        combined = "\n".join(texts)
        result["method"] = "pdf_text"

        # 文本过少 → OCR 补充
        if len(combined) < 200 and need_ocr_pages and self._llm and self._llm.available:
            ocr_texts = []
            for idx, page in need_ocr_pages[:8]:  # 最多 OCR 前8页
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                img_b64 = base64.b64encode(img_bytes).decode()
                try:
                    ocr = await self._llm.vision(
                        "请识别图片中的全部文字内容，按原文格式输出，保留表格结构。只输出文字，不加说明。",
                        img_b64, mime_type="image/png", max_tokens=3000,
                    )
                    ocr_texts.append(ocr)
                except Exception as e:
                    ocr_texts.append(f"[OCR失败 p{idx+1}: {e}]")
            combined = "\n".join(texts + ocr_texts)
            result["method"] = "pdf_ocr"

        result["text"] = combined[:max_chars]
        doc.close()

    # ── DOCX ────────────────────────────────────────────────────
    def _read_docx(self, p: Path, result: dict, max_chars: int):
        doc = DocxDocument(str(p))
        parts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        # 表格文本
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        result["text"] = "\n".join(parts)[:max_chars]
        result["pages"] = 1
        result["method"] = "docx"

    # ── DOC (旧格式) ────────────────────────────────────────────
    def _read_doc(self, p: Path, result: dict, max_chars: int):
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w") as tmp:
            tmp_path = tmp.name
        try:
            subprocess.run(
                ["textutil", "-convert", "txt", "-encoding", "UTF-8",
                 "-output", tmp_path, str(p)],
                capture_output=True, timeout=30,
            )
            text = Path(tmp_path).read_text(encoding="utf-8", errors="ignore")
            result["text"] = text[:max_chars]
            result["method"] = "textutil"
        except Exception as e:
            result["text"] = f"[DOC转换失败: {e}]"
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    # ── 图片 ────────────────────────────────────────────────────
    async def _read_image(self, p: Path, result: dict):
        if not self._llm or not self._llm.available:
            result["text"] = "[图片需 OCR 但 LLM 未配置]"
            result["method"] = "image_skip"
            return
        img_b64 = base64.b64encode(p.read_bytes()).decode()
        ext = p.suffix.lower().lstrip(".")
        mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg"}.get(ext, f"image/{ext}")
        try:
            text = await self._llm.vision(
                "请识别图片中的全部文字和数据，按原文格式输出。如果有表格请保留结构。只输出文字内容。",
                img_b64, mime_type=mime, max_tokens=3000,
            )
            result["text"] = text
            result["method"] = "image_ocr"
        except Exception as e:
            result["text"] = f"[图片OCR失败: {e}]"
            result["method"] = "image_error"
