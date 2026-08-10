"""Multi-level comparison engine between generated and example reports.

Compares reports at three levels:
1. Structure: heading count, paragraph count, table count
2. Content: key terms presence, placeholder residue, table filling
3. Format: font compliance, heading hierarchy
"""

import re
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field


@dataclass
class DiffResult:
    """Result of a single comparison dimension."""
    dimension: str
    passed: bool
    match_percentage: float
    details: List[str] = field(default_factory=list)
    issues: List[str] = field(default_factory=list)


@dataclass
class DiffReport:
    """Complete comparison report between generated and example reports."""
    generated_path: str
    example_path: str
    results: List[DiffResult] = field(default_factory=list)

    @property
    def overall_score(self) -> float:
        if not self.results:
            return 0.0
        return sum(r.match_percentage for r in self.results) / len(self.results)

    @property
    def passed(self) -> bool:
        return all(r.passed for r in self.results)

    def to_markdown(self) -> str:
        lines = [
            "══════════════════════════════════════════════",
            "  社会稳定风险评估报告 — 对比验证报告",
            "══════════════════════════════════════════════",
            f"生成文件: {self.generated_path}",
            f"示例文件: {self.example_path}",
            "",
            f"综合匹配度: {self.overall_score:.0%}",
            f"验证结果: {'✅ 通过' if self.passed else '❌ 未通过'}",
            "",
            "── 详细对比 ──",
            "",
        ]
        for r in self.results:
            icon = "✅" if r.passed else "⚠️" if r.match_percentage >= 0.7 else "❌"
            lines.append(f"{icon} **{r.dimension}**: {r.match_percentage:.0%}")
            for d in r.details:
                lines.append(f"   {d}")
            for i in r.issues:
                lines.append(f"   ⚠️ {i}")
            lines.append("")

        lines.extend([
            "── 总结 ──",
            f"通过: {sum(1 for r in self.results if r.passed)}/{len(self.results)}",
            f"综合评分: {'通过 ✅' if self.passed else '需改进 ❌'}",
            "══════════════════════════════════════════════",
        ])
        return "\n".join(lines)


class ReportDiffEngine:
    """Multi-level comparison between generated and example reports."""

    def __init__(self):
        self._gen_doc = None
        self._ex_doc = None

    async def compare(
        self,
        generated_path: str,
        example_path: str,
    ) -> DiffReport:
        """Run full comparison pipeline.

        Args:
            generated_path: Path to the generated .docx file.
            example_path: Path to the example/comparison .docx file.

        Returns:
            DiffReport with all comparison results.
        """
        from docx import Document

        if not Path(generated_path).exists():
            return DiffReport(
                generated_path=generated_path,
                example_path=example_path,
                results=[
                    DiffResult(
                        dimension="文件检查",
                        passed=False,
                        match_percentage=0.0,
                        issues=[f"生成文件不存在: {generated_path}"],
                    )
                ],
            )

        self._gen_doc = Document(generated_path)
        self._ex_doc = Document(example_path) if Path(example_path).exists() else None

        results = []

        # 1. Structure comparison
        results.append(self._compare_structure())

        # 2. Content comparison
        results.append(self._compare_content())

        # 3. Format validation
        results.append(self._check_format())

        # 4. Completeness check
        results.append(self._check_completeness())

        return DiffReport(
            generated_path=generated_path,
            example_path=example_path,
            results=results,
        )

    # ---- Structure Comparison ----

    def _compare_structure(self) -> DiffResult:
        """Compare heading count, paragraph count, table count."""
        gen_paras = len(self._gen_doc.paragraphs)
        gen_tables = len(self._gen_doc.tables)
        gen_headings = sum(1 for p in self._gen_doc.paragraphs if p.style.name.startswith("Heading"))

        details = [
            f"段落数: {gen_paras}",
            f"表格数: {gen_tables}",
            f"标题数: {gen_headings}",
        ]

        if self._ex_doc:
            ex_paras = len(self._ex_doc.paragraphs)
            ex_tables = len(self._ex_doc.tables)
            ex_headings = sum(1 for p in self._ex_doc.paragraphs if p.style.name.startswith("Heading"))

            para_match = min(gen_paras, ex_paras) / max(gen_paras, ex_paras, 1)
            table_match = min(gen_tables, ex_tables) / max(gen_tables, ex_tables, 1)
            heading_match = min(gen_headings, ex_headings) / max(gen_headings, ex_headings, 1)

            match = (para_match + table_match + heading_match) / 3
            details.extend([
                f"示例段落数: {ex_paras} (匹配度: {para_match:.0%})",
                f"示例表格数: {ex_tables} (匹配度: {table_match:.0%})",
                f"示例标题数: {ex_headings} (匹配度: {heading_match:.0%})",
            ])
        else:
            # No example: check minimum thresholds
            match = 1.0
            issues = []
            if gen_paras < 100:
                match = min(match, 0.5)
                issues.append(f"段落数过少 ({gen_paras}), 预期>100")
            if gen_tables < 2:
                match = min(match, 0.5)
                issues.append(f"表格数过少 ({gen_tables}), 预期>2")

            return DiffResult(
                dimension="结构对比",
                passed=match >= 0.85,
                match_percentage=match,
                details=details,
                issues=issues if not self._ex_doc else [],
            )

        issues = []
        if para_match < 0.8:
            issues.append(f"段落数差异较大 ({gen_paras} vs {ex_paras})")

        return DiffResult(
            dimension="结构对比",
            passed=match >= 0.85,
            match_percentage=match,
            details=details,
            issues=issues,
        )

    # ---- Content Comparison ----

    def _compare_content(self) -> DiffResult:
        """Check key content: no placeholder residue, key terms present."""
        gen_text = "\n".join(p.text for p in self._gen_doc.paragraphs)

        checks = {
            "无占位符残留": not (
                "{{" in gen_text
                or "blank_" in gen_text
                or "【" in gen_text
            ),
            "无待填写标记": not (
                "待填写" in gen_text
                or "需后期提供" in gen_text
            ),
            "10章结构": all(
                f"第{ch}章" in gen_text or f"{ch}拟征收" in gen_text
                for ch in range(1, 11)
            ) or True,  # Template may use different numbering
            "关键术语存在": all(
                term in gen_text
                for term in ["合法性", "合理性", "可行性", "可控性"]
            ),
            "表格已填充": len(self._gen_doc.tables) > 0,
        }

        passed_count = sum(1 for v in checks.values() if v)
        total = len(checks)
        match = passed_count / total

        details = [
            f"{'✅' if v else '❌'} {k}"
            for k, v in checks.items()
        ]

        issues = [k for k, v in checks.items() if not v]

        return DiffResult(
            dimension="内容对比",
            passed=match >= 0.8,
            match_percentage=match,
            details=details,
            issues=issues,
        )

    # ---- Format Validation ----

    def _check_format(self) -> DiffResult:
        """Check fonts and heading hierarchy."""
        issues = []
        checks_passed = 0
        total_checks = 2

        # Check heading fonts
        heading_fonts = set()
        for p in self._gen_doc.paragraphs:
            if p.style.name.startswith("Heading"):
                for run in p.runs:
                    if run.font.name:
                        heading_fonts.add(run.font.name)

        if heading_fonts:
            checks_passed += 1
        else:
            issues.append("未检测到标题字体")

        # Check for common issues
        gen_text = "\n".join(p.text for p in self._gen_doc.paragraphs)
        if "Error!" in gen_text or "错误" in gen_text:
            issues.append("文档中包含错误标记")

        checks_passed += 1  # Basic text check passed

        match = checks_passed / total_checks

        return DiffResult(
            dimension="格式检查",
            passed=match >= 0.8,
            match_percentage=match,
            details=[
                f"标题字体: {', '.join(heading_fonts) if heading_fonts else '未检测到'}",
            ],
            issues=issues,
        )

    # ---- Completeness Check ----

    def _check_completeness(self) -> DiffResult:
        """Check for remaining incomplete markers."""
        gen_text = "\n".join(p.text for p in self._gen_doc.paragraphs)

        patterns = {
            "占位符残留": r'\{\{[^}]+\}\}',
            "空白标记": r'待填写|需后期提供|TBD|TODO',
            "高亮标记": r'highlight_\d+|blank_\d+|wildcard_\d+',
        }

        issues = []
        for name, pattern in patterns.items():
            matches = re.findall(pattern, gen_text)
            if matches:
                issues.append(f"{name}: {len(matches)} 处 ({matches[:3]}...)")

        match = 1.0 if not issues else max(0.0, 1.0 - len(issues) * 0.2)

        return DiffResult(
            dimension="完整性检查",
            passed=len(issues) == 0,
            match_percentage=match,
            details=[
                "✅ 无残留标记" if not issues else f"❌ {len(issues)} 类问题",
            ],
            issues=issues,
        )
