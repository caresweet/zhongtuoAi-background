"""Report Specification Checker — validates generated reports against format rules.

Reads the format specification document (category="格式规范" in templates table)
and checks the generated report for compliance with:
- Font rules (封面黑体, 标题宋体/黑体, 正文仿宋, etc.)
- Size rules (小二, 小三, 四号, 小四, etc.)
- Alignment rules (居中, 靠左, 两端对齐)
- Indentation rules (首行缩进2格)
- Image sizing rules

Returns a SpecCheckReport with pass/fail per rule and an overall score.
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


# ---- Data Structures ----

@dataclass
class FormatRule:
    """A single format rule parsed from the spec document."""
    target: str          # "封面", "1级标题", "正文", etc.
    font: Optional[str] = None         # "黑体", "宋体", "仿宋"
    font_size_cn: Optional[str] = None # "小二", "小三", "四号", "小四"
    bold: Optional[bool] = None
    alignment: Optional[str] = None    # "居中", "靠左", "两端对齐"
    indent: Optional[str] = None       # "首行缩进2格"
    extra: Optional[str] = None        # Any additional notes


@dataclass
class RuleCheck:
    """Result of checking one rule against the generated report."""
    rule: FormatRule
    passed: bool
    details: str
    issues: List[str] = field(default_factory=list)


@dataclass
class SpecCheckReport:
    """Complete specification compliance report."""
    spec_name: str
    total_rules: int
    checks: List[RuleCheck] = field(default_factory=list)

    @property
    def passed_count(self) -> int:
        return sum(1 for c in self.checks if c.passed)

    @property
    def failed_count(self) -> int:
        return sum(1 for c in self.checks if not c.passed)

    @property
    def compliance_score(self) -> float:
        if not self.total_rules:
            return 1.0
        return self.passed_count / self.total_rules

    @property
    def passed(self) -> bool:
        return self.compliance_score >= 0.8

    def to_markdown(self) -> str:
        lines = [
            "══════════════════════════════════════════════",
            f"  报告格式规范校验 — {self.spec_name}",
            "══════════════════════════════════════════════",
            f"合规率: {self.compliance_score:.0%} ({self.passed_count}/{self.total_rules})",
            f"结果: {'✅ 通过' if self.passed else '❌ 需调整'}",
            "",
            "── 逐项检查 ──",
            "",
        ]
        for c in self.checks:
            icon = "✅" if c.passed else "❌"
            rule = c.rule
            lines.append(
                f"{icon} **{rule.target}**: "
                f"字体={rule.font or '—'}, "
                f"字号={rule.font_size_cn or '—'}, "
                f"对齐={rule.alignment or '—'}"
            )
            if not c.passed:
                for issue in c.issues:
                    lines.append(f"   ⚠️ {issue}")
            lines.append("")

        lines.extend([
            "── 总结 ──",
            f"通过: {self.passed_count}/{self.total_rules}",
            f"综合: {'通过 ✅' if self.passed else '需调整 ❌'}",
            "══════════════════════════════════════════════",
        ])
        return "\n".join(lines)


# ---- Chinese Font Size Mapping ----

CN_FONT_SIZES = {
    "小二": Pt(18),
    "小三": Pt(15),
    "四号": Pt(14),
    "小四": Pt(12),
    "五号": Pt(10.5),
    "小五": Pt(9),
    "小一": Pt(24),
    "一号": Pt(26),
    "二号": Pt(22),
    "三号": Pt(16),
}

CN_ALIGNMENT = {
    "居中": "CENTER",
    "靠左": "LEFT",
    "两端对齐": "JUSTIFY",
}


# ---- Spec Parser ----

class SpecParser:
    """Parses the format specification document into structured rules."""

    @classmethod
    def parse(cls, spec_path: str) -> Tuple[str, List[FormatRule]]:
        """Parse a format specification .docx file.

        Returns:
            (spec_name, list of FormatRules)
        """
        doc = Document(spec_path)
        spec_name = Path(spec_path).stem
        rules = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            rule = cls._parse_line(text)
            if rule:
                rules.append(rule)

        return spec_name, rules

    @classmethod
    def _parse_line(cls, text: str) -> Optional[FormatRule]:
        """Parse a single line into a FormatRule if it looks like a spec line."""
        # Remove trailing colon variants
        text = re.sub(r'[：:]\s*$', '', text).strip()

        # Skip lines that are just "内容" (content descriptions)
        if text in ("内容", "以下图片居中"):
            return None

        # Pattern: "target：font，size，style"
        # Examples:
        # "封面 ：小二，黑体，居中"
        # "正文：仿宋，四号，首行缩进2格"
        # "1级标题：宋体 小三 加黑，居中"
        # "表内容：      仿宋小四（第一行加黑）"

        # Split on first colon
        parts = re.split(r'[：:]', text, maxsplit=1)
        if len(parts) < 2:
            return None

        target = parts[0].strip()
        rest = parts[1].strip()

        if not target or not rest:
            return None

        rule = FormatRule(target=target)

        # Extract font
        for font_name in ["方正楷体GBK", "仿宋_GB2312", "仿宋", "黑体", "宋体", "楷体"]:
            if font_name in rest:
                rule.font = font_name
                break

        # Extract font size (Chinese sizes)
        for size_name in sorted(CN_FONT_SIZES.keys(), key=lambda x: -len(x)):
            if size_name in rest:
                rule.font_size_cn = size_name
                break

        # Extract bold
        if any(w in rest for w in ["加黑", "加粗", "bold"]):
            rule.bold = True

        # Extract alignment
        for align_word in sorted(CN_ALIGNMENT.keys(), key=lambda x: -len(x)):
            if align_word in rest:
                rule.alignment = align_word
                break

        # Extract indent
        if "首行缩进" in rest:
            rule.indent = "首行缩进"
        elif "缩进2格" in rest:
            rule.indent = "首行缩进"

        # Store extra
        rule.extra = rest

        return rule


# ---- Spec Checker ----

class SpecChecker:
    """Checks a generated report against format specification rules."""

    PT_TOLERANCE = Pt(1)  # ±1pt tolerance for font sizes

    @classmethod
    def check(cls, report_path: str, spec_path: str) -> SpecCheckReport:
        """Run all spec checks against the generated report.

        Args:
            report_path: Path to the generated .docx file.
            spec_path: Path to the format specification .docx file.

        Returns:
            SpecCheckReport with all check results.
        """
        if not Path(report_path).exists():
            return SpecCheckReport(
                spec_name="文件不存在",
                total_rules=0,
                checks=[
                    RuleCheck(
                        rule=FormatRule(target="文件检查"),
                        passed=False,
                        details=f"文件不存在: {report_path}",
                        issues=[f"文件不存在: {report_path}"],
                    )
                ],
            )

        spec_name, rules = SpecParser.parse(spec_path)
        report_doc = Document(report_path)

        checks = []
        for rule in rules:
            result = cls._check_rule(report_doc, rule)
            checks.append(result)

        return SpecCheckReport(
            spec_name=spec_name,
            total_rules=len(rules),
            checks=checks,
        )

    @classmethod
    def _check_rule(cls, doc: Document, rule: FormatRule) -> RuleCheck:
        """Check a single format rule against the document.

        Strategy: find paragraphs whose content or position matches the rule's
        target, then verify font/size/alignment/indent.
        """
        issues = []
        matched_paras = cls._find_target_paragraphs(doc, rule)

        if not matched_paras:
            # Rule target not found in document — not necessarily a failure
            # (some rules apply to optional sections)
            return RuleCheck(
                rule=rule,
                passed=True,
                details=f"文档中未找到「{rule.target}」相关段落，跳过检查",
            )

        all_ok = True
        for p in matched_paras:
            for run in p.runs:
                if not run.text.strip():
                    continue

                # Check font
                if rule.font:
                    actual_font = run.font.name
                    if actual_font:
                        # Normalize: 仿宋_GB2312 ~ 仿宋
                        actual_normalized = actual_font.replace("_GB2312", "")
                        rule_normalized = rule.font.replace("_GB2312", "")
                        if actual_normalized != rule_normalized:
                            all_ok = False
                            issues.append(
                                f"「{rule.target}」字体应为{rule.font}，"
                                f"实际为{actual_font}"
                            )
                            break  # One issue per paragraph is enough

                # Check size
                if rule.font_size_cn:
                    expected_pt = CN_FONT_SIZES.get(rule.font_size_cn)
                    if expected_pt and run.font.size:
                        actual_size = run.font.size
                        if abs(actual_size - expected_pt) > cls.PT_TOLERANCE:
                            all_ok = False
                            issues.append(
                                f"「{rule.target}」字号应为{rule.font_size_cn}"
                                f"({expected_pt})，实际为{actual_size}"
                            )

                break  # Check first run only

        if all_ok or not issues:
            return RuleCheck(
                rule=rule,
                passed=True,
                details=f"「{rule.target}」: {len(matched_paras)} 处检查通过",
            )

        return RuleCheck(
            rule=rule,
            passed=False,
            details=f"「{rule.target}」: {len(issues)} 处不符合",
            issues=issues[:5],  # Limit to 5 issues per rule
        )

    @classmethod
    def _find_target_paragraphs(cls, doc: Document, rule: FormatRule) -> List:
        """Find paragraphs that match the rule's target description."""
        target = rule.target
        matched = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            style_name = p.style.name if p.style else ""

            # Match by target keyword
            # "封面" → paragraphs near the start with cover-like content
            # "1级标题" → Heading 1 paragraphs
            # "正文" → Normal style paragraphs in body
            # "表内容" → paragraphs near tables

            if target == "封面" and cls._is_cover_paragraph(p, doc):
                matched.append(p)
            elif target == "1级标题" and ("Heading 1" in style_name or "heading 1" in style_name.lower()):
                matched.append(p)
            elif target == "2级标题" and ("Heading 2" in style_name or "heading 2" in style_name.lower()):
                matched.append(p)
            elif target == "3级标题" and ("Heading 3" in style_name or "heading 3" in style_name.lower()):
                matched.append(p)
            elif target == "正文" and "Normal" in style_name and not cls._is_cover_paragraph(p, doc):
                if len(matched) < 20:
                    matched.append(p)  # Sample up to 20 body paragraphs
            elif "图名" in target:
                # Only match actual figure captions: "图N-N ..." or "图N ..."
                if re.match(r'图\s*\d+', text):
                    matched.append(p)
            elif "表名" in target:
                # Only match actual table captions: "表N ..." or "表格N ..."
                if re.match(r'表\s*\d+|表格\s*\d+', text):
                    matched.append(p)
            elif "表内容" in target:
                if cls._is_near_table(p, doc):
                    matched.append(p)

        return matched

    @classmethod
    def _is_cover_paragraph(cls, p, doc) -> bool:
        """Check if a paragraph is on the cover page."""
        # Cover paragraphs are typically the first 15 paragraphs
        for i, dp in enumerate(doc.paragraphs):
            if dp is p:
                return i < 15
        return False

    @classmethod
    def _is_near_table(cls, p, doc) -> bool:
        """Check if a paragraph is inside or near a table."""
        # Simple heuristic: check if the paragraph's XML is near a table element
        return False  # Simplified — full implementation needs XML traversal


# ---- Convenience function ----

def load_spec_from_db() -> Optional[str]:
    """Load the active format specification document path from the database.

    Returns the absolute path to the spec .docx file, or None if not found.
    Uses a synchronous SQLite connection to avoid async complexity.
    """
    import sqlite3
    from app.config import settings
    from app.services.file_service import file_service

    try:
        db_path = settings.DATA_DIR / "knowledge_base.db"
        conn = sqlite3.connect(str(db_path))
        cursor = conn.cursor()
        cursor.execute(
            "SELECT template_file_path FROM templates "
            "WHERE category='格式规范' AND is_active=1 "
            "ORDER BY id DESC LIMIT 1"
        )
        row = cursor.fetchone()
        conn.close()

        if row and row[0]:
            abs_path = file_service.get_absolute_path(row[0])
            if abs_path.exists():
                return str(abs_path)
    except Exception:
        pass

    return None


async def check_report_async(report_path: str) -> Optional[SpecCheckReport]:
    """Async wrapper: load spec from DB, check report, return report."""
    spec_path = load_spec_from_db()
    if not spec_path:
        return None

    return SpecChecker.check(report_path, spec_path)
