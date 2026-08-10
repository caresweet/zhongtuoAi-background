#!/usr/bin/env python3
"""MCP Server: Report Format Validator with WPS integration.

Provides tools for validating generated .docx reports against
professional standards (DB32/T4013-2021) using python-docx deep
inspection and WPS Office for visual review.

Tools:
- validate_report: Full format validation against DB32/T4013-2021 standards
- check_fonts: Font compliance check (仿宋_GB2312, 黑体, 楷体)
- check_headings: Heading hierarchy validation
- check_tables: Table formatting check
- open_in_wps: Open report in WPS Office for visual review
- list_structure: Show document chapter/section structure
"""

import os
import sys
import json
import subprocess
from pathlib import Path
from typing import Optional

from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp.types import Tool, TextContent

# ── Server Setup ──────────────────────────────────────────────────────────────

server = Server("wps-report-validator")

# Storage path for generated reports
STORAGE_DIR = Path(__file__).resolve().parent.parent / "storage" / "generated"

# Professional standard requirements
REQUIRED_FONTS = {
    "body": "仿宋_GB2312",      # 正文
    "body_fallback": "仿宋",
    "heading1": "黑体",          # 一级标题
    "heading2": "黑体",          # 二级标题
    "heading3": "楷体_GB2312",  # 三级标题
    "heading3_fallback": "楷体",
    "cover_title": "方正小标宋简体",
}

REQUIRED_FONT_SIZES = {
    "body": 14,       # 小四 = 14pt
    "heading1": 18,   # 小二 = 18pt
    "heading2": 16,   # 三号 = 16pt
    "heading3": 14,   # 小四 = 14pt
    "cover_title": 26, # 一号 = 26pt
}

# ═══════════════════════════════════════════════════════════════════════════════
# Tool Implementation
# ═══════════════════════════════════════════════════════════════════════════════

def _find_latest_report() -> Optional[Path]:
    """Find the most recently generated report .docx file."""
    if not STORAGE_DIR.exists():
        return None
    docx_files = sorted(
        STORAGE_DIR.glob("*.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    return docx_files[0] if docx_files else None


def _get_report_path(file_path: Optional[str] = None) -> Optional[Path]:
    """Resolve a report file path."""
    if file_path:
        path = Path(file_path)
        if path.exists():
            return path
        # Try relative to storage
        path = STORAGE_DIR / file_path
        if path.exists():
            return path
        return None
    return _find_latest_report()


def _load_docx(file_path: str) -> tuple:
    """Load a .docx file and return (doc, paragraphs, styles, tables)."""
    from docx import Document
    from docx.shared import Pt

    doc = Document(file_path)
    paragraphs = doc.paragraphs
    tables = doc.tables
    sections = doc.sections

    return doc, paragraphs, tables, sections


def validate_report_impl(file_path: Optional[str] = None) -> str:
    """Comprehensive report format validation."""
    path = _get_report_path(file_path)
    if not path:
        return json.dumps({
            "error": "找不到报告文件",
            "hint": f"请将.docx文件放入 {STORAGE_DIR} 或提供完整路径",
        }, ensure_ascii=False)

    try:
        doc, paragraphs, tables, sections = _load_docx(str(path))
    except Exception as e:
        return json.dumps({"error": f"无法打开文件: {str(e)}"}, ensure_ascii=False)

    results = {
        "file": str(path.name),
        "file_size_kb": round(path.stat().st_size / 1024, 1),
        "total_paragraphs": len(paragraphs),
        "total_tables": len(tables),
        "total_sections": len(sections),
        "checks": {},
    }

    issues = []
    warnings = []
    passed = []

    # ── 1. Font Check ──
    font_issues = []
    for i, para in enumerate(paragraphs[:200]):  # Check first 200 paragraphs
        style_name = para.style.name if para.style else ""
        for run in para.runs:
            font_name = run.font.name
            if style_name.startswith("Heading") and font_name:
                # Headings should use 黑体
                if "heading 1" in style_name.lower() or "heading 2" in style_name.lower():
                    if "黑体" not in (font_name or ""):
                        font_issues.append(f"段落{i}: {style_name}使用字体'{font_name}'，应为'黑体'")

    if font_issues:
        issues.extend(font_issues[:10])
        results["checks"]["fonts"] = {"status": "warning", "issues": len(font_issues)}
    else:
        passed.append("字体")
        results["checks"]["fonts"] = {"status": "pass"}

    # ── 2. Heading Structure Check ──
    headings = []
    for i, para in enumerate(paragraphs):
        if para.style and para.style.name.startswith("Heading"):
            headings.append({
                "index": i,
                "level": int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1,
                "text": para.text[:80],
            })

    results["checks"]["headings"] = {
        "status": "pass" if headings else "warning",
        "count": len(headings),
        "list": headings[:20],
    }

    if not headings:
        warnings.append("未检测到标题样式（Heading），请确保使用了Word标题样式")
    elif len(headings) < 5:
        warnings.append(f"标题数量较少({len(headings)})，报告应包含至少10个一级标题（10个章节）")

    # Check for chapter headings (第X章 pattern)
    chapter_headings = [h for h in headings if "第" in h["text"] and "章" in h["text"]]
    if chapter_headings:
        passed.append(f"章节标题({len(chapter_headings)}章)")
    else:
        warnings.append("未找到'第X章'格式的章节标题，请检查章节命名")

    # ── 3. Table Format Check ──
    table_results = []
    for ti, table in enumerate(tables[:20]):
        rows = len(table.rows)
        cols = len(table.columns)
        has_borders = _check_table_borders(table)
        has_header = _check_table_header(table)

        table_results.append({
            "table_index": ti,
            "rows": rows,
            "cols": cols,
            "has_borders": has_borders,
            "has_header_style": has_header,
        })

        if not has_borders:
            warnings.append(f"表格{ti+1}({rows}行×{cols}列): 缺少边框（应使用'Table Grid'样式）")

    results["checks"]["tables"] = {
        "status": "pass" if not warnings else "warning",
        "count": len(tables),
        "details": table_results,
    }

    if tables:
        passed.append(f"表格({len(tables)}个)")
    else:
        warnings.append("未检测到表格，报告应包含数据表格")

    # ── 4. Page/Margin Check ──
    margin_issues = []
    for si, section in enumerate(sections):
        left_margin = section.left_margin
        right_margin = section.right_margin
        if left_margin and left_margin < 914400:  # 1 inch = 914400 EMU
            margin_issues.append(f"第{si+1}节左边距过小")
        if right_margin and right_margin < 914400:
            margin_issues.append(f"第{si+1}节右边距过小")

    results["checks"]["margins"] = {
        "status": "warning" if margin_issues else "pass",
        "issues": margin_issues,
    }

    if not margin_issues:
        passed.append("页边距")

    # ── 5. Content Completeness Check ──
    full_text = " ".join(p.text for p in paragraphs)
    text_length = len(full_text)
    results["checks"]["content"] = {
        "status": "pass" if text_length > 1000 else "warning",
        "total_chars": text_length,
    }

    if text_length > 5000:
        passed.append(f"内容({text_length}字)")
    elif text_length < 1000:
        warnings.append(f"正文内容过少({text_length}字)，正式报告通常需要数万字")

    # ── Summary ──
    total_checks = len(results["checks"])
    passed_count = sum(1 for c in results["checks"].values() if c.get("status") == "pass")

    return json.dumps({
        **results,
        "summary": {
            "passed": passed_count,
            "total": total_checks,
            "score": f"{passed_count}/{total_checks}",
            "passed_items": passed,
            "issues": issues[:10],
            "warnings": warnings[:10],
            "recommendation": _get_recommendation(passed_count, total_checks, warnings),
        },
        "validation_time": "",
    }, ensure_ascii=False, indent=2)


def _check_table_borders(table) -> bool:
    """Check if a table has borders defined."""
    if not table.rows:
        return True
    try:
        # Check first cell for border XML
        cell = table.rows[0].cells[0] if table.rows[0].cells else None
        if cell is None:
            return True
        tc_pr = cell._tc.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
        if tc_pr is not None:
            borders = tc_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
            if borders is not None:
                return True
    except Exception:
        pass
    return table.style is not None and "grid" in (table.style.name or "").lower()


def _check_table_header(table) -> bool:
    """Check if first row has bold or special formatting (header row)."""
    if not table.rows:
        return False
    first_row = table.rows[0]
    for cell in first_row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                if run.bold:
                    return True
    return False


def _get_recommendation(passed: int, total: int, warnings: list) -> str:
    """Generate a recommendation based on check results."""
    ratio = passed / total if total > 0 else 0
    if ratio >= 0.8 and not warnings:
        return "✅ 报告格式基本符合标准，建议在WPS中进行最终确认"
    elif ratio >= 0.6:
        return "⚠️ 报告存在格式问题，建议修正后重新检查"
    else:
        return "❌ 报告格式需要大幅调整，请参考DB32/T4013-2021标准检查"


def check_fonts_impl(file_path: Optional[str] = None) -> str:
    """Detailed font compliance check."""
    path = _get_report_path(file_path)
    if not path:
        return json.dumps({"error": "找不到报告文件"}, ensure_ascii=False)

    doc, paragraphs, tables, sections = _load_docx(str(path))

    font_usage = {}
    style_fonts = {}

    for para in paragraphs:
        style_name = para.style.name if para.style else "Normal"
        for run in para.runs:
            font_name = run.font.name or "默认字体"
            font_size = run.font.size
            size_pt = round(font_size / 12700, 1) if font_size else None  # EMU to pt

            if style_name not in style_fonts:
                style_fonts[style_name] = set()
            style_fonts[style_name].add(font_name)

            key = f"{font_name}({size_pt}pt)" if size_pt else font_name
            font_usage[key] = font_usage.get(key, 0) + 1

    # Sort by usage count
    sorted_fonts = sorted(font_usage.items(), key=lambda x: x[1], reverse=True)

    return json.dumps({
        "file": str(path.name),
        "font_usage": {k: v for k, v in sorted_fonts[:20]},
        "style_fonts": {k: list(v) for k, v in list(style_fonts.items())[:15]},
        "required": {
            "body": REQUIRED_FONTS["body"],
            "heading1": REQUIRED_FONTS["heading1"],
            "heading2": REQUIRED_FONTS["heading2"],
        },
    }, ensure_ascii=False, indent=2)


def check_headings_impl(file_path: Optional[str] = None) -> str:
    """Heading structure and hierarchy validation."""
    path = _get_report_path(file_path)
    if not path:
        return json.dumps({"error": "找不到报告文件"}, ensure_ascii=False)

    doc, paragraphs, tables, sections = _load_docx(str(path))

    headings = []
    current_chapter = 0
    hierarchy_ok = True
    prev_level = 0

    for i, para in enumerate(paragraphs):
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
            text = para.text[:100]

            # Track chapter number
            import re
            ch_match = re.match(r'第([一二三四五六七八九十\d]+)章', text)
            if ch_match and level == 1:
                current_chapter += 1

            # Check hierarchy: level should not jump more than 1
            if prev_level > 0 and level > prev_level + 1:
                hierarchy_ok = False

            prev_level = level

            headings.append({
                "para_index": i,
                "level": level,
                "chapter": current_chapter if current_chapter > 0 else None,
                "text": text,
            })

    return json.dumps({
        "file": str(path.name),
        "total_headings": len(headings),
        "hierarchy_valid": hierarchy_ok,
        "chapter_headings": [h for h in headings if h["chapter"] is not None],
        "all_headings": headings[:30],
        "issues": [] if hierarchy_ok else ["标题层级跳跃：存在不连续的标题级别（如H1→H3）"],
    }, ensure_ascii=False, indent=2)


def check_tables_impl(file_path: Optional[str] = None) -> str:
    """Table formatting detailed check."""
    path = _get_report_path(file_path)
    if not path:
        return json.dumps({"error": "找不到报告文件"}, ensure_ascii=False)

    doc, paragraphs, tables, sections = _load_docx(str(path))

    table_details = []
    for ti, table in enumerate(tables):
        rows = len(table.rows)
        cols = len(table.columns)
        has_header = _check_table_header(table)
        has_borders = _check_table_borders(table)

        # Check first row as potential header
        header_text = []
        if table.rows:
            for cell in table.rows[0].cells:
                header_text.append(cell.text[:50])

        # Check for [待填写] placeholders
        placeholder_cells = 0
        for row in table.rows:
            for cell in row.cells:
                if "[待填写]" in cell.text:
                    placeholder_cells += 1

        table_details.append({
            "index": ti,
            "dimensions": f"{rows}行 × {cols}列",
            "has_header": has_header,
            "has_borders": has_borders,
            "header_cells": header_text[:8],
            "placeholder_cells": placeholder_cells,
            "style": table.style.name if table.style else "无",
        })

    return json.dumps({
        "file": str(path.name),
        "total_tables": len(tables),
        "tables": table_details,
        "issues": [
            f"表格{t['index']+1}: {t['placeholder_cells']}个占位符待填写"
            for t in table_details if t["placeholder_cells"] > 0
        ],
    }, ensure_ascii=False, indent=2)


def open_in_wps_impl(file_path: Optional[str] = None) -> str:
    """Open the report in WPS Office for visual inspection."""
    path = _get_report_path(file_path)
    if not path:
        return json.dumps({"error": "找不到报告文件"}, ensure_ascii=False)

    try:
        # Use WPS Office to open the file
        subprocess.run(
            ["open", "-a", "WPS Office", str(path)],
            check=True,
            timeout=5,
        )
        return json.dumps({
            "success": True,
            "message": f"已在WPS Office中打开: {path.name}",
            "file": str(path.name),
        }, ensure_ascii=False)
    except subprocess.TimeoutExpired:
        return json.dumps({"error": "打开WPS Office超时"}, ensure_ascii=False)
    except Exception as e:
        # Fallback: use default app
        try:
            subprocess.run(["open", str(path)], check=True)
            return json.dumps({
                "success": True,
                "message": f"已用默认应用打开: {path.name}",
                "note": "WPS Office打开失败，已使用系统默认应用",
            }, ensure_ascii=False)
        except Exception as e2:
            return json.dumps({"error": f"无法打开文件: {str(e2)}"}, ensure_ascii=False)


def list_structure_impl(file_path: Optional[str] = None) -> str:
    """Show document chapter/section structure."""
    path = _get_report_path(file_path)
    if not path:
        return json.dumps({"error": "找不到报告文件"}, ensure_ascii=False)

    doc, paragraphs, tables, sections = _load_docx(str(path))

    import re

    structure = []
    current_chapter = {"number": 0, "title": "", "sections": [], "tables_in_chapter": 0}
    table_count = 0

    for para in paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        # Chapter heading detection
        ch_match = re.match(r'第([一二三四五六七八九十\d]+)章\s*(.*)', text)
        if ch_match and "Heading" in style_name:
            if current_chapter["number"] > 0:
                structure.append(current_chapter)
            current_chapter = {
                "number": ch_match.group(1),
                "title": ch_match.group(2) or text,
                "sections": [],
                "tables_in_chapter": 0,
            }
        elif style_name.startswith("Heading 2"):
            current_chapter["sections"].append({
                "level": 2,
                "text": text[:80],
            })
        elif style_name.startswith("Heading 3"):
            current_chapter["sections"].append({
                "level": 3,
                "text": text[:80],
            })

    if current_chapter["number"] > 0 or current_chapter["sections"]:
        structure.append(current_chapter)

    # Count tables per chapter (approximation based on position)
    for ti, table in enumerate(tables):
        table_count += 1

    return json.dumps({
        "file": str(path.name),
        "chapters": [
            {
                "number": ch["number"],
                "title": f"第{ch['number']}章 {ch['title']}",
                "subsections": len(ch["sections"]),
                "detail": ch["sections"][:5] if ch["sections"] else [],
            }
            for ch in structure if ch["number"]
        ],
        "total_tables": table_count,
        "total_paragraphs": len(paragraphs),
    }, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════════════════════════════════
# MCP Tool Registration
# ═══════════════════════════════════════════════════════════════════════════════

@server.list_tools()
async def list_tools():
    return [
        Tool(
            name="validate_report",
            description="全面验证生成报告的格式合规性。检查字体、标题层级、表格格式、页边距、内容完整性，对照DB32/T4013-2021标准。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "报告.docx文件的路径。留空则自动检查最新生成的报告。",
                    },
                },
            },
        ),
        Tool(
            name="check_fonts",
            description="详细检查报告中使用的字体是否符合标准（正文仿宋_GB2312，标题黑体，三级标题楷体）。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "报告.docx文件的路径。",
                    },
                },
            },
        ),
        Tool(
            name="check_headings",
            description="检查报告标题层级结构。验证一级标题(章节)、二级标题(小节)的层次是否正确，是否连续。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "报告.docx文件的路径。",
                    },
                },
            },
        ),
        Tool(
            name="check_tables",
            description="检查报告中的表格格式：边框、表头样式、占位符状态、行列数。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "报告.docx文件的路径。",
                    },
                },
            },
        ),
        Tool(
            name="open_in_wps",
            description="在WPS Office中打开报告文件，用于人工视觉审查格式。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "报告.docx文件的路径。留空则打开最新生成的报告。",
                    },
                },
            },
        ),
        Tool(
            name="list_structure",
            description="列出报告的文档结构：章节、小节、表格数、段落数。",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "报告.docx文件的路径。",
                    },
                },
            },
        ),
    ]


@server.call_tool()
async def call_tool(name: str, arguments: dict):
    """Route tool calls to implementations."""
    file_path = arguments.get("file_path")

    handlers = {
        "validate_report": validate_report_impl,
        "check_fonts": check_fonts_impl,
        "check_headings": check_headings_impl,
        "check_tables": check_tables_impl,
        "open_in_wps": open_in_wps_impl,
        "list_structure": list_structure_impl,
    }

    handler = handlers.get(name)
    if not handler:
        return [TextContent(type="text", text=json.dumps({"error": f"未知工具: {name}"}, ensure_ascii=False))]

    try:
        result = handler(file_path)
        return [TextContent(type="text", text=result)]
    except Exception as e:
        return [TextContent(type="text", text=json.dumps({"error": str(e)}, ensure_ascii=False))]


# ═══════════════════════════════════════════════════════════════════════════════
# Entry Point
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    """Run the MCP server via stdio."""
    import asyncio
    asyncio.run(_run())


async def _run():
    async with stdio_server() as (read_stream, write_stream):
        await server.run(read_stream, write_stream, server.create_initialization_options())


if __name__ == "__main__":
    main()
