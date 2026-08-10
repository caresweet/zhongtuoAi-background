"""Test Skill Integration — Quality Validation & Scoring.

Wraps the zhongtuo-report-test skill's four-dimensional testing:
- T1: 结构完整性 (Structure completeness)
- T2: 内容充实度 (Content richness)
- T3: 格式规范性 (Format compliance)
- T4: 数据准确性 (Data accuracy)

Integrates with existing validation/diff_engine.py and spec_checker.py.
"""

from dataclasses import dataclass, field
from typing import Dict, Any, Optional, List
import json


@dataclass
class TestResult:
    """Single test dimension result."""
    dimension: str          # "T1" | "T2" | "T3" | "T4"
    label: str              # Chinese label
    score: float            # 0-100
    max_score: float
    passed: bool
    details: List[str] = field(default_factory=list)
    suggestions: List[str] = field(default_factory=list)


@dataclass
class TestReport:
    """Complete test report with all four dimensions."""
    report_name: str
    overall_score: float         # 0-100
    overall_grade: str           # "A"/"B"/"C"/"D"
    passed: bool
    dimensions: List[TestResult] = field(default_factory=list)
    summary: str = ""
    recommendations: List[str] = field(default_factory=list)


class TestIntegration:
    """Integrates the test skill's quality validation into the backend.

    Provides:
    1. Four-dimensional quality scoring (T1-T4)
    2. Reference comparison against gold-standard reports
    3. Format compliance checking against DB32/T 4013-2021
    4. Auto-generated test reports
    """

    GRADE_THRESHOLDS = {
        "A": 85,   # >= 85: Excellent
        "B": 70,   # >= 70: Good
        "C": 55,   # >= 55: Needs improvement
        "D": 0,    # < 55: Failed
    }

    @classmethod
    def grade(cls, score: float) -> str:
        for grade, threshold in sorted(cls.GRADE_THRESHOLDS.items(),
                                        key=lambda x: -x[1]):
            if score >= threshold:
                return grade
        return "D"

    async def run_full_validation(
        self, generated_path: str,
        reference_path: str = "",
        template_path: str = "",
        context: Dict[str, Any] = None,
    ) -> TestReport:
        """Run all four test dimensions on a generated report.

        Args:
            generated_path: Path to the generated .docx report
            reference_path: Optional path to a gold-standard reference report
            template_path: Optional path to the template used
            context: Optional context (filled data, expected values)

        Returns:
            TestReport with all dimensions scored.
        """
        from pathlib import Path

        gen = Path(generated_path)
        if not gen.exists():
            return TestReport(
                report_name=gen.name,
                overall_score=0,
                overall_grade="D",
                passed=False,
                summary="Generated report file not found.",
            )

        report_name = gen.stem
        dimensions = []

        # T1: Structure completeness
        t1 = await self._check_structure_completeness(generated_path, template_path)
        dimensions.append(t1)

        # T2: Content richness
        t2 = self._check_content_richness(generated_path, context)
        dimensions.append(t2)

        # T3: Format compliance
        t3 = await self._check_format_compliance(generated_path)
        dimensions.append(t3)

        # T4: Data accuracy (if reference available)
        t4 = self._check_data_accuracy(generated_path, reference_path, context)
        dimensions.append(t4)

        # Calculate overall
        weights = {"T1": 0.25, "T2": 0.30, "T3": 0.25, "T4": 0.20}
        overall = sum(d.score * weights.get(d.dimension, 1/4) for d in dimensions)
        grade = self.grade(overall)

        # Build summary
        dim_summary = []
        for d in dimensions:
            icon = "✅" if d.passed else "⚠️"
            dim_summary.append(f"{icon} {d.label}: {d.score:.0f}/100")

        recommendations = []
        for d in dimensions:
            recommendations.extend(d.suggestions)

        return TestReport(
            report_name=report_name,
            overall_score=round(overall, 1),
            overall_grade=grade,
            passed=grade in ("A", "B"),
            dimensions=dimensions,
            summary=" | ".join(dim_summary),
            recommendations=recommendations,
        )

    # ═══════════════════════════════════════════════════════════════
    # T1: Structure Completeness
    # ═══════════════════════════════════════════════════════════════

    async def _check_structure_completeness(
        self, filepath: str, template_path: str = ""
    ) -> TestResult:
        """Check if the report has all required sections/chapters."""
        from docx import Document

        try:
            doc = Document(filepath)
        except Exception as e:
            return TestResult(
                dimension="T1", label="结构完整性", score=0, max_score=100,
                passed=False, details=[f"Cannot open document: {e}"],
            )

        # Required chapters for 稳评报告
        required = [
            "拟征收决策基本概况",
            "评估过程",
            "社会稳定风险因素调查",
            "决策综合分析",
            "风险因素识别",
            "措施前风险等级研判",
            "风险防范与化解措施",
            "措施后风险等级评估",
            "评估结论与建议",
            "应急预案",
        ]

        full_text = "\n".join(p.text for p in doc.paragraphs)
        full_text_lower = full_text.lower()

        found = 0
        details = []
        missing = []
        for ch in required:
            if ch in full_text:
                found += 1
            else:
                # Try partial match
                partial = ch[:4] if len(ch) >= 4 else ch
                if partial in full_text:
                    found += 0.7
                    details.append(f"⚠️ 部分匹配: {ch}")
                else:
                    missing.append(ch)

        score = (found / len(required)) * 100
        passed = score >= 80

        if missing:
            details.append(f"❌ 缺失章节 ({len(missing)}): {', '.join(missing)}")
        else:
            details.append(f"✅ 全部 {len(required)} 个必需章节均已覆盖")

        suggestions = []
        if missing:
            suggestions.append(f"补充缺失章节: {', '.join(missing[:3])}")

        return TestResult(
            dimension="T1", label="结构完整性",
            score=round(score, 1), max_score=100,
            passed=passed, details=details, suggestions=suggestions,
        )

    # ═══════════════════════════════════════════════════════════════
    # T2: Content Richness
    # ═══════════════════════════════════════════════════════════════

    def _check_content_richness(
        self, filepath: str, context: Dict[str, Any] = None
    ) -> TestResult:
        """Check content quality: word count, data presence, detail level."""
        from docx import Document

        try:
            doc = Document(filepath)
        except Exception:
            return TestResult(
                dimension="T2", label="内容充实度", score=0, max_score=100,
                passed=False, details=["无法打开文档"],
            )

        text = "\n".join(p.text for p in doc.paragraphs)
        tables = doc.tables
        char_count = len(text)

        details = []
        score = 0
        max_score = 100

        # Check 1: Minimum content length (稳评报告 typically 10000-30000 chars)
        if char_count >= 20000:
            score += 30
            details.append(f"✅ 内容长度充足: {char_count} 字符")
        elif char_count >= 10000:
            score += 20
            details.append(f"⚠️ 内容长度偏低: {char_count} 字符（建议>2万字符）")
        elif char_count >= 5000:
            score += 10
            details.append(f"⚠️ 内容长度不足: {char_count} 字符（严重偏低）")
        else:
            details.append(f"❌ 内容过短: {char_count} 字符")

        # Check 2: Data tables (稳评报告 should have risk tables, survey tables, etc.)
        if len(tables) >= 5:
            score += 25
            details.append(f"✅ 数据表格充足: {len(tables)} 个")
        elif len(tables) >= 3:
            score += 15
            details.append(f"⚠️ 数据表格偏少: {len(tables)} 个（建议≥5个）")
        else:
            score += 5
            details.append(f"⚠️ 数据表格不足: {len(tables)} 个")

        # Check 3: Contains key terms (法律条文引用)
        key_terms = ["土地管理法", "社会风险", "问卷调查", "评估单位", "风险等级"]
        term_count = sum(1 for t in key_terms if t in text)
        if term_count >= 4:
            score += 25
            details.append(f"✅ 关键术语覆盖: {term_count}/{len(key_terms)}")
        elif term_count >= 2:
            score += 15
            details.append(f"⚠️ 关键术语不足: {term_count}/{len(key_terms)}")
        else:
            score += 5
            details.append(f"❌ 关键术语缺失: {term_count}/{len(key_terms)}")

        # Check 4: Contains numbers/data (量化的风险评估应有具体数值)
        import re
        numbers = re.findall(r'\d+', text)
        if len(numbers) >= 50:
            score += 20
            details.append(f"✅ 数据点充足: {len(numbers)} 个数值")
        elif len(numbers) >= 20:
            score += 10
            details.append(f"⚠️ 数据点偏少: {len(numbers)} 个数值")
        else:
            details.append(f"⚠️ 缺少量化数据")

        passed = score >= 60
        suggestions = []
        if char_count < 10000:
            suggestions.append("增加各章节分析深度，目标≥2万字符")
        if len(tables) < 5:
            suggestions.append("增加风险因素表格、调查统计表等数据表格")
        if term_count < 4:
            suggestions.append("补充法律条文引用和专业术语")

        return TestResult(
            dimension="T2", label="内容充实度",
            score=round(score, 1), max_score=max_score,
            passed=passed, details=details, suggestions=suggestions,
        )

    # ═══════════════════════════════════════════════════════════════
    # T3: Format Compliance
    # ═══════════════════════════════════════════════════════════════

    async def _check_format_compliance(self, filepath: str) -> TestResult:
        """Check format compliance against DB32/T 4013-2021 standards.

        Uses the existing format_fixer and spec_checker where available.
        """
        details = []
        score = 0
        max_score = 100

        # Try existing spec checker first
        try:
            from app.validation.spec_checker import check_report_async
            spec_report = await check_report_async(filepath)
            if spec_report:
                score = spec_report.compliance_score * 100
                details.append(
                    f"格式规范校验: 合规率 {spec_report.compliance_score:.0%}")
                if spec_report.passed:
                    details.append("✅ 格式符合DB32/T 4013-2021标准")
                else:
                    for c in spec_report.checks:
                        if not c.passed:
                            for issue in c.issues[:2]:
                                details.append(f"⚠️ {issue}")
                return TestResult(
                    dimension="T3", label="格式规范性",
                    score=round(score, 1), max_score=max_score,
                    passed=spec_report.passed, details=details,
                )
        except Exception:
            pass

        # Fallback: basic format checks
        from docx import Document
        try:
            doc = Document(filepath)
        except Exception:
            return TestResult(
                dimension="T3", label="格式规范性", score=0, max_score=100,
                passed=False, details=["无法打开文档"],
            )

        # Check font styles
        para_count = len(doc.paragraphs)
        styled_count = sum(1 for p in doc.paragraphs if p.style and p.style.name)
        score += 20 if styled_count > para_count * 0.5 else 10

        # Check has page structure (sections/breaks)
        has_sections = len(doc.sections) > 0
        score += 15 if has_sections else 0

        # Check headers/footers
        has_header = any(s.header for s in doc.sections if s.header)
        has_footer = any(s.footer for s in doc.sections if s.footer)
        if has_header and has_footer:
            score += 15
            details.append("✅ 包含页眉和页脚")
        elif has_header or has_footer:
            score += 8
        else:
            details.append("⚠️ 缺少页眉/页脚")

        # Check images (稳评报告应包含位置图、照片等)
        from lxml import etree
        xml = etree.tostring(doc.element, encoding='unicode')
        image_count = xml.count('w:drawing') + xml.count('wp:inline')
        if image_count >= 3:
            score += 20
            details.append(f"✅ 包含 {image_count} 张图片")
        elif image_count >= 1:
            score += 10
            details.append(f"⚠️ 图片偏少: {image_count} 张（建议≥3张）")
        else:
            details.append("⚠️ 未检测到图片（建议添加位置图、照片等）")

        # General formatting score
        score += 30  # Base score if document opens
        passed = score >= 70

        return TestResult(
            dimension="T3", label="格式规范性",
            score=round(score, 1), max_score=max_score,
            passed=passed, details=details,
            suggestions=["建议通过 spec_checker 进行更详细格式校验"]
            if not passed else [],
        )

    # ═══════════════════════════════════════════════════════════════
    # T4: Data Accuracy
    # ═══════════════════════════════════════════════════════════════

    def _check_data_accuracy(
        self, generated_path: str,
        reference_path: str = "",
        context: Dict[str, Any] = None,
    ) -> TestResult:
        """Check if the correct data was used in key fields.

        Compares against reference if available, or verifies expected values.
        """
        details = []
        score = 0
        max_score = 100
        context = context or {}

        try:
            from docx import Document
            doc = Document(reference_path) if reference_path else None
            gen = Document(generated_path)
        except Exception:
            return TestResult(
                dimension="T4", label="数据准确性", score=0, max_score=100,
                passed=False, details=["无法读取文档进行数据比对"],
            )

        gen_text = "\n".join(p.text for p in gen.paragraphs)
        checks = []

        if reference_path and doc:
            ref_text = "\n".join(p.text for p in doc.paragraphs)
            # Compare structure similarity
            gen_paras = len(gen.paragraphs)
            ref_paras = len(doc.paragraphs)
            ratio = min(gen_paras, ref_paras) / max(gen_paras, ref_paras)
            if ratio >= 0.7:
                score += 30
                details.append(f"✅ 段落结构相似度: {ratio:.0%}")
            else:
                score += 15
                details.append(f"⚠️ 段落结构差异大: 对比 {gen_paras} vs {ref_paras}")
        else:
            score += 30  # No reference, assume OK

        # Check specific data from context
        filled = context.get("filled_data", context.get("basic_info", {}))
        if filled:
            data_found = 0
            for key, val in filled.items():
                if isinstance(val, str) and len(val) > 2 and val in gen_text:
                    data_found += 1
            if filled:
                match_rate = data_found / len(filled)
                if match_rate >= 0.8:
                    score += 40
                    details.append(f"✅ 填充数据匹配率: {match_rate:.0%} ({data_found}/{len(filled)})")
                else:
                    score += 20
                    details.append(f"⚠️ 数据匹配率偏低: {match_rate:.0%} ({data_found}/{len(filled)})")
        else:
            score += 30  # No context to check

        # Check no leftover template markers
        template_markers = ["XX", "点击此处", "需后期", "【", "】", "（", "）"]
        found_markers = [m for m in template_markers if m in gen_text and gen_text.count(m) > 5]
        if not found_markers:
            score += 30
            details.append("✅ 无模板占位符残留")
        else:
            score += 10
            details.append(f"⚠️ 检测到可能的模板残留: {found_markers}")

        passed = score >= 60
        suggestions = []
        if not passed:
            suggestions.append("建议提供参考报告进行比对校验")
            suggestions.append("检查填充数据是否完整写入报告")

        return TestResult(
            dimension="T4", label="数据准确性",
            score=round(score, 1), max_score=max_score,
            passed=passed, details=details, suggestions=suggestions,
        )

    # ═══════════════════════════════════════════════════════════════
    # API response helpers
    # ═══════════════════════════════════════════════════════════════

    def to_api_response(self, report: TestReport) -> Dict[str, Any]:
        """Convert TestReport to API-friendly dict."""
        return {
            "report_name": report.report_name,
            "overall_score": report.overall_score,
            "overall_grade": report.overall_grade,
            "passed": report.passed,
            "summary": report.summary,
            "dimensions": [
                {
                    "id": d.dimension,
                    "label": d.label,
                    "score": d.score,
                    "max_score": d.max_score,
                    "passed": d.passed,
                    "details": d.details,
                    "suggestions": d.suggestions,
                }
                for d in report.dimensions
            ],
            "recommendations": report.recommendations,
        }


# Singleton
test_integration = TestIntegration()
