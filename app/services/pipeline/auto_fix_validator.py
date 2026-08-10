"""Post-generation auto-fix validator.

Detects and auto-fixes common issues found during quality review:
- Broken dates, duplicate names, fabricated data
- Old area data from example report
- Old dates and survey numbers from example report
- Agent meta-text and notes
"""

import re
from typing import Dict, List
from pathlib import Path


class AutoFixValidator:
    """Auto-fix common report generation issues."""

    AUTO_FIX_PATTERNS = [
        # Broken dates
        (r'(?<![0-9])026年', '2026年'),
        (r'(?<![0-9])024年', '2024年'),
        # Duplicate names
        (r'朱坝街道办事处、朱坝街道办事处', '朱坝街道办事处'),
        (r'三圩社区、三圩社区', '三圩社区'),
        (r'高良涧街道三圩', '朱坝街道三圩'),
        # Old area data from example report
        (r'(?<!\d)108006(?!\d)', '326342'),
        (r'(?<!\d)105008(?!\d)', '326342'),
        (r'(?<!\d)323,?469(?!\d)', '326342'),  # 勘测定界报告实际面积
        (r'(?<!\d)2998(?!\d)', '0'),
        (r'约162\.01亩', '约489.51亩'),
        (r'约157\.51亩', '约489.51亩'),
        (r'约4\.5亩', '约0亩'),
        (r'(?<!\d)485\.20(?!\d)', '489.51'),  # 亩数修正 (326342/666.67)
        # Old dates
        (r'2024年6月13日', '2026年4月13日'),
        (r'2024年6月14日', '2026年4月13日'),
        (r'2024年6月27日', '2026年4月19日'),
        (r'6月14日至6月27日', '4月13日至4月19日'),
        (r'2024年(?=7月)', '2026年'),

        # 🔴 Notice period: 公告期限统一为 7 天（从4月13日起）
        (r'公示时间12天', '公示时间7天'),
        (r'公示期12天', '公示期7天'),
        (r'公告期限为2026年4月13日至2026年4月24日', '公告期限为2026年4月13日至2026年4月19日'),
        (r'4月13日至4月24日', '4月13日至4月19日'),
        (r'4月13日-2026年4月24日', '4月13日-2026年4月19日'),
        (r'2026年4月24日', '2026年4月19日'),
        # Wrong June/July internal work dates → correct April timeline
        (r'6月26日至7月3日', '4月13日至4月19日'),
        (r'2026年6月26日-2026年7月3日', '2026年4月13日-2026年4月19日'),
        (r'2026年6月26日', '2026年4月13日'),
        (r'2026年7月3日', '2026年4月19日'),
        # Internal work dates: after notice ends 4/19
        (r'2026年7月4日-2026年7月5日', '2026年4月20日-2026年4月21日'),
        (r'2026年7月6日-2026年7月7日', '2026年4月22日-2026年4月23日'),
        (r'2026年7月8日-2026年7月10日', '2026年4月24日-2026年4月26日'),
        # Single July dates
        (r'2026年7月4日', '2026年4月20日'),
        (r'2026年7月5日', '2026年4月21日'),
        (r'2026年7月6日', '2026年4月22日'),
        (r'2026年7月7日', '2026年4月23日'),
        (r'2026年7月8日', '2026年4月24日'),
        (r'2026年7月10日', '2026年4月26日'),
        # Month-only patterns
        (r'2026年7月', '2026年4月'),
        (r'2026年6月', '2026年4月'),
        # Day count
        (r'12天', '7天'),
        # Old project names
        (r'金湖县', '洪泽区'),
        (r'戴楼街道', '朱坝街道'),
        (r'大魏社区', '三圩社区'),
        (r'洪泽园三村社区', '三圩社区'),
        (r'洞庭湖路', ''),
        # Agent meta-text
        (r'好的[，,]?\s*作为稳评报告编制专家[，,]?\s*以下是针对该项目单元格的内容[：:]?\s*', ''),
        (r'好的[，,]?\s*根据您提供的要求和建议[，,]?\s*现完成对单元格内容的最终修订[：:]?\s*', ''),
        (r'[（(]注[：:][^）)]*[）)]', ''),
        (r'[（(]全文共\d+字[，,][^）)]*[）)]', ''),
        (r'[（(]较原稿删减[^）)]*[）)]', ''),
        (r'[（(]最终版本已[^）)]*[）)]', ''),
    ]

    FABRICATED_GROUPS = ['十一组', '十二组', '八组', '九组', '十组', '一组', '五组']

    @classmethod
    def fix(cls, doc_path: str) -> Dict[str, int]:
        """Run auto-fix on a generated report."""
        from docx import Document

        doc = Document(doc_path)
        counts = {'paragraph': 0, 'table': 0, 'fabricated': 0}

        # Fix paragraphs
        paras = list(doc.paragraphs)
        for pi, para in enumerate(paras):
            text = para.text
            if not text.strip():
                continue
            new_text = text
            fixed = False

            for pattern, replacement in cls.AUTO_FIX_PATTERNS:
                if re.search(pattern, new_text):
                    new_text = re.sub(pattern, replacement, new_text)
                    fixed = True

            for group in cls.FABRICATED_GROUPS:
                if group in new_text:
                    new_text = new_text.replace(f'、{group}', '')
                    new_text = new_text.replace(f'{group}、', '')
                    new_text = new_text.replace(group, '')
                    counts['fabricated'] += 1
                    fixed = True

            # Remove duplicate consecutive paragraphs
            if pi + 1 < len(paras):
                if text.strip() == paras[pi + 1].text.strip() and len(text.strip()) > 20:
                    for run in paras[pi + 1].runs:
                        run.text = ''
                    fixed = True

            if fixed and para.runs:
                for run in para.runs:
                    run.text = ''
                para.runs[0].text = new_text.strip()
                counts['paragraph'] += 1

        # Fix tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        new_text = text
                        fixed = False

                        for pattern, replacement in cls.AUTO_FIX_PATTERNS:
                            if re.search(pattern, new_text):
                                new_text = re.sub(pattern, replacement, new_text)
                                fixed = True

                        if fixed and para.runs:
                            for run in para.runs:
                                run.text = ''
                            para.runs[0].text = new_text.strip()
                            counts['table'] += 1

        # ── Add page breaks before each chapter ──
        cls._add_chapter_page_breaks(doc)
        counts['pagebreaks'] = cls._pagebreak_count

        doc.save(doc_path)
        return counts

    _pagebreak_count = 0

    @classmethod
    def _add_chapter_page_breaks(cls, doc):
        """Ensure each major chapter starts on a new page."""
        import re
        from lxml import etree
        from docx.oxml.ns import qn

        cls._pagebreak_count = 0
        first_chapter_found = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if not re.match(r'^第[一二三四五六七八九十\d]+章', text):
                continue

            if not first_chapter_found:
                first_chapter_found = True
                continue  # Skip first chapter (评审报告表)

            # Add pageBreakBefore to previous paragraph
            prev_elem = para._p.getprevious()
            if prev_elem is None:
                continue

            # Find or create pPr on previous paragraph
            pPr = prev_elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = etree.SubElement(prev_elem, qn('w:pPr'))

            if pPr.find(qn('w:pageBreakBefore')) is None:
                etree.SubElement(pPr, qn('w:pageBreakBefore'))
                cls._pagebreak_count += 1
