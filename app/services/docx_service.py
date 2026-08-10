"""Docx service — low-level python-docx operations for reading and filling templates."""
import re
import copy
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class SectionInfo:
    """Represents a document section (heading + its content)."""
    index: int
    title: str
    level: int  # heading level 1-6
    start_para_index: int
    end_para_index: Optional[int] = None
    placeholders: List["PlaceholderInfo"] = field(default_factory=list)


@dataclass
class PlaceholderInfo:
    """Represents a single placeholder in the document."""
    key: str
    display_name: str
    section_index: int
    section_title: str
    paragraph_index: int
    run_index: Optional[int] = None
    expected_type: str = "text"
    description: str = ""
    is_required: bool = True
    original_text: str = ""  # The full text of the run containing the placeholder
    placeholder_pattern: str = ""  # The actual pattern matched, e.g. "{{project_name}}"


# Regex patterns for detecting placeholders in docx text
PLACEHOLDER_PATTERNS = [
    (r'\{\{(\w+)\}\}', "double_brace"),           # {{project_name}}
    (r'\{\{([^}]+)\}\}', "double_brace_cn"),      # {{项目名称}}
    (r'[［【]([^］】]+)[］】]', "bracket_cn"),     # 【项目名称】
    (r'<([^>]+)>', "angle_bracket"),              # <project_name>
    (r'_{3,}', "underscore_blank"),               # ________
]


class DocxService:
    """Handles reading and writing docx files for template analysis and filling."""

    # ---- READING / ANALYSIS ----

    @classmethod
    def extract_structure(cls, docx_path: str) -> Dict[str, Any]:
        """
        Extract the structural map of a docx file.
        Returns a dict with sections and paragraphs suitable for AI analysis.
        """
        doc = Document(docx_path)
        sections = []
        current_section = None
        para_index = 0

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""

            # Detect headings
            if style_name.startswith("Heading") or style_name.startswith("heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1

                # Close previous section
                if current_section:
                    current_section["end_para_index"] = para_index - 1
                    sections.append(current_section)

                # Start new section
                current_section = {
                    "index": len(sections),
                    "title": paragraph.text.strip(),
                    "level": level,
                    "start_para_index": para_index,
                    "end_para_index": None,
                    "paragraphs": [],
                }

            # Add paragraph to current section
            if current_section:
                placeholders_found = cls._find_placeholders(paragraph.text)
                current_section["paragraphs"].append({
                    "para_index": para_index,
                    "style": style_name,
                    "text_preview": paragraph.text[:200] if paragraph.text else "",
                    "has_placeholder": len(placeholders_found) > 0,
                    "placeholders": placeholders_found,
                })

            para_index += 1

        # Close last section
        if current_section:
            current_section["end_para_index"] = para_index - 1
            sections.append(current_section)

        # Also extract tables
        tables_data = []
        for t_idx, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                cells_data = []
                for cell in row.cells:
                    cells_data.append(cell.text[:200])
                rows_data.append(cells_data)
            tables_data.append({
                "table_index": t_idx,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "preview": rows_data[:5],  # First 5 rows as preview
            })

        return {
            "sections": sections,
            "tables": tables_data,
            "total_paragraphs": para_index,
            "total_sections": len(sections),
        }

    @classmethod
    def find_all_placeholders(
        cls, docx_path: str, example_path: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Scan the entire document and return all placeholder locations with context.
        This is the comprehensive scan used for template analysis.
        If example_path is provided, also diff against the example report
        to find missing body text paragraphs.
        """
        doc = Document(docx_path)
        all_placeholders = []
        current_section_index = -1
        current_section_title = ""

        for p_idx, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name if paragraph.style else ""

            # Track heading changes
            if style_name.startswith("Heading") or style_name.startswith("heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                current_section_index += 1
                current_section_title = paragraph.text.strip()

            # Check each run in the paragraph
            for r_idx, run in enumerate(paragraph.runs):
                text = run.text
                if not text or not text.strip():
                    continue

                for pattern, pattern_type in PLACEHOLDER_PATTERNS:
                    matches = list(re.finditer(pattern, text))
                    for match in matches:
                        key = match.group(1) if match.lastindex else match.group(0)
                        all_placeholders.append({
                            "paragraph_index": p_idx,
                            "run_index": r_idx,
                            "section_index": current_section_index,
                            "section_title": current_section_title,
                            "key": key.strip() if key else "",
                            "display_name": key.strip() if key else "",
                            "original_text": text,
                            "matched_text": match.group(0),
                            "pattern_type": pattern_type,
                            "style": style_name,
                        })

        # Also detect highlighted (yellow etc.) prefilled text as placeholders
        highlighted = cls._find_highlighted_runs(docx_path)
        all_placeholders.extend(highlighted)

        # Detect table cells that need filling (with optional example comparison)
        table_placeholders = cls._find_table_placeholders(docx_path, example_path)
        all_placeholders.extend(table_placeholders)

        # Detect inline form fields (label + spaces, no value)
        form_fields = cls._find_form_fields(docx_path)
        all_placeholders.extend(form_fields)

        # Detect inline wildcard placeholders: "2026年*月", "受***委托", "**街道"
        wildcards = cls._find_wildcard_placeholders(doc)
        all_placeholders.extend(wildcards)

        # If an example report is provided, diff to find missing body text
        if example_path and Path(example_path).exists():
            diff_placeholders = cls.diff_with_example(docx_path, example_path)
            all_placeholders.extend(diff_placeholders)

            # ---- Cross-reference form fields with example report ----
            # Form fields (e.g. "稳评责任单位：      ") only have the label from the template.
            # Search the example report for the label within a range, since template and
            # example may have different paragraph structures (e.g. the example wraps
            # long text across multiple paragraphs, shifting all subsequent indices).
            example_doc = Document(example_path)
            example_paras = example_doc.paragraphs
            for p in all_placeholders:
                if p.get("pattern_type") == "form_field" and not p.get("example_content"):
                    label = p.get("display_name", "")
                    # Strip emoji prefix for matching (e.g. "📋 稳评责任单位：" → "稳评责任单位：")
                    clean_label = label
                    for prefix in ("📋 ", "📝 ", "🖼️ ", "🟡 "):
                        if clean_label.startswith(prefix):
                            clean_label = clean_label[len(prefix):]
                            break
                    form_pi = p.get("paragraph_index", -1)

                    # Search nearby paragraphs for the label (exact index, then ±30)
                    best_text = ""
                    search_range = [form_pi]  # Try exact match first
                    for offset in range(1, 31):
                        search_range.append(form_pi - offset)
                        search_range.append(form_pi + offset)

                    for pi in search_range:
                        if 0 <= pi < len(example_paras):
                            ex_text = example_paras[pi].text.strip()
                            # Check if this paragraph starts with the label and has content after
                            if ex_text.startswith(clean_label) and len(ex_text) > len(clean_label) + 2:
                                best_text = ex_text
                                break
                            # Also check if the label appears anywhere and there's content
                            if clean_label in ex_text and ex_text != clean_label and len(ex_text) > len(clean_label) + 2:
                                if not best_text:
                                    best_text = ex_text
                    # Fallback: use exact paragraph index content if it's different from label
                    if not best_text and 0 <= form_pi < len(example_paras):
                        ex_text = example_paras[form_pi].text.strip()
                        if ex_text and ex_text != clean_label and len(ex_text) > 3:
                            best_text = ex_text

                    if best_text:
                        p["example_content"] = best_text

            # ---- Wildcard example lookup ----
            # Build heading-aligned mapping: template para index → example para index.
            # Use the same heading-matching logic as diff_with_example.
            _tpl_paras = doc.paragraphs
            _tpl_headings = {}  # heading_text → template_para_index
            _ex_headings = {}   # heading_text → example_para_index
            for i, para in enumerate(_tpl_paras):
                s = para.style.name if para.style else ""
                if s.startswith("Heading") and para.text.strip():
                    _tpl_headings[para.text.strip()] = i
            for i, para in enumerate(example_paras):
                s = para.style.name if para.style else ""
                if s.startswith("Heading") and para.text.strip():
                    _ex_headings[para.text.strip()] = i

            # Build tpl_para → ex_para mapping for body paragraphs within
            # matching sections
            _tpl_to_ex = {}  # template para index → example para index (or None)
            for h_text, t_h_idx in _tpl_headings.items():
                if h_text not in _ex_headings:
                    continue
                e_h_idx = _ex_headings[h_text]
                # Find next heading in template
                t_next = len(_tpl_paras)
                for _, t_i in _tpl_headings.items():
                    if t_i > t_h_idx and t_i < t_next:
                        t_next = t_i
                e_next = len(example_paras)
                for _, e_i in _ex_headings.items():
                    if e_i > e_h_idx and e_i < e_next:
                        e_next = e_i
                # Map body paragraphs positionally
                t_body = list(range(t_h_idx + 1, t_next))
                e_body = list(range(e_h_idx + 1, e_next))
                for pos in range(max(len(t_body), len(e_body))):
                    if pos < len(t_body):
                        t_pi = t_body[pos]
                        e_pi = e_body[pos] if pos < len(e_body) else None
                        if e_pi is not None:
                            _tpl_to_ex[t_pi] = e_pi

            for p in all_placeholders:
                if (p.get("pattern_type") == "highlight"
                        and p.get("key", "").startswith("wildcard")
                        and not p.get("example_content")):
                    p_idx = p.get("paragraph_index", -1)
                    if p_idx < 0 or p_idx >= len(_tpl_paras):
                        continue
                    tpl_text = _tpl_paras[p_idx].text
                    match_text = p.get("matched_text", "")
                    wc_pos = tpl_text.find(match_text)
                    if wc_pos < 0:
                        continue
                    # Use heading-aligned example paragraph
                    ex_p_idx = _tpl_to_ex.get(p_idx)
                    if ex_p_idx is None:
                        continue
                    ex_text = example_paras[ex_p_idx].text if ex_p_idx < len(example_paras) else ""
                    wc_len = len(match_text)
                    # Try exact position first
                    ex_value = ""
                    if wc_pos + wc_len <= len(ex_text):
                        ex_value = ex_text[wc_pos:wc_pos + wc_len]
                    if (not ex_value or '*' in ex_value or ex_value == match_text):
                        # Try shifted positions (±8 chars)
                        for shift in range(-8, 9):
                            sp = wc_pos + shift
                            if 0 <= sp and sp + wc_len <= len(ex_text):
                                ev = ex_text[sp:sp + wc_len]
                                if ev and '*' not in ev and ev != match_text:
                                    ex_value = ev
                                    break
                    if ex_value and '*' not in ex_value and ex_value != match_text:
                        p["example_content"] = ex_value

            # ---- Merge wildcards into paragraph-level placeholders ----
            # Paragraphs with wildcards (e.g. "2026年*月，受***委托") should be
            # presented as a single "fill this paragraph" question instead of
            # individual "*" questions.  Create an example_diff for the whole para
            # and mark the individual wildcards so they won't be asked separately.
            _wc_paras = {}  # para_index → list of wildcard keys
            for p in all_placeholders:
                if (p.get("pattern_type") == "highlight"
                        and p.get("key", "").startswith("wildcard")):
                    pi = p.get("paragraph_index", -1)
                    if pi >= 0:
                        _wc_paras.setdefault(pi, []).append(p)

            for pi, wc_list in _wc_paras.items():
                # Check if this paragraph already has an example_diff/blank
                has_diff = any(
                    pp.get("paragraph_index") == pi
                    and pp.get("pattern_type") in ("example_diff", "blank")
                    for pp in all_placeholders
                )
                if has_diff:
                    continue  # Already covered by diff placeholder

                # Create a paragraph-level example_diff for the whole paragraph
                tpl_text = doc.paragraphs[pi].text.strip() if pi < len(doc.paragraphs) else ""
                section_title = wc_list[0].get("section_title", "")
                section_idx = wc_list[0].get("section_index", -1)

                # Find example paragraph text for reference
                ex_ref = ""
                if example_path and Path(example_path).exists():
                    ex_p_idx = _tpl_to_ex.get(pi)
                    if ex_p_idx is not None and ex_p_idx < len(example_paras):
                        ex_ref = example_paras[ex_p_idx].text.strip()

                # Collect wildcard keys to merge
                wc_keys = [w["key"] for w in wc_list]

                all_placeholders.append({
                    "paragraph_index": pi,
                    "run_index": 0,
                    "section_index": section_idx,
                    "section_title": section_title,
                    "key": f"wildcard_para_{pi}",
                    "display_name": f"📝 「{section_title}」评估过程内容",
                    "original_text": tpl_text,
                    "matched_text": "",
                    "pattern_type": "example_diff",
                    "style": doc.paragraphs[pi].style.name if pi < len(doc.paragraphs) else "",
                    "example_content": ex_ref,
                    "_merged_wildcards": wc_keys,  # Mark wildcards to skip
                    "_is_wildcard_para": True,
                })

            # Detect figure captions and attachment items that need image uploads
            fig_placeholders = cls._find_figure_image_placeholders(
                docx_path, example_path
            )
            all_placeholders.extend(fig_placeholders)

            # ---- Dedup: remove duplicate image placeholders ----
            # Collect known image paragraphs from figure_img_* and attachment_img_*
            known_img_paras = set()
            for p in fig_placeholders:
                pi = p.get("paragraph_index", -1)
                known_img_paras.add(pi)
                # Image may be immediately before or after caption (±1)
                known_img_paras.add(pi - 1)
                known_img_paras.add(pi + 1)

            if known_img_paras:
                all_placeholders = [
                    p for p in all_placeholders
                    if not (
                        p.get("pattern_type") == "image_placeholder"
                        and p.get("key", "").startswith("image_")
                        and not p.get("key", "").endswith("_inline")  # Keep inline body images
                        and not p.get("key", "").endswith("_extra")   # Keep extra images
                        and p.get("paragraph_index") in known_img_paras
                    )
                ]

            # Detect phase stage labels (第一阶段/第二阶段/第三阶段/第四阶段)
            # These have short label text but need time+content filled in
            import re as _re_phase
            doc_ex_phase = Document(example_path)
            for p_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                # Match "第X阶段" or "第X阶段："
                if (_re_phase.match(r'第[一二三四五六七八九十\d]+阶段[：:]?$', text)
                        and len(text) < 15):
                    # Extract phase number for matching
                    phase_num = _re_phase.search(r'第([一二三四五六七八九十\d]+)阶段', text)
                    phase_str = phase_num.group(1) if phase_num else ""
                    # Find corresponding example content for this specific phase
                    ex_content = ""
                    for j in range(max(0, p_idx - 5), min(len(doc_ex_phase.paragraphs), p_idx + 10)):
                        ex_text = doc_ex_phase.paragraphs[j].text.strip()
                        if (ex_text and len(ex_text) > 10
                                and phase_str in ex_text and '阶段' in ex_text):
                            ex_content = ex_text[:300]
                            break
                    all_placeholders.append({
                        "paragraph_index": p_idx,
                        "run_index": 0,
                        "section_index": -1,
                        "section_title": "",
                        "key": f"phase_{p_idx}",
                        "display_name": f"📅 {text}",
                        "original_text": text,
                        "matched_text": text,
                        "pattern_type": "example_diff",
                        "style": para.style.name if para.style else "",
                        "example_content": ex_content,
                    })

            # Detect empty paragraphs after numbered sub-section labels
            # e.g. "4.4.2决策宣传、公示" followed by empty para → needs content
            import re as _re_subsec
            for p_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                style_name = para.style.name if para.style else ""
                # Skip headings and TOC entries — they're not sub-section labels
                if style_name.startswith("Heading") or style_name.startswith("heading"):
                    continue
                if style_name.startswith("toc") or style_name.startswith("TOC"):
                    continue
                # Check if this is a numbered sub-section label (e.g. "4.4.2...", "7.2.3...")
                is_sub_label = (
                    _re_subsec.match(r'\d+\.\d+\.?\d*\s*\S', text)
                    and len(text) < 60
                )
                if not is_sub_label:
                    continue
                # Check if NEXT paragraph is empty
                next_p_idx = p_idx + 1
                if next_p_idx >= len(doc.paragraphs):
                    continue
                next_text = doc.paragraphs[next_p_idx].text.strip()

                # ---- Always-empty topics ----
                # Some sub-section labels always need user-provided content,
                # even if the next paragraph has text (which may be about a
                # different topic, e.g. "调查时间" label followed by
                # "调查形式" text).  These topics are project-specific.
                always_fill_topics = ['调查时间']
                is_always_fill = any(t in text for t in always_fill_topics)

                if next_text and not is_always_fill:
                    continue  # Next para has content, no placeholder needed

                # Extract sub-section number for matching in example
                num_prefix = _re_subsec.match(r'(\d+\.\d+\.?\d*)', text)
                num_prefix = num_prefix.group(1) if num_prefix else ""
                # Find corresponding content in example
                ex_content = ""
                for j in range(len(doc_ex_phase.paragraphs)):
                    ex_text = doc_ex_phase.paragraphs[j].text.strip()
                    # Match by the sub-section number (lower len threshold
                    # to also match short sub-section labels themselves)
                    if num_prefix and num_prefix in ex_text and len(ex_text) >= 5:
                        # For "always fill" topics (e.g. 调查时间), collect
                        # content from multiple following paragraphs.
                        if is_always_fill:
                            parts = []
                            for k in range(j + 1, min(len(doc_ex_phase.paragraphs), j + 10)):
                                k_text = doc_ex_phase.paragraphs[k].text.strip()
                                if k_text and len(k_text) > 10:
                                    parts.append(k_text[:200])
                                if len(' '.join(parts)) > 500:
                                    break
                            if parts:
                                ex_content = '\n\n'.join(parts)[:800]
                                break
                        # Normal case: use the NEXT paragraph
                        if j + 1 < len(doc_ex_phase.paragraphs):
                            next_ex = doc_ex_phase.paragraphs[j + 1].text.strip()
                            if next_ex and len(next_ex) > 10:
                                ex_content = next_ex[:300]
                                break
                        if not ex_content:
                            ex_content = doc_ex_phase.paragraphs[j].text.strip()[:300]
                            break

                all_placeholders.append({
                    "paragraph_index": next_p_idx,
                    "run_index": 0,
                    "section_index": -1,
                    "section_title": text,  # Use the sub-label as section title
                    "key": f"subsec_{next_p_idx}",
                    "display_name": f"「{text}」段落内容",
                    "original_text": "",
                    "matched_text": "",
                    "pattern_type": "example_diff",
                    "style": "",
                    "example_content": ex_content,
                })

        return all_placeholders

    @classmethod
    def _find_placeholders(cls, text: str) -> List[Dict[str, str]]:
        """Find placeholder patterns in a text string."""
        found = []
        for pattern, pattern_type in PLACEHOLDER_PATTERNS:
            matches = re.finditer(pattern, text)
            for match in matches:
                key = match.group(1) if match.lastindex else match.group(0)
                found.append({
                    "key": key.strip() if key else "",
                    "matched": match.group(0),
                    "type": pattern_type,
                })
        return found

    @classmethod
    def _find_form_fields(cls, docx_path: str) -> List[Dict[str, Any]]:
        """Detect inline form fields: label + colon + spaces, no actual value."""
        import re
        doc = Document(docx_path)
        result = []
        section_idx = -1
        section_title = ""

        for p_idx, para in enumerate(doc.paragraphs):
            style = para.style.name if para.style else ""
            if style.startswith("Heading") or style.startswith("heading"):
                section_idx += 1
                section_title = para.text.strip()

            text = para.text
            if not text:
                continue

            # "稳评责任单位：      " or "事 项  名 称：        "
            match = re.match(r'^(.{2,30}[：:])\s*$', text)
            if match:
                label = match.group(1)
                # Skip labels whose content is handled elsewhere:
                # - "附件清单" → content provided via attachment image items
                if '附件清单' in label:
                    continue
                result.append({
                    "paragraph_index": p_idx,
                    "run_index": 0,
                    "section_index": section_idx,
                    "section_title": section_title,
                    "key": f"form_{p_idx}",
                    "display_name": label,
                    "original_text": text,
                    "matched_text": "",
                    "pattern_type": "form_field",
                    "style": style,
                })
                continue

            # Date fields: "填 表  日 期：2026年4月10日       "
            # Label before colon, date value after — user should update the date
            date_match = re.match(r'^(.{2,20}(日\s*期|日\s+期)[：:])\s*(\S.*?)\s*$', text)
            if date_match and len(text) <= 80:
                label = date_match.group(1)
                existing_date = date_match.group(3)
                result.append({
                    "paragraph_index": p_idx,
                    "run_index": 0,
                    "section_index": section_idx,
                    "section_title": section_title,
                    "key": f"form_{p_idx}",
                    "display_name": label,
                    "original_text": text,
                    "matched_text": "",
                    "pattern_type": "form_field",
                    "style": style,
                    "expected_type": "date",
                    "example_content": existing_date,
                })
                continue

            if len(text) > 50:
                continue

        return result

    # ---- WILDCARD PLACEHOLDER DETECTION ----

    @classmethod
    def _find_wildcard_placeholders(cls, doc: Document) -> List[Dict[str, Any]]:
        """Detect inline wildcard placeholders: "2026年*月", "受***委托", "**街道".

        These are template runs containing one or more "*" characters used as
        placeholders for real values (dates, names, etc.) that the user should provide.
        """
        import re
        result = []
        section_idx = -1
        section_title = ""

        for p_idx, para in enumerate(doc.paragraphs):
            style = para.style.name if para.style else ""
            if style.startswith("Heading") or style.startswith("heading"):
                section_idx += 1
                section_title = para.text.strip()

            # Find all runs with "*" wildcards
            wildcard_runs = []
            for r_idx, run in enumerate(para.runs):
                text = run.text
                if not text:
                    continue
                # Match sequences of "*" characters (1 or more)
                if re.search(r'\*+', text):
                    wildcard_runs.append((r_idx, text))

            if not wildcard_runs:
                continue

            # Group adjacent wildcard runs in the same paragraph
            # into a single composite placeholder
            groups = []
            current_group = []
            last_idx = -2
            for r_idx, text in wildcard_runs:
                if r_idx == last_idx + 1 and current_group:
                    current_group.append((r_idx, text))
                else:
                    if current_group:
                        groups.append(current_group)
                    current_group = [(r_idx, text)]
                last_idx = r_idx
            if current_group:
                groups.append(current_group)

            # If wildcard runs are spread across the paragraph (>2 apart),
            # create individual placeholders. Otherwise merge them.
            for group in groups:
                run_indices = [g[0] for g in group]
                combined_text = "".join(g[1] for g in group)

                # Extract context: the full paragraph text around the wildcards
                para_text = para.text

                # Determine what the wildcard represents
                context_before = ""
                context_after = ""
                first_star = combined_text.find('*')
                if first_star >= 0:
                    # Get text before/after the first wildcard in the paragraph
                    star_pos_in_para = para_text.find('*')
                    if star_pos_in_para >= 0:
                        before = para_text[max(0, star_pos_in_para - 10):star_pos_in_para]
                        before = before.strip()
                        # Find preceding label (e.g. "受", "在")
                        context_before = before[-6:] if len(before) > 6 else before

                # Build a meaningful AND UNIQUE display name from context
                star_count = combined_text.count('*')
                # Show surrounding text to make each wildcard uniquely identifiable
                star_pos = para_text.find(combined_text)
                if star_pos < 0:
                    star_pos = para_text.find('*')
                snippet_start = max(0, star_pos - 8)
                snippet_end = min(len(para_text), star_pos + len(combined_text) + 8)
                snippet = para_text[snippet_start:snippet_end].replace('\n', ' ')

                if '年*月' in snippet or ('年' in snippet and '月' in snippet):
                    display_name = f"🟡 月份：{snippet}"
                    expected_type = "text"
                elif '月*日' in snippet:
                    display_name = f"🟡 日期：{snippet}"
                    expected_type = "text"
                elif star_count >= 3:
                    display_name = f"🟡 委托单位：{snippet}"
                    expected_type = "text"
                elif star_count == 2:
                    display_name = f"🟡 街道名称：{snippet}"
                    expected_type = "text"
                else:
                    display_name = f"🟡 {snippet}"
                    expected_type = "text"

                result.append({
                    "paragraph_index": p_idx,
                    "run_index": run_indices[0],
                    "run_indices": run_indices,
                    "section_index": section_idx,
                    "section_title": section_title,
                    "key": f"wildcard_{p_idx}_{run_indices[0]}",
                    # Use snippet + key to make each wildcard uniquely identifiable
                    # (prevents highlight dedup from merging e.g. 5 "*" in same para)
                    "display_name": f"{display_name} [{combined_text} #{run_indices[0]}]",
                    "original_text": combined_text,
                    "matched_text": combined_text,
                    "pattern_type": "highlight",
                    "expected_type": expected_type,
                    "style": style,
                    "is_required": True,
                })

        return result

    # ---- HIGHLIGHTED RUN DETECTION ----

    @classmethod
    def _find_highlighted_runs(cls, docx_path: str) -> List[Dict[str, Any]]:
        """
        Find runs with highlight formatting (yellow etc.) as prefilled placeholders.
        Adjacent highlighted runs in the same paragraph are merged into a single
        composite placeholder (e.g. "金湖县委政法委、" + "戴楼街道办事处" become one).
        """
        from docx.oxml.ns import qn
        doc = Document(docx_path)
        result = []

        # Track current section from headings
        current_section_index = -1
        current_section_title = ""

        for p_idx, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name if paragraph.style else ""

            # Track heading changes for section context
            if style_name.startswith("Heading") or style_name.startswith("heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                current_section_index += 1
                current_section_title = paragraph.text.strip()

            # Collect highlighted runs in this paragraph
            highlighted_runs = []
            for r_idx, run in enumerate(paragraph.runs):
                rpr = run._r.find(qn('w:rPr'))
                if rpr is None:
                    continue
                highlight = rpr.find(qn('w:highlight'))
                if highlight is None:
                    continue
                val = highlight.get(qn('w:val'))
                if val and val not in ('none',):
                    highlighted_runs.append((r_idx, run.text, val))

            if not highlighted_runs:
                continue

            # Merge adjacent runs with the same highlight color
            merged = []
            current = None
            for r_idx, text, color in highlighted_runs:
                if not text.strip():
                    continue  # skip empty highlighted runs
                if current and current['run_end'] == r_idx - 1 and current['color'] == color:
                    # Extend current group
                    current['run_end'] = r_idx
                    current['text'] += text
                    current['run_indices'].append(r_idx)
                else:
                    if current:
                        merged.append(current)
                    current = {
                        'run_start': r_idx,
                        'run_end': r_idx,
                        'run_indices': [r_idx],
                        'text': text,
                        'color': color,
                    }
            if current:
                merged.append(current)

            # Create placeholder entries
            for m in merged:
                result.append({
                    "paragraph_index": p_idx,
                    "run_index": m['run_start'],
                    "run_indices": m['run_indices'],
                    "section_index": current_section_index,
                    "section_title": current_section_title,
                    "key": f"highlight_{p_idx}_{m['run_start']}",
                    "display_name": m['text'][:40],
                    "original_text": m['text'],
                    "matched_text": m['text'],
                    "pattern_type": "highlight",
                    "highlight_color": m['color'],
                    "style": style_name,
                })

        return result

    @classmethod
    def diff_with_example(
        cls, template_path: str, example_path: str
    ) -> List[Dict[str, Any]]:
        """
        Compare template with a complete example report using heading-based
        section alignment. This handles paragraph index shifts caused by
        missing images/content in the template.

        For each section (between two headings), body paragraphs are compared
        positionally. Extra paragraphs in the example (images, captions) are
        detected as image_placeholder.
        """
        import re
        doc_tpl = Document(template_path)
        doc_ex = Document(example_path)

        result = []

        # ---- Step 1: Build heading index maps for both docs ----
        def _get_headings(doc):
            """Return list of (para_index, text, style, level)."""
            h = []
            for i, p in enumerate(doc.paragraphs):
                s = p.style.name if p.style else ""
                if s.startswith("Heading") or s.startswith("heading"):
                    try:
                        lv = int(s.split()[-1])
                    except (ValueError, IndexError):
                        lv = 1
                    h.append((i, p.text.strip(), s, lv))
            return h

        tpl_headings = _get_headings(doc_tpl)
        ex_headings = _get_headings(doc_ex)

        # Build mapping: heading_text → (tpl_para_index, ex_para_index)
        tpl_hmap = {h[1]: h[0] for h in tpl_headings if h[1]}
        ex_hmap = {h[1]: h[0] for h in ex_headings if h[1]}

        # ---- Step 2: Align sections by heading text ----
        section_idx = -1
        # Use template heading order, find matching example heading
        for h_idx, (t_pi, t_text, t_style, t_level) in enumerate(tpl_headings):
            if not t_text:
                continue
            section_idx += 1
            section_title = t_text

            # Find matching heading in example
            ex_pi = ex_hmap.get(t_text)
            if ex_pi is None:
                continue

            # Find next heading indices (section boundaries)
            t_next_pi = tpl_headings[h_idx + 1][0] if h_idx + 1 < len(tpl_headings) else len(doc_tpl.paragraphs)

            # Find matching example heading index
            ex_h_idx = None
            for j, (ep, et, es, el) in enumerate(ex_headings):
                if et == t_text:
                    ex_h_idx = j
                    break
            if ex_h_idx is None:
                continue
            ex_next_pi = ex_headings[ex_h_idx + 1][0] if ex_h_idx + 1 < len(ex_headings) else len(doc_ex.paragraphs)

            # ---- Step 3: Compare body paragraphs within this section ----
            tpl_body = list(range(t_pi + 1, t_next_pi))
            ex_body = list(range(ex_pi + 1, ex_next_pi))

            max_body = max(len(tpl_body), len(ex_body))
            for pos in range(max_body):
                t_para_idx = tpl_body[pos] if pos < len(tpl_body) else None
                e_para_idx = ex_body[pos] if pos < len(ex_body) else None

                t_text = doc_tpl.paragraphs[t_para_idx].text.strip() if t_para_idx is not None else ""
                e_text = doc_ex.paragraphs[e_para_idx].text.strip() if e_para_idx is not None else ""

                # Check if example paragraph has an image (drawing)
                e_has_image = False
                if e_para_idx is not None:
                    e_has_image = bool(
                        doc_ex.paragraphs[e_para_idx]._p.findall(
                            './/' + qn('w:drawing')
                        )
                    )

                # Case 1: Template paragraph exists and is empty, example has
                #          meaningful content → placeholder
                if t_para_idx is not None and not t_text and (e_text and len(e_text) > 3 or e_has_image):
                    # Figure captions → image placeholders
                    # (also catch empty paragraphs with images — e_has_image without text)
                    is_figure = (
                        (('图' in e_text or '照片' in e_text or '示意图' in e_text
                          or '平面图' in e_text or '效果图' in e_text or '红线图' in e_text)
                         and len(e_text) < 60)
                        or e_text.strip().startswith('图')
                    )
                    if is_figure or e_has_image:
                        result.append({
                            "paragraph_index": t_para_idx,
                            "run_index": 0,
                            "section_index": section_idx,
                            "section_title": section_title,
                            "key": f"image_{t_para_idx}",
                            "display_name": f"🖼️ 「{section_title}」图片占位",
                            "original_text": "",
                            "matched_text": "",
                            "pattern_type": "image_placeholder",
                            "expected_type": "image",
                            "style": "",
                            "example_content": e_text if e_text else "",
                        })
                        continue

                    # Skip TOC entries (e.g. "1.1决策名称	1")
                    if re.match(r'^\d+\.\d+.*\t\d+$', e_text):
                        continue
                    # Skip short labels / numbered sub-headings
                    if len(e_text) < 20:
                        continue
                    if re.match(r'^\d+\.\d+.*', e_text) and len(e_text) < 40:
                        continue
                    if re.match(r'^[\(\（]?\d+[\)\）]', e_text) and len(e_text) < 25:
                        continue

                    result.append({
                        "paragraph_index": t_para_idx,
                        "run_index": 0,
                        "section_index": section_idx,
                        "section_title": section_title,
                        "key": f"example_diff_{t_para_idx}",
                        "display_name": f"「{section_title}」段落内容",
                        "original_text": "",
                        "matched_text": "",
                        "pattern_type": "example_diff",
                        "style": "",
                        "example_content": e_text,
                    })
                    continue

                # Case 1b: Template has text, example has text + IMAGE embedded
                #          → create image placeholder alongside the text diff.
                #          (Moved OUTSIDE Case 1 — was dead code inside it since
                #           Case 1 requires not t_text but Case 1b needs t_text.)
                if t_para_idx is not None and t_text and e_has_image and len(t_text) > 3:
                    result.append({
                        "paragraph_index": t_para_idx,
                        "run_index": 0,
                        "section_index": section_idx,
                        "section_title": section_title,
                        "key": f"image_{t_para_idx}_inline",
                        "display_name": f"🖼️ 「{section_title}」图片占位",
                        "original_text": t_text,
                        "matched_text": "",
                        "pattern_type": "image_placeholder",
                        "expected_type": "image",
                        "style": "",
                        "example_content": "",
                    })
                    # Also check if example text is substantially different
                    # (longer/more detailed) → create text diff too
                    if e_text and len(e_text) > len(t_text) + 10:
                        # Example has richer content — ask user to provide it
                        result.append({
                            "paragraph_index": t_para_idx,
                            "run_index": 0,
                            "section_index": section_idx,
                            "section_title": section_title,
                            "key": f"example_diff_{t_para_idx}_inline",
                            "display_name": f"「{section_title}」段落内容",
                            "original_text": t_text,
                            "matched_text": "",
                            "pattern_type": "example_diff",
                            "style": "",
                            "example_content": e_text,
                        })

                # Case 2: Template has NO paragraph here but example does
                #          → extra image or content in example
                if t_para_idx is None and e_para_idx is not None and e_text:
                    if e_has_image or (
                        ('图' in e_text or '照片' in e_text) and len(e_text) < 80
                    ):
                        # Extra image in example — assign to nearest template para
                        anchor_pi = tpl_body[-1] if tpl_body else t_pi
                        result.append({
                            "paragraph_index": anchor_pi,
                            "run_index": 0,
                            "section_index": section_idx,
                            "section_title": section_title,
                            "key": f"image_{e_para_idx}_extra",
                            "display_name": f"🖼️ 「{section_title}」图片占位",
                            "original_text": "",
                            "matched_text": "",
                            "pattern_type": "image_placeholder",
                            "expected_type": "image",
                            "style": "",
                            "example_content": e_text,
                        })
                    elif len(e_text) > 20:
                        # Extra body text
                        anchor_pi = tpl_body[-1] if tpl_body else t_pi
                        result.append({
                            "paragraph_index": anchor_pi,
                            "run_index": 0,
                            "section_index": section_idx,
                            "section_title": section_title,
                            "key": f"example_diff_{e_para_idx}_extra",
                            "display_name": f"「{section_title}」段落内容",
                            "original_text": "",
                            "matched_text": "",
                            "pattern_type": "example_diff",
                            "style": "",
                            "example_content": e_text,
                        })

        return result

    @classmethod
    def _find_figure_image_placeholders(
        cls, template_path: str, example_path: str
    ) -> List[Dict[str, Any]]:
        """
        Find paragraphs in the template that are figure captions or
        attachment items but lack the actual embedded image.
        These need the user to upload the corresponding image file.

        Detects:
        - Figure captions: "图3-1决策评估公示内容", "图3-2公示照片", etc.
        - Attachment list: "1.拟征收土地公告、征地范围图", etc.

        Uses TEXT-BASED matching (not paragraph index) to handle index shifts
        between template and example documents.
        """
        import re
        doc_tpl = Document(template_path)
        doc_ex = Document(example_path)

        result = []
        seen_keys = set()

        # Build flat lists for text-based searching
        tpl_texts = [p.text.strip() for p in doc_tpl.paragraphs]
        ex_texts = [p.text.strip() for p in doc_ex.paragraphs]

        # ---- Find "附件清单：" anchor in template ----
        # Attachment items are only detected after the LAST "附件清单：".
        # (The TOC has one too, but the actual attachments section is at the end.)
        attachment_start_tpl = -1
        attachment_start_ex = -1
        for i, t in enumerate(tpl_texts):
            if t == "附件清单：" or t.startswith("附件清单"):
                attachment_start_tpl = i  # Keep last occurrence
        for i, t in enumerate(ex_texts):
            if t == "附件清单：" or t.startswith("附件清单"):
                attachment_start_ex = i  # Keep last occurrence

        # ---- Helper: check if example has images near a text-matched paragraph ----
        def _ex_has_images_near(ex_para_idx: int, look_ahead: int = 15) -> bool:
            """Check if there are drawings within range of ex_para_idx (both directions)."""
            start = max(0, ex_para_idx - 5)
            end = min(len(doc_ex.paragraphs), ex_para_idx + look_ahead)
            for j in range(start, end):
                if doc_ex.paragraphs[j]._p.findall('.//' + qn('w:drawing')):
                    return True
            return False

        # ---- Helper: find a paragraph in example by text matching ----
        def _find_in_example(text: str, start_from: int = 0) -> int:
            """Find paragraph index in example whose text contains `text`."""
            for j in range(start_from, len(doc_ex.paragraphs)):
                ex_text = doc_ex.paragraphs[j].text.strip()
                if not ex_text:
                    continue  # Skip empty paragraphs ('' in any_str == True)
                if text in ex_text or ex_text in text:
                    return j
            return -1

        for i, para in enumerate(doc_tpl.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Check if this paragraph is a figure caption
            # e.g. "图3-1决策评估公示内容", "图3-2公示照片"
            # Must start with "图" followed by number-dash-number pattern
            is_figure_caption = bool(
                re.match(r'图\d+[-–—]\d+', text)
            )

            # Check if this is an attachment list item
            # Only detected AFTER the "附件清单：" anchor paragraph,
            # and must look like a numbered list item with attachment keywords.
            is_attachment = (
                i > attachment_start_tpl >= 0
                and bool(re.match(r'^\d+\.\s*\S', text))
                and len(text) < 60
                and any(kw in text for kw in [
                    '照片', '问卷', '纪要', '公告', '范围图',
                    '公示', '现场', '座谈', '附件'
                ])
            )

            if not (is_figure_caption or is_attachment):
                continue

            # ---- Text-based matching in example ----
            # Search the ENTIRE example for matching text (not just ±3 index range)
            ex_match_idx = _find_in_example(text)
            if ex_match_idx < 0:
                # For figure captions, try searching with just the number prefix
                if is_figure_caption:
                    num_prefix = re.match(r'(图\d+[-–—]\d+)', text)
                    if num_prefix:
                        ex_match_idx = _find_in_example(num_prefix.group(1))
                if ex_match_idx < 0:
                    continue

            # Check if example has images after the matched paragraph
            ex_has_image = _ex_has_images_near(ex_match_idx)

            if not ex_has_image:
                continue

            if is_attachment:
                key = f"attachment_img_{i}"
            else:
                key = f"figure_img_{i}"

            if key in seen_keys:
                continue
            seen_keys.add(key)

            # Determine section context
            section_title = "附件" if is_attachment else "图片"
            # Find nearest heading
            for j in range(i, -1, -1):
                s = doc_tpl.paragraphs[j].style.name if doc_tpl.paragraphs[j].style else ""
                if s.startswith("Heading"):
                    section_title = doc_tpl.paragraphs[j].text.strip()
                    break

            display = text[:50] if text else f"图片_{i}"
            result.append({
                "paragraph_index": i,
                "run_index": 0,
                "section_index": -1,
                "section_title": section_title,
                "key": key,
                "display_name": f"🖼️ {display}",
                "original_text": text,
                "matched_text": text,
                "pattern_type": "image_placeholder",
                "expected_type": "image",
                "style": para.style.name if para.style else "",
                "example_content": "",
                "is_attachment_item": is_attachment,
            })

        return result

    @classmethod
    def _find_table_caption(cls, doc: Document, table_index: int) -> str:
        """
        Find the caption/title paragraph preceding a table.
        Looks for patterns like "表3-1..." or "表X..." before the table element.
        """
        import re as _re_tc
        # Find the table's XML element position in the document body
        body = doc.element.body
        tables_in_body = body.findall(qn('w:tbl'))
        if table_index >= len(tables_in_body):
            return ""

        tbl_element = tables_in_body[table_index]

        # Walk backwards through body children to find the preceding paragraph
        # that contains a table caption
        children = list(body)
        tbl_pos = None
        for i, child in enumerate(children):
            if child is tbl_element:
                tbl_pos = i
                break

        if tbl_pos is None:
            return ""

        # Look backwards up to 5 elements for a table caption paragraph
        for i in range(tbl_pos - 1, max(0, tbl_pos - 6), -1):
            child = children[i]
            if child.tag == qn('w:p'):
                # Extract paragraph text
                texts = child.findall('.//' + qn('w:t'))
                para_text = ''.join(t.text or '' for t in texts).strip()
                # Match "表X" or "表X-X" patterns (table captions)
                if _re_tc.match(r'表\d+', para_text):
                    return para_text[:80]
                # Also accept any non-empty text within 3 elements
                if i >= tbl_pos - 3 and para_text and len(para_text) > 3:
                    return para_text[:80]

        return ""

    @classmethod
    def _find_table_placeholders(
        cls, docx_path: str, example_path: str = None
    ) -> List[Dict[str, Any]]:
        """
        Scan tables for cells that need filling:
        - Cells containing only whitespace after a label
        - Empty cells in data rows (with optional example comparison)
        - Cells with old prefilled data (yellow-highlighted in tables)
        - Outline cells: numbered sub-headings with empty content paragraphs
          (e.g. "1.拟征地位置\\n \\n2.征收范围..." in decision overview table)
        """
        from docx.oxml.ns import qn
        import re as _re_table
        doc = Document(docx_path)
        result = []
        label_redirected_cells = set()  # (table, row, cell) already flagged

        # Load example doc for table cell comparison
        doc_ex = None
        if example_path and Path(example_path).exists():
            try:
                doc_ex = Document(example_path)
            except Exception:
                pass

        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()

                    # ---- OUTLINE CELL DETECTION ----
                    # Detect cells with numbered sub-headings but empty content paragraphs.
                    # Pattern: alternating heading paras ("1.xxx", "2.xxx") and
                    # content paras (whitespace). Each empty content para is a placeholder.
                    if 4 <= len(cell.paragraphs) <= 30:
                        for p_idx, para in enumerate(cell.paragraphs):
                            para_text = para.text.strip()
                            # Check if this is a numbered heading paragraph
                            if (_re_table.match(r'\d+\.\s*\S', para_text)
                                    and len(para_text) < 60
                                    and p_idx + 1 < len(cell.paragraphs)):
                                heading = para_text
                                content_para = cell.paragraphs[p_idx + 1]
                                content_text = content_para.text.strip()
                                # Content paragraph is empty/whitespace → needs filling
                                if not content_text or content_text.isspace():
                                    # Verify example also has content here
                                    # Skip if example is also empty (just formatting)
                                    ex_has_content = False
                                    if doc_ex is not None:
                                        try:
                                            if (t_idx < len(doc_ex.tables)
                                                    and r_idx < len(doc_ex.tables[t_idx].rows)
                                                    and c_idx < len(doc_ex.tables[t_idx].rows[r_idx].cells)):
                                                ex_cell = doc_ex.tables[t_idx].rows[r_idx].cells[c_idx]
                                                if p_idx + 1 < len(ex_cell.paragraphs):
                                                    ex_content = ex_cell.paragraphs[p_idx + 1].text.strip()
                                                    if ex_content and not ex_content.isspace():
                                                        ex_has_content = True
                                        except Exception:
                                            pass
                                    if not ex_has_content:
                                        continue  # Both template and example empty → skip
                                    key = (
                                        f"table_{t_idx}_{r_idx}_{c_idx}"
                                        f"_outline_{p_idx + 1}"
                                    )
                                    result.append({
                                        "table_index": t_idx,
                                        "row_index": r_idx,
                                        "cell_index": c_idx,
                                        "cell_para_index": p_idx + 1,
                                        "paragraph_index": -1,
                                        "run_index": 0,
                                        "section_index": -1,
                                        "section_title": f"表格{t_idx+1}",
                                        "key": key,
                                        "display_name": heading,
                                        "original_text": heading,
                                        "matched_text": heading,
                                        "pattern_type": "table_cell",
                                        "highlight_color": "",
                                        "style": "TableCell",
                                    })

                    # Skip fully empty cells (handled below via example comparison)
                    if not text:
                        # ---- EMPTY CELL WITH EXAMPLE CONTENT ----
                        # Skip if this cell is already flagged via label redirection
                        cell_key = (t_idx, r_idx, c_idx)
                        if cell_key in label_redirected_cells:
                            continue
                        # If template cell is empty but example has content → placeholder
                        if doc_ex is not None:
                            try:
                                if (t_idx < len(doc_ex.tables)
                                        and r_idx < len(doc_ex.tables[t_idx].rows)
                                        and c_idx < len(
                                            doc_ex.tables[t_idx].rows[r_idx].cells)):
                                    ex_text = (
                                        doc_ex.tables[t_idx]
                                        .rows[r_idx].cells[c_idx].text.strip()
                                    )
                                    if ex_text and len(ex_text) >= 2:
                                        # Find row context from cols 1/0 (col 1 is more specific)
                                        row_label = ""
                                        for col in [1, 0]:
                                            if col < len(row.cells) and col != c_idx:
                                                ltext = (
                                                    row.cells[col].text.strip()
                                                    .replace("\n", "")
                                                )
                                                if ltext and len(ltext) >= 2:
                                                    row_label = ltext
                                                    break
                                        # For disambiguation: count rows with same
                                        # row_label. If >1, append 调查对象.
                                        row_object = ""
                                        same_label_count = 0
                                        for rr in table.rows:
                                            for col in [1, 0]:
                                                if col < len(rr.cells):
                                                    rt = (
                                                        rr.cells[col].text.strip()
                                                        .replace("\n", "")
                                                    )
                                                    if rt == row_label:
                                                        same_label_count += 1
                                                        break
                                        if same_label_count > 1 and 2 < len(row.cells):
                                            obj_text = (
                                                row.cells[2].text.strip()
                                                .replace("\n", "")
                                            )
                                            if obj_text and obj_text != row_label:
                                                row_object = obj_text
                                        # Get column header from first row for context
                                        col_header = ""
                                        if (len(table.rows) > 0
                                                and c_idx < len(table.rows[0].cells)):
                                            col_header = (
                                                table.rows[0].cells[c_idx].text
                                                .strip().replace("\n", "")
                                            )
                                        # Build descriptive display name
                                        if row_label and col_header:
                                            display = f"{col_header} - {row_label}"
                                            if row_object:
                                                display += f"（{row_object}）"
                                        elif row_label:
                                            display = row_label
                                        elif col_header:
                                            display = f"{col_header}"
                                        else:
                                            display = (
                                                f"表格{t_idx+1}行{r_idx+1}列{c_idx+1}"
                                            )
                                        # Find table caption/title from surrounding paragraphs
                                        table_title = cls._find_table_caption(
                                            doc, t_idx
                                        )
                                        section_title = table_title or f"表格{t_idx+1}"
                                        result.append({
                                            "table_index": t_idx,
                                            "row_index": r_idx,
                                            "cell_index": c_idx,
                                            "paragraph_index": -1,
                                            "run_index": 0,
                                            "section_index": -1,
                                            "section_title": section_title,
                                            "key": f"table_{t_idx}_{r_idx}_{c_idx}",
                                            "display_name": display,
                                            "original_text": "",
                                            "matched_text": "",
                                            "pattern_type": "table_cell",
                                            "highlight_color": "",
                                            "style": "TableCell",
                                            "example_content": ex_text,
                                        })
                            except Exception:
                                pass
                        continue

                    # Check if cell has highlighted runs (prefilled data)
                    has_highlight = False
                    for para in cell.paragraphs:
                        for run in para.runs:
                            rpr = run._r.find(qn('w:rPr'))
                            if rpr is not None:
                                hl = rpr.find(qn('w:highlight'))
                                if hl is not None and hl.get(qn('w:val')) not in (None, 'none'):
                                    has_highlight = True
                                    break

                    # Check if cell text is just a label without a value
                    # Pattern: "单位名称" alone (no actual value filled)
                    is_label_only = (
                        len(text) < 20
                        and any(kw in text for kw in ['单位名称', '负责人', '联系人', '职务', '联络方式'])
                        and not any(c.isdigit() for c in text)  # no numbers = likely unfilled
                    )

                    # Check if cell is a signature/date boilerplate placeholder
                    # "负责人（签章）：\n年   月    日", "备案单位（签章）：..."
                    is_signature = any(
                        kw in text for kw in [
                            '签章', '年   月    日', '年    月    日',
                            '备案单位', '组长：\n'
                        ]
                    )

                    # Check if cell contains old project data (needs region replacement)
                    has_old_data = any(
                        kw in text for kw in ['金征预告', '金湖', '戴楼']
                    )

                    if is_label_only and not is_signature:
                        # Label cell → placeholder is the VALUE cell to its right
                        val_c_idx = c_idx + 1
                        if val_c_idx < len(row.cells):
                            val_text = row.cells[val_c_idx].text.strip()
                            # Only flag if value cell is empty (needs user input)
                            if not val_text:
                                key = f"table_{t_idx}_{r_idx}_{val_c_idx}"
                                label_redirected_cells.add((t_idx, r_idx, val_c_idx))
                                result.append({
                                    "table_index": t_idx,
                                    "row_index": r_idx,
                                    "cell_index": val_c_idx,
                                    "paragraph_index": -1,
                                    "run_index": 0,
                                    "section_index": -1,
                                    "section_title": f"表格{t_idx+1}",
                                    "key": key,
                                    "display_name": f"{text}",  # label as display name
                                    "original_text": "",
                                    "matched_text": "",
                                    "pattern_type": "table_cell",
                                    "highlight_color": "",
                                    "style": "TableCell",
                                })
                    elif has_highlight or has_old_data or is_signature:
                        key = f"table_{t_idx}_{r_idx}_{c_idx}"
                        # Build contextual display name for signature/date cells
                        if is_signature:
                            # Find row context from cols 0/1
                            row_label = ""
                            for col in [1, 0]:
                                if col < len(row.cells):
                                    ltext = row.cells[col].text.strip().replace("\n", "")
                                    if ltext and len(ltext) >= 2:
                                        row_label = ltext
                                        break
                            display = f"{row_label} - 结论内容" if row_label else f"表格{t_idx+1}行{r_idx+1}列{c_idx+1}"
                        else:
                            display = f"表格{t_idx+1} 行{r_idx+1}列{c_idx+1}"

                        result.append({
                            "table_index": t_idx,
                            "row_index": r_idx,
                            "cell_index": c_idx,
                            "paragraph_index": -1,
                            "run_index": 0,
                            "section_index": -1,
                            "section_title": f"表格{t_idx+1}",
                            "key": key,
                            "display_name": display,
                            "original_text": "" if is_signature else text[:200],
                            "matched_text": "",
                            "pattern_type": "table_cell",
                            "highlight_color": "yellow" if has_highlight else "",
                            "style": "TableCell",
                        })

        # Deduplicate merged cells (both vertical and horizontal).
        # Only dedup cells with actual content — empty cells are distinct placeholders.
        last_in_column = {}   # (table, cell_index) -> (row_index, text)
        seen_in_row = {}      # (table, row_index) -> set of seen texts
        seen_outline = set()  # (table, row, cell_para_index) for outline sub-items
        seen_cells = set()    # (table, row, cell_index) to dedup empty-cell placeholders
        deduped = []
        for p in result:
            cpi = p.get("cell_para_index")
            orig = p["original_text"].strip()
            cell_key = (p["table_index"], p["row_index"], p["cell_index"])

            if cpi is not None:
                # Outline sub-item: dedup by (table, row, cell_para_index)
                # since merged columns duplicate the same content paragraphs
                outline_key = (p["table_index"], p["row_index"], cpi)
                if outline_key in seen_outline:
                    continue
                seen_outline.add(outline_key)
            elif orig:
                # Vertical merge check
                col_key = (p["table_index"], p["cell_index"])
                prev_v = last_in_column.get(col_key)
                is_v_merged = (
                    prev_v is not None
                    and prev_v[0] == p["row_index"] - 1
                    and prev_v[1] == orig
                )
                last_in_column[col_key] = (p["row_index"], orig)

                # Horizontal merge check
                row_key = (p["table_index"], p["row_index"])
                is_h_merged = orig in seen_in_row.get(row_key, set())

                if is_v_merged or is_h_merged:
                    continue
                seen_in_row.setdefault(row_key, set()).add(orig)
            else:
                # Cells without original text: dedup by (table, row, pattern)
                # - Example-comparison cells: one per row
                # - Signature/empty cells without example: one per row
                row_key = (p["table_index"], p["row_index"])
                if p.get("example_content"):
                    if row_key in seen_cells:
                        continue
                    seen_cells.add(row_key)
                else:
                    if p.get("example_content"):
                        # Example-comparison cells: dedup by row
                        if row_key in seen_cells:
                            continue
                        seen_cells.add(row_key)
                    elif cell_key in label_redirected_cells:
                        # Label-redirected cells: each label is distinct
                        # Don't dedup — "职务" and "负责人" are different fields
                        pass
                    else:
                        # Signature/other empty cells: dedup by row
                        if row_key in seen_cells:
                            continue
                        seen_cells.add(row_key)

            deduped.append(p)
        return deduped

    # ---- WRITING / FILLING ----

    @classmethod
    def fill_template(
        cls,
        template_path: str,
        output_path: str,
        filled_data: Dict[str, str],
        placeholder_map: Optional[List[Dict[str, Any]]] = None,
    ) -> str:
        """
        Fill a template docx with provided data.
        - template_path: path to the original template
        - output_path: where to save the filled document
        - filled_data: {placeholder_key: value_to_fill}
        - placeholder_map: optional pre-computed placeholder locations

        CRITICAL RULES:
        1. NEVER modify any formatting (font, size, color, alignment, etc.)
        2. NEVER add or remove paragraphs, runs, or any content
        3. ONLY replace the placeholder text pattern with the user-provided value
        4. If a value is empty, fill with "需后期提供"
        """
        import shutil

        # Copy template — never modify the original
        shutil.copy2(template_path, output_path)
        doc = Document(output_path)

        # If no placeholder map provided, scan the document
        if placeholder_map is None:
            placeholder_map = cls.find_all_placeholders(output_path)

        # Build lookup: key -> location
        key_to_locations = {}
        for p in placeholder_map:
            key = p["key"]
            if key not in key_to_locations:
                key_to_locations[key] = []
            key_to_locations[key].append(p)

        # Pre-scan: group placeholders by paragraph_index.
        # When multiple example_diff/blank placeholders target the same paragraph,
        # concatenate their values with newlines instead of overwriting.
        # When a paragraph has both text AND image, process them together.
        _merged_keys = set()  # Keys already handled via merge
        _filled_grid_cells = set()  # (table_idx, row_idx, col_idx) already filled (handles gridSpan)
        _para_text_image = {}  # para_index -> {"text_keys": [...], "text_loc": ..., "image_key": ..., "image_loc": ...}
        for p in placeholder_map:
            pi = p.get("paragraph_index", -1)
            if pi < 0:
                continue
            pt = p.get("pattern_type", "")
            if pt in ("example_diff", "blank"):
                entry = _para_text_image.setdefault(pi, {})
                entry.setdefault("text_keys", []).append(p["key"])
                if "text_loc" not in entry:
                    entry["text_loc"] = p
            elif pt == "image_placeholder":
                _para_text_image.setdefault(pi, {})["image_key"] = p["key"]
                _para_text_image[pi]["image_loc"] = p

        # Fill each placeholder
        for key, value in filled_data.items():
            # Skip placeholders that should keep original template content
            # (conclusion tables, signature blocks, etc.)
            if value in ("__KEEP_ORIGINAL__", "__WITH_TEXT__"):
                continue

            # Actively remove image from the paragraph
            if value == "__REMOVE_IMAGE__":
                locations = key_to_locations.get(key)
                if locations:
                    for loc in locations:
                        para = doc.paragraphs[loc["paragraph_index"]]
                        # Collect all drawing elements first, then remove
                        to_remove = []
                        for elem in para._p.iter():
                            if elem.tag.endswith('}drawing') or elem.tag == 'drawing':
                                to_remove.append(elem)
                        for drawing in to_remove:
                            parent = drawing.getparent()
                            if parent is not None:
                                parent.remove(drawing)
                        # Clear any empty runs
                        for run in para.runs:
                            run.text = ""
                continue
            # Skip already-merged keys
            if key in _merged_keys:
                continue
            # For table cells and merged cells, empty value should not overwrite
            # (adjacent cells may already have content due to cell merging)
            fill_value = value if value else (
                "" if any(loc.get("pattern_type") == "table_cell"
                         for loc in (key_to_locations.get(key) or []))
                else "需后期提供"
            )

            locations = key_to_locations.get(key)
            if not locations:
                # Fuzzy match: try partial key matching
                for stored_key, locs in key_to_locations.items():
                    if key in stored_key or stored_key in key:
                        locations = locs
                        break

            if not locations:
                continue

            # ── Merged text+image at same paragraph ──
            # If this is a text placeholder whose paragraph also has other text
            # or image placeholders, handle them together to avoid overwrites.
            pi = locations[0].get("paragraph_index", -1) if locations else -1
            _merge_info = _para_text_image.get(pi, {}) if pi >= 0 else {}
            _text_keys = _merge_info.get("text_keys", [])
            _has_image = bool(_merge_info.get("image_key")
                              and _merge_info["image_key"] in filled_data
                              and filled_data[_merge_info["image_key"]] not in ("__KEEP_ORIGINAL__", "__WITH_TEXT__"))

            if pi >= 0 and len(_text_keys) > 1 and key == _text_keys[0]:
                # Multiple text placeholders at same para → concatenate
                _merged_keys.update(_text_keys)
                texts = []
                for tk in _text_keys:
                    tv = filled_data.get(tk, "")
                    if tv and tv not in ("__KEEP_ORIGINAL__", "__WITH_TEXT__"):
                        texts.append(tv)
                combined_text = "\n\n".join(texts)

                if _has_image:
                    _merged_keys.add(_merge_info["image_key"])
                    cls._fill_text_and_image(
                        doc, locations[0], combined_text,
                        _merge_info.get("image_loc"), filled_data[_merge_info["image_key"]],
                    )
                else:
                    # Fill with combined text
                    if len(doc.paragraphs[pi].runs) == 0:
                        from docx.oxml import OxmlElement
                        pPr = doc.paragraphs[pi]._p.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            doc.paragraphs[pi]._p.insert(0, pPr)
                        pStyle = pPr.find(qn('w:pStyle'))
                        if pStyle is None:
                            pStyle = OxmlElement('w:pStyle')
                            pPr.append(pStyle)
                            pStyle.set(qn('w:val'), 'Normal')
                        r_element = OxmlElement('w:r')
                        rPr = OxmlElement('w:rPr')
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        rPr.append(rFonts)
                        r_element.append(rPr)
                        t_element = OxmlElement('w:t')
                        t_element.text = combined_text
                        t_element.set(qn('xml:space'), 'preserve')
                        r_element.append(t_element)
                        doc.paragraphs[pi]._p.append(r_element)
                    else:
                        cls._replace_in_run(
                            doc.paragraphs[pi], locations[0].get("run_index", 0),
                            locations[0].get("matched_text", ""), combined_text,
                            remove_highlight=False,
                        )
                continue

            if (pi >= 0 and len(_text_keys) == 1 and _text_keys[0] == key
                    and _has_image):
                _merged_keys.add(key)
                _merged_keys.add(_merge_info["image_key"])
                img_key = _merge_info["image_key"]
                img_value = filled_data[img_key]
                # Process text+image together at this paragraph
                cls._fill_text_and_image(
                    doc, locations[0], fill_value,
                    _merge_info.get("image_loc"), img_value,
                )
                continue

            for loc in locations:
                # Handle table cell placeholders
                if loc.get("pattern_type") == "table_cell" and "table_index" in loc:
                    t_idx = loc["table_index"]
                    r_idx = loc["row_index"]
                    c_idx = loc["cell_index"]
                    cell_para_index = loc.get("cell_para_index")

                    if t_idx < len(doc.tables):
                        table = doc.tables[t_idx]
                        if r_idx < len(table.rows) and c_idx < len(table.rows[r_idx].cells):
                            row = table.rows[r_idx]
                            cell = row.cells[c_idx]

                            # Skip if this grid cell was already filled (handles gridSpan merged cells)
                            grid_cell_key = (t_idx, r_idx, c_idx)
                            if grid_cell_key in _filled_grid_cells:
                                continue

                            # Compute gridSpan to track merged cells
                            _tc_pr = cell._tc.find(qn('w:tcPr'))
                            _grid_span = 1
                            if _tc_pr is not None:
                                _gs = _tc_pr.find(qn('w:gridSpan'))
                                if _gs is not None:
                                    try:
                                        _grid_span = int(_gs.get(qn('w:val'), '1'))
                                    except (ValueError, TypeError):
                                        _grid_span = 1

                            # ---- OUTLINE SUB-ITEM FILLING ----
                            # Fill a specific paragraph within the cell (e.g. content
                            # after "1.拟征地位置" in the decision overview table).
                            if cell_para_index is not None:
                                if cell_para_index < len(cell.paragraphs):
                                    content_para = cell.paragraphs[cell_para_index]
                                    # Clear existing runs, keep first one for filling
                                    for run in content_para.runs:
                                        run.text = ""
                                    if content_para.runs:
                                        content_para.runs[0].text = fill_value
                                    elif content_para._p is not None:
                                        from docx.oxml import OxmlElement
                                        r = OxmlElement('w:r')
                                        t = OxmlElement('w:t')
                                        t.text = fill_value
                                        t.set(qn('xml:space'), 'preserve')
                                        r.append(t)
                                        content_para._p.append(r)
                                continue

                            # ---- WHOLE-CELL FILLING (existing logic) ----
                            original_cell_text = loc.get("original_text", "").strip()

                            # Check if this cell or its row has preserved content
                            PRESERVED_CELL_VALUES = {
                                '单位名称': '江苏众拓项目代理咨询有限公司',
                                '负责人': '陈春',
                                '职务': '董事长',
                                '联系人': '程诗茹',
                                '联络方式': '18252573739',
                            }

                            # Check if this cell or its row has preserved content
                            # Find the CLOSEST preceding label cell for value cells
                            row_preserved_key = None
                            # First check this cell itself
                            for kw in PRESERVED_CELL_VALUES:
                                if kw in original_cell_text.strip():
                                    row_preserved_key = kw
                                    break
                            # If not a label, find nearest preceding label cell in same row
                            if not row_preserved_key:
                                for scan_c in range(c_idx - 1, -1, -1):
                                    if scan_c < len(row.cells):
                                        scan_text = row.cells[scan_c].text.strip().replace('\n', '')
                                        for kw in PRESERVED_CELL_VALUES:
                                            if kw in scan_text:
                                                row_preserved_key = kw
                                                break
                                    if row_preserved_key:
                                        break

                            is_preserved_label = any(
                                kw in original_cell_text.strip()
                                for kw in PRESERVED_CELL_VALUES
                            )

                            # Check row context — only preserve for 稳评实施单位 section
                            # Scan ALL cells in the row (not just first 2) for reliable context
                            row_context = ''
                            for scan_c in range(len(row.cells)):
                                row_context += row.cells[scan_c].text.strip().replace('\n', '')
                            is_implement_section = ('实施' in row_context
                                                    and '责任' not in row_context.split('实施')[0]
                                                    if '实施' in row_context else False)
                            is_responsibility_section = ('责任单位' in row_context
                                                         and '实施' not in row_context)

                            # For 稳评责任单位 rows, do NOT use preserved company values.
                            # The fill_value from user data should be used as-is.
                            if is_responsibility_section and row_preserved_key and not is_preserved_label:
                                # Skip preserved content logic — use the user-provided fill_value
                                pass  # Fall through to normal cell filling below
                            elif row_preserved_key and not is_preserved_label and is_implement_section:
                                preserved_val = PRESERVED_CELL_VALUES.get(row_preserved_key, '')
                                if preserved_val and (not original_cell_text or original_cell_text.isspace() or len(original_cell_text) < 3):
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.text = ""
                                    if cell.paragraphs and cell.paragraphs[0].runs:
                                        cell.paragraphs[0].runs[0].text = preserved_val
                                    elif cell.paragraphs:
                                        from docx.oxml import OxmlElement
                                        r = OxmlElement('w:r')
                                        t = OxmlElement('w:t')
                                        t.text = preserved_val
                                        t.set(qn('xml:space'), 'preserve')
                                        r.append(t)
                                        cell.paragraphs[0]._p.append(r)
                                    continue

                            if is_preserved_label and is_implement_section:
                                preserved_val = PRESERVED_CELL_VALUES.get(
                                    # Find matching keyword
                                    next(kw for kw in PRESERVED_CELL_VALUES
                                         if kw in original_cell_text.strip()),
                                    ''
                                )
                                # Fill the label cell
                                label = original_cell_text.strip()
                                if '：' not in label and ':' not in label:
                                    label = label + '：'
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ""
                                if cell.paragraphs and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].text = label + preserved_val
                                elif cell.paragraphs:
                                    from docx.oxml import OxmlElement
                                    r = OxmlElement('w:r')
                                    t = OxmlElement('w:t')
                                    t.text = label + preserved_val
                                    t.set(qn('xml:space'), 'preserve')
                                    r.append(t)
                                    cell.paragraphs[0]._p.append(r)
                                continue  # Skip further processing for this cell

                            # Check if this is a label-only cell that needs user data
                            elif (
                                len(original_cell_text) < 20
                                and any(kw in original_cell_text for kw in ['事项', '名称', '适用'])
                                and not any(c.isdigit() for c in original_cell_text)
                            ):
                                label = original_cell_text
                                if '：' not in label and ':' not in label:
                                    label = original_cell_text + '：'
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ""
                                if cell.paragraphs and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].text = label + fill_value
                                elif cell.paragraphs:
                                    from docx.oxml import OxmlElement
                                    r = OxmlElement('w:r')
                                    t = OxmlElement('w:t')
                                    t.text = label + fill_value
                                    t.set(qn('xml:space'), 'preserve')
                                    r.append(t)
                                    cell.paragraphs[0]._p.append(r)
                            else:
                                # Non-label cell: clear and set fill value
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ""
                                if cell.paragraphs and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].text = fill_value
                                elif cell.paragraphs:
                                    from docx.oxml import OxmlElement
                                    r = OxmlElement('w:r')
                                    t = OxmlElement('w:t')
                                    t.text = fill_value
                                    t.set(qn('xml:space'), 'preserve')
                                    r.append(t)
                                    cell.paragraphs[0]._p.append(r)
                    # Mark grid cells as filled (handle gridSpan merged cells)
                    for _gs_off in range(_grid_span):
                        _filled_grid_cells.add((t_idx, r_idx, c_idx + _gs_off))
                    continue

                pattern_type = loc.get("pattern_type", "")
                run_indices = loc.get("run_indices", [loc.get("run_index", 0)])

                if pattern_type == "highlight" and len(run_indices) > 1:
                    # Multi-run highlight: fill first run, clear the rest
                    first_run_idx = run_indices[0]
                    cls._replace_in_run(
                        doc.paragraphs[loc["paragraph_index"]],
                        first_run_idx,
                        loc.get("matched_text", ""),
                        fill_value,
                        remove_highlight=True,
                    )
                    # Clear remaining highlighted runs
                    for extra_idx in run_indices[1:]:
                        try:
                            run = doc.paragraphs[loc["paragraph_index"]].runs[extra_idx]
                            run.text = ""
                            # Remove highlight
                            rpr = run._r.find(qn('w:rPr'))
                            if rpr is not None:
                                hl = rpr.find(qn('w:highlight'))
                                if hl is not None:
                                    rpr.remove(hl)
                        except IndexError:
                            pass
                elif pattern_type == "form_field":
                    # Form fields: "稳评责任单位：      " → keep label, fill after colon
                    # Multiple runs are common (e.g. "稳评" + "责任单位：" + " ")
                    # Fill first run with full label + value, clear remaining runs
                    para = doc.paragraphs[loc["paragraph_index"]]
                    # Find the colon position in the full paragraph text
                    # (not just at end — date fields like "填 表  日 期：2026年4月10日"
                    #  have existing values after the colon)
                    colon_match = re.search(r'[：:]', para.text)
                    if colon_match and para.runs:
                        label_part = para.text[:colon_match.end()]  # include colon
                        para.runs[0].text = label_part + fill_value
                        # Clear any remaining runs
                        for extra_run in para.runs[1:]:
                            extra_run.text = ""
                    elif para.runs:
                        para.runs[0].text = para.runs[0].text.rstrip() + fill_value
                        for extra_run in para.runs[1:]:
                            extra_run.text = ""
                    else:
                        para.add_run(fill_value)
                elif pattern_type == "image_placeholder":
                    # Image placeholder: insert the actual uploaded image
                    para = doc.paragraphs[loc["paragraph_index"]]
                    image_path = fill_value.strip()
                    is_attachment = loc.get("is_attachment_item", False)

                    # Resolve absolute path for uploaded images
                    abs_image_path = None
                    if image_path and image_path != "该地方为占位图":
                        from app.services.file_service import file_service as _fs
                        try:
                            resolved = _fs.get_absolute_path(image_path)
                            if resolved.exists():
                                abs_image_path = resolved
                        except Exception:
                            pass

                    # Check for multi-image value (JSON array of paths)
                    multi_images = None
                    if fill_value.startswith("[") and fill_value.endswith("]"):
                        try:
                            import json as _json
                            parsed = _json.loads(fill_value)
                            if isinstance(parsed, list) and all(isinstance(x, str) for x in parsed):
                                multi_images = parsed
                        except Exception:
                            pass

                    if abs_image_path or multi_images:
                        # Save original caption text (e.g. "图3-1决策评估公示内容"
                        # or attachment heading like "1.拟征收土地公告、征地范围图")
                        caption_text = loc.get("original_text", "").strip()

                        if is_attachment:
                            # ---- ATTACHMENT ITEM ----
                            # Keep heading text ABOVE the image(s):
                            #   1. First run: heading text
                            #   2. Line break (w:br)
                            #   3. Image(s) — one per run for multi-image
                            heading_text = caption_text

                            # Remove any existing drawings from the paragraph
                            drawing_ns = (
                                "http://schemas.openxmlformats.org/"
                                "wordprocessingml/2006/main"
                            )
                            for drawing in para._p.findall(
                                f"{{{drawing_ns}}}drawing"
                            ):
                                para._p.remove(drawing)

                            from docx.oxml import OxmlElement
                            from docx.shared import Inches
                            from docx.enum.text import WD_ALIGN_PARAGRAPH

                            # Clear all runs and repopulate in correct order
                            for run in para.runs:
                                run.text = ""
                            # Ensure at least one run for the heading text
                            if not para.runs:
                                para.add_run("")

                            # Place heading text in the first run
                            para.runs[0].text = heading_text

                            # Build list of image paths to render
                            img_paths = multi_images if multi_images else [image_path]
                            for ip in img_paths:
                                # Resolve each image path
                                ip_abs = None
                                if ip and ip != "该地方为占位图":
                                    try:
                                        resolved = _fs.get_absolute_path(ip)
                                        if resolved.exists():
                                            ip_abs = resolved
                                    except Exception:
                                        pass

                                # Add a line break before each image
                                br_run = para.add_run()
                                br = OxmlElement('w:br')
                                br.set(qn('w:type'), 'textWrapping')
                                br_run._r.append(br)

                                # Add the image in a new run
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                img_run = para.add_run()
                                if ip_abs:
                                    try:
                                        img_run.add_picture(
                                            str(ip_abs), width=Inches(5.5)
                                        )
                                    except Exception:
                                        img_run.text = f"\n[图片：{ip}]"
                                else:
                                    img_run.text = f"\n[图片：{ip}]"
                        else:
                            # ---- FIGURE CAPTION ----
                            # Clear existing paragraph runs and drawings
                            for run in para.runs:
                                run.text = ""
                            drawing_ns = (
                                "http://schemas.openxmlformats.org/"
                                "wordprocessingml/2006/main"
                            )
                            for drawing in para._p.findall(
                                f"{{{drawing_ns}}}drawing"
                            ):
                                para._p.remove(drawing)

                            from docx.shared import Inches
                            from docx.enum.text import WD_ALIGN_PARAGRAPH
                            from docx.oxml import OxmlElement
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            # Build list of image paths
                            img_paths = multi_images if multi_images else [image_path]

                            # Add caption text above first image
                            if caption_text and para.runs:
                                para.runs[0].text = caption_text
                            elif caption_text:
                                para.add_run(caption_text)

                            for ip in img_paths:
                                ip_abs = None
                                if ip and ip != "该地方为占位图":
                                    try:
                                        resolved = _fs.get_absolute_path(ip)
                                        if resolved.exists():
                                            ip_abs = resolved
                                    except Exception:
                                        pass

                                # Add line break before image (after caption/first image)
                                br_run = para.add_run()
                                br = OxmlElement('w:br')
                                br.set(qn('w:type'), 'textWrapping')
                                br_run._r.append(br)

                                img_run = para.add_run()
                                if ip_abs:
                                    try:
                                        img_run.add_picture(
                                            str(ip_abs), width=Inches(5.5)
                                        )
                                    except Exception:
                                        img_run.text = f"\n[图片：{ip}]"
                                else:
                                    img_run.text = f"\n[图片：{ip}]"
                    else:
                        # No valid image — insert placeholder text
                        placeholder_text = (
                            fill_value if fill_value else "该地方为占位图"
                        )
                        if is_attachment:
                            # Keep heading text, show placeholder below
                            heading_text = loc.get("original_text", "").strip()
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = (
                                    heading_text + "\n" + placeholder_text
                                )
                            else:
                                from docx.oxml import OxmlElement
                                r_element = OxmlElement('w:r')
                                t_element = OxmlElement('w:t')
                                t_element.text = (
                                    heading_text + "\n" + placeholder_text
                                )
                                t_element.set(qn('xml:space'), 'preserve')
                                r_element.append(t_element)
                                para._p.append(r_element)
                        else:
                            # Clear all runs and set text in first run
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = placeholder_text
                            elif para._p is not None:
                                from docx.oxml import OxmlElement
                                pPr = para._p.find(qn('w:pPr'))
                                if pPr is None:
                                    pPr = OxmlElement('w:pPr')
                                    para._p.insert(0, pPr)
                                pStyle = pPr.find(qn('w:pStyle'))
                                # Only set to Normal if no style is already defined
                                if pStyle is None:
                                    pStyle = OxmlElement('w:pStyle')
                                    pPr.append(pStyle)
                                    pStyle.set(qn('w:val'), 'Normal')
                                # If style already set (e.g. "Body Text"), leave it unchanged
                                r_element = OxmlElement('w:r')
                                rPr = OxmlElement('w:rPr')
                                rFonts = OxmlElement('w:rFonts')
                                rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                                rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                                rPr.append(rFonts)
                                r_element.append(rPr)
                                t_element = OxmlElement('w:t')
                                t_element.text = placeholder_text
                                t_element.set(qn('xml:space'), 'preserve')
                                r_element.append(t_element)
                                para._p.append(r_element)
                else:
                    para_idx = loc.get("paragraph_index", -1)
                    # Skip non-paragraph placeholders: tables (para_idx=-1),
                    # form fields with no paragraph mapping, etc.
                    if para_idx < 0 or para_idx >= len(doc.paragraphs):
                        continue
                    para = doc.paragraphs[para_idx]
                    if len(para.runs) == 0:
                        # Empty paragraph (no runs) — add a new run with the fill value
                        from docx.oxml import OxmlElement
                        # Ensure paragraph has a style element (preserve existing style)
                        pPr = para._p.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            para._p.insert(0, pPr)
                        pStyle = pPr.find(qn('w:pStyle'))
                        # Only set to Normal if no style is already defined
                        if pStyle is None:
                            pStyle = OxmlElement('w:pStyle')
                            pPr.append(pStyle)
                            pStyle.set(qn('w:val'), 'Normal')
                        # If style already set (e.g. "Body Text"), leave it unchanged

                        r_element = OxmlElement('w:r')
                        # Set run style to Normal
                        rPr = OxmlElement('w:rPr')
                        rFonts = OxmlElement('w:rFonts')
                        rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        rPr.append(rFonts)
                        r_element.append(rPr)

                        t_element = OxmlElement('w:t')
                        t_element.text = fill_value
                        t_element.set(qn('xml:space'), 'preserve')
                        r_element.append(t_element)
                        para._p.append(r_element)
                    else:
                        is_highlight = pattern_type == "highlight"
                        cls._replace_in_run(
                            para,
                            loc.get("run_index", 0),
                            loc.get("matched_text", ""),
                            fill_value,
                            remove_highlight=is_highlight,
                        )

        # Remove all yellow highlights and shading from the filled document
        cls._remove_all_highlights(doc)

        # Strip markdown formatting from ALL runs (AI-generated content may have ##, **, etc.)
        cls._strip_all_markdown(doc)

        # Remove AI instruction annotations (blue-font instruction content from template)
        cls._remove_instruction_content(doc)

        # Save
        doc.save(output_path)
        return output_path

    # ── Instruction Content Patterns ──
    # Regex patterns that identify AI instruction annotations in the template.
    # These are blue-font runs in the template that tell AI what to generate,
    # NOT actual report content. They must be removed from output.
    _INSTRUCTION_PATTERNS = [
        # "ai结合知识库完整报告的这部分内容进行编写，要合理，要数据准确..."
        r'ai.*(?:结合|根据).*(?:知识库|完整报告).*(?:编写|生成)',
        # "根据用户提供数据进行填充，时间和受哪个部门的委托..."
        r'用户提供.*(?:数据|填写|填充|基础数据)',
        # "需要ai对图片进行分析，所有的选项内容进行填充后修改..."
        r'需要(?:ai|AI).*(?:分析|填充|编写|生成|修改)',
        # "根据相关地区当地政策进行智能修改，建议上网查询..."
        r'根据(?:相关地区|当地).*(?:修改|调整|查询)',
        # "建议上网查询用户想要写的新地区的国土空间规划..."
        r'建议上网查询',
        # "得分需要ai重新编造，最终的数据一定要符合要求"
        r'(?:得分|数据)需要(?:ai|AI)重新',
        # "根据实际用户提供的数据进行合理构写"
        r'根据实际用户提供.*(?:构写|编写)',
        # "针对上述内容生成一个表格  表格格式要和原先这块的内容一模一样"
        r'针对上述内容生成',
        # "表格格式要和原先这块的内容一模一样"
        r'表格格式.*一模一样',
        # "这块内容分析比例那一列数据要改变是根据附件中的..."
        r'(?:这块|这部分).*(?:内容.*改变|需要.*填充后)',
        # "合理构写" / "智能修改"
        r'(?:合理|智能).*(?:构写|编写|修改)',
        # "根据用户提供数据进行填充"
        r'根据用户提供.*填充',
        # Generic: "不要胡编乱造" / "有理有据" / "数据准确"
        r'不要胡编乱造',
        # "用户想要生成的新报告的地方信息"
        r'(?:用户|你).*想要.*生成.*报告',
    ]
    _INSTRUCTION_RE = None  # Compiled lazily

    @classmethod
    def _get_instruction_re(cls):
        """Lazily compile instruction regex patterns."""
        if cls._INSTRUCTION_RE is None:
            import re as _re
            cls._INSTRUCTION_RE = _re.compile(
                '|'.join(cls._INSTRUCTION_PATTERNS),
                _re.IGNORECASE
            )
        return cls._INSTRUCTION_RE

    @classmethod
    def _remove_instruction_content(cls, doc: "Document"):
        """Remove AI instruction annotations from the document.

        Scans all runs with blue font (explicit RGB 0000FF or theme color)
        and:
        1. If the text matches instruction patterns → clear the run text entirely
        2. If the text is real content but blue → reset color to default (black)

        Also clears any remaining instruction-like text even without blue color
        (e.g., text that starts paragraphs with "🟡" emoji placeholders).
        """
        import re as _re

        instruction_re = cls._get_instruction_re()

        # ── Helper: check if a run has blue font color ──
        def _is_blue_run(run) -> bool:
            # Check python-docx color API
            try:
                color = run.font.color.rgb
                if color and str(color) in (
                    "0000FF", "0070C0", "1F4E79", "2F5496",
                    "0000CD", "000080", "4169E1", "1E90FF",
                ):
                    return True
            except Exception:
                pass
            # Check raw XML w:color element (handles theme colors)
            try:
                rPr = run._r.find(qn('w:rPr'))
                if rPr is not None:
                    color_el = rPr.find(qn('w:color'))
                    if color_el is not None:
                        val = color_el.get(qn('w:val'), '')
                        if val and val.upper() in (
                            "0000FF", "0070C0", "1F4E79", "2F5496",
                            "0000CD", "000080", "4169E1", "1E90FF",
                        ):
                            return True
            except Exception:
                pass
            return False

        # ── Helper: reset a run's font color to default ──
        def _reset_color_to_default(run):
            try:
                rPr = run._r.find(qn('w:rPr'))
                if rPr is not None:
                    color_el = rPr.find(qn('w:color'))
                    if color_el is not None:
                        # Remove explicit color — inherits from paragraph/theme
                        rPr.remove(color_el)
                # Also try the python-docx API
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black
            except Exception:
                pass

        # ── Clear highlight from a run (remove w:highlight element) ──
        def _clear_highlight(run):
            cls._strip_run_highlight(run)

        # ── Step 1: Process all paragraph runs ──
        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text
                if not text or not text.strip():
                    continue

                is_blue = _is_blue_run(run)

                if is_blue and instruction_re.search(text):
                    # Blue font + instruction pattern → clear text AND reset color
                    run.text = ""
                    _reset_color_to_default(run)
                elif is_blue:
                    # Blue font but real content → just reset color
                    _reset_color_to_default(run)
                elif instruction_re.search(text):
                    # Text matches instruction pattern (even without blue font)
                    # The instruction patterns are specific enough to avoid
                    # false positives. Clear the run text.
                    run.text = ""

        # ── Step 2: Process all table cell runs ──
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            text = run.text
                            if not text or not text.strip():
                                continue

                            is_blue = _is_blue_run(run)

                            if is_blue and instruction_re.search(text):
                                run.text = ""
                                _reset_color_to_default(run)
                            elif is_blue:
                                _reset_color_to_default(run)
                            elif instruction_re.search(text):
                                run.text = ""

        # ── Step 3: Strip instruction prefixes from mixed-content runs ──
        # Some runs have instruction text AND real content in the same run
        # (e.g. "根据用户提供数据进行填充，时间和受哪个部门的委托，稳评第三方机构...")
        # Strip the instruction prefix, keep the real content.
        _INSTRUCTION_PREFIX_RE = _re.compile(
            r'^\s*'
            r'(?:根据用户提供数据(?:进行)?填充[，,]\s*'
            r'(?:时间(?:和|及)受(?:哪个|什么)(?:部门|单位)的委托[，,]?\s*)?'
            r'|ai结合知识库完整报告的这(?:部分|块)内容进行编写[，,]\s*'
            r'(?:要合理[，,]?\s*(?:要)?数据准确[，,]?\s*)?'
            r'(?:根据用户想要生成的(?:新报告的)?地方信息进行编写[，,]?\s*)?'
            r'(?:不要胡编乱造[，,]?\s*(?:要)?有理有据[。.]?\s*)?'
            r')'
            r'\s*'
        )
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text:
                    original = run.text
                    cleaned = _INSTRUCTION_PREFIX_RE.sub('', original)
                    if cleaned != original:
                        run.text = cleaned if cleaned.strip() else ""

        # ── Step 4: Remove instruction-like emoji markers ──
        # These are placeholder markers from the template (🟡, 📝, etc.)
        # that appear as prefix annotations
        _EMOJI_INSTRUCTION_RE = _re.compile(
            r'^[🟡🔵📝📋🖼️📊💰📜🏢]\s*(?:本决策|图\s*\d|根据用户|ai结合|用户提供)'
        )
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and _EMOJI_INSTRUCTION_RE.search(run.text.strip()):
                    # Clear emoji instruction prefix, keep any trailing real content
                    cleaned = _EMOJI_INSTRUCTION_RE.sub('', run.text.strip())
                    run.text = cleaned if cleaned.strip() else ""

        # ── Step 5: Final pass — strip ALL remaining blue color attributes ──
        # This catches empty runs with residual blue color, theme-colored runs,
        # and any other blue font runs that slipped through earlier steps.
        # Iterates ALL runs (including empty-text ones) in paragraphs + tables.
        all_runs = []
        for para in doc.paragraphs:
            all_runs.extend(para.runs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        all_runs.extend(para.runs)

        _BLUE_COLORS = {
            "0000FF", "0070C0", "1F4E79", "2F5496",
            "0000CD", "000080", "4169E1", "1E90FF",
        }
        for run in all_runs:
            try:
                rPr = run._r.find(qn('w:rPr'))
                if rPr is not None:
                    color_el = rPr.find(qn('w:color'))
                    if color_el is not None:
                        val = (color_el.get(qn('w:val')) or '').upper()
                        theme = color_el.get(qn('w:themeColor')) or ''
                        if val in _BLUE_COLORS:
                            rPr.remove(color_el)
            except Exception:
                pass

    @classmethod
    def _remove_all_highlights(cls, doc: "Document"):
        """Remove ALL highlight and shading formatting from every run in the document.

        This is called after filling placeholders to ensure no yellow background
        remains in the final output — whether from replaced runs or from
        unfilled highlighted runs.
        """
        # 1. All paragraph runs
        for para in doc.paragraphs:
            for run in para.runs:
                cls._strip_run_highlight(run)

        # 2. All table cell runs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            cls._strip_run_highlight(run)

    @classmethod
    def _strip_all_markdown(cls, doc: "Document"):
        """Strip markdown formatting from ALL runs in the document.

        Removes: ## headings, **bold**, *italic*, |table| markers,
        - list markers, > blockquotes, === separators, numbered lists,
        emoji markers, and other markdown syntax that shouldn't appear in docx.

        This is more aggressive than strip_markdown_for_docx() because it
        operates on individual runs, where partial markdown fragments
        (like a standalone "##" or "**") may be split across runs.
        """
        import re as _re_md

        def clean_text(text: str) -> str:
            if not text:
                return text
            # Remove heading markers: ### text → text (multiline)
            text = _re_md.sub(r'^#{1,6}\s+', '', text, flags=_re_md.MULTILINE)
            # Handle ##text without space (common LLM output artifact)
            text = _re_md.sub(r'#{1,6}(?=[一-鿿A-Za-zĀ-ɏ])', '', text)
            # Handle standalone # markers (run-level: a run may have just "##" or "#")
            text = _re_md.sub(r'^#{1,6}$', '', text, flags=_re_md.MULTILINE)
            # Remove bold markers: **text** → text
            text = _re_md.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # Handle standalone ** (opening/closing split across runs)
            text = _re_md.sub(r'^\*\*$', '', text, flags=_re_md.MULTILINE)
            # Remove italic markers
            text = _re_md.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
            # Handle standalone * markers
            text = _re_md.sub(r'^\*$', '', text, flags=_re_md.MULTILINE)
            # Remove blockquote: > text
            text = _re_md.sub(r'^>\s+', '', text, flags=_re_md.MULTILINE)
            text = _re_md.sub(r'^>$', '', text, flags=_re_md.MULTILINE)
            # Collapse markdown table row separators: |---|---|
            text = _re_md.sub(r'\|?\s*[-:| ]{3,}\s*\|?', '', text)
            # Remove pipe-separated table cell markers
            # (only when clearly markdown: | cell1 | cell2 | with spaces around pipes)
            if text.count('|') >= 1:
                # Only strip pipes if the text looks like a table row
                if _re_md.search(r'\|\s*[^\|]+\s*\|', text):
                    cells = [c.strip() for c in text.split('|')]
                    text = '  '.join(c for c in cells if c)
            # Remove leading list markers: "- " or "1. " or "* "
            text = _re_md.sub(r'^[\-\*\d+]\.[ \t]+', '', text, flags=_re_md.MULTILINE)
            text = _re_md.sub(r'^[-]\s+', '', text, flags=_re_md.MULTILINE)
            # Remove horizontal rules: ---, ***, ===
            text = _re_md.sub(r'^[-*=_]{3,}\s*$', '', text, flags=_re_md.MULTILINE)
            # Remove standalone emoji that serves as markers
            text = _re_md.sub(
                r'^[📊📋✅⚠️📷💡🎯📝🖼️🔍🔗📄📐🧠🧪🔨📍📌💰📜🏢🟡🔵]\s*',
                '', text, flags=_re_md.MULTILINE
            )
            # Remove "##" and "**" even when embedded in Chinese text
            # (common LLM hallucination: using markdown inside Chinese paragraphs)
            text = text.replace('##', '').replace('**', '')
            # Remove "=== " and "--- " at line starts (restructured text artifacts)
            text = _re_md.sub(r'^(?:===|——-)\s*', '', text, flags=_re_md.MULTILINE)
            # Collapse multiple blank lines
            text = _re_md.sub(r'\n{3,}', '\n\n', text)
            # Remove lines that are entirely whitespace/separator
            text = _re_md.sub(r'^\s*[-=_]{2,}\s*$', '', text, flags=_re_md.MULTILINE)
            return text.strip()

        # Clean all paragraph runs
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text:
                    run.text = clean_text(run.text)

        # Clean all table cell runs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                run.text = clean_text(run.text)

    @staticmethod
    def _strip_run_highlight(run):
        """Remove w:highlight and w:shd (shading) from a single run's XML element."""
        rpr = run._r.find(qn('w:rPr'))
        if rpr is None:
            return
        # Remove <w:highlight w:val="yellow"/>
        highlight_elem = rpr.find(qn('w:highlight'))
        if highlight_elem is not None:
            rpr.remove(highlight_elem)
        # Remove <w:shd w:fill="FFFF00"/> etc. (some templates use shading for yellow bg)
        shd_elem = rpr.find(qn('w:shd'))
        if shd_elem is not None:
            fill_color = shd_elem.get(qn('w:fill'), '').upper() if shd_elem.get(qn('w:fill')) else ''
            if fill_color in ('FFFF00', 'FFEB3B', 'YELLOW', 'FFFF00', 'FFF000'):
                rpr.remove(shd_elem)

    @classmethod
    def _fill_text_and_image(
        cls,
        doc: "Document",
        text_loc: Dict[str, Any],
        text_value: str,
        image_loc: Optional[Dict[str, Any]],
        image_value: Optional[str],
    ):
        """Fill a paragraph with both text content and an image.

        Writes text in the first run, then appends the image as a new run.
        Prevents the two from overwriting each other when a paragraph has
        both an example_diff/blank and an image_placeholder.
        """
        para = doc.paragraphs[text_loc["paragraph_index"]]
        from docx.oxml import OxmlElement
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # ── Step 1: Write text content ──
        if len(para.runs) == 0:
            # Ensure paragraph has a style element
            pPr = para._p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                para._p.insert(0, pPr)
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is None:
                pStyle = OxmlElement('w:pStyle')
                pPr.append(pStyle)
                pStyle.set(qn('w:val'), 'Normal')
            r_element = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), '仿宋_GB2312')
            rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            rPr.append(rFonts)
            r_element.append(rPr)
            t_element = OxmlElement('w:t')
            t_element.text = text_value
            t_element.set(qn('xml:space'), 'preserve')
            r_element.append(t_element)
            para._p.append(r_element)
        else:
            cls._replace_in_run(
                para,
                text_loc.get("run_index", 0),
                text_loc.get("matched_text", ""),
                text_value,
                remove_highlight=False,
            )

        # ── Step 2: Append image ──
        if not image_value or not image_loc:
            return

        from app.services.file_service import file_service as _fs
        abs_image_path = None
        try:
            resolved = _fs.get_absolute_path(image_value)
            if resolved.exists():
                abs_image_path = resolved
        except Exception:
            pass

        if abs_image_path:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_run = para.add_run()
            try:
                img_run.add_picture(str(abs_image_path), width=Inches(5.5))
            except Exception:
                img_run.text = f"\n[图片：{image_value}]"
        else:
            para.add_run(f"\n[图片占位：{image_value}]")

    @classmethod
    def _replace_in_run(
        cls,
        paragraph,
        run_index: int,
        placeholder_text: str,
        new_value: str,
        remove_highlight: bool = False,
    ):
        """
        Replace placeholder text within a specific run.
        Sets paragraph style to Normal (正文) for filled content.
        If remove_highlight is True, also removes highlight/shading formatting.
        """
        try:
            run = paragraph.runs[run_index]
            original_text = run.text

            # Guard: empty placeholder_text would cause str.replace("", "x") to
            # insert between every character, corrupting the text.
            # Only replace entire run if it has highlight formatting (placeholder marker).
            if not placeholder_text:
                has_highlight = False
                rpr = run._r.find(qn('w:rPr'))
                if rpr is not None:
                    has_highlight = (rpr.find(qn('w:highlight')) is not None or
                                     rpr.find(qn('w:shd')) is not None)
                if has_highlight:
                    run.text = new_value
                    # Strip highlight after filling
                    if rpr is not None:
                        for tag in [qn('w:highlight'), qn('w:shd')]:
                            elem = rpr.find(tag)
                            if elem is not None:
                                rpr.remove(elem)
                return

            if placeholder_text in original_text:
                # Replace the placeholder pattern with the new value
                new_text = original_text.replace(placeholder_text, new_value)
                run.text = new_text
            else:
                # Try regex-based replacement
                for pattern, _ in PLACEHOLDER_PATTERNS:
                    if re.search(pattern, original_text):
                        new_text = re.sub(pattern, new_value, original_text, count=1)
                        run.text = new_text
                        break
                else:
                    # Last resort: replace the whole run text
                    run.text = new_value

            # Remove highlight and shading formatting after replacement
            if remove_highlight:
                rpr = run._r.find(qn('w:rPr'))
                if rpr is not None:
                    highlight_elem = rpr.find(qn('w:highlight'))
                    if highlight_elem is not None:
                        rpr.remove(highlight_elem)
                    shd_elem = rpr.find(qn('w:shd'))
                    if shd_elem is not None:
                        rpr.remove(shd_elem)
        except IndexError:
            # Run index out of range — skip gracefully
            pass

    @classmethod
    def fill_tables(
        cls,
        docx_path: str,
        table_data: Dict[int, List[List[str]]],  # {table_index: [[cell_value, ...], ...]}
    ) -> None:
        """
        Fill table cells with provided data. table_index matches the order in the document.
        """
        doc = Document(docx_path)
        for t_idx, rows_data in table_data.items():
            if t_idx < len(doc.tables):
                table = doc.tables[t_idx]
                for r_idx, row_data in enumerate(rows_data):
                    if r_idx < len(table.rows):
                        row = table.rows[r_idx]
                        for c_idx, cell_value in enumerate(row_data):
                            if c_idx < len(row.cells):
                                row.cells[c_idx].text = cell_value
        doc.save(docx_path)

    @classmethod
    def generate_placeholder_template(
        cls, docx_path: str, placeholders: list, output_path: str
    ) -> str:
        """Generate a placeholder-ized docx from a complete report.

        Uses content-based pattern matching to identify variable content
        (project names, locations, dates, area values, etc.) and replace
        them with {{placeholder_key}} markers across multi-run text spans.
        """
        import logging, re
        logger = logging.getLogger(__name__)

        doc = Document(docx_path)
        replace_count = 0

        # Build content patterns → placeholder keys
        # These regex patterns match common variable content in 稳评 reports
        CONTENT_PATTERNS = [
            # Project name — all variants
            (r'洞庭湖路\s*（S350[^）]*）\s*工程\s*土地征收\s*决策', 'project_name_full'),
            (r'洞庭湖路[^\n]{0,30}工程', 'project_name'),
            (r'洞庭湖路', 'project_name_short'),
            (r'S350', 'remove_s350'),
            (r'宁连一级路段', 'remove_ninglian'),
            # Location — all levels
            (r'金湖县', 'location_county'),
            (r'金湖(?!向)', 'location_city'),
            (r'淮安市', 'location_prefecture'),
            (r'淮安(?!经济技术)', 'location_pref_short'),
            (r'戴楼街道', 'location_street'),
            (r'戴楼', 'location_street_short'),
            # Dates
            (r'二〇二四年', 'report_year_cn'),
            (r'2024年', 'report_year'),
            (r'二〇二四年七月', 'report_date_full'),
            # Document references
            (r'金征预告\s*〔\d{4}〕\s*\d+\s*号', 'doc_reference'),
            # Land data — full paragraphs for complete replacement
            (r'拟征收[^\n]{0,30}面积\s*[约大约为]?\s*\d+\.?\d*\s*(?:亩|公顷|平方米)', 'land_area'),
            (r'征收范围[：:]\s*[^\n]{10,200}', 'land_scope_full'),
            (r'拟征收土地位于[^\n]{10,200}', 'land_scope_full'),
            # Decision info
            (r'决策事项[^\n]{0,30}名称[：:][^\n]{0,50}', 'decision_name'),
            (r'决策\s*单位[：:][^\n]{0,30}', 'decision_unit'),
            # Background/Purpose
            (r'为积极配合做好决策事项实施的各项工作[^\n]{0,80}', 'implementation_bg'),
            # Contact
            (r'联系人[：:]\s*\S{2,4}[；;]?\s*(?:电话[：:]\s*)?\d{3,4}[-\s]?\d{7,8}', 'contact_info'),
        ]

        # Chapter-level analysis markers — replace body text after headings
        CHAPTER_MARKERS = {
            "合法性": "{{legality_analysis}}",
            "合理性": "{{rationality_analysis}}",
            "可行性": "{{feasibility_analysis}}",
            "可控性": "{{controllability_analysis}}",
            "风险因素": "{{risk_factor_analysis}}",
            "风险防范": "{{risk_mitigation}}",
            "应急预案": "{{emergency_plan}}",
            "评估结论": "{{assessment_conclusion}}",
            "防范和化解": "{{risk_mitigation}}",
            "突发性事件": "{{emergency_plan}}",
        }

        # Survey data patterns to replace in tables and text
        SURVEY_PATTERNS = [
            (r'支持率\s*\d+\.?\d*\s*%', 'survey_support_rate'),
            (r'总样本[数数]\s*\d+\s*份?', 'survey_total_samples'),
            (r'支持\s*\d+\s*人', 'survey_support_count'),
            (r'反对\s*\d+\s*人', 'survey_oppose_count'),
            (r'条件支持\s*\d+\s*人', 'survey_conditional_count'),
            (r'调查日期\s*[^\n]{0,20}', 'survey_date'),
        ]

        # Expert opinion patterns
        EXPERT_PATTERNS = [
            (r'专家意见[：:][^\n]{0,100}', 'expert_opinion'),
            (r'评审结论[：:]\s*[^\n]{0,30}', 'review_conclusion'),
        ]

        def _replace_pattern(pattern, key, paragraphs):
            nonlocal replace_count
            regex = re.compile(pattern)
            marker = f"{{{{{key}}}}}"
            for para in paragraphs:
                if not para.runs:
                    continue
                runs = para.runs
                full_text = "".join(r.text for r in runs)
                m = regex.search(full_text)
                if not m:
                    continue

                # Build run boundary map
                run_bounds = []
                pos = 0
                for i, run in enumerate(runs):
                    start = pos
                    pos += len(run.text)
                    run_bounds.append((start, pos, i))

                match_start, match_end = m.start(), m.end()

                # Find affected runs
                affected = []
                for start, end, ri in run_bounds:
                    if end > match_start and start < match_end:
                        seg_start = max(start, match_start)
                        seg_end = min(end, match_end)
                        affected.append((ri, seg_start - start, seg_end - start))

                # Replace across runs
                for ai, (ri, local_start, local_end) in enumerate(affected):
                    run = runs[ri]
                    if ai == 0:
                        run.text = run.text[:local_start] + marker + run.text[local_end:]
                    else:
                        run.text = run.text[:local_start] + run.text[local_end:]

                replace_count += 1
                break  # One replacement per paragraph

        # Apply each pattern to all paragraphs and table cells
        all_paras = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)

        for pattern, key in CONTENT_PATTERNS:
            _replace_pattern(pattern, key, all_paras)

        # Apply survey and expert patterns to paragraphs AND table cells
        for pattern, key in SURVEY_PATTERNS + EXPERT_PATTERNS:
            _replace_pattern(pattern, key, all_paras)

        # Process TABLE CELLS — replace template data with placeholders
        table_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        runs = para.runs
                        if not runs: continue
                        full = "".join(r.text for r in runs)
                        modified = full
                        for pattern, key in CONTENT_PATTERNS:
                            m = re.search(pattern, modified)
                            if m:
                                marker = f"{{{{{key}}}}}"
                                modified = modified.replace(m.group(0), marker)
                                replace_count += 1
                        if modified != full:
                            # Rebuild runs
                            pos = 0; ri = 0; total = len(modified)
                            while pos < total and ri < len(runs):
                                rlen = min(len(runs[ri].text), total - pos)
                                runs[ri].text = modified[pos:pos + rlen]
                                pos += rlen; ri += 1
                            for rj in range(ri, len(runs)):
                                runs[rj].text = ""
                            table_count += 1
        if table_count > 0:
            logger.info(f"Table placeholders: {table_count} cells modified")

        # Add image placeholders — type-based, skip fixed content, no sequence numbers
        # Fixed content blocks — ALWAYS preserved (from KB, never user data)
        FIXED_BLOCK_MAP = {
            '公司营业执照': 'fixed_license',
            '营业执照': 'fixed_license',
            '稳评平台备案及人员证书': 'fixed_certificates',
            '平台备案': 'fixed_certificates',
            '人员证书': 'fixed_certificates',
            '图5-1 本决策风险评估流程图': 'fixed_risk_flowchart',
            '风险评估流程图': 'fixed_risk_flowchart',
            '公司简介': 'fixed_company_intro',
            '稳评工作组人员及分工情况': 'fixed_team_structure',
            '稳评负责人': 'fixed_team_lead',
            '陈春': 'fixed_team_lead',
            '江苏众拓项目代理咨询有限公司': 'fixed_company_name',
        }

        # Variable blocks — WHOLE paragraph replacement from PDF
        VARIABLE_BLOCK_MAP = {
            '决策事项名称': 'block_decision_name',
            '稳评责任单位': 'block_responsible_unit',
            '决策地理位置': 'block_location',
            '实施背景及地块用途': 'block_background',
            '征收目的': 'block_purpose',
            '拟征地位置': 'block_location',
        }
        TYPE_KEYWORDS = [
            (['公示', '公告栏', '张贴'], 'public_notice'),
            (['位置', '红线', '征地', '示意图', '范围图'], 'location_map'),
            (['现场', '施工', '用地', '地块'], 'site_photo'),
            (['座谈', '开会', '村民', '群众', '会议'], 'meeting_photo'),
            (['问卷', '调查表', '统计表', '问卷调查'], 'survey_scan'),
            (['专家', '评审', '意见', '签到'], 'expert_review'),
            (['附件', '附图', '公告', '勘测', '批文', '报告'], 'attachment'),
        ]

        def _get_type(text):
            for keywords, tname in TYPE_KEYWORDS:
                if any(kw in text for kw in keywords):
                    return tname
            return None

        # Iterate ALL images (check both runs AND paragraph-level drawings)
        WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        img_counters = {}
        img_idx = 0
        for i, para in enumerate(doc.paragraphs):
            # Collect all drawings: in runs + direct paragraph children
            all_drawings = []
            for run in para.runs:
                all_drawings.extend(run._element.findall(f'.//{{{WML_NS}}}drawing'))
            # Also check for drawings directly in paragraph (not in runs)
            all_drawings.extend(para._element.findall(f'.//{{{WML_NS}}}drawing'))

            if not all_drawings:
                continue

            # Check if in fixed content area
            is_fixed = False
            for j in range(max(0, i - 3), min(len(doc.paragraphs), i + 4)):
                if any(kw in doc.paragraphs[j].text for kw in FIXED_KEYWORDS):
                    is_fixed = True
                    break
            if is_fixed:
                continue

            # Find type from nearby text (±4 paragraphs)
            found_type = None
            for j in range(max(0, i - 4), min(len(doc.paragraphs), i + 5)):
                ct = doc.paragraphs[j].text.strip()
                if not ct or len(ct) > 80:
                    continue
                tp = _get_type(ct)
                if tp:
                    found_type = tp
                    break

            if not found_type:
                found_type = 'attachment'

            # Process EACH drawing individually (one marker per image)
            for d in all_drawings:
                cnt = img_counters.get(found_type, 0) + 1
                img_counters[found_type] = cnt
                marker = f"{{{{img_{found_type}_{cnt}}}}}"
                parent = d.getparent()
                if parent is not None:
                    parent.remove(d)
                # Add marker text where the image was
                if para.runs:
                    para.runs[0].text = marker + para.runs[0].text
                else:
                    para.add_run(marker)
                img_idx += 1

        # Also process table cell images
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        cell_drawings = []
                        for run in para.runs:
                            cell_drawings.extend(run._element.findall(f'.//{{{WML_NS}}}drawing'))
                        cell_drawings.extend(para._element.findall(f'.//{{{WML_NS}}}drawing'))
                        if not cell_drawings:
                            continue
                        ct = cell.text.strip()
                        if any(kw in ct for kw in FIXED_KEYWORDS):
                            continue
                        found_type = _get_type(ct) if (ct and len(ct) < 80) else 'attachment'
                        if not found_type:
                            found_type = 'attachment'
                        for d in cell_drawings:
                            cnt = img_counters.get(found_type, 0) + 1
                            img_counters[found_type] = cnt
                            marker = f"{{{{img_{found_type}_{cnt}}}}}"
                            parent = d.getparent()
                            if parent is not None:
                                parent.remove(d)
                        if para.runs:
                            para.runs[0].text = marker + para.runs[0].text
                        else:
                            para.add_run(marker)
                        img_idx += 1

        # Chapter 12 附件 — label each attachment item with a named placeholder
        ATTACHMENT_NAMES = [
            "征地红线图", "拟征地公告", "勘测定界报告",
            "建设项目用地预审与选址意见书", "法人证明及代码证",
            "座谈会签到表", "稳评问卷调查表",
            "专家评审会照片", "专家评审签到表", "专家评审意见",
        ]
        for i, para in enumerate(doc.paragraphs):
            t = para.text.strip()
            for name in ATTACHMENT_NAMES:
                if name in t and len(t) < 30:
                    marker = f"{{{{attachment_{name}}}}}"
                    if para.runs:
                        para.runs[0].text = marker
                    else:
                        para.add_run(marker)
                    break

        total = sum(img_counters.values())
        if total > 0:
            logger.info(f"Image placeholders: {total} slots, by type: {img_counters}")

        # Insert chapter body markers after chapter headings
        for i, para in enumerate(doc.paragraphs):
            t = para.text.strip()
            if not t:
                continue
            for keyword, marker in CHAPTER_MARKERS.items():
                if keyword in t and len(t) < 80 and re.search(r'第[一二三四五六七八九十\d]+章', t):
                    # Find the body paragraph after this heading
                    # and replace it with the analysis marker
                    for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                        body = doc.paragraphs[j]
                        body_text = body.text.strip()
                        if body_text and len(body_text) > 30:
                            # Replace body runs with marker
                            if body.runs:
                                body.runs[0].text = marker
                                for r in body.runs[1:]:
                                    r.text = ""
                            replace_count += 1
                            break
                    break  # One marker per heading

        doc.save(output_path)
        logger.info(
            f"Placeholder template: {replace_count} replacements → {output_path}"
        )
        return output_path

    @classmethod
    def get_document_metadata(cls, docx_path: str) -> Dict[str, Any]:
        """Get basic metadata about a docx file."""
        doc = Document(docx_path)
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        # Count headings
        heading_count = 0
        sections = []
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else ""
            if style_name.startswith("Heading") or style_name.startswith("heading"):
                heading_count += 1
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                sections.append({"level": level, "title": p.text.strip()[:100]})

        return {
            "paragraph_count": para_count,
            "table_count": table_count,
            "heading_count": heading_count,
            "sections": sections,
        }


docx_service = DocxService()
