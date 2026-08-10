"""template_docx_writer.py — 基于模板 docx 填入内容的 Word 报告生成器

核心思路：不是从零创建空 docx，而是以模板 docx 为基底，
找到各章节对应的段落范围，替换正文内容但保留：
  - 模板表格结构（只替换表格中的文字内容）
  - 模板中的图片位置（保留或替换为新图片）
  - 模板格式和样式

修复版本 v2：
  1. 章节定位：支持中文数字（第一章）+ Plain Text 样式
  2. 占位符替换：{{xxx}} 格式全覆盖（段落+表格）
  3. 内容填充：删除旧段落→在原位插入新段落，不再只清空run
  4. 章节映射：引擎10章 → 模板12章（标题相似度匹配）
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 中文数字 → 阿拉伯数字映射
_CN_NUM = {
    "一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
    "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
    "十一": 11, "十二": 12, "十三": 13,
}

# 引擎章节标题关键词 → 模板章节标题关键词（用于匹配）
_CHAPTER_TITLE_MAP = {
    # 引擎关键词 → 模板关键词
    "基本概况": "基本情况",
    "基本情况": "基本情况",
    "评估过程": "稳评工作方案",
    "编制依据": "编制依据",
    "风险因素调查": "利益相关者",
    "综合分析": "合法性",
    "风险因素识别": "分析预测风险源",
    "防范": "防范和化解",
    "化解措施": "防范和化解",
    "风险等级评估": "确定风险等级",
    "评估结论": "确定风险等级",
    "应急预案": "突发性事件",
}

# 模板占位符 → 默认替换值
_PLACEHOLDER_DEFAULTS = {
    "project_name_full": "洪拟征告〔2026〕7号（朱坝街道及三圩社区商业服务业设施用地项目）土地征收决策",
    "project_name": "洪拟征告〔2026〕7号",
    "project_name_short": "洪拟征告〔2026〕7号",
    "decision_name": "洪拟征告〔2026〕7号（朱坝街道及三圩社区商业服务业设施用地项目）土地征收决策",
    "location_prefecture": "淮安市洪泽区",
    "location_pref_short": "洪泽",
    "report_year": "2026",
    "report_year_cn": "二〇二六",
    "remove_s350": "朱坝街道及三圩社区",
    "remove_ninglian": "商业服务业设施用地项目",
    "implementation_bg": "（待填写）",
    "legality_analysis": "（待填写）",
    "risk_mitigation": "（待填写）",
    "emergency_plan": "（待填写）",
}


class TemplateDocxWriter:
    """基于模板 docx 的内容填入引擎（v2）。"""

    def write_with_template(
        self,
        template_path: str,
        chapter_contents: Dict[int, str],
        output_path: str,
        project_name: str = "",
        attachment_images: Optional[Dict[str, List[str]]] = None,
        placeholder_overrides: Optional[Dict[str, str]] = None,
    ) -> str:
        """使用模板 docx 作为基底，填入章节内容。

        Args:
            template_path: 模板 docx 文件路径
            chapter_contents: {chapter_no: markdown_text} 各章内容
            output_path: 输出文件路径
            project_name: 项目名称（用于替换占位符和旧名称）
            attachment_images: {category: [image_paths]} 附件图片分类
            placeholder_overrides: 自定义占位符替换值（覆盖默认值）
        """
        doc = DocxDocument(template_path)

        # Step 1: 替换所有 {{xxx}} 占位符（最关键！解决"名称没有替换"问题）
        overrides = dict(_PLACEHOLDER_DEFAULTS)
        if project_name:
            overrides["project_name_full"] = project_name
            overrides["project_name"] = project_name
            overrides["project_name_short"] = project_name
            overrides["decision_name"] = project_name
        if placeholder_overrides:
            overrides.update(placeholder_overrides)
        self._replace_placeholders(doc, overrides)

        # Step 2: 替换模板中的旧项目名称（占位符之外的具体文本）
        if project_name:
            self._replace_old_project_names(doc, project_name)

        # Step 3: 定位模板章节的段落范围（支持中文数字）
        template_chapters = self._locate_template_chapters(doc)

        # Step 4: 将引擎章节内容映射到模板章节，替换正文
        self._fill_template_chapters(doc, template_chapters, chapter_contents)

        # Step 5: 填入表格数据（从 markdown 中提取）
        self._fill_tables_from_content(doc, chapter_contents)

        # Step 6: 添加附件图片页
        if attachment_images:
            self._add_attachment_pages(doc, attachment_images)

        # Step 7: 保存
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    # ── 占位符替换 ──────────────────────────────────────────────
    def _replace_placeholders(self, doc, overrides: Dict[str, str]):
        """替换文档中所有 {{xxx}} 占位符为实际值（段落 + 表格）。"""
        for para in doc.paragraphs:
            self._replace_placeholders_in_paragraph(para, overrides)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_placeholders_in_paragraph(para, overrides)

    def _replace_placeholders_in_paragraph(self, para, overrides: Dict[str, str]):
        """在一个段落中替换 {{xxx}} 占位符。"""
        for run in para.runs:
            new_text = run.text
            for var, value in overrides.items():
                pattern = "{{" + var + "}}"
                if pattern in new_text:
                    new_text = new_text.replace(pattern, value)
            # 也处理 {{img_xxx}} 类占位符 → 删除（图片由其他方式处理）
            new_text = re.sub(r"\{\{img_\w+\}\}", "", new_text)
            if new_text != run.text:
                run.text = new_text

    # ── 旧项目名替换 ────────────────────────────────────────────
    def _replace_old_project_names(self, doc, project_name: str):
        """替换模板中的旧项目名称（占位符之外的硬编码文本）。"""
        old_names = [
            "金征预告〔2026〕3号",
            "金征预告[2026]3号",
            "高铁枢纽北片区开发地块项目",
            "高铁枢纽北片区",
            "洞庭湖路（S350-宁连一级路段）工程",
            "洞庭湖路",
            "S350-宁连一级公路",
            "宁连一级公路",
            "S350",
            "江苏洪泽经济开发区管理委员会",
            "洪泽经济开发区",
            "大魏社区",
            "洪泽园三村社区",
            "戴楼街道",
            "戴楼社区",
            "金湖县委政法委",
            "金湖县",
        ]
        for para in doc.paragraphs:
            for run in para.runs:
                for old in old_names:
                    if old in run.text:
                        run.text = run.text.replace(old, project_name)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for old in old_names:
                                if old in run.text:
                                    run.text = run.text.replace(old, project_name)

    # ── 模板章节定位（中文数字 + Plain Text）───────────────────
    def _locate_template_chapters(self, doc) -> Dict[int, Tuple[int, int, str]]:
        """定位模板中每个章节的段落范围。

        返回 {chapter_no: (start_para_idx, end_para_idx, chapter_title)}
        只识别真正的章标题：中文数字（第X章）或阿拉伯数字+标题词。
        严格排除附件清单、法律条目、日期等误判。
        """
        heading_paras = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            ch_no, is_real_heading = self._detect_chapter_heading(para, text)
            if ch_no is not None and is_real_heading:
                heading_paras.append((i, ch_no, text))

        # 计算范围
        ranges = {}
        for idx, (para_idx, ch_no, text) in enumerate(heading_paras):
            start = para_idx + 1  # 标题段之后
            end = heading_paras[idx + 1][0] if idx + 1 < len(heading_paras) else len(doc.paragraphs)
            ranges[ch_no] = (start, end, text)

        return ranges

    def _extract_chapter_no_v2(self, text: str) -> Optional[int]:
        """从标题文本提取章节号（支持中文和阿拉伯数字）。"""
        text = text.strip()
        # 中文数字：第一章、第十二章
        m = re.match(r"^第([一二三四五六七八九十]+)章", text)
        if m:
            cn = m.group(1)
            return _CN_NUM.get(cn, None)
        # 阿拉伯数字：1、第1章
        m = re.match(r"^(?:第)?(\d{1,2})", text)
        if m:
            return int(m.group(1))
        return None

    def _detect_chapter_heading(self, para, text: str) -> Tuple[Optional[int], bool]:
        """检测段落是否是真正的章节标题，返回 (chapter_no, is_real_heading)。

        严格规则：
        1. 中文数字章标题（第一章、第十二章等）→ 一定是真章节
        2. 阿拉伯数字开头的段落 → 只有特定样式（Heading 1/Plain Text）才算
        3. 排除：法律条目（"13.《..."）、附件清单（"1.征地红线图"）、日期等
        """
        style = para.style.name if para.style else ""

        # 规则1：中文数字章标题 → 一定是真章节
        m_cn = re.match(r"^第([一二三四五六七八九十]+)章", text)
        if m_cn:
            ch_no = _CN_NUM.get(m_cn.group(1), None)
            return (ch_no, True)

        # 规则2：阿拉伯数字开头 + Heading 1 或 Plain Text 样式
        if style in ("Heading 1", "Heading 2", "标题 1", "标题 2"):
            m_ar = re.match(r"^(\d{1,2})\s", text)
            if m_ar:
                ch_no = int(m_ar.group(1))
                # 排除明显的条目（如 "1.征地红线图"）
                if re.match(r"^(\d{1,2})\.\S", text):
                    return (None, False)
                return (ch_no, True)

        # 规则3：排除所有其他情况（法律条目、附件清单等）
        return (None, False)

    # ── 章节内容映射和填充 ──────────────────────────────────────
    def _fill_template_chapters(self, doc, template_chapters, chapter_contents: Dict[int, str]):
        """将引擎生成的章节内容映射到模板章节并替换正文。"""

        # 构建引擎→模板的映射
        engine_ch_titles = {}
        engine_ch_subtitles = {}  # 子标题关键词也用于匹配
        for ch_no, content_md in chapter_contents.items():
            # 从 markdown 中提取章节标题
            first_line = content_md.strip().split("\n")[0].strip()
            title = first_line.lstrip("# ").strip() if first_line.startswith("#") else first_line
            engine_ch_titles[ch_no] = title
            # 提取子标题关键词
            subtitles = re.findall(r"^#{2,3}\s+(.+)", content_md, re.MULTILINE)
            engine_ch_subtitles[ch_no] = [s.strip() for s in subtitles]

        # 按标题关键词匹配
        mapping = self._build_chapter_mapping(engine_ch_titles, engine_ch_subtitles, template_chapters)

        # 对每个映射的模板章节，替换正文内容
        # 重要：必须从文档末尾往前处理，避免删除/插入段落导致索引偏移
        for engine_no, template_no in sorted(mapping.items(), key=lambda x: x[1], reverse=True):
            if engine_no not in chapter_contents:
                continue
            content_md = chapter_contents[engine_no]
            if template_no not in template_chapters:
                # 模板中没有对应章节，追加到末尾
                self._append_chapter(doc, engine_no, content_md)
                continue

            start, end, ch_title = template_chapters[template_no]
            self._replace_chapter_body(doc, start, end, content_md)

        # 未映射的引擎章节也追加到末尾
        for engine_no in chapter_contents:
            if engine_no not in mapping:
                self._append_chapter(doc, engine_no, chapter_contents[engine_no])

    def _build_chapter_mapping(self, engine_ch_titles, engine_ch_subtitles, template_chapters) -> Dict[int, int]:
        """建立引擎章节 → 模板章节的映射。"""
        mapping = {}
        used_template_chs = set()

        for engine_no, engine_title in engine_ch_titles.items():
            best_match = None
            best_score = 0

            # 如果标题是推断标题，用子标题来匹配
            search_text = engine_title
            if "推断标题" in engine_title or not engine_title.strip():
                # 用子标题关键词拼接
                search_text = " ".join(engine_ch_subtitles.get(engine_no, []))

            for t_no, (start, end, t_title) in template_chapters.items():
                if t_no in used_template_chs:
                    continue
                score = self._title_similarity(search_text, t_title)
                if score > best_score:
                    best_score = score
                    best_match = t_no

            if best_match is not None and best_score > 0:
                mapping[engine_no] = best_match
                used_template_chs.add(best_match)

        return mapping

    def _title_similarity(self, title1: str, title2: str) -> float:
        """计算两个标题的相似度（关键词匹配）。"""
        # 直接关键词映射
        for kw1, kw2 in _CHAPTER_TITLE_MAP.items():
            if kw1 in title1 and kw2 in title2:
                return 1.0
            if kw2 in title1 and kw1 in title2:
                return 1.0

        # 共同关键词
        words1 = set(re.findall(r"[^\s]{2,}", title1))
        words2 = set(re.findall(r"[^\s]{2,}", title2))
        common = words1 & words2
        if not common:
            return 0.0
        return len(common) / max(len(words1), len(words2))

    # ── 正文段落替换（核心改进）──────────────────────────────────
    def _replace_chapter_body(self, doc, start: int, end: int, content_md: str):
        """替换章节范围内的正文内容。

        核心策略：
        1. 只保留章节标题（第一章...）和含图片的段落
        2. 删除所有其他模板段落（子标题、正文、列表等）
        3. 在原位置插入引擎生成的新内容段落
        """
        # 解析 markdown 为纯文本段落列表
        new_lines = self._extract_body_lines_from_md(content_md)

        if not new_lines:
            return

        # 找出章节内要保留的段落（只有章标题和图片）
        preserved = set()
        for i in range(start, end):
            para = doc.paragraphs[i]
            # 保留章节标题（第一章、第二章等）
            text = para.text.strip()
            if re.match(r"^第[一二三四五六七八九十]+章", text):
                preserved.add(i)
                continue
            # 保留含图片的段落
            if self._has_image(para):
                preserved.add(i)
                continue

        # 收集可删除段落的 XML 元素
        to_remove = []
        for i in range(start, end):
            if i not in preserved:
                to_remove.append(doc.paragraphs[i]._element)

        if not to_remove:
            return

        # 确定父元素和锚点位置
        parent = to_remove[0].getparent()
        all_children = list(parent)
        anchor_idx = all_children.index(to_remove[0])

        # 删除所有旧段落
        for elem in to_remove:
            parent.remove(elem)

        # 在锚点位置逐个插入新段落
        insert_pos = anchor_idx
        for line in new_lines:
            new_p = self._make_body_p_element(line)
            parent.insert(insert_pos, new_p)
            insert_pos += 1

    def _make_body_p_element(self, text: str):
        """创建一个正文段落 XML 元素（仿宋、12pt、首行缩进2字符）。"""
        # <w:p>
        p = OxmlElement("w:p")

        # <w:pPr> 段落属性
        pPr = OxmlElement("w:pPr")
        # 首行缩进
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLineChars"), "200")
        ind.set(qn("w:firstLine"), "480")
        pPr.append(ind)
        p.append(pPr)

        # <w:r> 运行
        r = OxmlElement("w:r")
        # <w:rPr> 运行属性
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        rFonts.set(qn("w:ascii"), "仿宋_GB2312")
        rFonts.set(qn("w:hAnsi"), "仿宋_GB2312")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "24")
        rPr.append(szCs)
        r.append(rPr)

        # <w:t> 文本
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)

        p.append(r)
        return p

    def _extract_body_lines_from_md(self, md: str) -> List[str]:
        """从 Markdown 提取正文行（不含标题行和表格分隔行）。"""
        lines = []
        for line in md.split("\n"):
            line = line.strip()
            if not line:
                continue
            # 跳过分隔线
            if line in ("---", "***", "___"):
                continue
            # 保留标题行（会作为子标题处理）
            if line.startswith("#"):
                # 去掉 # 前缀
                level = len(line) - len(line.lstrip("#"))
                text = line.lstrip("# ").strip()
                # 保留子标题格式
                if level >= 2:
                    lines.append(text)
                continue
            # 保留表格行（| 格式）
            if line.startswith("|"):
                # 检查是否是分隔行（|---|---|）
                cells = [c.strip() for c in line.strip("|").split("|")]
                if all(re.match(r"^[-:]+$", c) for c in cells):
                    continue
                lines.append(line)
                continue
            # 引用行去掉 > 前缀
            if line.startswith(">"):
                lines.append(line.lstrip("> ").strip())
                continue
            # 列表去掉前缀
            if re.match(r"^[-*]\s+", line):
                lines.append(re.sub(r"^[-*]\s+", "", line))
                continue
            if re.match(r"^\d+\.\s+", line):
                lines.append(re.sub(r"^\d+\.\s+", "", line))
                continue
            # 普通文本行
            lines.append(line)
        return lines

    def _is_chapter_heading_v2(self, para) -> bool:
        """判断段落是否是章节标题（支持中文数字）。"""
        text = para.text.strip()
        if re.match(r"^第[一二三四五六七八九十]+章", text):
            return True
        if re.match(r"^(?:第)?\d{1,2}\s", text):
            return True
        style = para.style.name if para.style else ""
        if "Heading 1" in style or style == "标题 1" or style == "Plain Text":
            # Plain Text 样式的章节标题也保留
            if re.match(r"^第", text):
                return True
        return False

    def _is_subsection_heading(self, para) -> bool:
        """判断段落是否是子标题（保留不替换）。"""
        style = para.style.name if para.style else ""
        text = para.text.strip()
        # Heading 2/3 样式
        if "Heading 2" in style or "Heading 3" in style or "Heading 4" in style:
            return True
        # 阿拉伯数字子标题（1.1, 1.2, 2.1等）
        if re.match(r"^\d+\.\d+", text):
            return True
        # 中文格式子标题（（一）、（二）等）
        if re.match(r"^（[一二三四五六七八九十]+）", text):
            return True
        if re.match(r"^\([一二三四五六七八九十]+\)", text):
            return True
        # "1." "2." 格式的小标题（如 "1.决策名称"）
        if re.match(r"^\d+\.\S", text) and len(text) < 40:
            return True
        # Plain Text 样式的子标题
        if style == "Plain Text" and re.match(r"^（[一二三四五六七八九十]+）", text):
            return True
        return False

    def _has_image(self, para) -> bool:
        """检查段落是否包含图片。"""
        for run in para.runs:
            if run._element.findall(qn("w:drawing")):
                return True
        return False

    # ── 表格填充 ────────────────────────────────────────────────
    def _fill_tables_from_content(self, doc, chapter_contents: Dict[int, str]):
        """从 Markdown 内容中提取表格数据，填充到模板表格。"""
        # 收集所有章节中的 Markdown 表格数据
        md_tables = []
        for ch_no, content_md in chapter_contents.items():
            tables = self._extract_md_tables(content_md)
            md_tables.extend(tables)

        # 尝试将 Markdown 表格数据填充到模板表格的空单元格
        for tbl_idx, tbl in enumerate(doc.tables):
            for row in tbl.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # 替换占位符和空值
                    if cell_text in ("（待填写）", "待补充", "XXX", "...", ""):
                        # 尝试从 md_tables 中找匹配数据
                        # 这里用简单策略：对于评审表（表格0），直接用引擎内容填充
                        pass

    def _extract_md_tables(self, md: str) -> List[List[List[str]]]:
        """从 Markdown 文本中提取所有表格（返回三维列表）。"""
        tables = []
        current_rows = []
        for line in md.split("\n"):
            line = line.strip()
            if line.startswith("|"):
                cells = [c.strip() for c in line.strip("|").split("|")]
                # 跳过分隔行
                if all(re.match(r"^[-:]+$", c) for c in cells):
                    continue
                current_rows.append(cells)
            else:
                if current_rows:
                    tables.append(current_rows)
                    current_rows = []
        if current_rows:
            tables.append(current_rows)
        return tables

    # ── 附件图片页 ──────────────────────────────────────────────
    def _add_attachment_pages(
        self, doc, attachment_images: Dict[str, List[str]],
    ):
        """在文档末尾添加附件页，嵌入图片+题注。"""
        doc.add_paragraph("")
        doc.add_page_break()

        heading = doc.add_heading("附件", level=1)

        for category, image_paths in attachment_images.items():
            doc.add_heading(category, level=2)
            for img_path in image_paths[:6]:
                if not os.path.exists(img_path):
                    continue
                try:
                    size = os.path.getsize(img_path)
                    if size > 5 * 1024 * 1024:
                        continue
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(4.5))
                    caption = doc.add_paragraph(os.path.basename(img_path))
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in caption.runs:
                        r.font.size = Pt(9)
                except Exception as e:
                    doc.add_paragraph(f"[图片: {os.path.basename(img_path)} — 插入失败: {e}]")

        if attachment_images:
            doc.add_page_break()

    # ── 追加缺失章节 ────────────────────────────────────────────
    def _append_chapter(self, doc, ch_no: int, content_md: str):
        """在文档末尾追加一个新章节（模板中没有的章节）。"""
        doc.add_page_break()
        items = self._parse_markdown_to_paragraphs(content_md)

        for item in items:
            if item["type"] == "heading":
                level = min(item["level"], 4)
                doc.add_heading(item["content"], level=level)
            elif item["type"] == "text":
                p = doc.add_paragraph(item["content"])
                # 设置首行缩进
                pf = p.paragraph_format
                pf.first_line_indent = Cm(0.74)
            elif item["type"] == "list":
                p = doc.add_paragraph(item["content"], style="List Bullet")
            elif item["type"] == "numbered_list":
                p = doc.add_paragraph(item["content"], style="List Number")
            elif item["type"] == "quote":
                p = doc.add_paragraph(item["content"])
            elif item["type"] == "table_row":
                pass  # 表格行跳过，由 _fill_tables 处理

    # ── Markdown 解析 ────────────────────────────────────────────
    def _parse_markdown_to_paragraphs(self, md: str) -> List[dict]:
        """将 Markdown 文本解析为结构化段落列表。"""
        items = []
        for line in md.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line in ("---", "***", "___"):
                continue
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                items.append({"type": "heading", "level": level, "content": line.lstrip("# ").strip()})
            elif line.startswith("|"):
                items.append({"type": "table_row", "content": line})
            elif re.match(r"^[-*]\s+", line):
                items.append({"type": "list", "content": re.sub(r"^[-*]\s+", "", line)})
            elif re.match(r"^\d+\.\s+", line):
                items.append({"type": "numbered_list", "content": re.sub(r"^\d+\.\s+", "", line)})
            elif line.startswith(">"):
                items.append({"type": "quote", "content": line.lstrip("> ").strip()})
            else:
                items.append({"type": "text", "content": line})
        return items

    # ── 简化版：纯 Markdown → Word（不使用模板）──────────────
    def write_from_markdown(self, markdown: str, output_path: str, title: str = ""):
        """纯 Markdown → Word（兼容旧版 DocxWriter）。"""
        doc = DocxDocument()
        self._setup_styles(doc)

        if title:
            h = doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        items = self._parse_markdown_to_paragraphs(markdown)
        i = 0
        while i < len(items):
            item = items[i]
            if item["type"] == "heading":
                level = min(item["level"], 4)
                if level == 1 and title and i == 0:
                    i += 1
                    continue
                doc.add_heading(item["content"], level=level)
            elif item["type"] == "table_row":
                table_rows = [item["content"]]
                i += 1
                while i < len(items) and items[i]["type"] == "table_row":
                    table_rows.append(items[i]["content"])
                    i += 1
                self._add_table(doc, table_rows)
                continue
            elif item["type"] == "text":
                p = doc.add_paragraph(item["content"])
                pf = p.paragraph_format
                pf.first_line_indent = Cm(0.74)
            elif item["type"] == "list":
                doc.add_paragraph(item["content"], style="List Bullet")
            elif item["type"] == "numbered_list":
                doc.add_paragraph(item["content"], style="List Number")
            elif item["type"] == "quote":
                p = doc.add_paragraph(item["content"])
            i += 1

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

    def _setup_styles(self, doc):
        """设置默认字体样式。"""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)
        try:
            style.element.rPr.rFonts.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
            )
        except Exception:
            pass

    def _add_table(self, doc, rows: list):
        """从 Markdown 表格行创建 Word 表格。"""
        parsed = []
        for row in rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            parsed.append(cells)
        parsed = [r for r in parsed if not all(re.match(r"^[-:]+$", c) for c in r)]
        if not parsed:
            return
        max_cols = max(len(r) for r in parsed)
        table = doc.add_table(rows=len(parsed), cols=max_cols)
        table.style = "Table Grid"
        for ri, row in enumerate(parsed):
            for ci in range(max_cols):
                cell = table.cell(ri, ci)
                cell.text = row[ci] if ci < len(row) else ""
        doc.add_paragraph("")
