#!/usr/bin/env python3
"""
模板解析器 — 解析空稳评模板，提取所有空白字段清单

用法:
  python template_parser.py <模板文件路径> [--output fields.json]
  
输出:
  JSON格式的字段清单，包含字段名称、类型（basic/survey/ai/attachment）、
  对应段落索引、当前值、是否为空。
"""
import os
import sys
import json
from typing import List, Dict, Tuple

try:
    from docx import Document
except ImportError:
    print("需要 python-docx: pip install python-docx")
    sys.exit(1)

# ============================================================
# 字段分类规则
# ============================================================

# A类：必填基础信息 — 段落完全为空且在关键位置
BASIC_INFO_SECTIONS = {
    "1.1": ("决策名称", "P182"),
    "1.2": ("稳评责任单位", "P184"),
    "1.3": ("拟征地位置", "P187"),
    "1.4": ("征收范围面积及地上附着物", "P189"),
    "1.5": ("资金筹措", "P191"),
    "1.6": ("决策实施周期", "P194"),
    "评审表-事项名称": ("事项名称", "P152首行"),
    "评审表-稳评责任单位": ("稳评责任单位", "P153首行"),
}

# B类：调研数据 — 日期、数字、百分比相关
SURVEY_SECTIONS = [
    "3.3", "3.4.1", "3.4", "3.5.2", "3.5.3", "3.5.4", "3.6"
]

# C类：AI生成段落 — 综合分析、风险识别、防范措施、结论、应急预案
AI_SECTIONS = [
    "4.2.1", "4.2.2", "4.2.3",
    "4.3.1", "4.3.2", "4.3.3",
    "4.4.1", "4.4.2", "4.4.3", "4.4.4",
    "5.2.1", "5.2.2", "5.2.3", "5.2.4",
    "7.2.1", "7.2.2", "7.2.3", "7.2.4", "7.2.5",
    "9.1", "9.2", "9.3",
    "10.1", "10.6",
]

# 附件标注关键词
ATTACHMENT_KEYWORDS = [
    "拟征收土地公告", "征地范围图", "公示照片",
    "现场照片", "座谈会相关资料", "会议纪要", "调查问卷"
]


def extract_section_number(text: str) -> str:
    """从段落文本提取章节号，如'4.2.1本决策符合...' → '4.2.1'"""
    import re
    m = re.match(r'^(\d+\.\d+(?:\.\d+)?)', text)
    return m.group(1) if m else ""


def analyze_template(docx_path: str) -> Dict:
    """分析模板文件，返回字段清单"""
    doc = Document(docx_path)
    
    result = {
        "template_path": docx_path,
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "fields": {
            "basic_info": [],       # A类
            "survey_data": [],      # B类
            "ai_sections": [],      # C类
            "attachments": [],      # D类
        },
        "paragraph_map": [],  # 每段的信息
        "table_summary": [],
    }
    
    # 1. 分析段落
    section_context = ""
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else "None"
        
        # 追踪当前所在章节
        sec_num = extract_section_number(text)
        if sec_num:
            section_context = sec_num
        
        para_info = {
            "index": idx,
            "style": style,
            "text": text[:200] if text else "",
            "text_preview": text[:100] if text else "",
            "is_empty": len(text) == 0,
            "section": section_context,
        }
        result["paragraph_map"].append(para_info)
        
        # 判断空白段落属于哪类
        if len(text) == 0 and style not in ("toc 1", "toc 2", None, ""):
            # 检查是否属于A类基础信息
            matched_a = False
            for key, (desc, _) in BASIC_INFO_SECTIONS.items():
                if key in section_context or (idx < 200 and "Normal" in style):
                    # 简化判断：前200段中Normal样式的空段落多为基础信息
                    pass
            
            # 检查是否属于B类调研数据
            matched_b = False
            for sec_prefix in SURVEY_SECTIONS:
                if section_context.startswith(sec_prefix):
                    matched_b = True
                    break
            
            # 检查是否属于C类AI生成
            matched_c = False
            for ai_sec in AI_SECTIONS:
                if section_context.startswith(ai_sec):
                    matched_c = True
                    break
            
            if matched_b:
                result["fields"]["survey_data"].append({
                    "paragraph_index": idx,
                    "section": section_context,
                    "type": "survey",
                    "description": f"调研数据 - {section_context}",
                })
            elif matched_c:
                result["fields"]["ai_sections"].append({
                    "paragraph_index": idx,
                    "section": section_context,
                    "type": "ai_generated",
                    "description": f"AI生成 - {section_context}",
                })
            elif "Heading" in style:
                # 标题级别的空段落，不处理
                pass
            else:
                # 其余非标题空段落默认为基础信息或待确认
                result["fields"]["basic_info"].append({
                    "paragraph_index": idx,
                    "section": section_context,
                    "type": "basic_or_ai",
                    "description": f"需确认 - P{idx:03d} ({section_context})",
                })
        
        # 检测附件相关段落
        for kw in ATTACHMENT_KEYWORDS:
            if kw in text and ("附件" in text or "见附件" in text or "清单" in text):
                result["fields"]["attachments"].append({
                    "paragraph_index": idx,
                    "section": section_context,
                    "keyword": kw,
                    "description": f"附件: {text[:80]}",
                })
                break
    
    # 2. 分析表格
    for ti, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns)
        empty_cells = 0
        total_cells = rows * cols
        
        for row in table.rows:
            for cell in row.cells:
                if not cell.text.strip():
                    empty_cells += 1
        
        result["table_summary"].append({
            "table_index": ti,
            "rows": rows,
            "cols": cols,
            "empty_cells": empty_cells,
            "total_cells": total_cells,
            "fill_rate": f"{(total_cells - empty_cells) / total_cells * 100:.1f}%",
            "first_row": [cell.text.strip()[:50] for cell in table.rows[0].cells],
        })
    
    # 3. 生成结构化字段清单（供用户填写）
    result["user_input_form"] = generate_input_form(result["paragraph_map"])
    
    # 统计
    result["summary"] = {
        "basic_info_fields": len(result["fields"]["basic_info"]),
        "survey_data_fields": len(result["fields"]["survey_data"]),
        "ai_sections_count": len(result["fields"]["ai_sections"]),
        "attachment_items": len(result["fields"]["attachments"]),
        "total_empty_basic": sum(1 for p in result["paragraph_map"] 
                                if p["is_empty"] and "Heading" not in p["style"] and "toc" not in p["style"]),
    }
    
    return result


def generate_input_form(paragraph_map: List[Dict]) -> Dict:
    """从段落映射生成用户输入表单结构"""
    form = {
        "basic_info": {
            "decision_name": {"label": "决策名称（完整）", "example": "金征预告〔2026〕3号（高铁枢纽北片区开发地块项目）土地征收决策", "value": ""},
            "bulletin_number": {"label": "公告文号", "example": "金征预告〔2026〕3号", "value": ""},
            "project_name": {"label": "项目名称（简短）", "example": "高铁枢纽北片区开发地块项目", "value": ""},
            "responsibility_unit": {"label": "稳评责任单位", "example": "金湖县戴楼街道办事处", "value": ""},
            "commission_month": {"label": "委托日期（月份）", "example": "2026年3月", "value": ""},
            "location_street": {"label": "拟征地街道", "example": "戴楼街道", "value": ""},
            "location_community": {"label": "拟征地社区/村组", "example": "戴楼社区四组、七组", "value": ""},
            "area_hectares": {"label": "征收面积（公顷）", "example": "32.1793", "value": ""},
            "area_mu": {"label": "征收面积（亩）", "example": "482.69", "value": ""},
            "land_type": {"label": "土地性质及地类", "example": "集体土地，水田、林地、坑塘水面等", "value": ""},
            "land_use": {"label": "征地用途", "example": "高铁枢纽北片区开发", "value": ""},
            "fund_per_mu": {"label": "亩均资金测算（万元）", "example": "18", "value": ""},
            "total_fund": {"label": "资金总额（万元）", "example": "8688", "value": ""},
            "fund_source": {"label": "资金来源", "example": "金湖县人民政府财政统筹", "value": ""},
            "bulletin_date": {"label": "公告日期", "example": "2026年3月23日", "value": ""},
        },
        "survey_data": {
            "survey_start": {"label": "调查开始日期", "example": "2026年4月7日", "value": ""},
            "survey_end": {"label": "调查结束日期", "example": "2026年4月24日", "value": ""},
            "phase1_dates": {"label": "第一阶段时间及内容", "example": "2026.4.7：走访责任单位收集资料", "value": ""},
            "phase2_dates": {"label": "第二阶段时间及内容", "example": "2026.4.8-4.9：实地勘查+座谈+问卷", "value": ""},
            "phase3_dates": {"label": "第三阶段时间及内容", "example": "2026.4.10-4.14：汇总材料+编制报告", "value": ""},
            "phase4_dates": {"label": "第四阶段时间及内容", "example": "2026.4.15-4.28：征求意见+专家评审+提交报告", "value": ""},
            "questionnaires_count": {"label": "发放问卷数", "example": "56", "value": ""},
            "support_rate": {"label": "群众支持率(%)", "example": "82.14", "value": ""},
            "awareness_rate": {"label": "知晓率(%)", "example": "89.29", "value": ""},
            "policy_understanding_rate": {"label": "了解补偿政策比例(%)", "example": "83.93", "value": ""},
            "grassroots_opinion": {"label": "基层组织意见", "example": "支持本决策实施，希望信息公开透明", "value": ""},
            "villager_demands": {"label": "村民主要诉求（多条用；分隔）", "example": "了解补偿标准；社保名单公平公正；资金及时到位", "value": ""},
            "online_opinion": {"label": "网络舆情结果", "example": "未发现负面消息", "value": ""},
            "opposition_rate_1": {"label": "反对率-肯定/强烈(%)", "example": "0", "value": "0"},
            "opposition_rate_2": {"label": "反对率-可能/较强(%)", "example": "17.86", "value": "17.86"},
            "opposition_rate_3": {"label": "反对率-概率小/一般(%)", "example": "82.14", "value": "82.14"},
        }
    }
    return form


def main():
    if len(sys.argv) < 2:
        print("用法: python template_parser.py <模板文件路径> [--output fields.json]")
        sys.exit(1)
    
    template_path = sys.argv[1]
    output_path = None
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]
    
    result = analyze_template(template_path)
    
    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"字段清单已保存到: {output_path}")
    
    # 打印摘要
    print("\n" + "="*60)
    print("模板分析结果")
    print("="*60)
    print(f"总段落数: {result['total_paragraphs']}")
    print(f"总表格数: {result['total_tables']}")
    print(f"\n字段统计:")
    print(f"  A类-基础信息: {result['summary']['basic_info_fields']} 个待填充字段")
    print(f"  B类-调研数据: {result['summary']['survey_data_fields']} 个待填充字段")
    print(f"  C类-AI生成:   {result['summary']['ai_sections_count']} 个段落需生成")
    print(f"  D类-附件标注: {result['summary']['attachment_items']} 项")
    print(f"\n用户需填写:")
    print(f"  基础信息: {len(result['user_input_form']['basic_info'])} 项")
    print(f"  调研数据: {len(result['user_input_form']['survey_data'])} 项")
    print(f"\n表格概况:")
    for t in result["table_summary"]:
        print(f"  表格{t['table_index']+1}: {t['rows']}x{t['cols']}, 填充率={t['fill_rate']}")
    print()
    
    return result


if __name__ == "__main__":
    main()
