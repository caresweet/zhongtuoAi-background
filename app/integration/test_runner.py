#!/usr/bin/env python3
"""
测试主控程序 — 串联所有测试，生成测试报告

用法:
  python test_runner.py <生成的报告.docx> [--reference 参考报告.docx] [--template 空模板.docx]
  
输出:
  测试报告.md + 评分
"""
import os
import sys
import json
import time
from typing import Dict, List, Tuple


def run_test(name: str, test_fn, *args, **kwargs) -> Tuple[bool, Dict]:
    """运行单个测试，返回(通过, 详情)"""
    print(f"\n{'='*50}")
    print(f"[测试] {name}")
    print(f"{'='*50}")
    try:
        start = time.time()
        result = test_fn(*args, **kwargs)
        elapsed = time.time() - start
        passed = result.get("passed", False)
        status = "✓ 通过" if passed else "✗ 未通过"
        print(f"  {status} ({elapsed:.1f}s)")
        if "details" in result and isinstance(result["details"], list) and result["details"]:
            for d in result["details"][:10]:
                print(f"    - {d}")
            if len(result["details"]) > 10:
                print(f"    ... 还有 {len(result['details']) - 10} 条")
        return passed, result
    except Exception as e:
        import traceback
        print(f"  ✗ 异常: {e}")
        traceback.print_exc()
        return False, {"passed": False, "error": str(e)}


# ============================================================
# T1: 字段覆盖率测试
# ============================================================

def test_field_coverage(docx_path: str, template_path: str = None) -> Dict:
    """检查报告中所有必填字段是否已填充"""
    try:
        from docx import Document
    except ImportError:
        return {"passed": False, "error": "需要 python-docx"}
    
    doc = Document(docx_path)
    
    details = []
    empty_count = 0
    total_text_paras = 0
    
    # 需要检查的关键段落索引
    critical_paras = {
        182: "1.1 决策名称",
        184: "1.2 决策单位",
        187: "1.3 拟征地位置",
        189: "1.4 征收范围",
        191: "1.5 资金筹措",
        194: "1.6 实施周期",
        200: "2.1.1 评估过程",
        252: "3.3 风险调查内容",
        263: "3.4.1 第一阶段",
        264: "3.4.1 第二阶段",
        293: "3.5.2 基层组织意见",
        295: "3.5.3 村民意见建议",
        297: "3.5.4 网络舆情",
        319: "4.2.1 经济社会发展",
        321: "4.2.2 投资环境",
        323: "4.2.3 群众利益",
        326: "4.3.1 资金保障",
        328: "4.3.2 政府支持",
        330: "4.3.3 群众认同",
        333: "4.4.1 安全可控",
        335: "4.4.2 宣传公示",
        337: "4.4.3 群体事件",
        339: "4.4.4 治安问题",
        353: "5.2.1 补偿方案风险",
        355: "5.2.2 资金分配风险",
        357: "5.2.3 社保名单风险",
        359: "5.2.4 信访舆论风险",
        393: "7.2.1 加强宣传",
        395: "7.2.2 补偿方案",
        397: "7.2.3 资金分配",
        399: "7.2.4 社会保障",
        401: "7.2.5 信访舆论",
        420: "9.1 评估结论",
        422: "9.2 社会稳定等级",
        424: "9.3 实施建议",
        427: "10.1 编制目的",
        442: "10.6 组织领导",
    }
    
    filled = 0
    unfilled = []
    for pidx, desc in critical_paras.items():
        if pidx < len(doc.paragraphs):
            text = doc.paragraphs[pidx].text.strip()
            if text and len(text) > 5:
                filled += 1
            else:
                unfilled.append(f"{desc} (P{pidx:03d})")
                details.append(f"✗ 空白: {desc} (P{pidx:03d}) - 需填充")
    
    # 检查附件标注
    attachment_check = check_attachment_annotations(doc)
    
    coverage_rate = filled / len(critical_paras) * 100 if critical_paras else 0
    passed = coverage_rate >= 90
    
    return {
        "passed": passed,
        "score": min(100, coverage_rate),
        "coverage_rate": f"{coverage_rate:.1f}%",
        "filled_fields": filled,
        "total_fields": len(critical_paras),
        "unfilled_fields": unfilled,
        "attachment_check": attachment_check,
        "details": details,
        "summary": f"字段覆盖率: {coverage_rate:.1f}% ({filled}/{len(critical_paras)})"
    }


def check_attachment_annotations(doc) -> Dict:
    """检查附件标注"""
    attachment_keywords = ["拟征收土地公告", "征地范围图", "公示照片", "现场照片",
                           "座谈会", "会议纪要", "调查问卷"]
    found = []
    missing = []
    for para in doc.paragraphs:
        text = para.text.strip()
        for kw in attachment_keywords:
            if kw in text and ("附件" in text or "需用户提供" in text or "【" in text):
                if kw not in found:
                    found.append(kw)
    
    for kw in attachment_keywords:
        if kw not in found:
            missing.append(kw)
    
    return {"found": found, "missing": missing, "all_present": len(missing) == 0}


# ============================================================
# T2: 内容质量测试
# ============================================================

def test_content_quality(docx_path: str) -> Dict:
    """内容质量检查"""
    try:
        from docx import Document
    except ImportError:
        return {"passed": False, "error": "需要 python-docx"}
    
    doc = Document(docx_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    
    checks = {}
    details = []
    
    # 四性术语检查
    four_properties = {
        "合法性": "合法性" in full_text,
        "合理性": "合理性" in full_text,
        "可行性": "可行性" in full_text,
        "可控性": "可控性" in full_text,
    }
    checks["四性术语"] = four_properties
    if all(four_properties.values()):
        details.append("✓ 四性术语完整")
    else:
        missing_terms = [k for k, v in four_properties.items() if not v]
        details.append(f"✗ 缺少术语: {missing_terms}")
    
    # 法规引用检查
    key_regulations = [
        "DB32/T4013", "中办发〔2021〕11号", "苏办发〔2021〕15号",
        "土地管理法", "城乡规划法", "苏政规〔2025〕5号", "苏政发〔2021〕87号"
    ]
    reg_found = sum(1 for r in key_regulations if r in full_text)
    checks["法规引用"] = f"{reg_found}/{len(key_regulations)}"
    if reg_found >= 5:
        details.append(f"✓ 法规引用: {reg_found}/{len(key_regulations)}")
    else:
        details.append(f"✗ 法规引用不足: {reg_found}/{len(key_regulations)}")
    
    # 风险等级检查
    has_risk_level = any(tag in full_text for tag in ["低风险", "中风险", "高风险"])
    checks["风险等级判定"] = has_risk_level
    if has_risk_level:
        details.append("✓ 包含风险等级判定")
    else:
        details.append("✗ 缺少风险等级判定")
    
    # 得分检查
    has_score = any(w in full_text for w in ["得分", "分", "W≤", "W≥"])
    checks["量化评分"] = has_score
    if has_score:
        details.append("✓ 包含量化评分")
    else:
        details.append("✗ 缺少量化评分")
    
    # 总长度检查
    text_len = len(full_text)
    checks["总字数"] = text_len
    if 10000 < text_len < 80000:
        details.append(f"✓ 报告字数: {text_len} (合理范围)")
    else:
        details.append(f"⚠ 报告字数: {text_len} (可能偏长偏短)")
    
    # 评分：布尔值和含/的字符串项
    score_items = {k: v for k, v in checks.items() if isinstance(v, bool) or (isinstance(v, str) and "/" in v)}
    total_passed = sum(1 for v in score_items.values() if v is True or (isinstance(v, str) and v.startswith("7")))
    max_checks = len(score_items)
    
    return {
        "passed": total_passed >= max(1, max_checks - 1),
        "score": total_passed / max(1, max_checks) * 100 if max_checks else 0,
        "checks": checks,
        "details": details,
        "summary": f"内容质量: {total_passed}/{max_checks} 项通过"
    }


# ============================================================
# T3: 格式一致性测试
# ============================================================

def test_format_consistency(docx_path: str) -> Dict:
    """格式一致性检查"""
    try:
        from docx import Document
    except ImportError:
        return {"passed": False, "error": "需要 python-docx"}
    
    doc = Document(docx_path)
    details = []
    
    # 检查章节结构
    headings = []
    for p in doc.paragraphs:
        if "Heading" in (p.style.name or ""):
            headings.append({"level": p.style.name, "text": p.text.strip()[:80]})
    
    # 检查关键章节是否存在
    required_chapters = [
        "拟征收决策基本概况",
        "评估过程",
        "风险因素调查",
        "综合分析",
        "风险因素识别",
        "风险等级",
        "风险防范",
        "风险等级评估",
        "评估结论",
        "应急预案",
    ]
    
    heading_texts = " ".join([h["text"] for h in headings])
    chapter_found = []
    chapter_missing = []
    for ch in required_chapters:
        if ch in heading_texts:
            chapter_found.append(ch)
        else:
            chapter_missing.append(ch)
    
    details.append(f"✓ 章节: 找到 {len(chapter_found)}/{len(required_chapters)} 个关键章节")
    if chapter_missing:
        details.append(f"⚠ 可能缺失: {chapter_missing}")
    
    # 表格检查
    table_count = len(doc.tables)
    if table_count >= 5:
        details.append(f"✓ 表格数量: {table_count} (充足)")
    else:
        details.append(f"✗ 表格数量: {table_count} (不足)")
    
    # 段落样式检查
    style_counts = {}
    for p in doc.paragraphs:
        name = p.style.name if p.style else "None"
        style_counts[name] = style_counts.get(name, 0) + 1
    
    has_heading1 = any("Heading 1" in s for s in style_counts)
    has_normal = "Normal" in style_counts
    details.append(f"✓ 样式: Heading1={has_heading1}, Normal={has_normal}")
    
    passed = len(chapter_missing) <= 2 and table_count >= 5
    
    return {
        "passed": passed,
        "score": 85 if passed else 60,
        "headings_count": len(headings),
        "table_count": table_count,
        "chapters_found": len(chapter_found),
        "chapters_required": len(required_chapters),
        "missing_chapters": chapter_missing,
        "details": details,
        "summary": f"格式: {len(chapter_found)}/{len(required_chapters)}章节, {table_count}表格"
    }


# ============================================================
# T4: 与参考报告对比
# ============================================================

def test_comparison(generated_path: str, reference_path: str) -> Dict:
    """与参考报告对比"""
    if not reference_path or not os.path.exists(reference_path):
        return {"passed": True, "skipped": True, "details": ["未提供参考报告，跳过对比"]}
    
    try:
        from docx import Document
    except ImportError:
        return {"passed": False, "error": "需要 python-docx"}
    
    gen_doc = Document(generated_path)
    ref_doc = Document(reference_path)
    
    gen_text = "\n".join([p.text for p in gen_doc.paragraphs])
    ref_text = "\n".join([p.text for p in ref_doc.paragraphs])
    
    details = []
    
    # 长度对比
    gen_len = len(gen_text)
    ref_len = len(ref_text)
    ratio = gen_len / max(1, ref_len) * 100
    details.append(f"生成报告: {gen_len} 字符 vs 参考报告: {ref_len} 字符 ({ratio:.0f}%)")
    
    if 50 < ratio < 200:
        details.append("✓ 长度比例在合理范围内")
    else:
        details.append(f"⚠ 长度偏差较大 ({ratio:.0f}%)")
    
    # 关键术语覆盖率
    key_terms = [
        "社会稳定风险评估", "合法性", "合理性", "可行性", "可控性",
        "低风险", "征收", "补偿", "社会保障", "应急预案",
        "DB32/T4013", "中办发〔2021〕11号", "群体性事件"
    ]
    gen_terms = sum(1 for t in key_terms if t in gen_text)
    details.append(f"关键术语: {gen_terms}/{len(key_terms)} 匹配")
    
    passed = ratio > 40 and gen_terms >= 8
    
    return {
        "passed": passed,
        "score": min(100, gen_terms / len(key_terms) * 100),
        "length_ratio": f"{ratio:.0f}%",
        "term_coverage": f"{gen_terms}/{len(key_terms)}",
        "details": details,
        "summary": f"对比: 长度{ratio:.0f}%, 术语{gen_terms}/{len(key_terms)}"
    }


# ============================================================
# 主测试流程
# ============================================================

def generate_test_report(results: Dict) -> str:
    """生成测试报告Markdown"""
    T1 = results.get("T1", {})
    T2 = results.get("T2", {})
    T3 = results.get("T3", {})
    T4 = results.get("T4", {})
    
    # 加权计算总分
    scores = {
        "T1_字段覆盖率": T1.get("score", 0) * 0.30,
        "T2_内容质量": T2.get("score", 0) * 0.30,
        "T3_格式一致性": T3.get("score", 0) * 0.20,
        "T4_对比一致性": T4.get("score", 0) * 0.20,
    }
    total_score = sum(scores.values())
    
    if total_score >= 90:
        grade = "A"
        grade_desc = "优秀，可直接交付"
    elif total_score >= 75:
        grade = "B"
        grade_desc = "良好，少量修改后交付"
    elif total_score >= 60:
        grade = "C"
        grade_desc = "合格，需人工审核修改"
    else:
        grade = "D"
        grade_desc = "不合格，需重新生成"
    
    report = f"""# 众拓稳评报告质量测试报告

**生成时间**: {time.strftime('%Y-%m-%d %H:%M:%S')}
**总体评分**: {total_score:.1f}/100
**等级**: {grade} — {grade_desc}

---

## 综合评分

| 测试维度 | 权重 | 得分 | 通过 |
|---------|------|------|------|
| T1 字段覆盖率 | 30% | {T1.get('score', 0):.0f}/100 | {'✓' if T1.get('passed') else '✗'} |
| T2 内容质量 | 30% | {T2.get('score', 0):.0f}/100 | {'✓' if T2.get('passed') else '✗'} |
| T3 格式一致性 | 20% | {T3.get('score', 0):.0f}/100 | {'✓' if T3.get('passed') else '✗'} |
| T4 对比一致性 | 20% | {T4.get('score', 0):.0f}/100 | {'✓' if T4.get('passed') else '✗'} |
| **总分** | **100%** | **{total_score:.1f}/100** | |

---

## T1 字段覆盖率

**{T1.get('summary', 'N/A')}**

"""
    if T1.get("unfilled_fields"):
        report += "### 未填充字段\n\n"
        for f in T1["unfilled_fields"]:
            report += f"- {f}\n"
    
    if T1.get("attachment_check"):
        ac = T1["attachment_check"]
        report += f"\n### 附件标注\n"
        report += f"- 已标注: {', '.join(ac.get('found', []))}\n"
        if ac.get("missing"):
            report += f"- 缺失: {', '.join(ac['missing'])}\n"
    
    report += f"""
---

## T2 内容质量

**{T2.get('summary', 'N/A')}**

"""
    for check_name, check_val in T2.get("checks", {}).items():
        report += f"- {check_name}: {check_val}\n"
    
    report += f"""
---

## T3 格式一致性

**{T3.get('summary', 'N/A')}**

- 标题总数: {T3.get('headings_count', 'N/A')}
- 表格数: {T3.get('table_count', 'N/A')}
- 章节覆盖: {T3.get('chapters_found', 'N/A')}/{T3.get('chapters_required', 'N/A')}

"""
    if T3.get("missing_chapters"):
        report += "### 可能缺失的章节\n"
        for ch in T3["missing_chapters"]:
            report += f"- {ch}\n"
    
    report += f"""
---

## T4 与参考报告对比

**{T4.get('summary', 'N/A')}**

"""
    for d in T4.get("details", []):
        report += f"- {d}\n"
    
    report += f"""
---

## 结论

总体评分 {total_score:.1f}/100，等级 **{grade}**：{grade_desc}

"""
    if grade == "A":
        report += "报告质量优秀，可直接交付使用。"
    elif grade == "B":
        report += "建议对未填充字段和内容偏差进行少量修改后交付。"
    elif grade == "C":
        report += "建议人工审核并修改报告中标注的缺失内容后再交付。"
    else:
        report += "报告质量不达标，建议检查用户输入数据完整性并重新生成。"
    
    return report


def main():
    if len(sys.argv) < 2:
        print("用法: python test_runner.py <生成的报告.docx> [--reference 参考.docx] [--template 空模板.docx]")
        sys.exit(1)
    
    generated_path = sys.argv[1]
    reference_path = None
    template_path = None
    
    if "--reference" in sys.argv:
        idx = sys.argv.index("--reference")
        if idx + 1 < len(sys.argv):
            reference_path = sys.argv[idx + 1]
    
    if "--template" in sys.argv:
        idx = sys.argv.index("--template")
        if idx + 1 < len(sys.argv):
            template_path = sys.argv[idx + 1]
    
    if not os.path.exists(generated_path):
        print(f"错误: 文件不存在: {generated_path}")
        sys.exit(1)
    
    print("=" * 70)
    print("  众拓稳评报告质量测试系统 v2.0")
    print("=" * 70)
    print(f"\n测试目标: {generated_path}")
    if reference_path:
        print(f"参考报告: {reference_path}")
    if template_path:
        print(f"空模板:   {template_path}")
    print()
    
    # 运行所有测试
    passed1, t1 = run_test("T1 字段覆盖率", test_field_coverage, generated_path, template_path)
    passed2, t2 = run_test("T2 内容质量", test_content_quality, generated_path)
    passed3, t3 = run_test("T3 格式一致性", test_format_consistency, generated_path)
    passed4, t4 = run_test("T4 参考报告对比", test_comparison, generated_path, reference_path)
    
    results = {"T1": t1, "T2": t2, "T3": t3, "T4": t4}
    
    # 生成测试报告
    report = generate_test_report(results)
    
    # 输出位置
    output_dir = os.path.dirname(generated_path)
    report_path = os.path.join(output_dir, "测试报告.md")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(report)
    
    print(f"\n{'='*70}")
    print(f"测试报告已生成: {report_path}")
    print(f"{'='*70}")
    
    # 打印摘要
    scores = {
        "T1": t1.get("score", 0) * 0.30,
        "T2": t2.get("score", 0) * 0.30,
        "T3": t3.get("score", 0) * 0.20,
        "T4": t4.get("score", 0) * 0.20,
    }
    total = sum(scores.values())
    print(f"\n总评分: {total:.1f}/100")


if __name__ == "__main__":
    main()
