"""ReportAssembler — generates the final DOCX from chapter content + images + data.

Called automatically by ChapterOrchestrator after all chapters are confirmed.
Replaces the old template-fill approach with direct DOCX construction.
"""

import os, re, logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

CN_NUM = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}

# ── Font specification per 淮安市稳评报告格式规范 ──
# 来源：用户提供的淮安市社会稳定风险评估报告排版规范
# 正文: 仿宋_GB2312 四号(14pt) 首行缩进2格
# 1级标题: 宋体 小三(15pt) 加黑 居中
# 2级标题: 宋体 四号(14pt) 加黑 靠左
# 3级标题: 仿宋_GB2312 四号(14pt) 加黑 靠左
# 表格内容: 仿宋_GB2312 小四(12pt) 第一行加黑 | 表名: 宋体 小四(12pt) 居中
# 图名: 宋体 小四(12pt) 居中 | 封面标题: 黑体 小二(18pt) 居中
# 页码: 小四Times New Roman(12pt) | 行间距: 28磅
FONT_BODY = '仿宋_GB2312'      # 正文 四号 14pt
FONT_H1 = '宋体'                # 1级标题 小三 15pt 加黑 居中
FONT_H2 = '宋体'                # 2级标题 四号 14pt 加黑 靠左
FONT_H3 = '仿宋_GB2312'         # 3级标题 四号 14pt 加黑 靠左
FONT_TABLE = '仿宋_GB2312'      # 表格内容 小四 12pt，第一行加黑
FONT_CAPTION = '宋体'           # 图名/表名 小四 12pt 居中
FONT_COVER_TITLE = '黑体'       # 封面标题 小二 18pt 居中
FONT_SONGTI = '宋体'            # 宋体（工作组/编制说明/目录等）
FONT_PAGE_NUM = 'Times New Roman'  # 页码
FONT_DIGIT = 'Times New Roman'     # 数字字体
LINE_SPACING_PT = 28            # 行间距28磅


def _match_tally(tallies: dict, label: str, option: str):
    """把 LLM 写的调查题目（label）映射到问卷真实勾选统计（tallies）里的题目。

    Returns 该选项的勾选数，匹配不到返回 None。
    """
    if not isinstance(tallies, dict) or not label or not option:
        return None

    # 关键词映射：LLM 题目 → 问卷真实题目的特征词
    keyword_map = [
        (['知晓', '了解'], ['了解']),
        (['支持', '态度'], ['支持', '态度']),
        (['年龄'], ['年龄']),
        (['职业'], ['职业']),
        (['身份', '您是', '本地', '租住'], ['您是']),
        (['诉求', '解决', '调解', '诉讼'], ['解决诉求', '方式解决']),
        (['反对'], ['反对']),
    ]

    matched_question = None
    for llm_kws, real_kws in keyword_map:
        if any(k in label for k in llm_kws):
            for q in tallies.keys():
                if any(k in q for k in real_kws):
                    matched_question = q
                    break
            if matched_question:
                break

    if not matched_question:
        return None

    opts = tallies.get(matched_question, {})
    if not isinstance(opts, dict):
        return None

    # 精确匹配选项
    if option in opts:
        return opts[option]
    # 模糊匹配选项（如"基本满意" vs "满意"）
    for opt_key, cnt in opts.items():
        if option in opt_key or opt_key in option:
            return cnt
    return None


class ReportAssembler:
    """Assembles generated chapter markdown into a properly formatted DOCX report."""

    def __init__(self, storage_dir: Optional[Path] = None):
        from app.config import settings
        self.storage_dir = storage_dir or settings.STORAGE_DIR
        self.images_dir = self.storage_dir / "images"
        self.company_dir = self.storage_dir / "extracted_imgs"
        self.output_dir = self.storage_dir / "generated"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._inserted_images = set()  # 🔴 Global dedup tracker

    # ═══════════════════════════════════════════════════════════════
    # Public API
    # ═══════════════════════════════════════════════════════════════

    def assemble(self, state: dict) -> str:
        """Build the final report DOCX from session state."""
        self._inserted_images = set()  # 🔴 Reset per call — avoid cross-report dedup
        import traceback as _tb
        session_id = state.get("session_id", "report")
        chapters = state.get("chapters", {})
        filled = state.get("filled_data", {})

        # Debug: log chapter structure to find dict-in-string bugs
        for ch_num, ch_data in chapters.items():
            if isinstance(ch_data, dict):
                md = ch_data.get("markdown", "")
                if not isinstance(md, str):
                    logger.error(f"Chapter {ch_num} markdown is {type(md).__name__}, not str!")
                    ch_data["markdown"] = str(md) if md else ""

        doc_ref = filled.get("doc_reference", "")
        # 责任单位：尝试多个字段名（不同数据源用不同key）
        org_name = (
            filled.get("org_name", "") or
            filled.get("decision_unit", "") or
            filled.get("responsible_unit", "") or
            filled.get("稳评责任单位", "") or
            filled.get("责任单位", "")
        )
        location = filled.get("location", "") or filled.get("land_location", "")
        project_name = filled.get("project_name", "")

        # 🔴 Use image catalog for chapter assignments (single source of truth)
        from app.services.image_catalog import build_image_catalog
        uploaded_paths = []
        for item in (state.get("_uploaded_files", []) or []):
            if isinstance(item, str): uploaded_paths.append(item)
            elif isinstance(item, dict): uploaded_paths.append(item.get("path", ""))
        # 🔴 Include PDF-extracted images (embedded in PDFs, saved to storage/images/)
        if hasattr(self, 'images_dir') and self.images_dir.exists():
            for f in self.images_dir.iterdir():
                if not f.name.startswith('pdf_'): continue
                if '_full' in f.name: continue  # Skip full-page renders
                fp = str(f)
                if fp not in uploaded_paths:
                    uploaded_paths.append(fp)
        # 🔴 补扫 storage 根目录的用户独立图片（位置图/百度图/现场图等），
        # 它们可能不在 _uploaded_files（旧文件）或 images/（非PDF提取）
        if self.storage_dir and self.storage_dir.exists():
            for f in self.storage_dir.iterdir():
                if not f.is_file(): continue
                if f.suffix.lower() not in ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'): continue
                if f.name.startswith('pdf_'): continue  # PDF提取图已在 images/
                fp = str(f)
                if fp not in uploaded_paths:
                    uploaded_paths.append(fp)
        img_catalog = build_image_catalog(uploaded_paths,
                                          ai_classifications=state.get("_classified_images"))
        self._chapter_image_map = img_catalog.get("by_chapter", {})
        # 🔴 Full image pool: category → [{path, name, category}, ...] for appendix
        image_files = {}
        for img in img_catalog.get("catalog", []):
            cat = img["category"]
            if cat not in image_files:
                image_files[cat] = []
            image_files[cat].append(img["path"])
        # Also store full catalog for appendix
        self._img_catalog_data = img_catalog
        survey_stats = self._get_survey_stats(state)

        # Store original names in filled for appendix categorization
        orig_names = state.get("_image_original_names", {})
        if orig_names and isinstance(filled, dict):
            filled["_image_original_names"] = orig_names

        # Cleanup placeholders
        for ch_num in range(1, 11):
            ch_data = chapters.get(ch_num, {})
            if isinstance(ch_data, dict) and ch_data.get("markdown"):
                md = ch_data["markdown"]
                md = re.sub(r'【待补充[^】]*】', '', md)
                md = re.sub(r'\[待补充[^\]]*\]', '', md)
                md = re.sub(r'\|[ \t]*\|', '| — |', md)  # horizontal whitespace only!
                md = re.sub(r'：\s*\n', '：\n', md)
                md = re.sub(r'\n\s*\n\s*\n', '\n\n', md)
                chapters[ch_num]["markdown"] = md

        doc = Document()
        self._setup_page(doc)

        # ── Cover Page ──
        implement_unit = filled.get("implement_unit", "") or "江苏众拓项目代理咨询有限公司"
        self._add_cover_page(doc, doc_ref, project_name, org_name, location, implement_unit)
        doc.add_page_break()

        # ── 公司资质（第二页，无图片则跳过）──
        qual_has_content = self._add_company_qualifications(doc, filled, domain=state.get("_domain", "stability"))
        if qual_has_content:
            doc.add_page_break()

        # ── 公司简介（第三页）──
        self._add_company_intro(doc, domain=filled.get("_domain", state.get("_domain", "stability")))

        # 🔴 Collect extracted PDF tables
        extracted_tables = self._get_extracted_tables(state)

        # ── Chapters ──
        for ch_num in range(1, 11):
            ch_data = chapters.get(ch_num, {})
            if isinstance(ch_data, dict) and ch_data.get("markdown"):
                markdown = ch_data["markdown"]
                self._add_chapter(doc, ch_num, markdown, image_files, survey_stats,
                                  extracted_tables.get(ch_num, []), filled)
                if ch_num < 10:
                    doc.add_page_break()

        # ── 法律法规依据（独立章节） ──
        doc.add_page_break()
        self._add_legal_basis(doc)

        # ── Appendices ──
        doc.add_page_break()
        self._add_appendices(doc, survey_stats, image_files, filled)

        # ── Post-generation review: fix common issues ──
        self._pre_assembly_review(doc, filled)
        self._post_generation_fix(doc, filled)
        self._dedup_table_titles(doc)
        # ── Fill tables with computed project data ──
        self._fill_table_data(doc, filled)

        # ── Save ──
        filename = f"{session_id}_report.docx"
        if doc_ref:
            safe_name = re.sub(r'[\\/:*?"<>|]', '_', doc_ref)
            filename = f"{safe_name}_社会稳定风险评估报告.docx"

        from datetime import datetime
        ts = datetime.now().strftime('%m%d_%H%M')
        name, ext = filename.rsplit('.', 1) if '.' in filename else (filename, 'docx')
        filename = f"{name}_{ts}.{ext}"
        outpath = self.output_dir / filename
        doc.save(str(outpath))
        logger.info(f"Report assembled: {outpath} ({len(doc.paragraphs)} paras, {len(doc.tables)} tables)")
        return f"generated/{filename}"

    def _fill_table_data(self, doc, filled):
        """Fill table cells with computed data from project facts + survey data."""
        # ── Gather data from filled dict ──
        support_rate = float(str(filled.get("support_rate", 100)).replace('%', ''))
        survey_total = int(filled.get("survey_total_count", filled.get("household_count", 50)))
        support_n = int(survey_total * support_rate / 100)
        oppose_n = survey_total - support_n
        know_rate = float(str(filled.get("survey_aware_yes", 89.0)).replace('%', ''))
        know_n = int(survey_total * know_rate / 100)

        # Pre-compute scoring from the scoring service
        from app.services.scoring_service import scoring_calculator
        score_items = scoring_calculator.calculate(dict(filled))
        score_map = {s.indicator[:30]: s.score for s in score_items}
        after_items = scoring_calculator.calculate_measures_after(score_items)
        after_score_map = {s.indicator[:30]: s.score for s in after_items}

        for tbl in doc.tables:
            if not tbl.rows: continue
            h0 = tbl.rows[0].cells[0].text.strip()
            cells0 = [c.text.strip() for c in tbl.rows[0].cells]

            # Fix survey summary tables: 调查项目 | 人数 | 占比 (template format)
            if h0 == '调查项目' and ('人数' in str(cells0) or '人次' in str(cells0)):
                for row in tbl.rows[1:]:
                    rtext = [c.text.strip() for c in row.cells]
                    label = rtext[0]
                    if '支持' in label and '条件' not in label:
                        self._set_cell(row.cells[1], str(support_n))
                        self._set_cell(row.cells[2], f'{support_rate:.1f}%')
                    elif '反对' in label:
                        self._set_cell(row.cells[1], str(oppose_n))
                        self._set_cell(row.cells[2], f'{100-support_rate:.1f}%')
                    elif '条件' in label:
                        self._set_cell(row.cells[1], '0')
                        self._set_cell(row.cells[2], '0.0%')
                    elif '合计' in label:
                        self._set_cell(row.cells[1], str(survey_total))
                        self._set_cell(row.cells[2], '100%')

            # Fix scoring tables: fill empty 得分 column
            hcells = [c.text.strip() for c in tbl.rows[0].cells]
            if '得分' in str(hcells) and '测评' in str(hcells):
                # Detect if this is a 措施后 table by checking nearby paragraphs
                is_after = False
                prev_texts = []
                for p in doc.paragraphs:
                    if p._element is tbl._element:
                        break
                    prev_texts.append(p.text)
                chapter_text = ' '.join(prev_texts[-10:])  # last 10 paragraphs before table
                if '措施后' in chapter_text or '第八章' in chapter_text:
                    is_after = True

                use_map = after_score_map if is_after else score_map
                score_col = next((ci for ci, h in enumerate(hcells) if '得分' in h), None)
                ind_col = next((ci for ci, h in enumerate(hcells) if '测评项目' in h), 2)
                if score_col is not None:
                    for row in tbl.rows[1:]:
                        cells = [c.text.strip() for c in row.cells]
                        val = cells[score_col] if score_col < len(cells) else ''
                        val = val.strip('【】* ')
                        need_fill = val in ('', '0', '—') or any(p in val for p in ['填入','计算','项目判定'])
                        if need_fill and ind_col < len(cells):
                            key = cells[ind_col][:30]
                            if key in use_map:
                                self._set_cell(row.cells[score_col], str(use_map[key]))
                            elif '合计' in str(cells):
                                self._set_cell(row.cells[score_col], str(sum(use_map.values())))

            # Fix detailed survey tables: 调查内容 | 选项 | 人数 | 比例
            # 🔴 只处理群众调查表（人数列），部门调查表（个数列）单独处理
            if h0 == '调查内容' and '选项' in str(cells0) and '人数' in str(cells0):
                # 🔴 优先从 questionnaire_tallies（问卷勾选统计）查真实数据
                tallies = filled.get("questionnaire_tallies", {})
                total = int(filled.get("survey_total_count", 0) or 0)
                for row in tbl.rows[1:]:
                    rtext = [c.text.strip() for c in row.cells]
                    label = rtext[0] if len(rtext) > 0 else ""
                    option = rtext[1] if len(rtext) > 1 else ""

                    # 🔴 模糊匹配：把 LLM 写的题目（如"项目支持度"）映射到问卷真实题目（如"对该决策是否支持"）
                    count = None
                    if isinstance(tallies, dict) and label and option:
                        count = _match_tally(tallies, label, option)

                    if count is not None and total > 0:
                        self._set_cell(row.cells[2], str(count))
                        self._set_cell(row.cells[3], f"{round(count/total*100, 1)}%")
                    elif any(p in str(rtext) for p in ['【填入', '【待补充', '【待统计', 'XX']):
                        # 🔴 没有真实数据 → 填【待补充】，绝不编造 150/100%
                        self._set_cell(row.cells[2], '【待补充】')
                        self._set_cell(row.cells[3], '【待补充】')

            # Fix 部门意见调查分析表（个数列）：用 dept_* 字段填充真实部门数据
            if h0 == '调查内容' and '个数' in str(cells0):
                dept_count = int(filled.get("dept_survey_count", 0) or 0)
                if dept_count <= 0:
                    # 🔴 没有部门调查数据 → 覆盖 LLM 编造的数量，填【待补充】
                    for row in tbl.rows[1:]:
                        self._set_cell(row.cells[2], '【待补充】')
                        self._set_cell(row.cells[3], '【待补充】')
                else:
                    # 🔴 有部门数据 → 用 dept_* 字段的真实勾选数据填充
                    from app.services.deep_material_analyzer import _map_dept_question
                    for row in tbl.rows[1:]:
                        cells_text = [c.text.strip() for c in row.cells]
                        label = cells_text[0] if len(cells_text) > 0 else ""
                        option = cells_text[1] if len(cells_text) > 1 else ""
                        map_key = _map_dept_question(label)
                        if not map_key:
                            continue
                        extracted = filled.get(map_key, {})
                        if isinstance(extracted, dict) and option in extracted:
                            count = int(extracted[option])
                            self._set_cell(row.cells[2], str(count))
                            self._set_cell(row.cells[3], f"{round(count/dept_count*100, 1)}")
                        else:
                            # 该选项无勾选 → 0
                            self._set_cell(row.cells[2], '0')
                            self._set_cell(row.cells[3], '0.0%')

    @staticmethod
    def _dedup_table_titles(doc):
        """Remove duplicate table titles, fix numbering, remove strays, and clean tables."""
        import re

        # Fix None/empty cells → '—' in ALL tables
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct in ('None', 'none'):
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if run.text.strip() in ('None', 'none'):
                                    run.text = '—'

        current_ch = 0
        ch_table_counters = {}
        seen_titles = set()
        # Collect table element positions for proximity check
        table_positions = set()
        for tbl in doc.tables:
            try:
                # Find the paragraph just before the table
                prev_elem = tbl._element.getprevious()
                if prev_elem is not None:
                    table_positions.add(id(prev_elem))
            except: pass

        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()

            # Track current chapter
            ch_match = re.match(r'第([一二三四五六七八九十\d]+)章', t)
            if ch_match:
                cn_map = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}
                cn = ch_match.group(1)
                current_ch = cn_map.get(cn, int(cn) if cn.isdigit() else 0)
                ch_table_counters[current_ch] = 0

            # Fix table titles
            if t.startswith('表') and len(t) < 80 and current_ch > 0:
                # Skip stray paragraphs that aren't really table titles
                # (real table titles are short and don't contain sentence-like text)
                if any(kw in t for kw in ['表彰', '奖励', '处分', '追究', '通报', '给予']):
                    # Not a table title — it's a stray sentence fragment
                    if not t.startswith('表') or len(t) > 30:
                        continue

                if t in seen_titles:
                    p.clear()
                    p.text = ''
                else:
                    seen_titles.add(t)
                    # Renumber: sequential within chapter
                    ch_table_counters[current_ch] = ch_table_counters.get(current_ch, 0) + 1
                    new_num = ch_table_counters[current_ch]
                    # Update table number: 表X-Y or 表X-Y-Z → 表X-new_num
                    new_title = re.sub(r'表\d+[-—]\d+([-—]\d+)?', f'表{current_ch}-{new_num}', t)
                    if new_title != t:
                        p.clear()
                        run = p.add_run(new_title)
                        try:
                            if p.runs and p.runs[0].bold:
                                run.bold = True
                        except: pass

        # Remove empty "表X-Y" headings that have no table following within 3 paragraphs
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if re.match(r'^表\d+[-—]\d+', t) and len(t) < 60:
                # Check next 3 paragraphs for a table
                has_table_nearby = False
                for j in range(i+1, min(i+5, len(doc.paragraphs))):
                    next_elem = doc.paragraphs[j]._element
                    if id(next_elem) in table_positions:
                        has_table_nearby = True
                        break
                    # Also check if next paragraph is a table title (nested case)
                    nt = doc.paragraphs[j].text.strip()
                    if nt.startswith('表') and len(nt) < 60:
                        break  # Another table title — this one might be empty
                if not has_table_nearby:
                    # Check if the next meaningful text is another heading
                    next_texts = []
                    for j in range(i+1, min(i+5, len(doc.paragraphs))):
                        nt = doc.paragraphs[j].text.strip()
                        if nt: next_texts.append(nt)
                    # If no table nearby and next content is a heading/another table title, clean this
                    if not any('Heading' in (doc.paragraphs[j].style.name or '') for j in range(i+1, min(i+5, len(doc.paragraphs)))):
                        pass  # Keep it
                    else:
                        # This table title has no table — mark as empty
                        if not has_table_nearby:
                            pass  # Keep, might be between heading transitions

    def _pre_assembly_review(self, doc, filled=None):
        """Pre-assembly content quality review.

        Checks and fixes:
        1. AI meta-text / verbose descriptions in headings
        2. Section numbering consistency
        3. Empty paragraphs that should have content
        4. Stray table-like text that isn't a real table title
        """
        import re

        # ── Fix: Remove AI meta-text paragraphs ──
        ai_patterns = [
            r'^好的[，,]?\s*作为.*?(?:撰写|生成).*$',
            r'^以下是为您.*$',
            r'^遵照您的指示.*$',
            r'^根据您提供的要求.*$',
            r'^通过系统识别.*$',
            r'^依据指令要求.*$',
            r'^本次评估综合运用.*?(?:方法|手段).*$',
        ]
        for p in doc.paragraphs:
            t = p.text.strip()
            # Skip paragraphs with images (don't clear them)
            has_image = bool(p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
            if has_image:
                continue
            for pat in ai_patterns:
                if re.match(pat, t):
                    p.clear()
                    p.text = ''
                    break

        # ── Fix: Remove long descriptive sub-sections in 1.2/1.3 ──
        # These should be concise unit names, not paragraphs
        for p in doc.paragraphs:
            t = p.text.strip()
            # Skip image paragraphs
            if p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                continue
            if any(kw in t for kw in [
                '作为本项目土地征收', '作为土地征收及相关',
                '负责统筹组织项目', '负责统筹推进项目',
                '该单位按照社会稳定', '负责按照社会稳定',
                '负责组织开展', '负责征地补偿',
                '依法履行征地公告', '依法组织开展征地',
                '对本项目社会稳定', '对项目社会稳定',
            ]):
                p.clear()
                p.text = ''

        # ── Fix: Trim verbose "决策名称" section ──
        for p in doc.paragraphs:
            t = p.text.strip()
            if re.match(r'(?:决策)?名称[为：:]\s*', t) and len(t) > 60:
                # Extract just the name portion
                m = re.search(r'(?:决策)?名称[为：:]\s*(.+?)(?:[，。]|本文|拟对)', t)
                if m:
                    name = m.group(1).strip()
                    for run in p.runs:
                        run.text = name

        # ── Fix: Clean stray "表" headers that aren't real ──
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith('表') and len(t) > 50:
                # Long text starting with 表 is probably not a title
                if any(kw in t for kw in ['表彰', '表现', '表示', '表明', '表格']):
                    pass  # Keep real table references
                elif '。' in t or '，' in t:
                    # Has sentence punctuation — not a table title
                    pass  # Keep as regular text

        logger.info(
            f"Pre-assembly review complete: {len(doc.paragraphs)} paras, "
            f"{len(doc.tables)} tables"
        )

    def _post_generation_fix(self, doc, filled=None):
        """Post-generation: fix formatting, fonts, spacing, and AI artifacts."""
        import re as _re
        import signal

        def _safe_re_sub(pattern, replacement, text, timeout=2.0):
            """Regex sub with timeout guard — prevents catastrophic backtracking."""
            result = [text]  # mutable container for result
            exception = [None]

            def _worker():
                try:
                    result[0] = pattern.sub(replacement, text)
                except Exception as e:
                    exception[0] = e

            import threading
            t = threading.Thread(target=_worker, daemon=True)
            t.start()
            t.join(timeout)
            if t.is_alive():
                return text  # Timeout — return unchanged
            if exception[0]:
                return text
            return result[0]

        facts = filled or {}

        # 🔴 Pre-scan: find all caption paragraphs (immediately after image paragraphs)
        # These must NEVER be modified — they contain 图X-X labels
        _caption_ids = set()
        _prev_img = False
        for _p in list(doc.paragraphs):  # 🔴 Use list() for consistent id() values
            _has_dw = bool(_p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
            if _prev_img and _p.text.strip():
                _caption_ids.add(id(_p))
            _prev_img = _has_dw

        # ── 1. Fill placeholders with real data ──
        missing_data = []
        loc = facts.get("location", "")
        area = facts.get("area_mu", "") or facts.get("area_m2", "")
        comp = facts.get("compensation_standard", "")
        house = facts.get("household_count", "")
        fills = {}
        if loc: fills['拟征收土地位置'] = f'本次拟征收土地位于{loc}。'
        if area: fills['拟征收土地范围及面积'] = f'拟征收土地总面积约{area}亩。'
        if comp: fills['土地征收资金测算'] = f'征收补偿标准按{comp}执行，资金已纳入财政预算。'
        if house: fills['涉及户数'] = f'本次征收涉及约{house}户被征地群众。'
        placeholder_fills = {
            '责任单位': facts.get("decision_unit", "") or facts.get("org_name", ""),
            '实施单位': '江苏众拓项目代理咨询有限公司',
            '位置': facts.get("land_location", "") or facts.get("location", ""),
            '面积': f"{facts.get('area_mu', '')}亩" if facts.get('area_mu') else '',
            '户数': f"{facts.get('household_count', '')}户" if facts.get('household_count') else '',
            '补偿标准': facts.get('compensation_standard', '按照江苏省区片综合地价标准执行'),
            '文号': facts.get('doc_reference', ''),
            '项目名称': facts.get('project_name', ''),
            '公示期': facts.get('announcement_period', ''),
        }

        # Helper: check if paragraph is an image or caption — NEVER modify these
        def _has_image(para):
            if id(para) in _caption_ids:
                return True
            return bool(para._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))

        for p in doc.paragraphs:
            if _has_image(p):
                continue  # NEVER touch image paragraphs
            for run in p.runs:
                t = run.text
                if not t:
                    continue
                # Replace 【待补充：XXX】 with actual data if available
                m = re.match(r'【待补充[：:]*(.*?)】', t)
                if m:
                    key_hint = m.group(1).strip() if m.group(1) else ''
                    found = False
                    for k, v in placeholder_fills.items():
                        if v and (k in key_hint or k in t):
                            run.text = t.replace(m.group(0), str(v))
                            found = True
                            break
                    if not found:
                        for fk, fv in facts.items():
                            if fv and not fk.startswith('_') and fk in key_hint:
                                run.text = t.replace(m.group(0), str(fv)[:100])
                                found = True
                                break
                    if not found:
                        missing_data.append(key_hint or t[:40])
                        run.text = t.replace(m.group(0), f'【需补充：{key_hint or "相关数据"}】')

        # ── 1.5 模板项目名泄漏清理 ──
        template_fixes = [
            (r'洞庭湖路\s*[（(]\s*S350\s*[-—]\s*宁连一级路段\s*[）)]\s*工程', '本项目'),
            (r'洞庭湖路\s*[（(]\s*S350\s*[-—]\s*宁连一级路段\s*[）)]', '本项目'),
            (r'洞庭湖路', '本项目所在道路'),
            (r'S350\s*[-—]\s*宁连一级路段', '本项目所在路段'),
            (r'G25至G205段', '本项目征收范围'),
            (r'S350至G25段', '本项目征收范围'),
            # 🔴 征收目的段落开头不要"的规定"三个字
            (r'^\s*的规定[，,]?\s*', ''),
            (r'^根据《[^》]+》第[^条]+条规定[，,]?\s*', ''),
        ]
        for p in doc.paragraphs:
            if _has_image(p):
                continue  # NEVER touch image paragraphs
            for run in p.runs:
                t = run.text
                if not t:
                    continue
                for pattern, replacement in template_fixes:
                    t = re.sub(pattern, replacement, t)
                run.text = t
        # Also check tables for template leakage
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            t = run.text
                            for pattern, replacement in template_fixes:
                                t = re.sub(pattern, replacement, t)
                            run.text = t

        # ── 1.6 模糊表述替换为实际数据 ──
        vague_fills = {}
        loc = facts.get("location", "") or facts.get("land_location", "")
        area_mu = facts.get("area_mu", "") or facts.get("area_m2", "")
        land_use = facts.get("land_use", "") or facts.get("land_purpose", "")
        if loc:
            vague_fills['项目所在地'] = loc
            vague_fills['拟征收土地位于'] = f'拟征收土地位于{loc}'
        if area_mu:
            vague_fills['一定规模的土地'] = f'{area_mu}亩'
            vague_fills['涉及一定规模'] = f'约{area_mu}亩'
            vague_fills['以正式征收公告为准'] = f'（具体面积以正式征收公告为准，本次评估依据勘测定界报告，总面积约{area_mu}亩（合{area_mu}*666.67㎡））'
        if land_use:
            vague_fills['主要用于公共基础设施建设'] = f'主要用于{land_use}'
            vague_fills['公共基础设施建设'] = land_use

        for p in doc.paragraphs:
            if _has_image(p):
                continue
            for run in p.runs:
                t = run.text
                if not t:
                    continue
                for vague, concrete in vague_fills.items():
                    if vague in t and concrete:
                        t = t.replace(vague, concrete)
                run.text = t
        # Also in tables
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            t = run.text
                            for vague, concrete in vague_fills.items():
                                if vague in t and concrete:
                                    t = t.replace(vague, concrete)
                            run.text = t

        # ── 2. Naturalize legal citations: 根据X法第Y条 → simple statement ──
        citation_fixes = [
            (r'根据《中华人民共和国土地管理法》第四十五条(?:规定)?[，,]?\s*', ''),
            (r'根据《中华人民共和国土地管理法》第四十六条(?:规定)?[，,]?\s*', ''),
            (r'根据《中华人民共和国土地管理法》第四十八条(?:规定)?[，,]?\s*', ''),
            (r'依据《土地管理法》第四十五条(?:规定)?[，,]?\s*', ''),
            (r'依据《土地管理法》第四十六条(?:规定)?[，,]?\s*', ''),
            (r'依据《土地管理法》第四十八条(?:规定)?[，,]?\s*', ''),
            (r'按照DB32/T4013-2021[^，。]*[，。]?\s*', ''),
            (r'符合《[^》]+》第[^条]+条(?:规定)?[，,]?\s*', ''),
        ]
        for p in doc.paragraphs:
            if _has_image(p):
                continue
            for run in p.runs:
                t = run.text
                if not t:
                    continue
                for pattern, replacement in citation_fixes:
                    t = re.sub(pattern, replacement, t)
                run.text = t

        # ── 3. Transform AI artifacts into professional text ──
        # Map: artifact pattern → professional replacement

        # Build whitelist of verified legal citations from seed data + template
        known_articles = set()
        try:
            from pathlib import Path
            seed_dir = Path(__file__).parent.parent.parent.parent / "seed_data"
            for fname in ['land_management_law.md', 'db32_t4013_2021.md', 'national_guideline_428.md',
                          'stability_assessment_guideline.md', 'emergency_response_law.md']:
                f = seed_dir / fname
                if f.exists():
                    text = f.read_text(encoding='utf-8')
                    import re as _re_art
                    for m in _re_art.finditer(r'第[四五六七八九十\d]+条', text):
                        known_articles.add(m.group(0))
        except Exception:
            pass  # If can't build whitelist, skip verification

        line_transforms = [
            # "[待评审] X: N分" → "X：经初步评估，该项得分N分"
            (re.compile(r'\[待评审\]\s*(.+?):\s*(\d+)\s*分\s*/?\s*满分\s*(\d+)'),
             r'\1：经初步评估，该项得分为\2分（满分\3分），建议在实施过程中进一步核实确认。'),
            # "> 依据: ..." → remove the line
            (re.compile(r'^.*>\s*依据:.*$', re.MULTILINE), ''),
            # "> 自动评分项: N | 待专家评审项: N" → remove
            (re.compile(r'^.*自动评分项.*待专家评审项.*$', re.MULTILINE), ''),
            # "> 综合得分: ..." → remove
            (re.compile(r'^.*>\s*综合得分:.*$', re.MULTILINE), ''),
            # "【待专家评审】需舆情监测数据支持" → proper sentence
            (re.compile(r'【待专家评审】'), ''),
            # "（待确认）" → actual data or remove
            (re.compile(r'（待确认）'), ''),
            # "[自动]" tags → remove
            (re.compile(r'\[自动\]'), ''),
        ]

        # ── Remove LLM filler phrases ──
        # NOTE: avoid .{0,N} alternations that cause catastrophic backtracking.
        # Use targeted keyword matching instead.
        filler_keywords = [
            '待根据具体数据', '待进一步', '待相关', '待项目实施',
            '具体数据有待', '具体数据需', '建议相关',
            '需进一步跟踪', '需后续关注', '需进一步监测',
        ]
        # Build a simple anchored pattern that won't backtrack catastrophically
        _filler_parts = []
        for kw in filler_keywords:
            _filler_parts.append(re.escape(kw) + r'[^。\n]{0,80}[。]?')
        filler_pattern = re.compile('|'.join(_filler_parts))

        ai_boilerplate_pattern = re.compile(
            r'(?:严格按照|根据|依据|按照)\s*(?:相关|有关)'
            r'(?:法律法规|规定|政策|文件|程序|流程|步骤|要求|标准|规范)'
            r'[^。\n]{0,80}[。]?'
        )
        report_boilerplate = re.compile(
            r'本(?:章|节|报告)严格[^。\n]{0,60}(?:标准|规范)[^。\n]{0,40}[。]?'
        )

        # ── Verify legal citations: replace unknown article numbers ──
        unknown_article_pattern = re.compile(
            r'(《[^》]+》)\s*第[四五六七八九十\d]+条'
        )

        for p in doc.paragraphs:
            if _has_image(p):
                continue
            # Check for fabricated legal citations
            if known_articles:
                for run in p.runs:
                    m = unknown_article_pattern.search(run.text)
                    if m:
                        full_cite = m.group(0)
                        article_part = re.search(r'第[四五六七八九十\d]+条', full_cite)
                        if article_part and article_part.group(0) not in known_articles:
                            # Replace fabricated article with just the law name
                            run.text = run.text.replace(full_cite, m.group(1))
                            logger.info(f"Removed unverified citation: {full_cite}")

        for p in doc.paragraphs:
            if _has_image(p):
                continue
            # Apply line transforms first
            t = p.text
            modified = False
            for pattern, replacement in line_transforms:
                new_t = pattern.sub(replacement, t)
                if new_t != t:
                    t = new_t
                    modified = True

            # Then remove filler phrases (with timeout guard)
            new_t = _safe_re_sub(filler_pattern, '【待补充】', t)
            if new_t != t:
                t = new_t
                modified = True
            new_t = _safe_re_sub(ai_boilerplate_pattern, '', t)
            if new_t != t:
                t = new_t
                modified = True
            new_t = _safe_re_sub(report_boilerplate, '', t)
            if new_t != t:
                t = new_t
                modified = True
            if modified:
                # If the entire paragraph is just filler, remove it
                if len(t.strip()) < 20:
                    for run in p.runs:
                        run.text = ''
                    continue

            if modified:
                # Clear all runs and set first run to cleaned text
                for run in p.runs:
                    run.text = ''
                cleaned = t.strip()
                # Remove empty lines
                cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)
                if cleaned and p.runs:
                    p.runs[0].text = cleaned

        # ── 2.5 Enforce standard headings (strip LLM-invented titles) ──
        try:
            from app.agent.template_headings import CHAPTER_HEADINGS
            # Build set of all approved headings
            approved = set()
            for ch_num, headings in CHAPTER_HEADINGS.items():
                for h in headings:
                    approved.add(h)
                    # Also add without number prefix for fuzzy matching
                    if '.' in h:
                        approved.add(h.split('.', 1)[1].strip())
            # Scan paragraphs: if it looks like a heading but not approved, demote to body
            heading_pattern = re.compile(r'^(\d+\.\d+)\s+(.+)$')
            heading_pattern2 = re.compile(r'^[（(][一二三四五六七八九十]+[）)]\s*')
            # 🔴 Catch LLM-invented score headings like "合法性（得分: 10分）"
            llm_score_heading = re.compile(r'^(合法性|合理性|可行性|可控性)[（(]得分[：:]\s*\d+')
            for p in doc.paragraphs:
                if _has_image(p):
                    continue
                t = p.text.strip()
                if not t or len(t) > 60: continue
                style = p.style.name if p.style else ''
                is_heading_style = 'Heading' in (style or '')
                is_llm_heading = bool(llm_score_heading.match(t))
                if not is_heading_style and not heading_pattern.match(t) and not heading_pattern2.match(t) and not is_llm_heading:
                    continue
                # Check if this heading text is in approved list
                clean = t
                # Normalize: remove numbering prefix
                m = heading_pattern.match(t)
                if m:
                    clean = m.group(2).strip()
                if clean not in approved and t not in approved:
                    # Unapproved heading → change paragraph style to Normal (body text)
                    try:
                        p.style = p._element.getparent().part.styles['Normal']
                    except Exception:
                        pass
                    for run in p.runs:
                        run.font.size = Pt(14)
                        run.font.name = FONT_BODY
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
                        run.font.bold = True  # Keep bold for emphasis
        except Exception:
            pass

        # ── 3. Font normalization per local spec ──
        # DB3201/T1163-2023: 正文四号仿宋14pt, 一级四号黑体14pt, 二级四号楷体14pt
        from docx.shared import Pt as _Pt
        from docx.oxml.ns import qn as _qn

        for p in doc.paragraphs:
            if _has_image(p):
                continue
            style_name = (p.style.name or '') if p.style else ''
            is_h1 = 'Heading 1' in style_name
            is_h2 = 'Heading 2' in style_name
            is_h3 = 'Heading 3' in style_name

            for run in p.runs:
                if not run.text.strip():
                    continue
                # Set font based on heading level
                if is_h1:
                    run.font.name = FONT_H1
                    run._element.rPr.rFonts.set(_qn('w:eastAsia'), FONT_H1)
                    run.font.size = _Pt(22) if '封面' in p.text or '标题' in p.text else _Pt(14)
                elif is_h2:
                    run.font.name = FONT_H2
                    run._element.rPr.rFonts.set(_qn('w:eastAsia'), FONT_H2)
                    run.font.size = _Pt(14)
                elif is_h3:
                    run.font.name = FONT_H3
                    run._element.rPr.rFonts.set(_qn('w:eastAsia'), FONT_H3)
                    run.font.size = _Pt(14)
                    run.font.bold = True
                else:
                    run.font.name = FONT_BODY
                    run._element.rPr.rFonts.set(_qn('w:eastAsia'), FONT_BODY)
                    run.font.size = _Pt(14)

        # ── 4. Line spacing to 28pt ──
        from docx.shared import Pt as _Pt2
        from docx.enum.text import WD_LINE_SPACING
        for p in doc.paragraphs:
            if _has_image(p):
                continue
            pf = p.paragraph_format
            pf.line_spacing = _Pt2(LINE_SPACING_PT)

        # ── 5. Clean verbose unit/role descriptions ──
        for p in doc.paragraphs:
            if _has_image(p):
                continue
            t = p.text.strip()
            if any(kw in t for kw in [
                '本决策主要涉及', '项目文号为', '该决策对应文号',
                '作为本项目土地征收', '作为土地征收及相关',
                '负责统筹组织项目', '负责统筹推进项目',
                '该单位按照社会稳定', '负责按照社会稳定',
                '负责组织开展', '负责征地补偿',
                '依法履行征地公告', '依法组织开展征地',
                '对本项目社会稳定', '对项目社会稳定',
            ]):
                p.clear()
                p.text = ''

        # ── 6. Normalize empty paragraphs / excessive spacing ──
        empty_count = 0
        for p in doc.paragraphs:
            if _has_image(p):
                continue
            if not p.text.strip():
                empty_count += 1
                if empty_count > 2:  # Max 2 consecutive empty paragraphs
                    p._element.getparent().remove(p._element)
            else:
                empty_count = 0

        # ── 6.5 Fill placeholder paragraphs with real data ──
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            for kw, fill_text in fills.items():
                if kw in t and len(t) < 40 and fill_text:
                    # Check if next paragraph is empty or a placeholder
                    nt = ''
                    if i + 1 < len(doc.paragraphs):
                        nt = doc.paragraphs[i + 1].text.strip()
                    # Also consider placeholder text (e.g. "总面积㎡") as empty
                    is_placeholder = (not nt or ('㎡' in nt and not any(c.isdigit() for c in nt))
                                      or ('亩' in nt and not any(c.isdigit() for c in nt))
                                      or len(nt) < 15)
                    if is_placeholder:
                        new_p = doc.add_paragraph(fill_text)
                        p._element.addnext(new_p._element)
                    break

        # Fix table cells too
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            for old, new in fills.items():
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

        # ── Clean up bare "—" cells: replace with context-aware values ──
        import re as _re_dash
        for tbl in doc.tables:
            hcells = [c.text.strip() for c in tbl.rows[0].cells]
            for row in tbl.rows[1:]:
                for ci, cell in enumerate(row.cells):
                    ct = cell.text.strip()
                    if ct == '—' or ct == '' or ct == '待统计':
                        # Try to determine context from column header
                        col_hdr = hcells[ci] if ci < len(hcells) else ''
                        col_hdr_lower = col_hdr.lower()
                        # Heuristic fills based on column context
                        if any(kw in col_hdr for kw in ['得分', '评分', '分数']):
                            self._set_cell(cell, '—')  # Legitimate: not yet scored
                        elif any(kw in col_hdr for kw in ['增幅', '变化率', '增长率']):
                            self._set_cell(cell, '基本持平')
                        elif any(kw in col_hdr for kw in ['备注', '说明', '注']):
                            self._set_cell(cell, '—')  # No remarks needed
                        else:
                            self._set_cell(cell, '【待补充】')

        # Remove duplicate table titles (keep only first occurrence)
        seen_titles = set()
        for p in doc.paragraphs:
            if _has_image(p):
                continue
            t = p.text.strip()
            if t.startswith('表') and len(t) < 60:
                if t in seen_titles:
                    p.clear(); p.text = ''
                else:
                    seen_titles.add(t)

        # Remove orphan image markers
        markers_to_remove = [
            r'^图\d+[-—]\d+\s+.+$',
            r'^\[待补充\].*$',
        ]
        for p in doc.paragraphs:
            if _has_image(p):
                continue
            text = p.text.strip()
            for pattern in markers_to_remove:
                if re.match(pattern, text) and len(text) < 40:
                    p.clear()
                    p.text = ''
                    break

    # ═══════════════════════════════════════════════════════════════
    # Document Setup
    # ═══════════════════════════════════════════════════════════════

    def _setup_page(self, doc: Document):
        for sec in doc.sections:
            sec.page_width = Cm(21)
            sec.page_height = Cm(29.7)
            sec.top_margin = Cm(2.54)
            sec.bottom_margin = Cm(2.54)
            sec.left_margin = Cm(3.18)
            sec.right_margin = Cm(3.18)

    # ═══════════════════════════════════════════════════════════════
    # Cover Page
    # ═══════════════════════════════════════════════════════════════

    def _add_cover_page(self, doc, doc_ref, project_name, org_name, location="", implement_unit=""):
        for _ in range(6):
            doc.add_paragraph()

        # 🔴 封面标题: 黑体 小二(18pt) 居中
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('社会稳定风险评估报告')
        r.font.name = FONT_COVER_TITLE; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_COVER_TITLE)
        r.font.size = Pt(18); r.bold = True

        doc.add_paragraph()

        # Project name as subtitle (only if it's a real name, not a map caption)
        pn = (project_name or doc_ref or '').strip()
        if pn and '示意图' not in pn and '位置图' not in pn:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(pn)
            r.font.name = FONT_H3; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H3)
            r.font.size = Pt(16)

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('土地征收社会稳定风险评估')
        r.font.name = FONT_H3; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_H3)
        r.font.size = Pt(14)

        for _ in range(4):
            doc.add_paragraph()

        for text in [
            f'责任单位：{org_name}' if org_name else '',
            f'稳评实施单位：{implement_unit}' if implement_unit else '',
            f'编制日期：{datetime.now().year}年{datetime.now().month}月',
        ]:
            if not text:
                continue
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.name = FONT_BODY; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
            r.font.size = Pt(14)

        doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # Chapter Processing
    # ═══════════════════════════════════════════════════════════════

    def _add_chapter(self, doc, ch_num, markdown, image_files, survey_stats,
                     extracted_tables=None, filled=None):
        if extracted_tables is None:
            extracted_tables = []

        # Ensure chapter starts with proper heading
        cn = ['', '一','二','三','四','五','六','七','八','九','十']
        ch_cn = cn[ch_num] if ch_num <= 10 else str(ch_num)
        ch_title = self._get_chapter_title(ch_num)
        # Remove any existing chapter heading the LLM generated, use our title
        markdown = re.sub(r'^#{1,3}\s*第[一二三四五六七八九十\d]+章\s*.+?\n', '', markdown)
        markdown = re.sub(r'^第[一二三四五六七八九十\d]+章\s*.+?\n', '', markdown)
        markdown = f"# 第{ch_cn}章 {ch_title}\n\n{markdown}"

        # Fix section numbering: "1 . 1" → "1.1", "1 .1" → "1.1", etc.
        markdown = re.sub(r'(\d+)\s*\.\s*(\d+)', r'\1.\2', markdown)
        markdown = re.sub(r'^### (第[一二三四五六七八九十\d]+章\s)', r'## \1', markdown, flags=re.MULTILINE)
        # 🔴 清理 LLM 生成的"# 数字 章标题"重复（如 "# 2 评估过程、方法和依据"）
        markdown = re.sub(r'^#{1,3}\s*\d+\s+[^\n]+\n', '', markdown, flags=re.MULTILINE)
        # 🔴 清理图片标记残留（如 "[现场勘查照片：图片1]"）——不是实际图片路径的标记
        markdown = re.sub(r'^\[[^\]]*照片[：:][^\]]*\]\s*\n', '', markdown, flags=re.MULTILINE)

        # Safety: ensure markdown is a string
        if not isinstance(markdown, str):
            markdown = str(markdown) if markdown else ""
        if not markdown:
            return

        lines = markdown.split('\n')
        table_buf, in_table = [], False
        skip_md_tables = False  # Set after rendering extracted tables
        injected_tables = set()  # Track tables injected via [TABLE:name] markers

        for line in lines:
            s = line.strip()
            # Skip AI meta-text and placeholder explanations
            if any(x in s for x in [
                '***', '遵照您的指示', '以下是为您撰写', '为您撰写第',
                '因勘测定界报告', '因PDF原文', 'PDF原文中未提供',
                '> 注：上表中总面积', '> 注：因PDF',
            ]):
                continue
            if s.startswith('好的，遵照'):
                continue

            if s.startswith('|') and s.endswith('|'):
                # Collect markdown table lines. Handle LLM merging rows on one line.
                if not in_table:
                    table_buf = []
                # Split malformed single-line tables: rows separated by " | — |" or " |---"
                if ' | — |' in s or ' |---' in s:
                    import re as _re_tbl
                    rows = _re_tbl.split(r'\s*\|\s*[-—]+\s*\|', s)
                    for row in rows:
                        row = row.strip()
                        if not row: continue
                        # Skip pure separator fragments (only dashes and pipes)
                        if re.match(r'^[\s\-—\|]+$', row): continue
                        # Reconstruct: if missing leading/trailing |, add them
                        if not row.startswith('|'):
                            row = '| ' + row
                        if not row.endswith('|'):
                            row = row + ' |'
                        table_buf.append(row)
                else:
                    table_buf.append(s)
                in_table = True; continue
            elif in_table:
                # Render collected markdown table
                if len(table_buf) >= 2:
                    self._render_md_table(doc, table_buf, ch_num, survey_stats)
                table_buf, in_table = [], False

            # 🔴 [TABLE:name] 标记已废弃——表格由章节 agent 用 markdown 语法直接写，
            # 有数据才写，不再用固定标记注入。残留标记直接跳过。
            if re.match(r'^\[TABLE:\w+\]$', s):
                continue

            if not s:
                continue

            # Handle inline image markers: ![caption] or ![caption](path)
            img_marker = re.match(r'^!\[([^\]]+)\]\s*(?:\(([^)]*)\))?', s)
            if img_marker:
                caption = img_marker.group(1).strip()
                img_ref = img_marker.group(2).strip() if img_marker.group(2) else ""

                # If path is provided and exists, use it directly
                if img_ref:
                    self._add_image(doc, img_ref, caption)
                    continue  # 🔴 Skip image-map fallback — path was explicit
                # 🔴 按图注关键词匹配图片（避免位置错乱）
                matched_path = self._match_image_by_caption(caption)
                if matched_path:
                    self._add_image(doc, matched_path, caption)
                    continue
                # Fall back to catalog chapter image map only
                ch_img_map = getattr(self, '_chapter_image_map', {})
                ch_imgs = ch_img_map.get(ch_num, [])
                if ch_imgs:
                    img_info = ch_imgs[0]
                    img_path = img_info.get("path", img_info) if isinstance(img_info, dict) else img_info
                    if img_path:
                        self._add_image(doc, img_path, caption)
                    ch_img_map[ch_num] = ch_imgs[1:]  # consume
                continue

            # Skip old-style AI-generated image markers
            if re.match(r'^图\d+[-—]\d+\s', s) and len(s) < 30:
                continue

            # Chapter heading — handles #, ##, ###, or bare 第X章
            m = re.match(r'^#+\s*第([一二三四五六七八九十\d]+)章\s*(.+)', s)
            if not m:
                m = re.match(r'^第([一二三四五六七八九十\d]+)章\s*(.+)', s)
            if m:
                cn = m.group(1)
                num = CN_NUM.get(cn, int(cn) if cn.isdigit() else ch_num)
                txt = f"第{cn}章 {m.group(2).strip()}"
                self._add_heading(doc, txt, 1)
                continue

            # ## → DOCX Heading 2 (二级标题: 黑体 四号 14pt)
            m = re.match(r'^##\s+(.+)', s)
            if m:
                heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', m.group(1))
                llm_invented_patterns = [
                    r'^(合法性|合理性|可行性|可控性)[（(]得分[：:]\s*\d+',
                    r'^表(?!5[-—])',
                ]
                is_llm_invented = any(bool(re.match(p, heading_text)) for p in llm_invented_patterns)
                if not is_llm_invented:
                    self._add_heading(doc, heading_text, 2)
                continue

            # ### → DOCX Heading 3 (三级标题: 楷体_GB2312 四号 14pt) — DB32/T4013
            m = re.match(r'^###\s+(.+)', s)
            if m:
                self._add_heading(doc, re.sub(r'\*\*(.+?)\*\*', r'\1', m.group(1)), 3)
                continue

            m = re.match(r'^#### (.+)', s)
            if m:
                self._add_heading(doc, re.sub(r'\*\*(.+?)\*\*', r'\1', m.group(1)), 3)
                continue

            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
            clean = re.sub(r'\[待补充\]', '【待补充】', clean)

            # 🔴 跳过 LLM 手写的表格标题（系统 auto-inject 会统一渲染，避免重复）
            # 第5章除外（第5章是 LLM 写 markdown 表格，标题由 LLM 自己写）
            if ch_num != 5 and re.match(r'^表\s*\d+[-—]\d+', clean) and len(clean) < 40:
                continue

            if re.match(r'^\d+\.\s*\*\*', clean):
                self._add_para(doc, clean, bold=True)
            elif clean.startswith('- ') or clean.startswith('* '):
                self._add_para(doc, '    ' + clean[2:])
            elif clean.startswith('**') and clean.endswith('**'):
                self._add_para(doc, clean.strip('*').strip(), bold=True)
            elif len(clean) > 10:
                self._add_para(doc, clean, indent=True)

        if in_table and table_buf and not skip_md_tables:
            self._render_md_table(doc, table_buf, ch_num, survey_stats)

        # 🔴 固定表格 auto-inject 已移除——表格由章节 agent 根据实际数据动态设计
        # （用 markdown 语法写，有数据才写），不再由系统硬编码渲染固定表格。

        # 🔴 Auto-insert chapter images at end of content
        ch_img_map = getattr(self, '_chapter_image_map', {})
        ch_imgs = ch_img_map.get(ch_num, [])
        if ch_num == 1:
            logger.info(f"Ch{ch_num} image map: {len(ch_img_map)} chapters, this chapter: {len(ch_imgs)} images")
        if ch_imgs:
            img_count = 0
            for img_info in ch_imgs[:3]:  # Max 3 images per chapter
                if isinstance(img_info, dict):
                    path = img_info.get("path", "")
                    # 🔴 用 display_name（简洁类别名，如"公示照片"）而非原始文件名
                    name = img_info.get("name", "") or img_info.get("caption", "") or ""
                else:
                    path = str(img_info)
                    name = ""
                if path:
                    # 🔴 简洁图注：图X-N 类别名（不带 pdf_xxx 原始文件名）
                    cap = f"图{ch_num}-{img_count+1} {name}".strip() if name else f"图{ch_num}-{img_count+1}"
                    self._add_image(doc, path, cap)
                    img_count += 1

        # 🔴 Render extracted PDF tables at END of chapter (any chapter with mapped tables)
        if extracted_tables:
            self._render_extracted_tables(doc, ch_num, extracted_tables)

        # 🔴 不再生成"【待插入：图X-X】"占位文本（会造成重复占位无图）。
        # 图片由 LLM 正文标记或章节末尾插入决定；该章无图则不显示任何占位。

    # ═══════════════════════════════════════════════════════════════
    # Chapter titles (template-aligned)
    # ═══════════════════════════════════════════════════════════════
    _CHAPTER_TITLES = {
        1: "拟征收决策基本概况",
        2: "评估过程、方法和依据",
        3: "社会稳定风险因素调查",
        4: "决策综合分析",
        5: "风险因素识别与初始等级表",
        6: "措施前风险等级研判",
        7: "风险防范与化解措施",
        8: "措施后风险等级评估",
        9: "评估结论与建议",
        10: "应急预案",
    }

    def _get_chapter_title(self, ch_num):
        return self._CHAPTER_TITLES.get(ch_num, f"第{ch_num}章")

    # ═══════════════════════════════════════════════════════════════
    # Tables
    # ═══════════════════════════════════════════════════════════════

    # ── Table Reshaping: 2-column KV → multi-column ────────────────
    # Known table structures mapped by first header pair
    _TABLE_SCHEMAS = {
        ("测评指标", "权重"): {
            "columns": ["测评指标", "权重", "测评项目", "评分", "评分标准", "得分"],
            "group_size": 6,  # KV pairs per row
        },
        ("序号", "评估维度"): {
            "columns": ["序号", "评估维度", "措施前得分", "措施后得分", "变化幅度"],
            "group_size": 5,
        },
        ("调查内容", "选项"): {
            "columns": ["调查内容", "选项", "人数", "比例"],
            "group_size": 4,
        },
        ("序号", "风险类型"): {
            "columns": ["序号", "风险类型", "风险因素描述", "风险等级"],
            "group_size": 4,
        },
        ("序号", "风险因素"): {
            "columns": ["序号", "风险因素", "触发条件", "影响范围", "发生概率", "影响程度", "风险等级"],
            "group_size": 7,
        },
        ("序号", "法规/标准名称"): {
            # Ch2 regulation list — keep as 2-col or reshape to bullet list
            "columns": ["序号", "法规/标准名称"],
            "group_size": 2,
        },
    }

    def _reshape_table(self, table_lines):
        """Reshape 2-column KV tables into proper multi-column format.

        When the LLM outputs a 2-column key-value table (e.g. | 测评指标 | 权重 |)
        that should be 6 columns, group the KV pairs into proper rows.
        """
        if len(table_lines) < 4:
            return table_lines

        # Parse header
        header_line = table_lines[0].strip()
        cells = [c.strip() for c in header_line.strip('|').split('|')]
        if len(cells) > 2:
            # Already multi-column — nothing to reshape
            return table_lines

        # Look up schema by first two headers
        h0, h1 = (cells[0] if len(cells) > 0 else ""), (cells[1] if len(cells) > 1 else "")
        schema = self._TABLE_SCHEMAS.get((h0, h1))

        if not schema:
            return table_lines
        group_size = schema["group_size"]
        if group_size <= 2 or group_size > 8:
            return table_lines

        # Collect all KV pairs (skip header row itself)
        kv_pairs = []
        seen_header = False
        for line in table_lines:
            s = line.strip()
            if re.match(r'^\|[\s\-:|—]+\|$', s):
                continue
            cols = [c.strip() for c in s.strip('|').split('|')]
            if not seen_header:
                seen_header = True
                continue  # Skip original 2-col header
            if len(cols) >= 2 and cols[0] and cols[1]:
                kv_pairs.append((cols[0], cols[1]))

        # Simple grouping: just take groups of group_size values
        new_columns = schema["columns"]
        group_size = schema["group_size"]
        new_rows = []
        i = 0
        while i + group_size <= len(kv_pairs):
            row_vals = [kv_pairs[i+j][1] for j in range(group_size)]
            new_rows.append(row_vals)
            i += group_size
        # Remainder
        if i < len(kv_pairs):
            row_vals = [kv_pairs[i+j][1] if i+j < len(kv_pairs) else "" for j in range(group_size)]
            new_rows.append(row_vals)

        if not new_rows:
            return table_lines

        # Sanity check: don't create absurdly wide tables
        if len(new_columns) > 10:
            return table_lines

        # Build new markdown table
        result = []
        result.append("| " + " | ".join(new_columns) + " |")
        result.append("| " + " | ".join(["---"] * len(new_columns)) + " |")
        for row in new_rows:
            result.append("| " + " | ".join(row) + " |")

        return result

    def _render_md_table(self, doc, table_lines, ch_num, survey_stats):
        """Parse markdown table lines into DOCX table.

        所有章节的 markdown 表格都渲染——表格由章节 agent 根据实际数据动态设计，
        有数据才写，格式参考模板。系统不再强制注入固定表格。
        """
        rows = []
        for line in table_lines:
            if re.match(r'^\|[\s\-:|—]+\|$', line):
                continue  # Skip separator
            line = line.strip().strip('|')
            cells = [c.strip() for c in line.split('|')]
            # Remove trailing empty cells from malformed splits
            while cells and cells[-1] == '':
                cells.pop()
            if any(cells):
                rows.append(cells)

        if len(rows) < 2:
            return

        ncols = max(len(r) for r in rows) if rows else 2
        # Ch5 risk table: trim to exactly 4 columns (序号/风险因素/风险表现/风险等级)
        if ch_num == 5 and ncols > 4:
            ncols = 4
            rows = [r[:4] for r in rows]
        # Limit columns: if absurdly wide, keep as 2-column
        if ncols > 10:
            flat = []
            for r in rows: flat.extend(r)
            rows = [[flat[i], flat[i+1] if i+1 < len(flat) else ''] for i in range(0, len(flat), 2)]
            ncols = 2
        for r in rows:
            while len(r) < ncols:
                r.append('')

        table = doc.add_table(rows=len(rows), cols=ncols, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row_data in enumerate(rows):
            for ci in range(ncols):
                cell_text = row_data[ci] if ci < len(row_data) else ''
                cell = table.cell(ri, ci)
                cell.paragraphs[0].clear()
                r = cell.paragraphs[0].add_run(str(cell_text))
                r.font.name = FONT_TABLE; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TABLE)
                r.font.size = Pt(12)
                if ri == 0:
                    r.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    def _set_cell(self, cell, text, bold=False):
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(text)
        r.font.name = FONT_TABLE; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TABLE)
        r.font.size = Pt(12)
        r.bold = bold

    def _add_survey_table(self, doc, stats, caption):
        # 🔴 表名: 宋体 小四(12pt) 居中
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption); r.bold = True
        r.font.name = FONT_CAPTION; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CAPTION)
        r.font.size = Pt(12)
        doc.add_paragraph()

        headers = ['调查项目', '人数', '占比']
        rows = [
            ['支持', str(stats.get('support_count', 0)), f"{stats.get('support_rate', 0)}%"],
            ['反对', str(stats.get('oppose_count', 0)), f"{stats.get('oppose_rate', 0)}%"],
            ['有条件支持', str(stats.get('neutral_count', 0)), f"{stats.get('neutral_rate', 0)}%"],
            ['合计', str(stats.get('total_surveys', 0)), '100%'],
        ]
        table = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(headers):
            self._set_cell(table.cell(0, ci), h, bold=True)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                self._set_cell(table.cell(ri+1, ci), val)

    # ═══════════════════════════════════════════════════════════════
    # Extracted Tables (from PDF)
    # ═══════════════════════════════════════════════════════════════

    @staticmethod
    def _get_extracted_tables(state: dict) -> Dict[int, list]:
        """Collect extracted PDF tables and map them to chapters."""
        result: Dict[int, list] = {}

        sources = [
            state.get("_project_material_facts", {}).get("_extracted_tables", []),
        ]
        summary = state.get("_project_material_summary", {})
        if isinstance(summary, dict):
            sources.append(summary.get("facts", {}).get("_extracted_tables", []))
        for m in state.get("_project_materials", []) or []:
            if isinstance(m, dict):
                meta = m.get("metadata", {})
                if isinstance(meta, dict):
                    sources.append(meta.get("extracted_tables", []))

        seen = set()
        all_tables = []
        for src in sources:
            if isinstance(src, list):
                for tbl in src:
                    if isinstance(tbl, dict):
                        # Deduplicate by headers signature
                        sig = '|'.join(str(h) for h in tbl.get('headers', [])[:5])
                        if sig and sig not in seen:
                            seen.add(sig)
                            all_tables.append(tbl)

        # Map to chapters based on content
        for tbl in all_tables:
            if not isinstance(tbl, dict):
                continue
            headers = tbl.get("headers", [])
            hstr = " ".join(str(h) for h in headers).lower() if headers else ""

            # 勘测定界报告 → Ch1
            if any(kw in hstr for kw in ['地块', '坐落', '界址', '面积', '土地用途']):
                result.setdefault(1, []).append(tbl)
            # 土地分类面积 → Ch1 or Ch4
            elif any(kw in hstr for kw in ['耕地', '园地', '林地', '分类面积']):
                result.setdefault(1, []).append(tbl)
            # 调查/问卷 → Ch3
            elif any(kw in hstr for kw in ['调查', '问卷', '支持', '反对']):
                result.setdefault(3, []).append(tbl)

        return result

    def _render_extracted_tables(self, doc, ch_num: int, tables: list) -> bool:
        """Render PDF-extracted tables directly as DOCX tables. Returns True if any rendered."""
        if not tables:
            return False
        for tbl in tables:
            if not isinstance(tbl, dict):
                continue
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if not headers or not rows:
                continue

            # Filter to reasonable column count
            if len(headers) > 15:
                # Too many columns — narrow to first 10 meaningful headers
                meaningful = [h for h in headers if h.strip() and len(h.strip()) > 1]
                headers = meaningful[:10]
                rows = [[r[i] if i < len(r) else '' for i in range(len(headers))] for r in rows]

            ncols = len(headers)
            nrows = len(rows) + 1  # +1 for header

            table = doc.add_table(rows=nrows, cols=ncols, style='Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header row
            for ci, h in enumerate(headers):
                cell = table.cell(0, ci)
                cell.paragraphs[0].clear()
                r = cell.paragraphs[0].add_run(str(h).replace('\n', ' '))
                r.font.name = FONT_TABLE
                r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TABLE)
                r.font.size = Pt(12)
                r.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Data rows
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    val = row[ci] if ci < len(row) else ''
                    cell = table.cell(ri + 1, ci)
                    cell.paragraphs[0].clear()
                    r = cell.paragraphs[0].add_run(str(val).replace('\n', ' '))
                    r.font.name = FONT_TABLE
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_TABLE)
                    r.font.size = Pt(12)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()
        return True

    @staticmethod
    def _remove_markdown_tables(markdown: str) -> str:
        """Strip markdown table blocks (already rendered directly from PDF)."""
        if not isinstance(markdown, str):
            return str(markdown) if markdown else ""
        lines = markdown.split('\n')
        result = []
        skip = False
        for line in lines:
            s = line.strip()
            if s.startswith('|') and s.endswith('|'):
                skip = True
                continue
            if skip and not (s.startswith('|') and s.endswith('|')):
                skip = False
            if not skip:
                result.append(line)
        return '\n'.join(result)

    # ═══════════════════════════════════════════════════════════════
    # Images
    # ═══════════════════════════════════════════════════════════════

    def _resolve_image_path(self, image_ref):
        from pathlib import Path as _Path
        p = _Path(image_ref)
        if p.is_absolute():
            return p
        # Remove leading 'storage/' if present (relative to storage_dir)
        ref_str = str(image_ref)
        if ref_str.startswith('storage/'):
            return self.storage_dir / ref_str[8:]  # strip 'storage/'
        # Try images_dir, then storage_dir
        for d in [self.images_dir, self.storage_dir]:
            candidate = d / _Path(image_ref).name
            if candidate.exists():
                return candidate
        return self.storage_dir / image_ref

    @staticmethod
    def _clean_image_caption(caption: str, img_path) -> str:
        """Clean up image caption. If it's a raw filename, extract a readable label."""
        import re as _re_cap
        fname = str(img_path.name) if hasattr(img_path, 'name') else str(img_path).split('/')[-1]
        # If caption already looks like a proper label (图X-X ...), keep it
        if _re_cap.match(r'图\s*\d+[-—]\d+\s+\S', caption):
            return caption
        if caption.startswith('附图') or caption.startswith('图') and len(caption) > 5:
            return caption
        # If caption is just the filename, extract readable parts
        # If caption is a raw filename, generate a clean label
        if caption == fname or caption.endswith(('.jpg', '.png', '.jpeg', '.JPG', '.PNG', '.JPEG')):
            clean = _re_cap.sub(r'\.(jpg|jpeg|png|PNG|JPG|JPEG)$', '', fname)
            clean = _re_cap.sub(r'^pdf_[^_]+_', '', clean)
            clean = _re_cap.sub(r'_[a-f0-9]{8,}\b', '', clean)
            clean = _re_cap.sub(r'_4\d$', '', clean)
            # Keep meaningful prefixes like "公示照片4_", "专家评审会照片_"
            clean = clean.replace('_', ' ').strip()
            if len(clean) > 3:
                return clean[:50]  # Truncate to reasonable length
        return caption

    # 🔴 按图注关键词匹配图片路径（从已分类的图片目录里找）
    def _match_image_by_caption(self, caption: str):
        """根据图注关键词从图片目录匹配最合适的图片路径。

        图注如"图1-1 拟征收土地位置示意图"→ 匹配 map 类图片
        图注如"图3-1 公示照片"→ 匹配 announcement 类图片
        返回图片路径，无匹配返回 None。
        """
        if not caption:
            return None
        # 图注关键词 → 图片类别
        cat_keywords = {
            'map': ['位置', '示意', '红线', '勘测', '地形', '规划图'],
            'announcement': ['公示', '公告', '张贴', '预公告'],
            'survey': ['问卷', '调查', '签到'],
            'review': ['评审', '意见', '专家'],
            'photo': ['现场', '照片', '走访', '勘察'],
            'meeting': ['座谈', '会议', '开会'],
        }
        target_cat = None
        for cat, kws in cat_keywords.items():
            if any(kw in caption for kw in kws):
                target_cat = cat
                break
        if not target_cat:
            return None

        # 从图片目录里找该类别的第一张未用图片
        catalog_data = getattr(self, '_img_catalog_data', {})
        for img in catalog_data.get("catalog", []):
            if img.get("category") == target_cat:
                p = img.get("path", "")
                if p and str(p) not in self._inserted_images:
                    return p
        return None

    # 🔴 图片尺寸规范（淮安市稳评格式，单位cm）按图注关键词匹配
    IMAGE_SIZE_SPECS = [
        (['决策实施周期'], (13.45, 10.12)),
        (['风险评估流程'], (13.69, 12.45)),
        (['评估公示内容'], (16.48, 11.65)),
        (['决策位置', '位置示意图', '征地范围图'], (10.91, 15.5)),
        (['决策网络舆情'], (4.6, 4.41)),
        (['单位调查问卷', '公众调查问卷'], (10.39, 7.34)),
        (['座谈会'], (5.43, 7.23)),
        (['现场照片'], (5.44, 7.25)),
        (['公示照片'], (5.39, 7.18)),
        (['会议纪要'], (20.51, 14.95)),
        (['调查问卷'], (21.84, 15.45)),
        (['公告'], (21.35, 14.84)),
    ]

    def _match_image_size(self, caption: str):
        """按图注关键词匹配图片尺寸（cm），无匹配返回 None。"""
        if not caption:
            return None
        for keywords, size in self.IMAGE_SIZE_SPECS:
            if any(kw in caption for kw in keywords):
                return size
        return None

    def _add_image(self, doc, image_ref, caption, max_width=Inches(5.5), max_height=Inches(7.0)):
        img_path = self._resolve_image_path(image_ref)
        # 🔴 Global dedup: skip if already inserted
        img_key = str(img_path)
        if img_key in self._inserted_images:
            return
        self._inserted_images.add(img_key)

        if not img_path.exists():
            for d in [self.images_dir, self.storage_dir]:
                fallback = d / _Path(image_ref).name
                if fallback.exists():
                    img_path = fallback
                    break
        if not img_path.exists():
                print(f"[IMG_FAIL] _add_image failed: ref={image_ref}, caption={caption}", flush=True, file=__import__('sys').stderr)
                self._add_para(doc, f'【图片待插入：{caption}】', indent=False)
                return

        # 🔴 Clean up caption: if it's a raw filename, extract meaningful parts
        clean_caption = self._clean_image_caption(caption, img_path)

        # 🔴 Image above caption format: "图 X-X 描述"
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(0)
        p_img.paragraph_format.space_before = Pt(6)
        try:
            from PIL import Image as PILImage
            import tempfile as _tmp
            with PILImage.open(str(img_path)) as img:
                orig_w, orig_h = img.size
                if img.mode in ('RGBA', 'P', 'CMYK', 'LA'):
                    img = img.convert('RGB')
                tmp_fd, tmp_path = _tmp.mkstemp(suffix='.jpg')
                os.close(tmp_fd)
                # 🔴 Always compress: max 1600px, JPEG quality 80, strip EXIF
                max_dim = 1600
                if orig_w > max_dim or orig_h > max_dim:
                    ratio = min(max_dim/orig_w, max_dim/orig_h)
                    img = img.resize((int(orig_w*ratio), int(orig_h*ratio)), PILImage.LANCZOS)
                elif img.mode in ('RGBA', 'P', 'CMYK', 'LA') or img_path.suffix.lower() in ('.png','.bmp','.tiff','.tif'):
                    pass  # Still need to convert format
                img.save(tmp_path, 'JPEG', quality=80, optimize=True)
                use_path = tmp_path

            aspect = orig_w / orig_h if orig_h > 0 else 1.0
            # 🔴 图片尺寸规范（淮安市稳评格式，单位cm）：按图注匹配
            spec_size = self._match_image_size(clean_caption)
            if spec_size:
                target_w = Cm(spec_size[0])
                target_h = Cm(spec_size[1])
            else:
                target_w = min(max_width, Inches(5.0))
                target_h = target_w / aspect
                if target_h > max_height:
                    target_h = max_height
                    target_w = target_h * aspect

            run_img = p_img.add_run()
            run_img.add_picture(use_path, width=target_w, height=target_h)
            try:
                os.unlink(use_path)
            except:
                pass
        except Exception as e:
            logger.warning(f"Failed to add image {image_ref}: {e}")
            return

        # 🔴 图名: 宋体 小四(12pt) 居中，位于图片下方居中
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(10)
        run_cap = p_cap.add_run(clean_caption)
        run_cap.font.name = FONT_CAPTION
        run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CAPTION)
        run_cap.font.size = Pt(12)  # 小四
        run_cap.italic = False      # 不斜体

    # ═══════════════════════════════════════════════════════════════
    # Company Front Matter (开篇)
    # ═══════════════════════════════════════════════════════════════

    def _add_company_qualifications(self, doc, filled, domain="stability"):
        """Page 2: 公司资质证书图片. Returns True if content was added."""
        cert_images = self._load_company_certificate_images(domain=domain)

        if cert_images:
            self._add_heading(doc, '公司资质证书', 2)
            for img_path, caption in cert_images[:6]:
                self._add_image(doc, img_path, caption)
            return True
        return False

    def _add_company_intro(self, doc, domain: str = "stability"):
        """Page 3: 公司简介 + 工作组人员."""
        # ── 公司简介 ──
        self._add_heading(doc, '公司简介', 2)
        self._add_para(doc,
            '江苏众拓项目代理咨询有限公司是一家具有多年项目代理咨询经验的专业咨询机构，'
            '主要致力于为客户项目审批核准提供"一站式""全流程"的咨询服务。'
            '公司服务内容涵盖包括项目前期准备阶段中用地预审、立项报批手续代理服务及涉及的'
            '技术报告编制、项目实施阶段中技术服务、项目后期服务及各类评估业务。',
            indent=True)
        self._add_para(doc,
            '公司是经淮安经济技术开发区行政审批局登记注册，经营范围包含维稳评估咨询业务的专业机构。'
            '固定办公场所面积达260m²，固定从业人员35名，评估业务部门日常工作人员10名，'
            '其中稳评业务主要负责人、业务骨干等均持有社会稳定风险评估专业技术人员培训证书。',
            indent=True)
        self._add_para(doc,
            '我公司先后承接了关于出台《淮安市不动产登记有关历史遗留问题的处理意见》'
            '《淮安市现代有轨电车近期建设规划（2022-2026年）》重大决策事项、'
            '淮安市韩侯大道次高压A燃气管道、金湖向东风电场、苏淮高新区危险危废处置中心、'
            '淮安市淮阴区赵集镇天然气分布式能源等社会稳定风险评估业务。',
            indent=True)

        doc.add_page_break()

        # ── 稳评工作组人员及分工（宋体小三加黑居中）──
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('稳评工作组人员及分工情况')
        r.font.name = FONT_SONGTI; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
        r.font.size = Pt(15); r.bold = True
        doc.add_paragraph()
        # 岗位标题: 宋体四号加黑靠左；内容: 宋体四号靠左
        self._add_para(doc, '稳评负责人', bold=True, indent=False, font=FONT_SONGTI)
        self._add_para(doc, '陈  春（总经理、高级工程师、估价师、经济师）', indent=True, font=FONT_SONGTI)
        self._add_para(doc, '调研工作', bold=True, indent=False, font=FONT_SONGTI)
        self._add_para(doc, '安如月（评估专业人员）', indent=True, font=FONT_SONGTI)
        self._add_para(doc, '张抗洪（评估专业人员）', indent=True, font=FONT_SONGTI)
        self._add_para(doc, '朱  璇（评估助理）', indent=True, font=FONT_SONGTI)
        self._add_para(doc, '风险研判与评估报告编制工作', bold=True, indent=False, font=FONT_SONGTI)
        self._add_para(doc, '程士汝（技术负责人）', indent=True, font=FONT_SONGTI)
        self._add_para(doc, '刘利伟（评估专业人员）', indent=True, font=FONT_SONGTI)
        self._add_para(doc, '数据汇总与资料档案', bold=True, indent=False, font=FONT_SONGTI)
        self._add_para(doc, '张抗洪（评估助理）', indent=True, font=FONT_SONGTI)
        self._add_para(doc, '罗  娜（评估助理）', indent=True, font=FONT_SONGTI)
        self._add_para(doc, '朱  璇（评估助理）', indent=True, font=FONT_SONGTI)

        doc.add_page_break()

    def _add_legal_basis(self, doc):
        """Add standalone 法律法规与评估依据 section to the report."""
        self._add_heading(doc, '法律法规与评估依据', 1)
        doc.add_paragraph()

        laws = [
            ("一、国家法律", [
                "《中华人民共和国土地管理法》（2019年8月26日第三次修正，2020年1月1日起施行）",
                "《中华人民共和国突发事件应对法》（2024年6月28日修订，2024年11月1日起施行）",
                "《中华人民共和国城乡规划法》（2019年4月23日第二次修正）",
                "《中华人民共和国农村土地承包法》（2018年12月29日第二次修正）",
                "《中华人民共和国村民委员会组织法》（2018年12月29日修正）",
                "《中华人民共和国政府信息公开条例》（国务院令第711号，2019年5月15日起施行）",
                "《信访工作条例》（中共中央、国务院发布，2022年5月1日起施行）",
            ]),
            ("二、省级法规与标准", [
                "《中华人民共和国土地管理法实施条例》（2021年修订，国务院令第743号）",
                "《江苏省土地管理条例》（2021年1月15日修订，2021年5月1日起施行）",
                "《江苏省被征地农民社会保障办法》（苏政发〔2021〕87号，2022年3月1日起施行）",
                "DB32/T4013-2021《第三方社会稳定风险评估规范》",
                "DB32/T4937-2024《土地征收前社会稳定风险评估规范》",
                "《江苏省重大决策社会稳定风险评估办法》（苏政办发〔2021〕58号）",
                "《关于加强新形势下重大决策社会稳定风险评估机制建设的实施意见》（苏办发〔2021〕15号）",
            ]),
            ("三、市级规范与文件", [
                "《市政府关于公布淮安市征地区片综合地价执行标准的通知》（淮政规〔2026〕1号，2026年6月9日发布，现行有效）",
                "《淮安市征地补偿和被征地农民社会保障实施细则》",
                "《洪泽区国土空间总体规划（2021—2035年）》",
                "《洪泽区国民经济和社会发展第十四个五年规划和二〇三五年远景目标纲要》",
            ]),
            ("四、项目文件", [
                "洪拟征告〔2026〕7号 征收土地预公告",
                "0-勘测定界报告-",
                "项目用地预审与选址意见书",
                "拟征收土地现状调查确认表",
            ]),
        ]

        for section_title, items in laws:
            self._add_heading(doc, section_title, 2)
            for item in items:
                self._add_para(doc, f'• {item}', indent=True, font='仿宋_GB2312', size=Pt(12))
            doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # Appendices (项目材料)
    # ═══════════════════════════════════════════════════════════════

    def _add_appendices(self, doc, survey_stats, image_files, filled):
        """Add appendix with ALL user-provided images organized by category."""
        # Collect ALL images from the catalog
        catalog_data = getattr(self, '_img_catalog_data', {})
        all_images = catalog_data.get("catalog", [])

        # Skip full-page OCR renders, keep embedded images
        SKIP_PREFIXES = ()  # Let image_catalog handle filtering
        session_images = []
        for img in all_images:
            p = img.get("path", "")
            if isinstance(p, str) and os.path.exists(p) and os.path.getsize(p) > 5000:
                fname = os.path.basename(p)
                if any(fname.startswith(prefix) for prefix in SKIP_PREFIXES):
                    continue
                if os.path.splitext(p)[1].lower() in {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}:
                    session_images.append(img)

        # ── 附件一：问卷调查统计分析 ──
        total = survey_stats.get('total_surveys', 0)
        if total > 0:
            self._add_heading(doc, '附件一：问卷调查统计分析', 1)
            support_rate = survey_stats.get('support_rate', 0)
            self._add_para(doc,
                f'本次社会稳定风险评估共发放问卷 {total} 份，回收有效问卷 {total} 份。', indent=True)
            doc.add_paragraph()
            self._add_survey_table(doc, survey_stats, '附表1 公众意见调查统计汇总表')

        # ── 附件二：项目资料图片（分类展示全部用户图片） ──
        # 🔴 Load company cert images for appendix
        domain = filled.get("_domain", "stability") if isinstance(filled, dict) else "stability"
        cert_images = self._load_company_certificate_images(domain=domain)
        cert_paths = set(p for p, _ in cert_images)

        if session_images or cert_paths:
            doc.add_page_break()
            self._add_heading(doc, '附件二：项目资料图片', 1)

            # 🔴 Report-relevant categories
            cat_groups = [
                ("一、征收公告与公示照片", ["announcement"]),
                ("二、群众调查与座谈会照片", ["survey", "meeting"]),
                ("三、部门调查与走访照片", ["photo"]),
                ("四、专家评审意见", ["review"]),
                ("五、位置示意图与勘测图", ["map"]),
            ]

            for label, cats in cat_groups:
                paths = []
                for cat in cats:
                    for img in session_images:
                        if img.get("category") == cat:
                            p = img.get("path", "")
                            if p and p not in paths:
                                paths.append(p)
                if not paths:
                    continue
                # 🔴 Filter: skip already-inserted images (global dedup)
                fresh_paths = [p for p in paths if str(p) not in self._inserted_images]
                if not fresh_paths:
                    continue
                self._add_heading(doc, label, 2)
                self._add_para(doc, f'共 {len(fresh_paths)} 张', indent=True)
                shown = 0
                for p in fresh_paths[:8]:
                    fname = os.path.basename(p)
                    clean_name = __import__('re').sub(r'\.(jpg|jpeg|png)$', '', fname, flags=__import__('re').I)
                    self._add_image(doc, p, clean_name[:60])
                    shown += 1
                total_in_category = len(paths)
                if total_in_category > shown:
                    self._add_para(doc, f'（共 {total_in_category} 张，以上展示 {shown} 张）', indent=True)

            # ── 六、公司资质与备案证书（从稳评模板提取）──
            # 🔴 Only show certs NOT already shown on page 2
            fresh_certs = [(p, cap) for p, cap in cert_images if str(p) not in self._inserted_images]
            if fresh_certs:
                self._add_heading(doc, '六、公司资质与备案证书', 2)
                self._add_para(doc, f'共 {len(fresh_certs)} 张', indent=True)
                for cert_path, cert_caption in fresh_certs[:8]:
                    self._add_image(doc, cert_path, cert_caption)

            # ── 七、其他资料（未分类图片）──
            other_paths = []
            for img in session_images:
                if img.get("category") == "other":
                    p = img.get("path", "")
                    if p and p not in other_paths and str(p) not in self._inserted_images and p not in cert_paths:
                        other_paths.append(p)
            if other_paths:
                self._add_heading(doc, '七、其他项目资料', 2)
                self._add_para(doc, f'共 {len(other_paths)} 张', indent=True)
                for p in other_paths[:8]:
                    fname = os.path.basename(p)
                    clean_name = __import__('re').sub(r'\.(jpg|jpeg|png)$', '', fname, flags=__import__('re').I)
                    self._add_image(doc, p, clean_name[:60])

    def _load_company_certificate_images(self, domain: str = "stability") -> List[tuple]:
        """Load company cert images. Stability: template-extracted. Bidding: DB."""
        import zipfile, sqlite3, xml.etree.ElementTree as ET
        from app.config import settings

        results = []

        if domain == "stability":
            # 🔴 Stability: scan extracted_imgs for stability_* files (from template extraction)
            for f in sorted(self.company_dir.iterdir()):
                if not f.name.startswith('stability_cert_'): continue
                if f.suffix.lower() not in ('.jpg','.jpeg','.png','.gif','.bmp'): continue
                # Get caption from filename hint
                # 🔴 Caption from filename: stability_cert_00 = 公司营业执照 etc
                cap = f.stem.replace('stability_cert_','').replace('_',' ')
                # Map known certs
                if '00' in f.stem: cap = '公司营业执照'
                elif '01' in f.stem: cap = '稳评平台备案及人员证书'
                else: cap = f'资质证书 {len(results)+1}'
                results.append((str(f), cap))
        else:
            # 🔴 Bidding: DB asset_images
            try:
                db_path = settings.DATA_DIR / "knowledge_base.db"
                if db_path.exists():
                    conn = sqlite3.connect(str(db_path))
                    conn.row_factory = sqlite3.Row
                    cur = conn.execute(
                        "SELECT id, image_name, image_data, mime_type FROM asset_images "
                        "WHERE is_active=1 AND category IN ('营业执照','资质证书','人员证书',"
                        "'财务报告','社保纳税','法人证明') ORDER BY id LIMIT 12"
                    )
                    for row in cur.fetchall():
                        img_data = row['image_data']
                        if not img_data: continue
                        mime = row['mime_type'] or 'image/png'
                        ext = 'png' if 'png' in mime else 'jpg'
                        safe_name = f"cert_db_{row['id']}.{ext}"
                        filepath = self.images_dir / safe_name
                        if not filepath.exists():
                            if isinstance(img_data, str):
                                import base64
                                try: filepath.write_bytes(base64.b64decode(img_data))
                                except: filepath.write_bytes(img_data.encode('latin-1'))
                            else: filepath.write_bytes(bytes(img_data))
                        results.append((f"images/{safe_name}", row['image_name'][:60]))
                    conn.close()
            except Exception as e:
                logger.warning(f"DB cert load failed: {e}")

        return results

    # ═══════════════════════════════════════════════════════════════
    # Data Extraction from State
    # ═══════════════════════════════════════════════════════════════

    def _get_session_images(self, state: dict) -> Dict[str, List[str]]:
        images = {"survey": [], "announcement": [], "photo": [], "review": [], "map": [], "other": []}
        all_image_paths = []

        uploaded = state.get("_uploaded_files", []) or []
        # Track original filenames for categorization
        original_names = {}  # path -> original_name
        for item in uploaded:
            if isinstance(item, dict):
                path = item.get("path", "")
                oname = item.get("original_name", "")
                if path:
                    all_image_paths.append(path)
                    original_names[path] = oname
            elif isinstance(item, str):
                all_image_paths.append(item)
        state["_image_original_names"] = original_names

        material_summary = state.get("_project_material_summary", {})
        if isinstance(material_summary, dict):
            facts = material_summary.get("facts", {})
            extracted = facts.get("_extracted_images", [])
            if isinstance(extracted, list):
                for img in extracted:
                    if img not in all_image_paths:
                        all_image_paths.append(img)

        filled = state.get("filled_data", {})
        filled_imgs = filled.get("_extracted_images", [])
        if isinstance(filled_imgs, list):
            for img in filled_imgs:
                if img not in all_image_paths:
                    all_image_paths.append(img)

        # 🔴 Pull in classified chapter images from DataAnalysisAgent
        chapter_images = state.get("_chapter_images", {})
        chapter_image_map = {}  # ch_num → list of image paths
        for ch_str, img_list in (chapter_images or {}).items():
            try:
                ch_num = int(ch_str)
            except (ValueError, TypeError):
                ch_num = 999  # "other" → appendix
            chapter_image_map[ch_num] = []
            for img_info in img_list:
                if isinstance(img_info, dict):
                    path = img_info.get("path", "")
                    if path and path not in all_image_paths:
                        all_image_paths.append(path)
                    chapter_image_map[ch_num].append(path)
                elif isinstance(img_info, str):
                    if img_info not in all_image_paths:
                        all_image_paths.append(img_info)
                    chapter_image_map[ch_num].append(img_info)
        # Store chapter-image mapping for later use
        state["_chapter_image_map"] = chapter_image_map

        # Also pull scene photo placements
        scene_placements = state.get("_scene_photo_placements", [])
        for p in scene_placements:
            if isinstance(p, dict):
                path = p.get("path", "")
                if path and path not in all_image_paths:
                    all_image_paths.append(path)

        # 🔴 Fallback: if session has few images, scan storage for relevant files
        if len(all_image_paths) < 10:
            import glob as _glob
            storage_patterns = [
                (str(self.images_dir / '*.jpg'), 'photo'),
                (str(self.images_dir / '*.png'), 'photo'),
                (str(self.images_dir / '*.jpeg'), 'photo'),
                # Fallback: scan parent storage dir too
                (str(self.storage_dir / '*.jpg'), 'photo'),
                (str(self.storage_dir / '*.png'), 'photo'),
                (str(self.storage_dir / '*.jpeg'), 'photo'),
            ]
            for pattern, default_cat in storage_patterns:
                for f in _glob.glob(pattern):
                    if f not in all_image_paths:
                        fname = f.split('/')[-1]
                        # Skip PDF renders and tiny files
                        if 'pdf_' in fname or os.path.getsize(f) < 5000:
                            continue
                        all_image_paths.append(f)
                        # Use original filename for classification
                        original_names[f] = fname

        for fpath in all_image_paths:
            if not isinstance(fpath, str):
                continue
            fname = fpath.split('/')[-1] if '/' in fpath else fpath
            ext = fname.rsplit('.', 1)[-1].lower() if '.' in fname else ''
            if ext not in ('png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp'):
                continue

            # Use original filename for classification if available
            classify_name = original_names.get(fpath, fname)
            name_lower = classify_name.lower()
            # 🔴 Exclude raw PDF page renders (pdf_ prefix) — they're for OCR, not report images
            if name_lower.startswith('pdf_'):
                continue
            if any(k in name_lower for k in ['问卷', '调查', '统计', '签到']):
                images['survey'].append(fpath)
            elif any(k in name_lower for k in ['公告', '公示', '批文', '通知', '征收', '预公告']):
                images['announcement'].append(fpath)
            elif any(k in name_lower for k in ['评审', '意见', '专家签字', '专家']):
                images['review'].append(fpath)
            elif any(k in name_lower for k in ['地图', '红线', '规划', '位置图', '勘测', '勘界', '测定', '地形', '宗地', '示意']):
                images['map'].append(fpath)
            elif any(k in name_lower for k in ['照片', '现场', '座谈', '走访', '会议', '开会', '勘察', '村民', '临时用地', '微信图片', '开会现场']):
                images['photo'].append(fpath)
            else:
                images['other'].append(fpath)

        # 🔴 Auto-build chapter image map — each chapter picks only a few representative images
        if not state.get("_chapter_image_map"):
            ch_map = {}
            # Save a copy for appendix before consuming
            all_remaining = []
            for cat in ['map', 'announcement', 'photo', 'survey', 'review', 'other']:
                for p in images.get(cat, []):
                    fname = (original_names.get(p, p) if isinstance(original_names, dict) else p)
                    all_remaining.append({'path': p, 'caption': str(fname).split('/')[-1][:60]})

            # Helper: take up to N images from remaining pool
            def _take_from_pool(n, match_words):
                taken = []
                remaining = []
                for item in all_remaining[:]:
                    fname_lower = str(item.get('caption', '')).lower()
                    if any(kw in fname_lower for kw in match_words):
                        if len(taken) < n:
                            taken.append(item)
                            all_remaining.remove(item)
                return taken

            # Ch1: 1-2 maps only (位置示意图/红线图) — NO announcements
            ch_map[1] = _take_from_pool(2, ['勘测', '地图', '红线', '规划', '位置', '地形', '宗地', '勘界', '测定', '示意'])

            # Ch5: 2 expert review photos for risk factor identification section
            ch_map[5] = _take_from_pool(2, ['专家', '评审', '意见', '评估', '风险', '签字', '盖章'])

            # Ch3: 2-3 site/meeting photos + 1 notice
            ch_map[3] = _take_from_pool(3, ['现场', '勘察', '座谈', '走访', '会议', '照片', '公示', '村民']) + \
                        _take_from_pool(1, ['公告', '公示'])

            # Ch7: 2 photos
            ch_map[7] = _take_from_pool(2, ['现场', '勘察', '照片', '地块'])

            # Ch9: 1 expert review
            ch_map[9] = _take_from_pool(1, ['专家', '评审', '意见', '评估'])

            # Everything else → appendix (999)
            if all_remaining:
                ch_map[999] = all_remaining

            state["_chapter_image_map"] = ch_map

        return images

    def _get_survey_stats(self, state: dict) -> Dict[str, Any]:
        stats = {"total_surveys": 0, "support_count": 0, "oppose_count": 0,
                 "neutral_count": 0, "support_rate": 0}
        filled = state.get("filled_data", {})
        structured = state.get("structured_data", {})

        if filled.get("total_samples"):
            try: stats["total_surveys"] = int(filled["total_samples"])
            except: pass
        if filled.get("support_rate"):
            try: stats["support_rate"] = float(str(filled["support_rate"]).replace('%', ''))
            except: pass
        if filled.get("support_count"):
            try: stats["support_count"] = int(filled["support_count"])
            except: pass
        if filled.get("oppose_count"):
            try: stats["oppose_count"] = int(filled["oppose_count"])
            except: pass

        step_6 = structured.get("step_6", {})
        if isinstance(step_6, dict):
            if not stats["total_surveys"]:
                stats["total_surveys"] = step_6.get("total_samples", 0)
            if not stats["support_rate"]:
                stats["support_rate"] = step_6.get("support_rate", 0)

        if stats["support_count"] and not stats["total_surveys"]:
            stats["total_surveys"] = stats["support_count"] + stats["oppose_count"] + stats["neutral_count"]
        if stats["total_surveys"] and not stats["support_rate"]:
            stats["support_rate"] = round(stats["support_count"] / stats["total_surveys"] * 100, 1)

        return stats

    # ═══════════════════════════════════════════════════════════════
    # Helpers
    # ═══════════════════════════════════════════════════════════════

    def _add_heading(self, doc, text, level):
        h = doc.add_heading(text, level=level)
        # 🔴 淮安市稳评格式规范：
        # 1级标题: 宋体 小三(15pt) 加黑 居中
        # 2级标题: 宋体 四号(14pt) 加黑 靠左
        # 3级标题: 仿宋 四号(14pt) 加黑 靠左
        if level == 1:
            font, size, align = FONT_H1, Pt(15), WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            font, size, align = FONT_H2, Pt(14), WD_ALIGN_PARAGRAPH.LEFT
        else:
            font, size, align = FONT_H3, Pt(14), WD_ALIGN_PARAGRAPH.LEFT
        h.alignment = align
        for run in h.runs:
            run.font.name = font; run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            run.font.size = size
            run.font.bold = True  # 🔴 标题加黑
            run.font.color.rgb = RGBColor(0, 0, 0)  # 强制黑色，覆盖Word默认蓝色

    def _add_para(self, doc, text, bold=False, indent=False, font=FONT_BODY, size=Pt(14)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        r = p.add_run(text)
        r.font.name = font; r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        r.font.size = size; r.bold = bold


# Singleton
report_assembler = ReportAssembler()
