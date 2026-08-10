#!/usr/bin/env python3
"""Extract images from knowledge-base bidding docx files and store in DB.

For each image in the document:
1. Derive a humane name from the surrounding paragraph text (caption, heading, label)
2. Classify into a category (营业执照/资质证书/人员/业绩/设备/承诺函/...)
3. Extract binary data, resize if needed
4. Store in the asset_images table

Usage:
    cd backend && python extract_kb_images.py
    cd backend && python extract_kb_images.py --template-id 18   # single template
"""

import asyncio
import base64
import io
import logging
import re
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from PIL import Image as PILImage

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)

# ── Category keyword mapping ──
CATEGORY_MAP: Dict[str, List[str]] = {
    "营业执照": ["营业执照", "统一社会信用代码", "注册资本", "经营范围", "成立日期", "license"],
    "资质证书": ["资质", "证书", "测绘", "资格", "备案", "cert", "qualification"],
    "人员": ["人员", "身份证", "职称", "建造师", "执业", "团队", "简历", "personnel"],
    "业绩": ["业绩", "合同", "验收", "中标", "项目名称", "委托", "performance"],
    "承诺函": ["承诺", "声明", "函", "无重大", "中小", "信用", "独立", "commitment"],
    "法人证明": ["法定", "法人", "代表", "身份证明", "legal"],
    "授权委托": ["授权", "委托", "托书", "authorization"],
    "财务报告": ["财务", "审计", "报表", "资产", "负债", "financial"],
    "社保纳税": ["社保", "纳税", "税收", "完税", "保障", "tax", "social"],
    "设备": ["设备", "仪器", "车辆", "软件", "equipment"],
    "办公场所": ["办公", "场所", "租赁", "地点", "office"],
    "技术方案": ["技术", "方案", "路线", "流程", "架构", "技术路线图"],
    "公司信息": ["公司", "单位", "概况", "组织", "简介", "company"],
    "其他": [],
}

# ── Keywords that indicate stability-assessment content (NOT bidding) ──
_STABILITY_KEYWORDS = [
    "社会稳定", "稳评", "风险评估", "地质灾害", "征地补偿",
    "被征地", "农民社会保障", "补偿登记", "征地报批",
    "补偿安置方案", "土地征收", "勘测定界",
]

# ── Name extraction helpers ──

def _safe_text(para) -> str:
    """Return the paragraph text, stripping extra whitespace."""
    try:
        return para.text.strip()
    except Exception:
        return ""


def _find_caption_text(doc: Document, image_para_index: int) -> str:
    """Find the caption/text nearest to an image paragraph.

    Tries in order:
    1. The image paragraph's own text (if inline with text)
    2. The paragraph immediately before the image
    3. The preceding Heading 2/3 title
    4. A numbered list item above
    """
    paras = doc.paragraphs

    # 1. Own text (for inline images)
    own = _safe_text(paras[image_para_index]) if image_para_index < len(paras) else ""
    if own and len(own) >= 2:
        return own[:200]

    # 2. Previous paragraph
    if image_para_index > 0:
        prev = _safe_text(paras[image_para_index - 1])
        if prev:
            # Is it a figure caption? E.g. "图3-1 决策评估公示内容"
            if re.match(r'图\d+[-–—]\d+|照片|示意图|平面图|效果图|红线图|附图', prev):
                return prev[:200]
            # Short text is likely a label
            if 2 <= len(prev) <= 100:
                return prev[:200]

    # 3. Heading above (within 5 paragraphs)
    for offset in range(2, 7):
        idx = image_para_index - offset
        if idx < 0:
            break
        text = _safe_text(paras[idx])
        style = paras[idx].style.name if paras[idx].style else ""
        if style.startswith("Heading") and text:
            return text[:200]

    # 4. Any non-empty text within 3 paragraphs
    for offset in range(1, 4):
        idx = image_para_index - offset
        if idx < 0:
            break
        text = _safe_text(paras[idx])
        if text and len(text) >= 2 and len(text) <= 200:
            return text

    return f"图片_{image_para_index + 1}"


def _classify_image(name: str, context_text: str = "") -> str:
    """Classify image into a category based on its name and context."""
    combined = (name + " " + context_text).lower()
    for cat, keywords in CATEGORY_MAP.items():
        if cat == "其他":
            continue
        for kw in keywords:
            if kw.lower() in combined:
                return cat
    return "其他"


def _derive_search_keywords(name: str, category: str) -> str:
    """Generate search keywords from name and category."""
    kws = set()
    # Category-based
    for kw in CATEGORY_MAP.get(category, []):
        kws.add(kw)

    # Chinese 2-char substrings from name
    for i in range(len(name) - 1):
        chunk = name[i:i + 2]
        if re.search(r'[一-龥]', chunk):
            kws.add(chunk)

    return " ".join(sorted(kws, key=len, reverse=True)[:20])


def _get_image_size(data: bytes) -> Tuple[Optional[int], Optional[int]]:
    """Get image dimensions in pixels."""
    try:
        with PILImage.open(io.BytesIO(data)) as img:
            return img.size
    except Exception:
        return None, None


# ── Main extraction logic ──

async def extract_from_docx(
    docx_path: str,
    source_name: str = "",
    template_id: Optional[int] = None,
) -> int:
    """Extract all images from a docx file into the asset_images table.

    Returns the number of images extracted.
    """
    from app.database.knowledge_db import async_session
    from app.models.knowledge import AssetImage
    from sqlalchemy import select

    doc = Document(docx_path)
    total = 0

    # Scan for image paragraphs and their surrounding text
    # Use the docx internals: iterate over blip fills
    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue

        target = rel.target_ref
        image_data = rel.target_part.blob
        ext = Path(target).suffix.lower()
        mime_map = {".png": "image/png", ".jpg": "image/jpeg",
                    ".jpeg": "image/jpeg", ".gif": "image/gif",
                    ".bmp": "image/bmp", ".webp": "image/webp"}
        mime = mime_map.get(ext, "image/png")

        # Find which paragraph this image is in
        image_para_index = -1
        for i, para in enumerate(doc.paragraphs):
            blips = para._p.findall('.//' + qn('a:blip'))
            for b in blips:
                embed = b.get(qn('r:embed'))
                if embed == rel_id:
                    image_para_index = i
                    break
            if image_para_index >= 0:
                break

        # Derive name from surrounding text
        caption = _find_caption_text(doc, image_para_index) if image_para_index >= 0 else ""
        # Clean up the caption to make a good name
        name = caption.strip() if caption else f"image_{rel_id}"
        name = re.sub(r'\s+', ' ', name)[:500]

        # Get context from broader area (for classification)
        context = ""
        if image_para_index >= 0:
            start = max(0, image_para_index - 3)
            end = min(len(doc.paragraphs), image_para_index + 2)
            context = " ".join(_safe_text(doc.paragraphs[j]) for j in range(start, end))[:500]

        # Classify
        category = _classify_image(name, context)
        keywords = _derive_search_keywords(name, category)
        width, height = _get_image_size(image_data)

        # ── Skip stability-assessment images ──
        combined_check = (name + " " + context + " " + category).lower()
        if any(kw in combined_check for kw in _STABILITY_KEYWORDS):
            logger.info(f"  ⊘ 跳过(稳评内容): {name[:60]}")
            continue

        # Base64 encode for SQLite storage
        b64_data = base64.b64encode(image_data).decode("ascii")

        # Upsert into DB
        async with async_session() as db:
            existing = (
                await db.execute(
                    select(AssetImage).where(
                        AssetImage.image_name == name,
                        AssetImage.source_file == source_name,
                    )
                )
            ).scalar_one_or_none()

            if existing:
                existing.image_data = b64_data
                existing.mime_type = mime
                existing.width_px = width
                existing.height_px = height
                existing.category = category
                existing.search_keywords = keywords
                existing.is_active = True
                logger.info(f"  ↻ 更新: {name[:60]} ({category})")
            else:
                db.add(AssetImage(
                    image_name=name,
                    category=category,
                    image_data=b64_data,
                    mime_type=mime,
                    width_px=width,
                    height_px=height,
                    source_file=source_name,
                    source_template_id=template_id,
                    search_keywords=keywords,
                ))
                logger.info(f"  + 新增: {name[:60]} ({category})")
            await db.commit()

        total += 1

    logger.info(f"  ✅ 共提取 {total} 张图片")
    return total


async def main():
    """Main entry point: extract from all bidding templates in the knowledge base."""
    from app.database.knowledge_db import async_session
    from sqlalchemy import text

    # Parse CLI args
    target_id = None
    args = sys.argv[1:]
    for i, a in enumerate(args):
        if a == "--template-id" and i + 1 < len(args):
            target_id = int(args[i + 1])

    # Find bidding templates
    async with async_session() as db:
        if target_id:
            result = await db.execute(
                text("SELECT id, name, template_file_path FROM templates WHERE id=:tid"),
                {"tid": target_id},
            )
        else:
            result = await db.execute(
                text("SELECT id, name, template_file_path FROM templates WHERE is_active=1 AND (category LIKE '%招标%' OR category LIKE '%投标%' OR name LIKE '%投标%' OR name LIKE '%招标%')"),
            )
        templates = [{"id": r[0], "name": r[1], "path": r[2]} for r in result]

    if not templates:
        logger.warning("No bidding templates found!")
        # Fallback: try all active templates with docx files
        async with async_session() as db:
            result = await db.execute(
                text("SELECT id, name, template_file_path FROM templates WHERE is_active=1 AND template_file_path LIKE '%.docx'"),
            )
            templates = [{"id": r[0], "name": r[1], "path": r[2]} for r in result]

    logger.info(f"Found {len(templates)} templates to process")

    storage_dir = Path(__file__).resolve().parent / "storage"
    total_extracted = 0

    for tpl in templates:
        docx_path = storage_dir / tpl["path"]
        if not docx_path.exists():
            logger.warning(f"  ⚠️ 文件不存在: {docx_path}")
            continue

        logger.info(f"\n📄 处理: {tpl['name']} (id={tpl['id']})")
        logger.info(f"   文件: {docx_path} ({docx_path.stat().st_size / 1024:.0f} KB)")

        try:
            n = await extract_from_docx(
                str(docx_path),
                source_name=tpl["name"],
                template_id=tpl["id"],
            )
            total_extracted += n
        except Exception as e:
            logger.exception(f"  ❌ 提取失败: {e}")

    logger.info(f"\n{'='*50}")
    logger.info(f"总计: {total_extracted} 张图片入库")


if __name__ == "__main__":
    asyncio.run(main())
