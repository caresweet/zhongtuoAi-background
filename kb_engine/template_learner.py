"""template_learner.py — 模板学习

从 docx 模板/用例中学习：
1. 章节大纲（标题层级结构）
2. 每章子节列表
3. 每章表格结构（行/列/表头/内容概要）
4. 每章图片位置（所在段落文本/章节位置）
5. 每章写作指引（AI 生成：写什么/要点/必备数据/句式风格）
6. 整体写作风格备注

学习结果存入对应知识库的 templates + learned_chapters 表。
"""

import json
import re
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from .db import DualKB
from .llm import LLMClient


class TemplateLearner:
    """解析模板 docx → 提取大纲+表格+图片 → AI 生成逐章写作指引 → 存库。"""

    def __init__(self, db: DualKB, llm: Optional[LLMClient] = None):
        self.db = db
        self.llm = llm

    # ── 主流程 ──────────────────────────────────────────────────
    async def learn_from_docx(
        self, domain: str, file_path: str, name: str,
        doc_role: str = "template", category: str = "",
    ) -> dict:
        """从 docx 模板学习，返回 {template_id, outline, chapters}。"""
        doc = DocxDocument(file_path)
        outline = self._extract_outline(doc)
        style_notes = self._extract_style_notes(doc)

        # 提取表格结构 → 按章节归属
        table_map = self._extract_tables(doc)
        # 提取图片位置 → 按章节归属
        image_map = self._extract_images(doc)

        # 存模板
        template_id = self.db.upsert_template(
            domain, name=name, file_path=file_path,
            doc_role=doc_role, category=category,
            outline=outline, style_notes=style_notes,
        )

        # AI 生成逐章写作指引（包含表格/图片信息）
        chapters_with_guide = await self._generate_chapter_guides(
            domain, outline, style_notes, doc, table_map, image_map
        )

        self.db.save_learned_chapters(domain, template_id, chapters_with_guide)
        return {
            "template_id": template_id,
            "outline": outline,
            "style_notes": style_notes,
            "chapters": chapters_with_guide,
        }

    # ── 大纲提取 ────────────────────────────────────────────────
    def _extract_outline(self, doc) -> List[dict]:
        """从 docx 段落样式中提取标题大纲。"""
        chapters: List[dict] = []
        current_h1: Optional[dict] = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""

            level = self._heading_level(style_name, text)
            if level == 0:
                continue

            if level == 1:
                ch_no = self._parse_chapter_no(text)
                current_h1 = {
                    "chapter_no": ch_no or len(chapters) + 1,
                    "title": text,
                    "level": 1,
                    "subsections": [],
                }
                chapters.append(current_h1)
            elif level >= 2:
                sec_no = self._parse_section_no(text)
                # 推断缺失的 H1
                if sec_no and "." in sec_no:
                    inferred_ch = int(sec_no.split(".")[0])
                    existing = {c.get("chapter_no") for c in chapters}
                    if inferred_ch not in existing:
                        current_h1 = {
                            "chapter_no": inferred_ch,
                            "title": f"{inferred_ch}（推断标题）",
                            "level": 1,
                            "subsections": [],
                            "_inferred": True,
                        }
                        chapters.append(current_h1)
                if current_h1 is not None:
                    current_h1["subsections"].append({
                        "no": sec_no,
                        "title": text,
                    })

        return chapters

    def _heading_level(self, style_name: str, text: str) -> int:
        """判断段落标题级别。"""
        sn = style_name.lower()
        if "heading 1" in sn or style_name == "标题 1":
            return 1
        if "heading 2" in sn or style_name == "标题 2":
            return 2
        if "heading 3" in sn or style_name == "标题 3":
            return 3
        if "heading 4" in sn or style_name == "标题 4":
            return 4
        # 无样式但以数字章节号开头的
        if re.match(r"^\d{1,2}\s+\S", text) and not re.match(r"^\d+\.\d+", text):
            return 1
        if re.match(r"^\d+\.\d+\s+\S", text):
            return 2
        if re.match(r"^\d+\.\d+\.\d+\s+\S", text):
            return 3
        return 0

    def _parse_chapter_no(self, text: str) -> Optional[int]:
        m = re.match(r"^(\d{1,2})", text)
        return int(m.group(1)) if m else None

    def _parse_section_no(self, text: str) -> str:
        m = re.match(r"^(\d+(?:\.\d+)*)", text)
        return m.group(1) if m else ""

    # ── 表格提取 ────────────────────────────────────────────────
    def _extract_tables(self, doc) -> Dict[int, List[dict]]:
        """提取模板中的所有表格，按其所属章节归类。

        每张表格包含：
        - skeleton: 表头+行分类（固定骨架，不随项目变化）
        - fill_map: 每列的填充类型（fixed/project_data/calculated）
        - example_values: 模板原始值（仅参考，生成时不应照抄）
        """
        if not doc.tables:
            return {}

        # 遍历 body XML 元素，确定表格在文档流中的位置
        body = doc.element.body
        current_ch = 0  # 0 = 前置区（封面/目录等）
        tbl_global_idx = 0

        table_map: Dict[int, List[dict]] = {}

        for child in body:
            if child.tag == qn('w:p'):
                full_text = ''.join([
                    t.text for t in child.findall('.//' + qn('w:t'))
                    if t.text
                ]).strip()
                pPr = child.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        val = pStyle.get(qn('w:val'), '')
                        if 'Heading' in val or val.startswith('Heading') or '标题' in val:
                            m = re.match(r'^(\d{1,2})', full_text)
                            if m:
                                current_ch = int(m.group(1))
                if current_ch == 0:
                    m = re.match(r'^(\d{1,2})\s+\S', full_text)
                    if m and not re.match(r'^\d+\.\d+', full_text):
                        current_ch = int(m.group(1))
            elif child.tag == qn('w:tbl'):
                tbl_global_idx += 1
                tbl = doc.tables[tbl_global_idx - 1]
                rows_count = len(tbl.rows)
                cols_count = len(tbl.columns)

                # 提取完整表头（清洗换行）
                headers = []
                if tbl.rows:
                    headers = [
                        c.text.strip().replace('\n', ' ')[:50]
                        for c in tbl.rows[0].cells
                    ]

                # 提取骨架行分类（每行的左列标签）
                row_labels = []
                for row in tbl.rows[1:]:
                    first_cell = row.cells[0].text.strip().replace('\n', ' ')[:30]
                    row_labels.append(first_cell)

                # 提取所有行的完整内容（作为 example_values）
                example_rows = []
                for ri, row in enumerate(tbl.rows):
                    row_data = [
                        c.text.strip().replace('\n', ' ')[:40]
                        for c in row.cells
                    ]
                    example_rows.append(row_data)

                # 分析每列的填充类型
                fill_map = self._classify_columns(tbl, headers, example_rows)

                # 判断此表格属于哪个章节
                ch_no = current_ch if current_ch > 0 else 0
                if ch_no == 0:
                    ch_no = self._infer_table_chapter(headers, example_rows[:4])

                table_info = {
                    "idx": tbl_global_idx,
                    "rows": rows_count,
                    "cols": cols_count,
                    "headers": headers,
                    "row_labels": row_labels,          # 骨架：每行左列标签
                    "fill_map": fill_map,              # 每列填充类型
                    "example_values": example_rows[:4], # 模板原始值（仅参考）
                    "chapter_loc": ch_no,
                }

                if ch_no not in table_map:
                    table_map[ch_no] = []
                table_map[ch_no].append(table_info)

        return table_map

    def _classify_columns(self, tbl, headers: List[str], example_rows: List[List[str]]) -> List[dict]:
        """分析每列的填充类型。

        返回 [{col_idx, header, fill_type, data_source, description}]
        fill_type:
          - fixed: 法规/标准内容，不随项目变化（如评分标准、风险情形）
          - project_data: 必须根据用户资料填写（如事项名称、调查范围、反对率）
          - calculated: 从其他数据计算得出（如得分、综合评分）
          - structural: 行号/分类标签，保持模板结构不变
        """
        fill_map = []
        for col_idx, header in enumerate(headers):
            # 收集该列所有数据行的文本
            col_texts = []
            for row in tbl.rows[1:]:
                text = row.cells[col_idx].text.strip().replace('\n', ' ')[:40]
                col_texts.append(text)

            header_clean = header.strip()[:30] if header else f"列{col_idx+1}"
            empty_count = sum(1 for t in col_texts if not t)
            unique_values = set(t for t in col_texts if t)

            # 判断填充类型
            if col_idx == 0:
                # 第一列通常是行号或行分类标签 → structural
                fill_type = "structural"
                data_source = "模板骨架"
                description = "行分类/序号，保持模板原样"
            elif empty_count >= len(col_texts) * 0.7:
                # 大部分空 → project_data
                fill_type = "project_data"
                data_source = "用户资料"
                description = f"需从用户资料中提取{header_clean}相关数据"
            elif all(t in ['0', '0.00', '0/0', '', '√', '0.00%'] for t in col_texts if t):
                # 全是占位数字 → calculated 或 project_data
                if '得分' in header_clean or '评分' in header_clean or '权重' in header_clean:
                    fill_type = "calculated"
                    data_source = "计算/评分"
                    description = f"{header_clean}，根据评分标准和项目情况计算"
                elif '百分比' in header_clean or '率' in header_clean:
                    fill_type = "project_data"
                    data_source = "用户资料+计算"
                    description = f"{header_clean}，需从调查问卷数据计算"
                else:
                    fill_type = "project_data"
                    data_source = "用户资料"
                    description = f"需填入{header_clean}的实际值"
            elif any(kw in header_clean for kw in ['评分标准', '情形', '结论', '测评项目', '测评指标', '方法', '对象', '工作事项', '工作内容', '危险源', '危害说明', '现有措施', '环境因素', '环境影响']):
                # 法规/标准类内容 → fixed
                fill_type = "fixed"
                data_source = "模板/法规"
                description = f"{header_clean}，为标准/法规内容，保持模板原样"
            elif any(kw in header_clean for kw in ['名称', '单位', '范围', '项目', '合同', '时间', '人员', '姓名', '学历', '专业', '职称', '证书', '岗位']):
                # 项目相关 → project_data（需替换模板示例数据）
                fill_type = "project_data"
                data_source = "用户资料/公司固定资料"
                description = f"{header_clean}，需替换为实际项目/公司数据"
            elif any(kw in header_clean for kw in ['权重', '序号', '编号', '天数', '备注']):
                # 权重/编号 → structural 或 fixed
                fill_type = "fixed"
                data_source = "模板骨架"
                description = f"{header_clean}，数值固定，保持模板原样"
            else:
                # 默认：有内容但不确定 → 标注为 project_data
                # LLM 写作时应根据资料决定是否替换
                fill_type = "project_data"
                data_source = "用户资料"
                description = f"{header_clean}，需根据项目数据判断是否替换模板值"

            fill_map.append({
                "col_idx": col_idx,
                "header": header_clean,
                "fill_type": fill_type,
                "data_source": data_source,
                "description": description,
            })

        return fill_map

    def _infer_table_chapter(self, headers: List[str], content_rows: List[List[str]]) -> int:
        """从表头内容推断表格应属于哪个章节。"""
        all_text = " ".join(headers) + " " + " ".join(
            " ".join(row) for row in content_rows
        )
        # 关键词映射
        chapter_keywords = {
            1: ["决策名称", "决策单位", "征收范围", "征地位置", "基本概况", "事项名称"],
            2: ["评估过程", "评估方法", "评估依据", "社会稳定风险评估依据"],
            3: ["调查层级", "调查范围", "调查对象", "调查方法", "风险调查", "意见建议"],
            4: ["合法性", "合理性", "可行性", "可控性", "综合分析"],
            5: ["风险因素识别", "测评指标", "风险识别"],
            7: ["风险防范", "化解措施", "防范措施"],
            8: ["风险等级", "风险分析", "措施后", "反对率", "测评"],
            9: ["评估结论", "社会稳定等级"],
            10: ["应急预案", "编制目的", "保障措施"],
        }
        for ch, keywords in chapter_keywords.items():
            for kw in keywords:
                if kw in all_text:
                    return ch
        return 0

    def _build_para_chapter_map(self, doc) -> Dict[int, int]:
        """建立段落索引 → 章节号的映射。"""
        current_ch = 0
        para_map = {}
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else ""
            if "Heading 1" in style or style == "标题 1":
                m = re.match(r'^(\d{1,2})', text)
                if m:
                    current_ch = int(m.group(1))
            elif "Heading 2" in style or style == "标题 2":
                pass  # 保持当前章节
            elif re.match(r'^\d{1,2}\s+\S', text) and not re.match(r'^\d+\.\d+', text):
                m = re.match(r'^(\d{1,2})', text)
                if m:
                    current_ch = int(m.group(1))
            para_map[i] = current_ch
        return para_map

    # ── 图片提取 ────────────────────────────────────────────────
    def _extract_images(self, doc) -> Dict[int, List[dict]]:
        """提取模板中的所有图片位置，按章节归类。

        返回 {chapter_no: [{idx, para_text, caption, chapter_loc}]}
        """
        # 建立段落索引 → 章节映射
        para_ch_map = self._build_para_chapter_map(doc)
        current_ch = 0
        img_idx = 0
        image_map: Dict[int, List[dict]] = {}

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else ""

            # 更新当前章节号
            if "Heading 1" in style or style == "标题 1":
                m = re.match(r'^(\d{1,2})', text)
                if m:
                    current_ch = int(m.group(1))
            elif re.match(r'^\d{1,2}\s+\S', text) and not re.match(r'^\d+\.\d+', text):
                m = re.match(r'^(\d{1,2})', text)
                if m:
                    current_ch = int(m.group(1))

            # 检查是否有嵌入图片
            for run in para.runs:
                drawings = run._element.findall(qn('w:drawing'))
                if drawings:
                    img_idx += 1
                    # 获取题注（图片段落的文本或前后段落）
                    caption = text[:80] if text else ""
                    if not caption and i + 1 < len(doc.paragraphs):
                        next_text = doc.paragraphs[i + 1].text.strip()
                        if next_text and len(next_text) < 100:
                            caption = next_text
                    # 如果没题注，检查前一段
                    if not caption and i > 0:
                        prev_text = doc.paragraphs[i - 1].text.strip()
                        if prev_text and len(prev_text) < 100:
                            caption = prev_text

                    ch_no = current_ch if current_ch > 0 else para_ch_map.get(i, 0)

                    image_info = {
                        "idx": img_idx,
                        "para_text": text[:80],
                        "caption": caption,
                        "chapter_loc": ch_no,
                    }
                    if ch_no not in image_map:
                        image_map[ch_no] = []
                    image_map[ch_no].append(image_info)

        return image_map

    # ── 风格备注 ────────────────────────────────────────────────
    def _extract_style_notes(self, doc) -> str:
        """提取文档整体的格式/风格特征。"""
        notes = []
        notes.append(f"模板共 {len(doc.paragraphs)} 段落、{len(doc.tables)} 个表格")
        try:
            for p in doc.paragraphs[:5]:
                for run in p.runs:
                    if run.font.name:
                        notes.append(f"正文字体: {run.font.name}")
                        break
                break
        except Exception:
            pass
        h1_texts = [p.text.strip() for p in doc.paragraphs
                    if p.style and "Heading 1" in (p.style.name or "")]
        if h1_texts:
            notes.append(f"一级标题共 {len(h1_texts)} 个，编号风格: 「{h1_texts[0][:20]}...」")
        return "；".join(notes)

    # ── AI 逐章写作指引 ─────────────────────────────────────────
    async def _generate_chapter_guides(
        self, domain: str, outline: List[dict], style_notes: str, doc,
        table_map: Dict[int, List[dict]] = None,
        image_map: Dict[int, List[dict]] = None,
    ) -> List[dict]:
        """为每章生成写作指引，包含表格结构和图片位置。"""
        if not self.llm or not self.llm.available:
            return [self._rule_based_guide(ch, table_map, image_map) for ch in outline]

        chapter_texts = self._extract_chapter_bodies(doc, outline)

        results = []
        for ch in outline:
            ch_no = ch["chapter_no"]
            ch_title = ch["title"]
            subsections = ch.get("subsections", [])
            sub_titles = "；".join(s["title"] for s in subsections) if subsections else "无"
            template_body = chapter_texts.get(ch_no, "")[:3000]

            # 表格信息
            ch_tables = (table_map or {}).get(ch_no, [])
            table_desc = self._format_table_desc(ch_tables)

            # 图片信息
            ch_images = (image_map or {}).get(ch_no, [])
            image_desc = self._format_image_desc(ch_images)

            system = (
                "你是报告模板分析专家。根据模板中该章节的标题结构、表格结构和图片位置，"
                "生成一份简明的「写作指引」，帮助 AI 按模板风格撰写该章节内容。"
                "特别注意：如果该章有表格，写作指引必须说明需要填写哪些表格以及表格内容。"
                "如果该章有图片位置，说明需要插入什么类型的图片。"
            )
            prompt = (
                f"报告类型：{domain}\n"
                f"章节编号：第{ch_no}章\n"
                f"章节标题：{ch_title}\n"
                f"子节列表：{sub_titles}\n"
                f"模板表格结构：{table_desc}\n"
                f"模板图片位置：{image_desc}\n"
                f"模板原文片段（参考）：\n{template_body}\n\n"
                f"模板风格备注：{style_notes}\n\n"
                "请生成该章的写作指引，返回 JSON：\n"
                '{"writing_guide": "该章应写什么内容、分几个方面、每个方面要点、'
                '如有表格需说明填什么数据、如有图片需说明插入什么照片", '
                '"required_data": ["该章必须包含的数据字段1", "字段2"]}'
            )
            try:
                resp = await self.llm.chat_json(
                    messages=[{"role": "user", "content": prompt}],
                    system=system, max_tokens=1500, temperature=0.3,
                )
                guide = resp.get("writing_guide", "")
                req_data = resp.get("required_data", [])
            except Exception as e:
                guide = f"[写作指引生成失败: {e}]，请参照模板子节结构撰写"
                req_data = []

            # 合并前置区(ch_no=0)的表格到第1章
            actual_ch_tables = ch_tables
            if ch_no == 1 and (table_map or {}).get(0):
                actual_ch_tables = (table_map or {}).get(0, []) + ch_tables
            # 合并前置区的图片到第1章
            actual_ch_images = ch_images
            if ch_no == 1 and (image_map or {}).get(0):
                actual_ch_images = (image_map or {}).get(0, []) + ch_images

            results.append({
                "chapter_no": ch_no,
                "title": ch_title,
                "level": ch.get("level", 1),
                "subsections": subsections,
                "writing_guide": guide,
                "required_data": req_data,
                "tables": actual_ch_tables,
                "images": actual_ch_images,
            })
        return results

    def _format_table_desc(self, tables: List[dict]) -> str:
        """格式化表格结构描述，包含骨架和填充映射，供 AI 参考。

        关键：明确标注哪些列是固定的（照抄模板）、哪些需根据用户资料替换。
        """
        if not tables:
            return "本章无模板表格"
        parts = []
        for t in tables:
            header_str = " | ".join(t.get("headers", [])[:6])
            parts.append(f"表格{t['idx']}: {t['rows']}行×{t['cols']}列, 表头: [{header_str}]")

            # 填充映射：哪些列照抄模板，哪些需替换
            fill_map = t.get("fill_map", [])
            if fill_map:
                parts.append("  各列填充规则：")
                for fm in fill_map:
                    fill_type_cn = {
                        "fixed": "【固定】保持模板原样",
                        "structural": "【骨架】行标签/序号保持不变",
                        "project_data": "【需替换】根据用户资料填写",
                        "calculated": "【计算】根据评分标准计算",
                    }.get(fm["fill_type"], "【未知】")
                    parts.append(
                        f"    {fm['header']}: {fill_type_cn} — "
                        f"来源: {fm['data_source']}, 说明: {fm['description']}"
                    )

            # 骨架行标签
            row_labels = t.get("row_labels", [])
            if row_labels:
                parts.append(f"  行分类: {row_labels[:5]}")

            # 模板原始值（标注为示例，不应照抄）
            example = t.get("example_values", [])
            if example:
                parts.append(f"  模板示例值（仅供参考，project_data列不应照抄）: {example[:2]}")

        return "\n".join(parts)

    def _format_image_desc(self, images: List[dict]) -> str:
        """格式化图片位置描述。"""
        if not images:
            return "本章无模板图片"
        parts = []
        for img in images:
            parts.append(
                f"图片{img['idx']}: 题注=\"{img['caption']}\""
            )
        return "\n".join(parts)

    def _extract_chapter_bodies(self, doc, outline: List[dict]) -> Dict[int, str]:
        """按章节切分模板原文，供 AI 参考。"""
        bodies: Dict[int, List[str]] = {ch["chapter_no"]: [] for ch in outline}
        current_ch = None
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            sn = (para.style.name or "") if para.style else ""
            if "Heading 1" in sn or re.match(r"^\d{1,2}\s+\S", text):
                m = re.match(r"^(\d{1,2})", text)
                if m:
                    current_ch = int(m.group(1))
            if current_ch and current_ch in bodies:
                bodies[current_ch].append(text)
        return {k: "\n".join(v[:30]) for k, v in bodies.items()}

    def _rule_based_guide(self, ch: dict, table_map=None, image_map=None) -> dict:
        sub_titles = [s["title"] for s in ch.get("subsections", [])]
        guide = f"按模板结构撰写「{ch['title']}」，包含以下子节：" + "、".join(sub_titles)
        ch_no = ch.get("chapter_no", 0)
        ch_tables = (table_map or {}).get(ch_no, [])
        if ch_no == 1 and (table_map or {}).get(0):
            ch_tables = (table_map or {}).get(0, []) + ch_tables
        ch_images = (image_map or {}).get(ch_no, [])
        if ch_no == 1 and (image_map or {}).get(0):
            ch_images = (image_map or {}).get(0, []) + ch_images
        if ch_tables:
            guide += f"。本章含 {len(ch_tables)} 个模板表格，需按表头结构填写数据。"
        if ch_images:
            guide += f"。本章含 {len(ch_images)} 个图片位置，需插入相应照片。"
        return {
            "chapter_no": ch_no,
            "title": ch["title"],
            "level": ch.get("level", 1),
            "subsections": ch.get("subsections", []),
            "writing_guide": guide,
            "required_data": [],
            "tables": ch_tables,
            "images": ch_images,
        }
