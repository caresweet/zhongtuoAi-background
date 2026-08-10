#!/usr/bin/env python3
"""
报告填充器 — 将生成的内容写入模板.docx，保留原始格式

用法:
  python report_filler.py <模板路径> <数据JSON路径> <输出路径>
  
数据JSON格式:
{
  "basic_info": { ... },
  "survey_data": { ... },
  "ai_content": {
    "P319": "生成的第4.2.1段内容...",
    "P321": "生成的第4.2.2段内容...",
    ...
  }
}
"""
import os
import sys
import json
import shutil
import copy
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("需要 python-docx: pip install python-docx")
    sys.exit(1)

# ============================================================
# 段落映射：章节 → 需要填充的字段
# ============================================================

# 基于对模板的精确分析，建立章节→段落索引→填充内容的映射
SECTION_FILL_MAP = {
    # === 封面/评审表 ===
    "cover_title": {
        "paragraph_index": 2,  # P002: 报告标题
        "field": "decision_name",
        "format": "{decision_name}",
    },
    "cover_subtitle": {
        "paragraph_index": 3,  # P003: 报告副标题
        "field": None,
        "format": "土地征收决策社会稳定风险评估报告",
    },
    "review_name": {
        "paragraph_index": 152,  # 事项名称
        "field": "decision_name",
        "format": "事 项  名 称：{decision_name}",
    },
    "review_unit": {
        "paragraph_index": 153,  # 稳评责任单位
        "field": "responsibility_unit",
        "format": "稳评责任单位：{responsibility_unit}",
    },

    # === 1.1 决策名称 ===
    "sec1_1": {
        "paragraph_index": 182,
        "field": "decision_name",
        "format": '{decision_name}（以下简称为\u201c本决策\u201d）',
    },

    # === 1.2 决策单位 ===
    "sec1_2": {
        "paragraph_index": 184,
        "field": "responsibility_unit",
        "format": '稳评责任单位：{responsibility_unit}（以下简称为\u201c{short_unit}\u201d）',
    },

    # === 1.3 拟征地位置 ===
    "sec1_3": {
        "paragraph_index": 187,
        "field": None,
        "template": "{decision_name}拟征收土地位于{location_community}范围内。地理位置如图所示。",
    },

    # === 1.4 征收范围 ===
    "sec1_4": {
        "paragraph_index": 189,
        "field": None,
        "template": ('本次征收{num_plots}个地块，面积{area_hectares}公顷（{area_mu}亩），'
                   '土地权属为{location_community}农民集体所有，土地性质为{land_type}。'
                   '征地用途为{land_use}，符合《土地管理法》规定的'
                   '\u201c在土地利用总体规划确定的城镇建设用地范围内，经省级以上人民政府'
                   '批准由县级以上地方人民政府组织实施的成片开发建设需要用地的\u201d'
                   '可以依法实施征收的情形。地块上有青苗、沟渠，不占用永久基本农田。'),
    },

    # === 1.5 资金筹措 ===
    "sec1_5": {
        "paragraph_index": 191,
        "field": None,
        "template": "本决策资金暂按平均{fund_per_mu}万元/亩标准进行测算，则本决策相关资金约"
                   "{total_fund}万元（后续需根据进社保人员及青苗补偿等实际情况进行核算），"
                   "本决策资金全部由{fund_source}统筹。",
    },

    # === 1.6 实施周期 ===
    "sec1_6": {
        "paragraph_index": 194,
        "field": None,
        "template": "自拟征地公告公示之日（{bulletin_date}）起至征地工作结束。",
    },

    # === 2.1.1 评估过程 ===
    "sec2_1_1": {
        "paragraph_index": 200,
        "field": None,
        "template": "{commission_month}，{responsibility_unit}委托稳评第三方机构"
                   "江苏众拓项目代理咨询有限公司开展决策社会稳定风险评估工作，"
                   "确定本公司为决策事项社会稳定风险评估实施主体。随后成立了本决策"
                   "社会稳定风险评估工作组，稳评工作组制定了社会稳定风险评估工作方案。"
                   "方案经{responsibility_unit}审查通过后开始实施，"
                   "本公司开始对本决策事项社会稳定风险进行正式评估。",
    },

    # === 3.3 风险调查内容 ===
    "sec3_3": {
        "paragraph_index": 252,
        "field": None,
        "template": "评估工作组根据搜集的资料情况，于{survey_start}-{survey_end}，"
                   "在{street_name}的支持下，对决策各相关方进行了调查，风险调查包括以下主要内容：",
    },

    # === 3.4.1 调查时间 第一阶段 ===
    "sec3_4_1_phase1": {
        "paragraph_index": 263,
        "field": "phase1_dates",
        "format": "第一阶段（{phase1_dates}）",
    },
    # 第二阶段
    "sec3_4_1_phase2": {
        "paragraph_index": 264,
        "field": "phase2_dates",
        "format": "第二阶段（{phase2_dates}）",
    },
}

# C类AI生成段落的索引列表（需要AI生成内容的段落）
AI_PARAGRAPH_INDICES = [
    # 4.2 合理性分析
    319, 321, 323,
    # 4.3 可行性分析
    326, 328, 330,
    # 4.4 可控性分析
    333, 335, 337, 339,
    # 5.2 风险识别结果
    353, 355, 357, 359,
    # 7.2 风险防范措施
    393, 395, 397, 399, 401,
    # 9 评估结论
    420, 422, 424,
    # 10 应急预案
    427, 442,
]

# 附件标注段落
ATTACHMENT_INDICES = [494, 495, 496, 497, 498, 499, 500]
ATTACHMENT_LABELS = [
    "拟征收土地公告、征地范围图",
    "决策公示照片",
    "现场照片",
    "决策座谈会相关资料",
    "会议纪要",
    "调查问卷",
]


def smart_truncate(s: str, max_len: int) -> str:
    """智能截断字符串，在中文句号处截断"""
    if len(s) <= max_len:
        return s
    truncated = s[:max_len]
    last_period = max(truncated.rfind("。"), truncated.rfind("；"), truncated.rfind("，"))
    if last_period > max_len * 0.6:
        return truncated[:last_period + 1]
    return truncated


def fill_template(template_path: str, data: Dict, ai_content: Dict, output_path: str) -> str:
    """
    将数据填充到模板中，生成完整报告
    
    Args:
        template_path: 空模板路径
        data: 用户输入数据（basic_info + survey_data）
        ai_content: AI生成内容 {"paragraph_index": "content"}
        output_path: 输出路径
    
    Returns:
        输出文件路径
    """
    # 合并数据
    merged_data = {}
    if "basic_info" in data:
        merged_data.update(data["basic_info"])
    if "survey_data" in data:
        merged_data.update(data["survey_data"])
    
    doc = Document(template_path)
    total_filled = 0
    
    # 1. 填充精确映射的字段
    for sec_key, mapping in SECTION_FILL_MAP.items():
        pidx = mapping["paragraph_index"]
        if pidx >= len(doc.paragraphs):
            continue
        
        if "field" in mapping and mapping["field"] and mapping["field"] in merged_data:
            # 单个字段替换
            value = merged_data[mapping["field"]]
            text = mapping["format"].replace(f"{{{mapping['field']}}}", str(value))
            doc.paragraphs[pidx].text = ""
            run = doc.paragraphs[pidx].add_run(text)
            total_filled += 1
        elif "template" in mapping:
            # 多字段模板
            try:
                text = mapping["template"].format(**merged_data)
            except KeyError as e:
                text = mapping["template"] + f"\n[缺少字段: {e}]"
            doc.paragraphs[pidx].text = ""
            run = doc.paragraphs[pidx].add_run(text)
            total_filled += 1
    
    # 2. 填充AI生成内容
    for pidx_str, content in ai_content.items():
        try:
            pidx = int(pidx_str)
        except (ValueError, TypeError):
            continue
        if pidx < len(doc.paragraphs) and content:
            # 保留原有段落样式
            para = doc.paragraphs[pidx]
            style = para.style
            # 清除原有运行
            for run in para.runs:
                run.text = ""
            # 添加新内容
            if para.runs:
                para.runs[0].text = content
            else:
                para.add_run(content)
            total_filled += 1
    
    # 3. 处理附件标注
    for idx, label in zip(ATTACHMENT_INDICES, ATTACHMENT_LABELS):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            text = para.text.strip()
            if label in text or ("附件" in text and not "【需用户提供】" in text):
                para.text = ""
                para.add_run(f"【需用户提供：{label}】")
    
    # 4. 保存
    doc.save(output_path)
    print(f"报告已生成: {output_path}")
    print(f"共填充 {total_filled} 个字段 + {len(ai_content)} 个AI生成段落")
    
    return output_path


def load_data_json(json_path: str) -> Dict:
    """加载数据JSON"""
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def generate_empty_data_template(template_path: str) -> Dict:
    """生成空数据模板（供用户填写）"""
    from template_parser import analyze_template, generate_input_form
    
    result = analyze_template(template_path)
    form = generate_input_form(result["paragraph_map"])
    
    return {
        "template_path": template_path,
        "basic_info": {k: v.get("example", "") for k, v in form["basic_info"].items()},
        "survey_data": {k: v.get("example", "") for k, v in form["survey_data"].items()},
        "ai_paragraphs_to_generate": AI_PARAGRAPH_INDICES,
        "instructions": "请修改 basic_info 和 survey_data 中的值为实际数据，然后保存。",
    }


def main():
    if len(sys.argv) < 2:
        print("用法:")
        print("  python report_filler.py <模板.docx> <数据.json> <输出.docx>")
        print("  python report_filler.py --generate-template <模板.docx> > data_template.json")
        sys.exit(1)
    
    if sys.argv[1] == "--generate-template":
        if len(sys.argv) < 3:
            print("用法: python report_filler.py --generate-template <模板.docx>")
            sys.exit(1)
        template_path = sys.argv[2]
        template = generate_empty_data_template(template_path)
        print(json.dumps(template, ensure_ascii=False, indent=2))
        return
    
    if len(sys.argv) < 4:
        print("用法: python report_filler.py <模板.docx> <数据.json> <输出.docx>")
        sys.exit(1)
    
    template_path = sys.argv[1]
    data_json_path = sys.argv[2]
    output_path = sys.argv[3]
    
    data = load_data_json(data_json_path)
    ai_content = data.get("ai_content", {})
    
    fill_template(template_path, data, ai_content, output_path)


if __name__ == "__main__":
    main()
