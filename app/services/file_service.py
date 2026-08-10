"""File storage service — handles upload, retrieval, and cleanup of docx files."""
import os
import shutil
import uuid
from pathlib import Path
from typing import Optional, BinaryIO
from fastapi import UploadFile, HTTPException
import aiofiles

from app.config import settings


class FileService:
    """Manages file storage for templates, examples, and generated reports."""

    # Storage subdirectories
    TEMPLATES = "templates"
    EXAMPLES = "examples"
    GENERATED = "generated"
    TEMP = "temp"
    IMAGES = "images"
    KNOWLEDGE_DOCS = "knowledge_docs"

    # Allowed extensions for knowledge documents（支持 PDF/DOC/图片/文本等）
    KNOWLEDGE_EXTENSIONS = [
        ".docx", ".doc", ".pdf", ".txt", ".md",
        ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp",
        ".xlsx", ".xls", ".csv",
    ]

    @classmethod
    def _get_storage_path(cls, subdir: str, filename: Optional[str] = None) -> Path:
        """Get absolute path within storage directory."""
        path = settings.STORAGE_DIR / subdir
        path.mkdir(parents=True, exist_ok=True)
        if filename:
            path = path / filename
        return path

    @classmethod
    def _generate_filename(cls, original_name: str) -> str:
        """Generate a unique filename preserving the original extension."""
        ext = Path(original_name).suffix.lower()
        if ext not in settings.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {ext}。仅支持: {', '.join(settings.ALLOWED_EXTENSIONS)}",
            )
        return f"{uuid.uuid4().hex}{ext}"

    @classmethod
    async def save_upload(cls, upload_file: UploadFile, subdir: str) -> tuple[str, str]:
        """
        Save an uploaded file to storage.
        Returns (relative_path, original_filename).
        """
        if not upload_file.filename:
            raise HTTPException(status_code=400, detail="文件名不能为空")

        # Validate extension
        ext = Path(upload_file.filename).suffix.lower()
        if ext not in settings.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {ext}。仅支持: {', '.join(settings.ALLOWED_EXTENSIONS)}",
            )

        # Check size
        content = await upload_file.read()
        size_mb = len(content) / (1024 * 1024)
        if size_mb > settings.MAX_UPLOAD_SIZE_MB:
            raise HTTPException(
                status_code=400,
                detail=f"文件大小({size_mb:.1f}MB)超过限制({settings.MAX_UPLOAD_SIZE_MB}MB)",
            )

        # Generate unique filename and save
        unique_name = cls._generate_filename(upload_file.filename)
        file_path = cls._get_storage_path(subdir, unique_name)

        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)

        relative_path = f"{subdir}/{unique_name}"
        return relative_path, upload_file.filename

    @classmethod
    def get_absolute_path(cls, relative_path: str) -> Path:
        """Convert a relative storage path to absolute filesystem path."""
        return settings.STORAGE_DIR / relative_path

    # ── Content hash mapping for dedup ──

    _hash_index_path = settings.STORAGE_DIR / "file_hashes.json"

    @classmethod
    def _load_hash_index(cls) -> dict:
        import json
        if cls._hash_index_path.exists():
            try:
                with open(cls._hash_index_path) as f:
                    return json.load(f)
            except Exception:
                pass
        return {}

    @classmethod
    def _save_hash_index(cls, index: dict):
        import json
        cls._hash_index_path.parent.mkdir(parents=True, exist_ok=True)
        with open(cls._hash_index_path, "w") as f:
            json.dump(index, f)

    @classmethod
    def _find_existing_by_hash(cls, content_hash: str) -> str | None:
        """Return existing relative_path if hash exists, else None."""
        index = cls._load_hash_index()
        existing = index.get(content_hash)
        if existing:
            abs_path = settings.STORAGE_DIR / existing
            if abs_path.exists():
                return existing
        return None

    @classmethod
    def _store_hash_mapping(cls, content_hash: str, relative_path: str):
        index = cls._load_hash_index()
        index[content_hash] = relative_path
        cls._save_hash_index(index)

    @classmethod
    def delete_file(cls, relative_path: str) -> bool:
        """Delete a file from storage. Returns True if successful."""
        file_path = settings.STORAGE_DIR / relative_path
        if file_path.exists() and file_path.is_file():
            file_path.unlink()
            return True
        return False

    @classmethod
    def copy_template_for_generation(cls, template_path: str, session_id: str) -> Path:
        """
        Copy a template to the temp directory for filling.
        Returns the absolute path of the working copy.
        """
        src_path = settings.STORAGE_DIR / template_path
        if not src_path.exists():
            raise HTTPException(status_code=404, detail=f"模板文件不存在: {template_path}")

        temp_dir = cls._get_storage_path(cls.TEMP)
        dest_name = f"{session_id}.docx"
        dest_path = temp_dir / dest_name

        shutil.copy2(src_path, dest_path)
        return dest_path

    @classmethod
    def move_temp_to_generated(cls, temp_path: Path, session_id: str) -> str:
        """Move a filled temp file to the generated directory. Returns relative path."""
        gen_dir = cls._get_storage_path(cls.GENERATED)
        dest_name = f"{session_id}.docx"
        dest_path = gen_dir / dest_name

        shutil.move(str(temp_path), str(dest_path))

        return f"{cls.GENERATED}/{dest_name}"

    @classmethod
    def cleanup_temp(cls, session_id: str):
        """Remove temp files for a session."""
        temp_path = cls._get_storage_path(cls.TEMP, f"{session_id}.docx")
        if temp_path.exists():
            temp_path.unlink()

    @classmethod
    def get_file_size(cls, relative_path: str) -> Optional[int]:
        """Get file size in bytes, or None if file doesn't exist."""
        file_path = settings.STORAGE_DIR / relative_path
        if file_path.exists():
            return file_path.stat().st_size
        return None

    @classmethod
    async def save_image(cls, upload_file: UploadFile) -> dict:
        """
        Save an uploaded image file to the images subdirectory.
        Returns {relative_path, original_name, url, size_bytes}.
        """
        if not upload_file.filename:
            raise HTTPException(status_code=400, detail="文件名不能为空")

        ext = Path(upload_file.filename).suffix.lower()
        if ext not in settings.ALLOWED_IMAGE_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的图片类型: {ext}。仅支持: {', '.join(settings.ALLOWED_IMAGE_EXTENSIONS)}",
            )

        content = await upload_file.read()
        size_mb = len(content) / (1024 * 1024)
        if size_mb > settings.MAX_IMAGE_SIZE_MB:
            raise HTTPException(
                status_code=400,
                detail=f"图片大小({size_mb:.1f}MB)超过限制({settings.MAX_IMAGE_SIZE_MB}MB)",
            )

        unique_name = f"{uuid.uuid4().hex}{ext}"
        file_path = cls._get_storage_path(cls.IMAGES, unique_name)

        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)

        relative_path = f"{cls.IMAGES}/{unique_name}"
        return {
            "relative_path": relative_path,
            "original_name": upload_file.filename,
            "url": f"/api/v1/files/{relative_path}",
            "size_bytes": len(content),
        }

    @classmethod
    def get_file_url(cls, relative_path: str) -> str:
        """Get the API URL for accessing a stored file."""
        return f"/api/v1/files/{relative_path}"

    @classmethod
    async def save_knowledge_document(cls, upload_file: UploadFile) -> dict:
        """
        Save an uploaded knowledge document.
        Supports .docx, .pdf, .txt, .md, .png, .jpg, .xlsx and more.
        Returns {relative_path, original_name, file_size, file_type}.
        """
        if not upload_file.filename:
            raise HTTPException(status_code=400, detail="文件名不能为空")

        ext = Path(upload_file.filename).suffix.lower()
        if ext not in cls.KNOWLEDGE_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {ext}。仅支持: {', '.join(cls.KNOWLEDGE_EXTENSIONS)}",
            )

        content = await upload_file.read()
        size_mb = len(content) / (1024 * 1024)

        # Different size limits for images vs documents
        is_image = ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
        max_size = settings.MAX_IMAGE_SIZE_MB if is_image else settings.MAX_UPLOAD_SIZE_MB
        if size_mb > max_size:
            raise HTTPException(
                status_code=400,
                detail=f"文件大小({size_mb:.1f}MB)超过限制({max_size}MB)",
            )

        unique_name = f"{uuid.uuid4().hex}{ext}"
        file_path = cls._get_storage_path(cls.KNOWLEDGE_DOCS, unique_name)

        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)

        relative_path = f"{cls.KNOWLEDGE_DOCS}/{unique_name}"
        return {
            "relative_path": relative_path,
            "original_name": upload_file.filename,
            "file_size": len(content),
            "file_type": ext.lstrip("."),
        }

    @classmethod
    def read_text_file(cls, relative_path: str) -> str:
        """Read a text-based file (.txt, .md) from storage."""
        file_path = settings.STORAGE_DIR / relative_path
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="文件不存在")
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    @classmethod
    def extract_docx_text(cls, relative_path: str) -> str:
        """Extract text from .docx or .doc files in storage.

        For .doc (old Word format), auto-converts to .docx via:
        1. macOS: textutil (built-in)
        2. Linux: libreoffice --headless (installed in Dockerfile.prod)
        """
        from docx import Document
        file_path = settings.STORAGE_DIR / relative_path
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="文件不存在")

        # .doc → convert to .docx first
        if file_path.suffix.lower() == '.doc':
            import subprocess, tempfile, os, shutil
            converted = tempfile.mktemp(suffix='.docx')
            converters = [
                # macOS
                ['textutil', '-convert', 'docx', '-output', converted, str(file_path)],
                # Linux (libreoffice)
                ['libreoffice', '--headless', '--convert-to', 'docx',
                 '--outdir', os.path.dirname(converted), str(file_path)],
            ]
            converted_ok = False
            last_error = ""
            for cmd in converters:
                try:
                    if shutil.which(cmd[0]) is None:
                        continue
                    result = subprocess.run(
                        cmd, capture_output=True, text=True, timeout=60,
                    )
                    # libreoffice output filename differs: it creates <basename>.docx
                    if cmd[0] == 'libreoffice':
                        lo_output = os.path.join(
                            os.path.dirname(converted),
                            file_path.stem + '.docx',
                        )
                        if os.path.exists(lo_output):
                            if lo_output != converted:
                                os.rename(lo_output, converted)
                    if result.returncode == 0 and os.path.exists(converted):
                        converted_ok = True
                        break
                    else:
                        last_error = result.stderr or result.stdout or "unknown error"
                except FileNotFoundError:
                    continue
                except Exception as e:
                    last_error = str(e)
                    continue

            if not converted_ok:
                raise ValueError(
                    f"无法转换.doc文件（{file_path.name}），"
                    f"请使用.docx格式上传。错误: {last_error or '无可用的转换工具'}"
                )

            try:
                doc = Document(converted)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return "\n\n".join(paragraphs)
            finally:
                if os.path.exists(converted):
                    os.unlink(converted)

        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    @classmethod
    def extract_pdf_text_fast(cls, relative_path: str) -> str:
        """Fast PDF text extraction — pdfplumber only, NO OCR fallback.

        Used during upload for instant text availability. Scanned PDFs will
        return empty string; OCR happens in background task.
        """
        import logging
        logger = logging.getLogger(__name__)

        file_path = settings.STORAGE_DIR / relative_path
        if not file_path.exists():
            return ""

        all_pages_text = []
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        all_pages_text.append(page_text)
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table:
                                cells_list = [[c or "" for c in row] for row in table if row]
                                all_pages_text.append("\n".join(" | ".join(row) for row in cells_list))
        except Exception as e:
            logger.warning(f"pdfplumber failed for '{relative_path}': {e}")

        return "\n\n".join(all_pages_text).strip()

    @classmethod
    def extract_pdf_text(cls, relative_path: str) -> str:
        """Extract text from a PDF file in storage.

        Tries pdfplumber first, then falls back to automatic OCR via vision API
        for scanned/image-based PDFs (transparent to the user).
        """
        import logging
        logger = logging.getLogger(__name__)

        file_path = settings.STORAGE_DIR / relative_path
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="文件不存在")

        all_pages_text = []
        try:
            import pdfplumber
            with pdfplumber.open(str(file_path)) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        all_pages_text.append(page_text)

                    # Also extract tables if any
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            if table:
                                rows = []
                                for row in table:
                                    if row:
                                        cells = [c or "" for c in row]
                                        rows.append(" | ".join(cells))
                                if rows:
                                    all_pages_text.append("\n".join(rows))
        except Exception as e:
            logger.warning(f"pdfplumber failed for '{relative_path}': {e}")

        result = "\n\n".join(all_pages_text).strip()
        if result:
            return result

        # ---- No text extracted → scanned PDF, auto OCR via vision API ----
        logger.info(f"No text extracted from '{relative_path}', auto-OCR via vision API")
        return cls._ocr_pdf_pages(relative_path, file_path)

    @classmethod
    def _ocr_pdf_pages(cls, relative_path: str, file_path: Path) -> str:
        """Auto OCR a scanned PDF by rendering pages and using the vision model.

        Renders pages via pypdfium2, then calls the vision API (sync HTTP with retry).
        For large PDFs, samples strategically up to ~20 pages.
        """
        import logging
        import base64
        import httpx
        import time
        from io import BytesIO
        logger = logging.getLogger(__name__)

        try:
            import pypdfium2 as pdfium
        except ImportError:
            logger.error("pypdfium2 not available for PDF OCR")
            return ""

        pdf = None
        try:
            pdf = pdfium.PdfDocument(str(file_path))
            total_pages = len(pdf)

            if total_pages <= 10:
                page_indices = list(range(total_pages))
            elif total_pages <= 30:
                page_indices = list(range(total_pages))
            else:
                step = max(1, total_pages // 15)
                page_indices = list(range(0, 3))
                page_indices += list(range(step, total_pages - 2, step))
                page_indices += [total_pages - 2, total_pages - 1]
                page_indices = sorted(set(page_indices))

            page_indices = page_indices[:20]
            logger.info(
                f"PDF OCR: processing {len(page_indices)}/{total_pages} pages"
            )

            from app.config import settings as s
            # Config-driven vision model/endpoint (VISION_* → falls back to LLM_*).
            _vbase = (s.VISION_BASE_URL or s.LLM_BASE_URL or s.ANTHROPIC_BASE_URL
                      or "https://dashscope.aliyuncs.com/compatible-mode/v1").rstrip("/")
            if _vbase.endswith("/anthropic"):
                _vbase = _vbase[:-9]
            vision_api_url = f"{_vbase}/chat/completions"
            api_key = s.VISION_API_KEY or s.ANTHROPIC_API_KEY or s.EMBEDDING_API_KEY
            vision_model = s.VISION_MODEL or s.LLM_MODEL or "qwen-vl-max"

            all_text = []
            for idx in page_indices:
                page_text = ""
                for attempt in range(3):  # Retry up to 3 times
                    try:
                        page = pdf[idx]
                        bitmap = page.render(scale=2)
                        img = bitmap.to_pil()
                        buf = BytesIO()
                        img.save(buf, format="PNG")
                        img_b64 = base64.b64encode(buf.getvalue()).decode()

                        prompt = (
                            f"PDF第{idx + 1}/{total_pages}页扫描件。"
                            "请提取本页所有文字内容，保留标题、数字、日期、单位名称。"
                            "按原文格式输出，不要添加解释说明。"
                        )

                        with httpx.Client(timeout=90.0) as client:
                            resp = client.post(
                                vision_api_url,
                                headers={
                                    "Content-Type": "application/json",
                                    "Authorization": f"Bearer {api_key}",
                                },
                                json={
                                    "model": vision_model,
                                    "messages": [{
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": prompt},
                                            {
                                                "type": "image_url",
                                                "image_url": {
                                                    "url": f"data:image/png;base64,{img_b64}",
                                                },
                                            },
                                        ],
                                    }],
                                    "max_tokens": 1024,
                                    "temperature": 0.2,
                                },
                            )

                        if resp.status_code == 200:
                            data = resp.json()
                            choices = data.get("choices", [])
                            if choices:
                                page_text = choices[0].get("message", {}).get("content", "")
                                if page_text and len(page_text) > 10:
                                    break  # Success
                        elif resp.status_code == 429:
                            time.sleep(2 * (attempt + 1))  # Rate limit backoff
                        else:
                            logger.debug(
                                f"  Page {idx + 1} HTTP {resp.status_code} (attempt {attempt + 1})"
                            )
                            time.sleep(1)
                    except Exception as e:
                        logger.debug(f"  Page {idx + 1} attempt {attempt + 1}: {e}")
                        time.sleep(1)

                if page_text and len(page_text) > 10:
                    all_text.append((idx, page_text))
                    logger.info(f"  Page {idx + 1}: {len(page_text)} chars OCR'd")
                else:
                    logger.warning(f"  Page {idx + 1}: failed after retries")

            pdf.close()
            pdf = None

            all_text.sort(key=lambda x: x[0])
            result = "\n\n".join(
                f"[第{idx + 1}页]\n{text}" for idx, text in all_text
            )
            logger.info(f"Auto OCR complete: {len(result)} chars from {len(all_text)} pages")
            return result

        except Exception as e:
            logger.error(f"PDF OCR pipeline failed: {e}")
            return ""
        finally:
            if pdf:
                try:
                    pdf.close()
                except Exception:
                    pass

    @classmethod
    async def save_attachment(cls, upload_file) -> dict:
        """Save an uploaded attachment (image, PDF, Word, Excel, text, archive) to storage.

        Returns {relative_path, original_name, url, size_bytes, file_type}.
        """
        from fastapi import UploadFile

        if not upload_file.filename:
            raise HTTPException(status_code=400, detail="文件名不能为空")

        ext = Path(upload_file.filename).suffix.lower()
        if ext not in settings.ALLOWED_ATTACHMENT_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {ext}。仅支持: {', '.join(settings.ALLOWED_ATTACHMENT_EXTENSIONS)}",
            )

        content = await upload_file.read()
        size_mb = len(content) / (1024 * 1024)

        # Determine file type and size limit
        is_pdf = ext == ".pdf"
        is_image = ext in settings.ALLOWED_IMAGE_EXTENSIONS
        is_docx = ext in (".docx", ".doc")
        is_xlsx = ext in (".xlsx", ".xls", ".csv")
        is_pptx = ext in (".ppt", ".pptx")
        is_text = ext in (".txt", ".md")
        is_archive = ext in (".zip", ".rar")

        if is_image:
            max_mb = settings.MAX_IMAGE_SIZE_MB
        elif is_pdf:
            max_mb = settings.MAX_PDF_SIZE_MB
        else:
            max_mb = getattr(settings, 'MAX_ATTACHMENT_SIZE_MB', 100)

        if size_mb > max_mb:
            raise HTTPException(
                status_code=400,
                detail=f"文件大小({size_mb:.1f}MB)超过限制({max_mb}MB)",
            )

        # Dedup: compute content hash, skip save if identical file exists
        import hashlib
        content_hash = hashlib.sha256(content).hexdigest()
        existing_path = cls._find_existing_by_hash(content_hash)
        if existing_path:
            relative_path = existing_path
        else:
            # Use original filename (sanitized) instead of UUID for human readability
            base_name = Path(upload_file.filename).stem
            # Sanitize: keep Chinese chars, letters, digits, -, _, .
            safe_name = ''.join(c if c.isalnum() or c in '.-_ ' or '一' <= c <= '鿿' else '_' for c in base_name)
            safe_name = safe_name.strip()[:80]  # limit length
            if not safe_name:
                safe_name = "unnamed"
            # Add short hash suffix to avoid collisions
            short_hash = content_hash[:8]
            unique_name = f"{safe_name}_{short_hash}{ext}"
            file_path = cls._get_storage_path(cls.IMAGES, unique_name)
            async with aiofiles.open(file_path, "wb") as f:
                await f.write(content)
            cls._store_hash_mapping(content_hash, f"{cls.IMAGES}/{unique_name}")
            relative_path = f"{cls.IMAGES}/{unique_name}"

        # Determine file_type for frontend display
        if is_image:
            file_type = "image"
        elif is_pdf:
            file_type = "pdf"
        elif is_docx:
            file_type = "docx"
        elif is_xlsx:
            file_type = "xlsx"
        elif is_pptx:
            file_type = "pptx"
        elif is_text:
            file_type = "text"
        elif is_archive:
            file_type = "archive"
        else:
            file_type = "file"

        return {
            "relative_path": relative_path,
            "original_name": upload_file.filename,
            "url": f"/api/v1/files/{relative_path}",
            "size_bytes": len(content),
            "file_type": file_type,
        }

    @classmethod
    def extract_attachment_text(cls, relative_path: str) -> str:
        """Extract text from an attachment based on its file extension.

        - .pdf → extract_pdf_text()
        - .docx/.doc → extract_docx_text()
        - .txt/.md → read_text_file()
        - Images → empty string (handled by ImageAnalysisAgent)
        """
        ext = Path(relative_path).suffix.lower()

        if ext == ".pdf":
            return cls.extract_pdf_text(relative_path)
        elif ext in (".docx", ".doc"):
            return cls.extract_docx_text(relative_path)
        elif ext in (".txt", ".md"):
            return cls.read_text_file(relative_path)
        else:
            # Image files — no text to extract
            return ""


file_service = FileService()
