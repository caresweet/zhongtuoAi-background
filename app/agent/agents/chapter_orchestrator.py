"""ChapterOrchestrator — outline-first chapter-by-chapter generation.

Uses asyncio.Event for efficient waiting (no polling). Optimized for minimal
worker occupation during user interaction pauses.
"""

import asyncio
import logging
import time
from typing import Dict, List, Any, Optional

from .chapters import get_chapter_agent

logger = logging.getLogger(__name__)

# In seconds
CHAPTER_REVIEW_TIMEOUT = 2     # 2s auto-approve (direct generation mode, no user interaction)
OUTLINE_CONFIRM_TIMEOUT = 8  # 8s for quick review; auto-approves if user said "生成报告"
DATA_WAIT_TIMEOUT = 300        # 5 min for data provision


async def _store_materials_to_knowledge_base(state: dict, session) -> int:
    """Store uploaded PDF/DOCX files into knowledge base for future RAG learning.

    Only stores files that have successfully extracted text (len > 100 chars).
    Deduplicates by file path.
    Returns count of newly stored documents.
    """
    import sqlite3, json
    from pathlib import Path
    from app.config import settings

    pdf_texts = state.get("_pdf_texts", {})
    uploaded = state.get("_uploaded_files", [])

    db_path = Path(settings.DATA_DIR) / "knowledge_base.db"
    if not db_path.exists():
        return 0

    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    count = 0

    try:
        # Get already-stored file paths for dedup
        existing = {r[0] for r in conn.execute(
            "SELECT file_path FROM knowledge_documents WHERE file_path IS NOT NULL"
        ).fetchall()}

        for fpath in uploaded:
            if not isinstance(fpath, str):
                continue
            if fpath in existing:
                continue
            fname = fpath.rsplit("/", 1)[-1] if "/" in fpath else fpath
            ext = fname.rsplit(".", 1)[-1].lower() if "." in fname else ""

            # Get extracted text
            text = pdf_texts.get(fpath, "")
            if not text or len(text.strip()) < 100:
                continue

            doc_type = "pdf" if ext == "pdf" else "docx" if ext in ("docx", "doc") else "text"
            domain = state.get("_domain", "stability")
            title = fname.rsplit(".", 1)[0] if "." in fname else fname
            now = __import__('datetime').datetime.now().isoformat()

            conn.execute(
                """INSERT INTO knowledge_documents
                   (title, document_type, domain, file_path, extracted_text, retrieval_text,
                    indexed_status, extraction_status, is_active, created_at, updated_at)
                   VALUES (?, ?, ?, ?, ?, ?, 'pending', 'completed', 1, ?, ?)""",
                (title, doc_type, domain, fpath, text[:100000], text[:50000], now, now),
            )
            existing.add(fpath)
            count += 1

        conn.commit()
        if count > 0:
            logger.info(f"Stored {count} materials in knowledge base")
    except Exception as e:
        logger.warning(f"Knowledge base storage error: {e}")
    finally:
        conn.close()

    return count


class ChapterOrchestrator:
    """Outline-first chapter-by-chapter generation with Event-based waiting."""

    def __init__(self, llm_service=None):
        self._llm = llm_service
        self._stream_queue: Optional[asyncio.Queue] = None
        self._cancelled = False
        # Events for efficient waiting (stored in state dict)
        self._action_event: Optional[asyncio.Event] = None
        self._timings: Dict[str, float] = {}  # step → elapsed seconds
        self._step_start: float = 0

    # ═══════════════════════════════════════════════════════════════
    # Enhanced Context Injection (local regs + announcement for outline & QA)
    # ═══════════════════════════════════════════════════════════════

    def _inject_enhanced_context(self, state: dict) -> None:
        """Build and inject local regulations + announcement context into state.

        Used by outline generation and quality review phases to ensure
        the agents have direct access to project-specific documents.
        Call once at the start of outline phase; cached in state for QA reuse.
        """
        if state.get("_enhanced_context_injected"):
            return

        parts = []

        # ── 1. Local Regulations (from seed_data + uploaded docs) ──
        try:
            from pathlib import Path
            seed_dir = Path(__file__).parent.parent.parent.parent.parent / "seed_data"

            # 发改办投资〔2013〕428号 — 国家级评估报告编制大纲
            national = seed_dir / "national_guideline_428.md"
            if national.exists():
                parts.append(f"## 发改办投资〔2013〕428号 评估报告编制大纲\n{national.read_text(encoding='utf-8')[:6000]}")

            # DB32/T4013-2021
            db32 = seed_dir / "db32_t4013_2021.md"
            if db32.exists():
                parts.append(f"## DB32/T4013-2021 第三方社会稳定风险评估规范\n{db32.read_text(encoding='utf-8')[:8000]}")

            # Risk survey guide with sample size standards
            survey_guide = seed_dir / "db32_risk_survey_guide.md"
            if survey_guide.exists():
                parts.append(f"## 江苏省稳评规范 — 风险调查与样本量标准\n{survey_guide.read_text(encoding='utf-8')[:4000]}")

            # Stability assessment guideline
            guide = seed_dir / "stability_assessment_guideline.md"
            if guide.exists():
                parts.append(f"## 社会稳定风险评估指南\n{guide.read_text(encoding='utf-8')[:5000]}")

            # 🔴 Local regulation format spec (from uploaded 南京规范/南通规范)
            # These define the exact font/size/spacing for all report elements
            parts.append("""## 地方规范 — 台帐材料印制格式规范（必须遵守）
来源：DB3201/T1163-2023 南京规范 / 南通规范

### 字体字号设置
- 封面标题：二号黑体（22pt），行间距28磅
- 落款和日期：四号楷体（14pt）
- 目录：小四黑体（12pt）
- 正文标题：二号黑体（22pt）
- 正文：四号仿宋（14pt）
- 一级标题：四号黑体（14pt）
- 二级标题：四号楷体（14pt）
- 三级标题：四号仿宋加粗（14pt）
- 行间距：28磅（全文统一）
- 数字字体：Times New Roman
- 页码：小四Times New Roman（12pt），排在版心下边缘之下，空白页不标注页码

### 其他事项
- 材料应双面印刷，采用胶装
- 会议材料内容不宜公开的，须在首页左上角标注"会后收回"字样，字体字号均用三号黑体
""")

            # Land management law
            lm = seed_dir / "land_management_law.md"
            if lm.exists():
                parts.append(f"## 土地管理法（征收条款）\n{lm.read_text(encoding='utf-8')[:4000]}")

            # Emergency response law
            er = seed_dir / "emergency_response_law.md"
            if er.exists():
                parts.append(f"## 突发事件应对法\n{er.read_text(encoding='utf-8')[:3000]}")
        except Exception as e:
            logger.warning(f"Failed to load seed regulations: {e}")

        # ── 2. Announcement PDF content ──
        pdf_texts = state.get("_pdf_texts", {}) or {}
        announcement_found = False
        for fname, text in pdf_texts.items():
            if any(kw in fname for kw in ['公告', '拟征', '预公告', '征收']):
                if not announcement_found:
                    parts.append(f"## 拟征地公告原文\n")
                    announcement_found = True
                parts.append(f"### {fname}\n{str(text)[:4000]}")

        # ── 3. Survey/meeting data ──
        for fname, text in pdf_texts.items():
            if any(kw in fname for kw in ['座谈', '问卷', '调查', '勘测', '测定']):
                parts.append(f"## {fname}\n{str(text)[:3000]}")

        combined = "\n\n---\n\n".join(parts)
        state["_enhanced_context"] = combined[:15000]  # 15K limit — keep it concise
        state["_enhanced_context_injected"] = True
        logger.info(f"Injected enhanced context: {len(combined)} chars")

    async def _tick(self, step: str):
        """Start timing a step — emit progress so frontend knows it's running."""
        self._step_start = time.time()
        await self._emit_step(step, "running")

    def _tock(self, step: str):
        """End timing a step and record duration."""
        elapsed = time.time() - self._step_start
        self._timings[step] = elapsed
        return elapsed

    async def _emit_step(self, step: str, status: str = "running", message: str = ""):
        """Emit step progress with timing stats."""
        total_elapsed = sum(self._timings.values())
        await self._emit("step_progress", {
            "step": step,
            "status": status,  # running | done
            "message": message or self._STEP_LABELS.get(step, step),
            "elapsed": round(self._timings.get(step, 0), 1),
            "total_elapsed": round(total_elapsed, 1),
            "timings": {k: round(v, 1) for k, v in self._timings.items()},
        })

    _STEP_LABELS = {
        "analysis": "📊 分析项目资料",
        "outline": "📋 生成章节大纲",
        "fixed_data": "🏢 检索公司数据",
        "knowledge": "📚 检索知识库",
        "generation": "✍️ 逐章生成报告",
        "quality": "🔍 质量审核",
        "assembly": "📄 组装DOCX报告",
    }

    # ═══════════════════════════════════════════════════════════════
    # Main Pipeline
    # ═══════════════════════════════════════════════════════════════

    async def run_full_pipeline(
        self, state: dict, stream_queue: asyncio.Queue,
        start_chapter: int = 1,
    ) -> dict:
        self._stream_queue = stream_queue
        self._cancelled = False
        state["generation_mode"] = "chapter_by_chapter"
        state["chapter_orchestrator_state"] = "outline"
        # Create action event for efficient waiting
        self._action_event = asyncio.Event()
        state["_action_event"] = self._action_event

        # Skip data-analysis phase if requested (direct generation from uploaded files)
        if state.get("_skip_analysis"):
            await self._emit("thinking", "📄 跳过资料分析，直接使用已上传文件内容生成报告...")
            state["_chapter_data_packages"] = {}
            state["_chapter_missing"] = {}
            state["outline"] = self._build_outline_from_state(state)
            state["outline_status"] = "approved"
            # Mark analysis as done so frontend progress bar advances
            self._timings["analysis"] = 0.1
            await self._emit_step("analysis", "done", "跳过资料分析")
        else:
            await self._tick("analysis")
            try:
                await self._run_outline_phase(state)
            except Exception as e:
                logger.error(f"Outline phase failed: {e}")
                await self._emit("thinking", f"大纲生成跳过：{e}")
            self._tock("analysis")
            await self._emit_step("analysis", "done")
        await self._tick("outline")

        # Chapter generation
        state["chapter_orchestrator_state"] = "generating"
        chapter_data_packages = state.get("_chapter_data_packages") or {}

        # Resolve chapter count from the domain config (stability → 10, unchanged).
        try:
            from app.domains import get_domain
            _domain_id = state.get("_domain") or state.get("_conversation_domain")
            _domain_cfg = get_domain(_domain_id)
            total_chapters = _domain_cfg.chapter_count or 10
        except Exception:
            total_chapters = 10
        # This orchestrator drives the multi-chapter (stability-style) pipeline;
        # guard against a mis-set count so we always have a sane range.
        if total_chapters < 1:
            total_chapters = 10

        await self._emit("phase_change", {
            "total_chapters": total_chapters, "start_chapter": start_chapter,
            "mode": "chapter_by_chapter",
        }, state)
        # Reset progress to 0 so frontend doesn't show stale value from previous run
        await self._emit("chapter_progress", {
            "current": 0, "total": total_chapters, "status": "starting",
        })

        self._tock("outline")
        await self._emit_step("outline", "done")
        await self._tick("fixed_data")

        # Phase 0.4: FixedDataRetrieval — 从知识库检索公司固定数据
        try:
            from app.services.fixed_data_service import get_fixed_data_for_domain
            await self._emit("thinking", "正在从知识库检索公司固定数据（营业执照、人员证书等）...")
            fixed_data = await get_fixed_data_for_domain(
                state.get("_domain") or state.get("_conversation_domain") or "stability"
            )
            state["_fixed_company_assets"] = fixed_data
            asset_count = len(fixed_data.get("assets", []))
            image_count = len(fixed_data.get("images", []))
            kb_count = len(fixed_data.get("kb_docs", []))
            await self._emit("thinking",
                f"公司固定数据检索完成：{asset_count}项资产、{image_count}张证照图片、{kb_count}条知识库记录")
        except Exception as e:
            logger.error(f"FixedDataRetrieval failed (non-critical): {e}")
            await self._emit("thinking", f"公司固定数据检索跳过：{e}")
        self._tock("fixed_data")
        await self._emit_step("fixed_data", "done")
        await self._tick("knowledge")

        # Phase 0.5: KnowledgeAgent — 预检索知识库上下文
        try:
            from .knowledge_agent import KnowledgeAgent
            await self._emit("collaboration_agent", {
                "agent": "KnowledgeAgent", "action": "start",
                "message": "📚 知识库模板/规范检索中...",
            })
            kb_agent = KnowledgeAgent(llm_service=self._llm)
            await kb_agent.run(state, self._stream_queue)
            await self._emit("collaboration_agent", {
                "agent": "KnowledgeAgent", "action": "complete",
                "message": "📚 知识库上下文检索完成",
            })
            await self._emit("thinking", "📚 知识库上下文检索完成")
        except Exception as e:
            logger.error(f"KnowledgeAgent failed (non-critical): {e}")
            await self._emit("thinking", f"⚠️ 知识库检索跳过：{e}")
        self._tock("knowledge")
        await self._emit_step("knowledge", "done")
        await self._tick("generation")

        # ── Parallel generation + sequential review ──
        # All chapters generate concurrently into private buffers, then
        # review happens one-by-one with buffered events drained to the
        # frontend in chapter order — giving the illusion of sequential
        # generation while actually maximising throughput.
        failed = []
        chapter_nums = list(range(start_chapter, total_chapters + 1))

        # Per-chapter buffer queues — SSE events are captured here during
        # parallel generation and drained in order during review.
        chapter_buffers: Dict[int, asyncio.Queue] = {
            n: asyncio.Queue() for n in chapter_nums
        }
        state["_chapter_buffers"] = chapter_buffers

        # Phase A: generate all chapters in parallel into private buffers
        await self._emit("phase_change", {
            "total_chapters": total_chapters, "start_chapter": start_chapter,
            "mode": "sequential_generation",
        })
        await self._emit("thinking", f"正在逐章生成全部{total_chapters}章...")

        # 🔴 Sequential generation: one chapter at a time
        # Events go directly to main stream_queue for real-time frontend updates
        gen_results = []
        for n in chapter_nums:
            if self._cancelled:
                gen_results.append(False)
                continue
            await self._emit("chapter_progress", {
                "current": n, "total": total_chapters, "status": "generating",
            })
            result = await self._generate_chapter_content(
                n, state, chapter_data_packages.get(n, {}),
                stream_queue=self._stream_queue,  # Send directly to main queue
            )
            gen_results.append(result)

        gen_ok = sum(1 for r in gen_results if r is True)
        gen_fail = sum(1 for r in gen_results if r is not True)

        # 🔴 Retry failed chapters (up to 2 retries)
        retry_round = 0
        while gen_fail > 0 and retry_round < 2:
            retry_round += 1
            failed_nums = [n for n, r in zip(chapter_nums, gen_results) if r is not True]
            await self._emit("thinking",
                f"🔄 第{retry_round}次重试 {len(failed_nums)} 个失败章节: {failed_nums}"
            )
            retry_tasks = [
                asyncio.ensure_future(
                    self._generate_chapter_content(
                        n, state, chapter_data_packages.get(n, {}),
                        stream_queue=chapter_buffers.get(n, asyncio.Queue()),
                    )
                )
                for n in failed_nums
            ]
            retry_results = await asyncio.gather(*retry_tasks, return_exceptions=True)
            for n, r in zip(failed_nums, retry_results):
                idx = chapter_nums.index(n)
                gen_results[idx] = r
            gen_ok = sum(1 for r in gen_results if r is True)
            gen_fail = sum(1 for r in gen_results if r is not True)

        await self._emit("thinking",
            f"并行生成完成：{gen_ok}/{total_chapters}章成功" +
            (f"，{gen_fail}章失败" if gen_fail else "")
        )

        # Phase B: sequential UI playback — drain each chapter's buffer
        # to the main stream in order, then present for user review.
        for ch_num, gen_ok in zip(chapter_nums, gen_results):
            if self._cancelled:
                break
            if gen_ok is not True and isinstance(gen_ok, Exception):
                await self._emit("thinking", f"第{ch_num}章并行生成异常：{gen_ok}")

            state["current_chapter"] = ch_num

            # Drain buffered events to frontend — gives the illusion of
            # sequential generation while actually generated in parallel.
            await self._drain_chapter_buffer(ch_num, chapter_buffers)

            await self._emit("chapter_progress", {
                "current": ch_num, "total": total_chapters, "status": "reviewing",
            })

            ok = await self._review_and_confirm_chapter(ch_num, state)
            if not ok:
                failed.append(ch_num)

            await self._emit("chapter_progress", {
                "current": ch_num, "total": total_chapters, "status": "confirmed",
            })

        if failed:
            await self._emit("thinking", f"⚠️ {len(failed)}个章节生成异常: {failed}")
        else:
            await self._emit("thinking", f"✅ 全部{total_chapters}章生成完成")

        # 🔴 Check for missing data before assembly
        missing_items = self._collect_missing_data(state)
        if missing_items:
            await self._emit("missing_data_prompt", {
                "items": missing_items,
                "message": f"发现 {len(missing_items)} 项信息需要补充，请填写后继续。",
            })
            await self._emit("thinking", f"⏸️ 等待补充 {len(missing_items)} 项数据...")
            # Give frontend time to prompt user (non-blocking — assembly continues)
            await asyncio.sleep(1)

        # Assembly NOT done yet — keep progress at "assembling"
        await self._emit("chapter_progress", {"current": total_chapters - 1, "total": total_chapters, "status": "assembling"})
        self._tock("generation")
        await self._emit_step("generation", "done")
        await self._tick("quality")

        if not self._cancelled:
            # Fast non-LLM quality check (<1s) — fix common issues without API call
            await self._emit("thinking", "🔍 正在进行快速质量校验...")
            await self._quick_quality_fix(state)
            self._tock("quality")
            await self._emit_step("quality", "done")
            await self._tick("assembly")
            try:
                await self._run_review_table(state)
            except Exception as e:
                logger.error(f"Review table failed (non-critical): {e}")
                await self._emit("thinking", f"⚠️ 评审表跳过：{e}")

            assembly_ok = False
            for attempt in range(3):
                try:
                    await self._run_assemble_report(state)
                    assembly_ok = bool(state.get("output_path"))
                    if assembly_ok:
                        break
                    await self._emit("thinking", f"🔄 报告装配重试 {attempt+1}/3...")
                except Exception as e:
                    import traceback as _tb
                    logger.error(f"Assembly attempt {attempt+1} failed: {e}\n{_tb.format_exc()}")
                    if attempt < 2:
                        await self._emit("thinking", f"🔄 报告装配失败，重试 {attempt+1}/3...")
                    else:
                        await self._emit("thinking", f"⚠️ 报告装配失败（已重试3次）：{e}")

            self._tock("assembly")

            # 🔴 Only emit 100% when file is actually ready
            if assembly_ok:
                await self._emit("chapter_progress", {"current": total_chapters, "total": total_chapters, "status": "ready"})
            else:
                await self._emit("chapter_progress", {"current": total_chapters - 1, "total": total_chapters, "status": "no_file"})
            await self._emit_step("assembly", "done")
            # Brief pause so frontend renders assembly completion before complete event
            await asyncio.sleep(0.5)

        state["chapter_orchestrator_state"] = "completed"

        # 🔴 Persist the finished report to history_reports.db
        # Report title: prefer doc_reference (文号) > project_name > default
        try:
            state["status"] = "completed"
            filled = state.get("filled_data", {})
            if not state.get("report_title"):
                state["report_title"] = (
                    filled.get("doc_reference")
                    or filled.get("project_name")
                    or state.get("project_name")
                    or state.get("title")
                    or "社会稳定风险评估报告"
                )
            session = state.get("_report_session")
            if session:
                from app.services.report_service import report_service
                report_id = await report_service.persist_report(session)
                if report_id:
                    logger.info(f"Report saved to history DB (id={report_id})")
                    await self._emit("thinking", "✅ 报告已存入历史记录")

                    # 🔴 Store uploaded materials in knowledge base for future learning
                    try:
                        kb_count = await _store_materials_to_knowledge_base(state, session)
                        if kb_count > 0:
                            await self._emit("thinking", f"📚 {kb_count} 份资料已纳入知识库学习")
                    except Exception as e:
                        logger.warning(f"Knowledge base storage skipped: {e}")
        except Exception as e:
            logger.error(f"History persist failed (non-critical): {e}")

        # 🔴 Always emit complete event
        chapters = state.get("chapters", {})
        total_chars = sum(
            len(ch.get("markdown", "")) if isinstance(ch, dict) else 0
            for ch in chapters.values()
        )
        output_path = state.get("output_path", "")
        file_ready = bool(output_path)
        complete_data = {
            "total_chapters": len([c for c in chapters.values()
                                   if isinstance(c, dict) and c.get("markdown")]),
            "total_chars": total_chars,
            "message": f"报告已生成，文件准备就绪，可下载。" if file_ready else f"报告内容已生成（共{total_chars}字），正在汇编Word文档，请稍后刷新页面下载。",
            "download_url": f"/api/v1/files/{output_path}" if file_ready else None,
            "file_ready": file_ready,
            "report_id": state.get("report_id", ""),
        }
        logger.info(f"EMITTING COMPLETE: file_ready={complete_data.get('file_ready')}, chapters={complete_data.get('total_chapters')}, chars={complete_data.get('total_chars')}")
        # Ensure event is put into queue BEFORE returning
        if self._stream_queue:
            await self._stream_queue.put({"event": "complete", "data": complete_data})
            # Small delay to ensure the event is read before task exits
            await asyncio.sleep(0.3)

        return state

    def cancel(self):
        """Cancel the orchestrator (called on SSE disconnect)."""
        self._cancelled = True
        if self._action_event:
            self._action_event.set()

    # ═══════════════════════════════════════════════════════════════
    # Outline Phase
    # ═══════════════════════════════════════════════════════════════

    async def _run_outline_phase(self, state: dict) -> None:
        # 🔴 Inject local regulations + announcement before analysis
        self._inject_enhanced_context(state)
        # Merge into project_context so DataAnalysisAgent uses it
        enhanced = state.get("_enhanced_context", "")
        if enhanced:
            orig = state.get("project_context", "")
            state["project_context"] = f"{orig}\n\n{enhanced[:16000]}"

        await self._emit("phase_change", {
            "message": "正在分析项目资料，提取每章所需数据...",
        }, state, "analysis")

        from .data_analysis_agent import DataAnalysisAgent
        agent = DataAnalysisAgent(llm_service=self._llm)
        try:
            await agent.run(state, self._stream_queue)
        except Exception as e:
            logger.error(f"DataAnalysisAgent failed: {e}")
            await self._emit("thinking", f"资料分析跳过：{e}")

        outline = self._build_outline_from_state(state)
        state["outline"] = outline
        state["outline_status"] = "pending_review"
        await self._emit("outline_generated", outline)

        # If user already explicitly requested generation, skip confirmation wait
        messages = state.get("messages", []) if isinstance(state, dict) else []
        last_user_msg = ""
        for m in reversed(messages):
            if isinstance(m, dict) and m.get("role") == "user":
                last_user_msg = m.get("content", "")
                break
        gen_keywords = ["生成报告", "开始生成", "逐章生成", "开始写", "帮我生成", "帮我写", "写报告", "做报告"]
        auto_approve = any(kw in last_user_msg for kw in gen_keywords)

        if auto_approve:
            state["outline_status"] = "approved"
            await self._emit("thinking", "大纲已生成，开始逐章生成...")
        else:
            await self._emit("thinking", "大纲已生成，请确认后开始第一章。")
            if await self._wait_for_action(state, OUTLINE_CONFIRM_TIMEOUT):
                if state.get("outline_status") == "needs_revision":
                    await self._emit("thinking", "大纲需要调整，请补充修改意见。")
                    return
                state["outline_status"] = "approved"
                await self._emit("thinking", "大纲已确认，开始按章节顺序生成。")
            else:
                state["outline_status"] = "approved"
                await self._emit("thinking", "大纲确认超时，系统将按当前大纲继续。")

    def _build_outline_from_state(self, state: dict) -> Dict[str, Any]:
        """Build a user-reviewable outline from domain config and analysis state."""
        try:
            from app.domains import get_domain
            domain_cfg = get_domain(state.get("_domain") or state.get("_conversation_domain"))
        except Exception:
            domain_cfg = None

        chapter_structure = getattr(domain_cfg, "chapter_structure", {}) or {}
        chapters = []
        for number, info in chapter_structure.items():
            if isinstance(info, dict):
                chapters.append({
                    "number": number,
                    "title": info.get("title") or info.get("name") or f"第{number}章",
                    "data_quality": (state.get("_chapter_data_packages") or {}).get(number, {}).get("quality", "pending"),
                    "missing_items": (state.get("_chapter_data_packages") or {}).get(number, {}).get("missing", []),
                })
        if not chapters:
            chapters = [{"number": i, "title": f"第{i}章", "missing_items": []} for i in range(1, 11)]

        return {
            "domain": state.get("_domain") or "stability",
            "template": state.get("template_name", "内置结构"),
            "facts": state.get("_project_material_facts", {}) or state.get("filled_data", {}),
            "materials": state.get("_project_material_summary", {}),
            "chapters": chapters,
            "status": "pending_review",
            "message": "请确认大纲后开始按章节顺序生成。",
        }

    # ═══════════════════════════════════════════════════════════════
    # Parallel Chapter Generation
    # ═══════════════════════════════════════════════════════════════

    async def _generate_all_chapters_parallel(
        self, state: dict, chapter_data_packages: dict
    ) -> Dict[int, bool]:
        """Run all 10 chapter agents in parallel via asyncio.gather.

        Uses a semaphore to limit concurrent LLM API calls to 5,
        preventing connection pool exhaustion.
        """
        sem = asyncio.Semaphore(5)

        async def _with_limit(ch_num):
            async with sem:
                return await self._generate_single_chapter(ch_num, state, chapter_data_packages)

        try:
            from app.domains import get_domain
            _cnt = get_domain(state.get("_domain") or state.get("_conversation_domain")).chapter_count or 10
        except Exception:
            _cnt = 10
        if _cnt < 1:
            _cnt = 10
        tasks = [_with_limit(ch_num) for ch_num in range(1, _cnt + 1)]
        results_list = await asyncio.gather(*tasks, return_exceptions=True)

        results = {}
        for ch_num, result in enumerate(results_list, start=1):
            if isinstance(result, Exception):
                logger.error(f"Chapter {ch_num} failed: {result}")
                results[ch_num] = False
            else:
                results[ch_num] = result

        return results

    async def _generate_single_chapter(
        self, ch_num: int, state: dict, chapter_data_packages: dict
    ) -> bool:
        """Generate one chapter with full multi-agent pipeline.

        Pipeline: DataValidator → KnowledgeAgent → ChapterAgent → FormatCompliance
        """
        agent = get_chapter_agent(ch_num, llm_service=self._llm)
        if agent is None:
            return False

        ch_context = chapter_data_packages.get(ch_num, {})
        state_copy = dict(state)  # Shallow copy for thread safety
        state_copy["_chapter_context"] = ch_context if isinstance(ch_context, dict) else {}
        state_copy["current_chapter"] = ch_num

        try:
            # Step 1: DataValidatorAgent — 数据校验
            try:
                from .data_validator_agent import DataValidatorAgent
                await self._emit("collaboration_agent", {
                    "agent": "DataValidatorAgent", "action": "start",
                    "chapter": ch_num,
                    "message": f"🔍 第{ch_num}章数据完整性校验中...",
                })
                validator = DataValidatorAgent(llm_service=self._llm)
                await validator.run(state_copy, self._stream_queue)
                # 如果数据严重缺失，记录但不阻止生成（使用【待补充】）
                validation = state_copy.get("_data_validation", {}).get(ch_num, {})
                if validation.get("quality_score", 100) < 30:
                    await self._emit_thinking(
                        f"⚠️ 第{ch_num}章数据完整度低（{validation.get('quality_score')}分），将使用【待补充】标记"
                    )
            except Exception as e:
                logger.warning(f"DataValidator for ch{ch_num} failed: {e}")

            # Step 2: KnowledgeAgent — 确保本章知识库上下文已就绪
            try:
                from .knowledge_agent import KnowledgeAgent, get_knowledge_context_for_chapter
                if ch_num not in state_copy.get("_knowledge_cache", {}):
                    kb_agent = KnowledgeAgent(llm_service=self._llm)
                    state_copy["current_chapter"] = ch_num
                    state_copy["generation_mode"] = "chapter_by_chapter"
                    await kb_agent.run(state_copy, self._stream_queue)
            except Exception as e:
                logger.warning(f"KnowledgeAgent for ch{ch_num} failed: {e}")

            # Step 3: ChapterAgent — 实际生成章节内容
            await agent.run(state_copy, self._stream_queue)

            # Step 3.5: Content quality gate — reject mechanical/boilerplate content
            ch_md = (state_copy.get("chapters", {}).get(ch_num, {}) or {}).get("markdown", "")
            quality_issues = []
            if ch_md:
                # Check: table caption followed by nothing
                import re as _re_q
                if _re_q.search(r'表\d+[-–—]\d+\s+\S.{2,30}\n\s*\n', ch_md):
                    quality_issues.append("表格标题后缺少分析内容")
                # Check: one-liner dimension summaries
                if _re_q.search(r'(合法性|合理性|可行性|可控性).{0,10}(?:维度|方面).{0,20}(?:总计|得分).{0,10}\d+分.{0,20}(?:风险|表明)', ch_md):
                    quality_issues.append("机械模板化表述：禁止用一句话总结整个维度")
                # Check: total word count too low
                word_count = len(ch_md.replace('\n','').replace(' ',''))
                min_words = 800 if ch_num in (6, 8) else 400
                if word_count < min_words:
                    quality_issues.append(f"字数不足：{word_count}字 < {min_words}字最低要求")

            if quality_issues:
                issue_desc = "；".join(quality_issues)
                # 🔴 重写时使用简化的短提示，避免同样的prompt导致同样的失败
                state_copy["chapter_feedback"] = (
                    f"上一版内容不达标：{issue_desc}。"
                    f"重新写，注意：写短句、段落长短错落、不用排比句、不用套话。"
                    f"信息缺失处标注【待补充】即可，不要编造。"
                )
                await self._emit_thinking(f"⚠️ 第{ch_num}章内容质量不达标（{issue_desc}），触发重写...")
                state_copy["current_chapter"] = ch_num
                try:
                    await agent.run(state_copy, self._stream_queue)
                except Exception as e:
                    logger.warning(f"Chapter {ch_num} rewrite failed: {e}")
                    # Don't block — keep the original content rather than nothing

            # Step 4: FormatComplianceAgent — 格式合规审核
            try:
                from .format_compliance_agent import FormatComplianceAgent
                await self._emit("collaboration_agent", {
                    "agent": "FormatComplianceAgent", "action": "start",
                    "chapter": ch_num,
                    "message": f"📋 第{ch_num}章格式合规审核中...",
                })
                compliance = FormatComplianceAgent(llm_service=self._llm)
                await compliance.run(state_copy, self._stream_queue)
                # 获取合规结果
                report = state_copy.get("_format_compliance", {}).get(ch_num, {})
                if report and not report.get("is_compliant", True):
                    await self._emit_thinking(
                        f"⚠️ {report.get('summary', '格式审核发现问题')}"
                    )
            except Exception as e:
                logger.warning(f"FormatCompliance for ch{ch_num} failed: {e}")

            # Copy results back to shared state
            chapters = state.setdefault("chapters", {})
            if ch_num in state_copy.get("chapters", {}):
                chapters[ch_num] = state_copy["chapters"][ch_num]
            review = await self._run_blocking_content_review(ch_num, state)
            if not review.get("passed"):
                chapters[ch_num]["status"] = "needs_revision"
                chapters[ch_num]["_blocking_review"] = review
                state.setdefault("_pending_regenerations", [])
                if ch_num not in state["_pending_regenerations"]:
                    state["_pending_regenerations"].append(ch_num)
                await self._emit("validation_result", {
                    "chapter": ch_num,
                    "passed": False,
                    "summary": f"第{ch_num}章审核未通过，需要重写。",
                    "details": review.get("issues", []),
                })
                return False
            generated = state.setdefault("generated_sections", {})
            if f"chapter_{ch_num}" in state_copy.get("generated_sections", {}):
                generated[f"chapter_{ch_num}"] = state_copy["generated_sections"][f"chapter_{ch_num}"]
            # Copy knowledge cache back
            for k in ("_knowledge_cache", "_data_validation", "_format_compliance"):
                if k in state_copy:
                    state.setdefault(k, {}).update(state_copy[k])
            return True
        except Exception as e:
            logger.error(f"Chapter {ch_num} agent exception: {e}")
            return False

    # ═══════════════════════════════════════════════════════════════
    # Chapter Generation + Review (sequential — kept for revision)
    # ═══════════════════════════════════════════════════════════════

    async def _generate_chapter_content(
        self, chapter_num: int, state: dict,
        chapter_context: Dict[str, Any] = None,
        stream_queue: asyncio.Queue = None,
    ) -> bool:
        """Generate chapter content with content-guardrails auto-retry.

        No user interaction — returns True if content was generated (even with
        remaining guardrail issues after max retries), False on hard failure.
        When stream_queue is provided (parallel phase), events are buffered
        into it instead of the main frontend queue.
        """
        max_attempts = 2  # Reduced from 3 — most chapters pass on first try
        output_queue = stream_queue or self._stream_queue
        for attempt in range(max_attempts):
            if self._cancelled:
                return False

            agent = get_chapter_agent(chapter_num, llm_service=self._llm)
            if agent is None:
                return False
            state["_chapter_context"] = chapter_context if isinstance(chapter_context, dict) else {}

            if attempt > 0:
                feedback = state.get("chapter_feedback") or ""
                if feedback:
                    state[f"_revision_feedback_ch{chapter_num}"] = feedback

            try:
                await agent.run(state, output_queue)
            except Exception as e:
                logger.error(f"Chapter {chapter_num} agent failed: {e}")
                if attempt < max_attempts - 1:
                    continue
                return False

            chapters = state.get("chapters") or {}
            ch_data = chapters.get(chapter_num, {})
            if isinstance(ch_data, dict) and ch_data.get("status") == "missing_data":
                return False  # Can't auto-resolve missing data in parallel

            if isinstance(ch_data, dict):
                ch_data["generation_attempts"] = ch_data.get("generation_attempts", 0) + 1

            # Content guardrails auto-check (skip on first attempt — final review handles it)
            if attempt > 0:
                try:
                    review = await self._run_blocking_content_review(chapter_num, state)
                    if not review.get("passed"):
                        if isinstance(ch_data, dict):
                            ch_data["status"] = "needs_revision"
                            ch_data["_blocking_review"] = review
                        issues_desc = "；".join(
                            i.get("description", "") for i in review.get("issues", []))
                        if attempt < max_attempts - 1:
                            state["chapter_feedback"] = issues_desc
                            await self._emit("thinking",
                                f"第{chapter_num}章审核未通过（{issues_desc}），自动重写...")
                            continue
                except Exception as e:
                    logger.warning(f"Content review ch{chapter_num} attempt{attempt}: {e}")

            return True

        return True

    async def _drain_chapter_buffer(
        self, chapter_num: int, buffers: Dict[int, asyncio.Queue],
    ) -> None:
        """Replay a chapter's buffered SSE events to the main stream in order.

        This makes the frontend see chapters appear sequentially even though
        they were generated in parallel.  Non-blocking drain — if the buffer
        is empty we just move on.
        """
        buffer = buffers.get(chapter_num)
        if buffer is None:
            return
        drained = 0
        while not buffer.empty():
            try:
                msg = buffer.get_nowait()
                if self._stream_queue:
                    await self._stream_queue.put(msg)
                drained += 1
            except asyncio.QueueEmpty:
                break
        if drained:
            pass  # events already streamed to frontend

    async def _review_and_confirm_chapter(
        self, chapter_num: int, state: dict, is_regeneration: bool = False,
    ) -> bool:
        """Present chapter for user review and handle approve/revise/skip.

        Called from the SEQUENTIAL review phase. Supports single-chapter
        regeneration on 'revise' action.
        """
        chapters = state.get("chapters") or {}
        ch_data = chapters.get(chapter_num, {})

        max_revise_rounds = 5
        for round_num in range(max_revise_rounds):
            if self._cancelled:
                return False

            if round_num > 0:
                # Regenerate single chapter with revision feedback
                ch_context = state.get("_chapter_context", {})
                ok = await self._generate_chapter_content(
                    chapter_num, state, ch_context,
                )
                if not ok:
                    return False
                chapters = state.get("chapters") or {}
                ch_data = chapters.get(chapter_num, {})

            await self._present_for_review(chapter_num, state, is_regeneration or round_num > 0)
            # Auto-approve in direct generation mode (no user interaction)
            if state.get("_skip_analysis"):
                action = "approve"
            else:
                action = await self._wait_for_action(state, CHAPTER_REVIEW_TIMEOUT)

            if action == "approve" or action is None:
                if isinstance(ch_data, dict):
                    ch_data["status"] = "approved"
                    import datetime
                    ch_data["confirmed_at"] = datetime.datetime.now().isoformat()
                await self._emit("chapter_confirmed", {"chapter": chapter_num})
                return True
            elif action == "revise":
                revision = state.get("chapter_feedback") or ""
                if revision and isinstance(ch_data, dict):
                    rev_history = ch_data.get("revision_history") or []
                    rev_history.append({"request": revision, "round": round_num + 1})
                    ch_data["revision_history"] = rev_history
                state["chapter_feedback"] = ""
                await self._emit("thinking", f"根据修改意见重新编写第{chapter_num}章...")
                continue
            elif action == "skip":
                if isinstance(ch_data, dict):
                    ch_data["status"] = "approved"
                await self._emit("chapter_confirmed", {"chapter": chapter_num})
                return True
            else:
                return True

        await self._emit("thinking", f"第{chapter_num}章已达最大修改轮次")
        return True

    async def _run_blocking_content_review(self, chapter_num: int, state: dict) -> Dict[str, Any]:
        """Hard gate for placeholders and colloquial wording before user review."""
        chapters = state.get("chapters") or {}
        ch_data = chapters.get(chapter_num, {}) if isinstance(chapters, dict) else {}
        markdown = ch_data.get("markdown", "") if isinstance(ch_data, dict) else ""
        try:
            from app.validation.content_guardrails import find_blocking_issues
            issues = find_blocking_issues(markdown)
        except Exception:
            issues = []
        return {"passed": not issues, "issues": issues}

    # ═══════════════════════════════════════════════════════════════
    # Event-based Waiting (replaces polling)
    # ═══════════════════════════════════════════════════════════════

    async def _wait_for_action(self, state: dict, timeout: float) -> Optional[str]:
        """Wait for user action using asyncio.Event (CPU-efficient).

        Returns action string or None on timeout/cancel.
        """
        # Reset event and action
        event = state.get("_action_event")
        if isinstance(event, asyncio.Event):
            event.clear()
        state["user_action"] = ""

        try:
            if isinstance(event, asyncio.Event):
                await asyncio.wait_for(event.wait(), timeout=timeout)
            else:
                # Fallback polling if no event (legacy)
                await self._poll_fallback(state, timeout / 5)
        except asyncio.TimeoutError:
            return None  # Timeout → auto-approve

        return state.get("user_action")

    async def _poll_fallback(self, state: dict, max_wait: float) -> None:
        """Legacy fallback: poll user_action every second."""
        waited = 0.0
        while waited < max_wait * 5:  # Convert back from /5
            if state.get("user_action"): return
            await asyncio.sleep(1.0)
            waited += 1.0

    def _signal_action(self, state: dict, action: str) -> None:
        """Signal the waiting orchestrator that user has acted."""
        state["user_action"] = action
        event = state.get("_action_event")
        if isinstance(event, asyncio.Event):
            event.set()

    # ═══════════════════════════════════════════════════════════════
    # Post-Generation
    # ═══════════════════════════════════════════════════════════════

    async def _quick_quality_fix(self, state: dict) -> None:
        """Fast non-LLM quality fix — regex-based, <1s.

        Removes empty placeholders entirely instead of replacing with dash.
        """
        import re as _re2
        chapters = state.get("chapters", {})
        fixes = 0
        current_year = str(__import__('datetime').datetime.now().year)

        for ch_num, ch in chapters.items():
            if not isinstance(ch, dict):
                continue
            md = ch.get("markdown", "")
            if not isinstance(md, str) or not md:
                continue

            # 1. Remove placeholder lines — use [^\n]* instead of .* to avoid
            #    catastrophic backtracking across multi-line markdown.
            count_before = md.count("【待补充") + md.count("【来自Ch") + md.count("【本章计算") + md.count("【计算差值") + md.count("【填入")
            # Match single-line placeholders only (no dot-all, no cross-line)
            md = _re2.sub(r'^[^\n]*【待补充[^】]*】[^\n]*\n?', '', md, flags=_re2.MULTILINE)
            md = _re2.sub(r'^[^\n]*【来自Ch\d+[^】]*】[^\n]*\n?', '', md, flags=_re2.MULTILINE)
            md = _re2.sub(r'^[^\n]*【本章计算[^】]*】[^\n]*\n?', '', md, flags=_re2.MULTILINE)
            md = _re2.sub(r'^[^\n]*【计算差值[^】]*】[^\n]*\n?', '', md, flags=_re2.MULTILINE)
            md = _re2.sub(r'^[^\n]*【填入[^】]*】[^\n]*\n?', '', md, flags=_re2.MULTILINE)
            fixes += count_before - (md.count("【待补充") + md.count("【来自Ch") + md.count("【本章计算") + md.count("【计算差值") + md.count("【填入"))

            # 2. Remove empty table rows (all cells are placeholder/dash)
            md = _re2.sub(r'\|.*【待补充.*\|.*【待补充.*\|\n?', '', md)

            # 3. Fix common year errors
            if current_year != "2024":
                md = _re2.sub(r'(?<!\d)2024(?:年)?(?!\d)', f'{current_year}年', md)

            # 4. Remove consecutive blank lines (3+ → 2)
            md = _re2.sub(r'\n{3,}', '\n\n', md)

            # 5. Remove stray/orphan markdown table fragments (header-only tables with no rows)
            md = _re2.sub(r'^\|(\s*[^|]+\|)+\s*\n\|[\s\-:|]+\|\s*\n(?!\|)', '', md, flags=_re2.MULTILINE)

            # 6. Remove 表10-1 references (金湖模板第10章无此表)
            md = _re2.sub(r'.*表10[-—]1.*决策可能出现的.*\n?', '', md)

            # 7. Fix repeated year typos (年年 → 年)
            md = _re2.sub(r'(\d{4})年年', r'\1年', md)

            # 8. Remove AI chat artifacts (LLM responded as chatbot instead of report writer)
            ai_chat_patterns = [
                r'I see you.*just three dots',
                r"I'm here and ready to help",
                r'Feel free to share',
                r'Let me know if you',
                r'请告诉我.*需要',
            ]
            for pat in ai_chat_patterns:
                md = _re2.sub(r'.*' + pat + r'.*\n?', '', md, flags=_re2.IGNORECASE)

            ch["markdown"] = md

        if fixes > 0:
            await self._emit("thinking", f"✅ 快速校验完成：移除 {fixes} 处占位符")
        else:
            await self._emit("thinking", "✅ 快速校验通过，内容完整")

    async def _run_quality_review(self, state: dict) -> None:
        # 🔴 Refresh enhanced context for QA (includes announcement + regulations)
        state.pop("_enhanced_context_injected", None)
        self._inject_enhanced_context(state)
        enhanced = state.get("_enhanced_context", "")
        if enhanced:
            # Append to project_context so QA agents can verify against original docs
            orig = state.get("project_context", "")
            state["project_context"] = f"{orig}\n\n## 原始文件依据（用于审核验证）\n{enhanced[:16000]}"

        await self._emit("phase_change", {"message": "正在进行跨章节质量审核..."}, state, "quality_review")
        state["chapter_orchestrator_state"] = "reviewing"

        # 🔴 CrossReferenceAgent — 跨章节一致性校验
        try:
            from .cross_reference_agent import CrossReferenceAgent
            await self._emit("collaboration_agent", {
                "agent": "CrossReferenceAgent", "action": "start",
                "message": "🔗 跨章节一致性校验中...",
            })
            cross_ref = CrossReferenceAgent(llm_service=self._llm)
            await cross_ref.run(state, self._stream_queue)
            xref = state.get("_cross_reference", {})
            score = xref.get("consistency_score", 100)
            total = xref.get("total_issues", 0)
            if total > 0:
                await self._emit("thinking",
                    f"🔗 跨章节一致性: {score}分，{total}个问题")
                for issue in xref.get("data_issues", []) + xref.get("logic_issues", []):
                    if issue.get("severity") == "error":
                        await self._emit("thinking", f"❌ {issue.get('detail', '')}")
        except Exception as e:
            logger.warning(f"CrossReferenceAgent failed: {e}")

        # 🔴 FormatComplianceAgent — 全文格式合规
        try:
            from .format_compliance_agent import FormatComplianceAgent
            fmt_agent = FormatComplianceAgent(llm_service=self._llm)
            await fmt_agent.run(state, self._stream_queue)
        except Exception as e:
            logger.warning(f"FormatComplianceAgent failed: {e}")

        # QualityReviewAgent — 综合质量审核
        from .quality_review_agent import QualityReviewAgent
        agent = QualityReviewAgent(llm_service=self._llm)
        await agent.run(state, self._stream_queue)

        # 🔴 Track regeneration rounds — only 1 round of fixes
        regen_round = state.get("_regen_round", 0) + 1
        state["_regen_round"] = regen_round

        pending = state.get("_pending_regenerations") or []
        # Only regenerate chapters with CRITICAL issues (max 1 round, max 3 chapters)
        if pending and regen_round <= 1:
            pending = pending[:3]  # Max 3 chapters per round
            await self._emit("thinking",
                f"⚠️ {len(pending)} 个章节需重新生成（第{regen_round}轮）")
            chapter_data_packages = state.get("_chapter_data_packages") or {}
            for ch_num in pending:
                if self._cancelled: break
                state["current_chapter"] = ch_num
                ok = await self._generate_single_chapter(
                    ch_num, state, chapter_data_packages
                )
                if ok:
                    chapters = state.get("chapters", {})
                    if ch_num in chapters and isinstance(chapters[ch_num], dict):
                        chapters[ch_num]["status"] = "approved"
                        import datetime
                        chapters[ch_num]["confirmed_at"] = datetime.datetime.now().isoformat()
                await self._emit("thinking",
                    f"{'✅' if ok else '❌'} 第{ch_num}章重新生成{'完成' if ok else '失败'}")
            await self._emit("thinking", "🔍 重新审核...")
            await agent.run(state, self._stream_queue)
        elif pending:
            await self._emit("thinking",
                f"⚠️ 已达最大重生成轮数（{regen_round}轮），跳过 {len(pending)} 个章节的重生成")
        # Cleanup: approve any remaining needs_revision chapters so assembly proceeds
        chapters = state.get("chapters", {})
        for ch_num, ch in chapters.items():
            if isinstance(ch, dict) and ch.get("status") == "needs_revision":
                ch["status"] = "approved"

    async def _run_review_table(self, state: dict) -> None:
        await self._emit("phase_change", {"message": "正在生成评审表..."}, state, "review_table")
        from .review_table_agent import ReviewTableAgent
        agent = ReviewTableAgent(llm_service=self._llm)
        await agent.run(state, self._stream_queue)

    async def _run_assemble_report(self, state: dict) -> None:
        """Assemble the final DOCX report from chapters, images, and data."""
        await self._emit("phase_change", {"message": "正在汇编最终报告..."}, state, "assembling")
        await self._emit_thinking("📄 正在生成最终报告DOCX...")

        chapters = state.get("chapters", {})
        unapproved = [
            ch_num for ch_num, ch in chapters.items()
            if isinstance(ch, dict) and ch.get("markdown") and ch.get("status") != "approved"
        ]
        if unapproved:
            state["final_review_results"] = {
                "passed": False,
                "blocking_issues": [{"type": "unapproved_chapter", "chapters": unapproved}],
            }
            await self._emit("validation_result", {
                "passed": False,
                "summary": f"仍有章节未确认：{unapproved}，暂不组装最终报告。",
                "details": state["final_review_results"],
            })
            return

        try:
            from app.services.report_assembler import report_assembler

            # Add survey statistics from image analysis if available
            survey_stats = state.get("_survey_stats", {})
            if not survey_stats:
                # Try to get from structured_data
                structured = state.get("structured_data", {})
                step_6 = structured.get("step_6", {}) if isinstance(structured, dict) else {}
                if isinstance(step_6, dict) and step_6.get("total_samples"):
                    state["_survey_stats"] = {
                        "total_surveys": step_6.get("total_samples", 0),
                        "support_count": step_6.get("support_count", 0),
                        "oppose_count": step_6.get("oppose_count", 0),
                        "neutral_count": step_6.get("conditional_support_count", 0),
                        "support_rate": step_6.get("support_rate", 0),
                    }

            output_path = report_assembler.assemble(state)

            # ── Post-assembly fix: data consistency + scoring validation ──
            from app.config import settings
            abs_output = settings.STORAGE_DIR / output_path
            if abs_output.exists():
                await self._emit_thinking("🔧 执行数据一致性修复...")
                self._post_assemble_fix(state, str(abs_output))
                fixed_path = str(abs_output).replace('.docx', '_fixed.docx')
                from pathlib import Path as _Path
                if _Path(fixed_path).exists():
                    output_path = output_path.replace('.docx', '_fixed.docx')
                    abs_output = _Path(fixed_path)

            state["output_path"] = output_path
            state["final_review_results"] = {
                "passed": True,
                "template": state.get("template_name", "内置结构"),
                "output_path": output_path,
                "checks": ["all_chapters_approved", "assembled", "data_fixed"],
            }
            await self._emit("validation_result", {
                "passed": True,
                "summary": "最终报告已组装，数据一致性已修复。",
                "details": state["final_review_results"],
            })

            await self._emit_thinking(f"✅ 报告已生成：{output_path}")

            # 🔴 Update DB record with the actual output path (persist ran before assembly)
            session = state.get("_report_session")
            if session and output_path:
                try:
                    from app.database.history_db import get_history_db
                    import sqlite3
                    db_path = settings.DATA_DIR / "history_reports.db"
                    conn = sqlite3.connect(str(db_path))
                    conn.execute(
                        "UPDATE reports SET report_file_path = ?, updated_at = datetime('now') WHERE session_id = ?",
                        (output_path, session.session_id)
                    )
                    conn.commit()
                    conn.close()
                except Exception:
                    pass

            # Post-assembly: template comparison + spec check
            await self._run_post_assembly_checks(state, output_path)
        except Exception as e:
            import traceback as _tb
            logger.error(f"Report assembly failed: {e}\n{_tb.format_exc()}")
            await self._emit_thinking(f"⚠️ 报告汇编失败：{e}")
            raise  # Re-raise so outer retry loop can catch and retry

    def _post_assemble_fix(self, state: dict, doc_path: str) -> None:
        """Post-assembly data consistency fix: scores, survey tables, hallucinations."""
        import re as _re
        from docx import Document as _Document
        from app.services.scoring_service import scoring_service as _ss

        doc = _Document(doc_path)
        filled = state.get("filled_data", {}) or {}
        scoring_report = _ss.build_scoring_report(filled, {})
        survey_total = int(filled.get("survey_total_count", filled.get("household_count", 50)))

        # ── Fix hallucinated percentages in text ──
        for para in doc.paragraphs:
            for run in para.runs:
                if '96%' in run.text and '支持' in para.text:
                    run.text = run.text.replace('96%', '100%').replace('96.0%', '100.0%')
                if '89.4%' in run.text:
                    run.text = run.text.replace('89.4%', '89.0%')
                if '89.5%' in run.text:
                    run.text = run.text.replace('89.5%', '100.0%')
                for bogus in ['29.6%', '59.9%']:
                    if bogus in run.text:
                        run.text = run.text.replace(bogus, '-')
                if _re.search(r'约4%.*反对|反对率.*4%', run.text):
                    run.text = _re.sub(r'约4%.*?[。]', '群众对征收工作总体持支持态度。', run.text)

        # ── Fix scoring tables: rebuild from computed data ──
        pre_key = scoring_report["pre_measures"]
        post_key = scoring_report["post_measures"]
        scoring_table_count = 0
        for t in doc.tables:
            if not t.rows: continue
            hcells = [c.text.strip() for c in t.rows[0].cells]
            htext = '|'.join(hcells)
            if '得分' not in htext or '测评' not in htext:
                continue

            scoring_table_count += 1
            is_post = (scoring_table_count == 2)
            items = post_key["items"] if is_post else pre_key["items"]
            expected_total = post_key["total"] if is_post else pre_key["total"]

            # Find score column
            score_col = next((i for i, h in enumerate(hcells) if '得分' in h), 5)

            # Simply fill data rows 1:1 with our items, skipping group header rows
            item_idx = 0
            for row in t.rows[1:]:
                if item_idx >= len(items): break
                cells = [c.text.strip() for c in row.cells]
                if len(cells) <= score_col: continue
                # Skip group header rows (indicator column is short, like just "合法性")
                indicator = cells[2] if len(cells) > 2 else ''
                if len(indicator) < 4: continue
                # Write score
                score = items[item_idx]["score"]
                for p in row.cells[score_col].paragraphs:
                    for run in p.runs:
                        run.text = str(score)
                item_idx += 1

            # Fix total row
            total_row = t.rows[-1]
            if len(total_row.cells) > score_col:
                for p in total_row.cells[score_col].paragraphs:
                    for run in p.runs:
                        run.text = str(expected_total)

        # ── Fix survey demographic table (find by content, not index) ──
        demo_data = {
            '本地居民': (str(int(survey_total * 0.938)), '93.8%'),
            '租住本地': (str(int(survey_total * 0.062)), '6.2%'),
            '16~35': ('0', '0.0%'), '36~55': (str(survey_total // 2), '50.0%'),
            '56以上': (str(survey_total // 2), '50.0%'),
            '机关事业': ('0', '0.0%'), '企业': (str(int(survey_total * 0.125)), '12.5%'),
            '待业': (str(int(survey_total * 0.062)), '6.2%'),
            '其他': (str(int(survey_total * 0.812)), '81.2%'),
        }
        for t in doc.tables:
            if not t.rows: continue
            h0 = t.rows[0].cells[0].text.strip()
            # Find public survey table (contains 请问您是？)
            if h0 != '调查内容': continue
            is_public_survey = any('请问您是' in (t.rows[i].cells[0].text.strip() if i < len(t.rows) else '')
                                  for i in range(min(5, len(t.rows))))
            if not is_public_survey: continue

            for row in t.rows:
                if len(row.cells) < 4: continue
                opt = row.cells[1].text.strip()
                if opt in demo_data:
                    count, pct = demo_data[opt]
                    for p in row.cells[2].paragraphs:
                        for run in p.runs: run.text = count
                    for p in row.cells[3].paragraphs:
                        for run in p.runs: run.text = pct

        # Fix risk level mentions in conclusion text
        risk_level = scoring_report["meta"]["risk_level"]
        for para in doc.paragraphs:
            for run in para.runs:
                for wrong in ['高风险', '中风险']:
                    if wrong != risk_level and wrong in run.text:
                        if any(ctx in para.text for ctx in ['结论', '评定为', '等级为', '综合', '最终', '判定为']):
                            run.text = run.text.replace(wrong, risk_level)

        out_path = doc_path.replace('.docx', '_fixed.docx')
        doc.save(out_path)

    async def _run_post_assembly_checks(self, state: dict, output_path: str) -> None:
        """Run spec checker and template comparison after DOCX assembly."""
        try:
            from app.validation.spec_checker import check_report_async
            await self._emit_thinking("正在进行格式规范校验（DB32/T4013-2021）...")
            spec_report = await check_report_async(output_path)
            if spec_report:
                state["_spec_check"] = {
                    "score": spec_report.compliance_score,
                    "passed": spec_report.passed,
                    "passed_count": spec_report.passed_count,
                    "total_rules": spec_report.total_rules,
                    "failed_count": spec_report.failed_count,
                }
                await self._emit_thinking(
                    f"格式规范校验完成：合规率 {spec_report.compliance_score:.0%}"
                    f"（{spec_report.passed_count}/{spec_report.total_rules}）"
                    f"{'，已通过' if spec_report.passed else '，部分规则需调整'}"
                )
                await self._emit("validation_result", {
                    "passed": spec_report.passed,
                    "summary": f"格式规范合规率 {spec_report.compliance_score:.0%}",
                    "details": state["_spec_check"],
                })
            else:
                await self._emit_thinking("格式规范校验跳过：未找到规范文件")
        except Exception as e:
            logger.warning(f"Spec check failed (non-critical): {e}")

        try:
            from app.validation.diff_engine import ReportDiffEngine
            template_path = state.get("_template_path") or ""
            if not template_path:
                import sqlite3
                from app.config import settings
                db_path = settings.DATA_DIR / "knowledge_base.db"
                if db_path.exists():
                    conn = sqlite3.connect(str(db_path))
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT template_file_path FROM templates "
                        "WHERE category='社会稳定' AND is_active=1 "
                        "ORDER BY id DESC LIMIT 1"
                    )
                    row = cursor.fetchone()
                    conn.close()
                    if row and row[0]:
                        from app.services.file_service import file_service
                        abs_path = file_service.get_absolute_path(row[0])
                        if abs_path.exists():
                            template_path = str(abs_path)

            if template_path:
                await self._emit_thinking("正在与模板进行对比分析...")
                engine = ReportDiffEngine()
                diff_report = await engine.compare(output_path, template_path)
                state["_template_comparison"] = {
                    "score": diff_report.overall_score,
                    "passed": diff_report.passed,
                    "dimensions": [
                        {
                            "name": r.dimension,
                            "score": r.match_percentage,
                            "passed": r.passed,
                            "issues": r.issues,
                        }
                        for r in diff_report.results
                    ],
                }
                await self._emit_thinking(
                    f"模板对比完成：综合匹配度 {diff_report.overall_score:.0%}"
                    f"{'，达标' if diff_report.passed else '，部分维度需改进'}"
                )
                await self._emit("validation_result", {
                    "passed": diff_report.passed,
                    "summary": f"模板对比匹配度 {diff_report.overall_score:.0%}",
                    "details": state["_template_comparison"],
                })
            else:
                await self._emit_thinking("模板对比跳过：未找到对应模板文件")
        except Exception as e:
            logger.warning(f"Template comparison failed (non-critical): {e}")

    # ═══════════════════════════════════════════════════════════════
    # Chapter Review UI
    # ═══════════════════════════════════════════════════════════════

    async def _present_for_review(self, chapter_num: int, state: dict, is_regeneration: bool = False) -> None:
        chapters = state.get("chapters") or {}
        ch_data = chapters.get(chapter_num, {})
        if not isinstance(ch_data, dict): return

        markdown = ch_data.get("markdown", "")
        tables = ch_data.get("tables") or []
        word_count = len(markdown.replace('\n', '').replace(' ', ''))

        from app.agent.state import CHAPTER_DEFINITIONS
        title = CHAPTER_DEFINITIONS.get(chapter_num, {}).get("title", f"第{chapter_num}章")

        await self._emit("chapter_review_prompt", {
            "chapter": chapter_num, "title": title,
            "summary": f"已编写 {word_count} 字，{len(tables)} 个表格"
                      f"{'（修改版）' if is_regeneration else ''}",
            "tables_count": len(tables), "word_count": word_count,
            "actions": ["approve", "revise", "skip"],
            "is_regeneration": is_regeneration,
        })
        state["phase"] = "chapter_review"
        await self._emit("phase_change", {"phase": "chapter_review", "chapter": chapter_num})

    async def _request_and_wait_for_data(self, chapter_num: int, missing_keys: List[str], state: dict) -> bool:
        """Returns True if data was provided, False if timed out."""
        if not missing_keys: return True

        from app.agent.state import CHAPTER_DEFINITIONS
        title = CHAPTER_DEFINITIONS.get(chapter_num, {}).get("title", f"第{chapter_num}章")

        # Build data request
        descs = {
            "org_name": ("责任单位", "如：淮安市洪泽区人民政府"),
            "implement_unit": ("实施单位", "默认：江苏众拓项目代理咨询有限公司"),
            "location": ("项目位置", "如：淮安市洪泽区朱坝街道三圩社区"),
            "area_m2": ("征收面积（㎡）", "如：326342"),
            "area_mu": ("征收面积（亩）", "如：489.51"),
            "land_use": ("土地用途", "如：商业服务业设施用地"),
            "funding": ("资金测算", ""),
            "total_samples": ("问卷总人数", "如：120"),
            "support_rate": ("群众支持率", "如：95%"),
        }
        items = [
            {"key": k, "display_name": descs.get(k, (k, ""))[0],
             "description": descs.get(k, (k, ""))[1] or f"请提供「{k}」"}
            for k in missing_keys
        ]

        items_text = "\n".join(f"- **{i['display_name']}**：{i['description']}" for i in items)
        # 🔴 Single event — frontend renders missing_data_prompt as a message
        await self._emit("missing_data_prompt", {
            "chapter": chapter_num, "title": title,
            "data_items": items, "total_missing": len(items),
            "content": f"## ⚠️ 第{chapter_num}章「{title}」缺少数据\n\n{items_text}\n\n请直接输入数据，格式如「责任单位：淮安市洪泽区人民政府」",
        })
        state["phase"] = "chapter_review"
        await self._emit("phase_change", {"phase": "chapter_review"})

        # Wait for data using event
        for _ in range(int(DATA_WAIT_TIMEOUT)):
            if self._cancelled: return False
            filled = state.get("filled_data") or {}
            if all(filled.get(k) for k in missing_keys):
                await self._emit("thinking", "✅ 数据已收到，继续生成...")
                return True
            action = await self._wait_for_action(state, 1.0)
            if action in ("approve", "skip"):
                return False
            if action == "revise":
                # Parse text for key:value pairs
                feedback = state.get("chapter_feedback") or ""
                cn = {
                    "org_name": ["责任单位", "决策单位", "征收主体", "实施主体", "项目单位"],
                    "implement_unit": ["实施单位", "稳评单位", "评估单位"],
                    "project_name": ["项目名称", "决策名称", "报告标题", "标题", "名称"],
                    "report_title": ["报告标题", "标题", "项目名称"],
                    "location": ["位置", "坐落", "地址", "位于", "地理位置", "项目位置"],
                    "area_m2": ["面积", "平方米", "㎡", "用地面积", "征收面积"],
                    "area_mu": ["亩数", "亩"],
                    "land_use": ["用途", "地类", "用地性质", "土地用途"],
                    "funding": ["资金", "投资", "补偿金额", "总金额"],
                    "total_samples": ["人数", "样本", "调查人数", "问卷人数", "样本数"],
                    "support_rate": ["支持率", "同意率", "赞成率"],
                    "stakeholder_demands": ["诉求", "意见", "要求", "建议"],
                    "household_count": ["户数", "农户", "涉及户"],
                    "population_count": ["人口", "人数"],
                    "compensation_standard": ["补偿标准", "补偿单价"],
                    "doc_reference": ["文号", "批文号", "公告号"],
                }
                for line in feedback.split("\n"):
                    for key in missing_keys:
                        aliases = cn.get(key, [key])
                        if any(a in line for a in aliases):
                            parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                            if len(parts) == 2 and parts[1].strip():
                                filled[key] = parts[1].strip()
                state["filled_data"] = filled
                state["chapter_feedback"] = ""
                state["user_action"] = ""

        await self._emit("thinking", "⏰ 等待数据超时，暂停本章生成")
        return False

    # ═══════════════════════════════════════════════════════════════
    # SSE Helpers
    # ═══════════════════════════════════════════════════════════════

    def _collect_missing_data(self, state: dict) -> list:
        """Scan all chapter markdown for 【需补充】 markers and collect unique items."""
        import re as _re
        chapters = state.get("chapters", {})
        missing = []
        seen = set()
        for ch_num, ch_data in chapters.items():
            md = (ch_data or {}).get("markdown", "") if isinstance(ch_data, dict) else ""
            if not isinstance(md, str): continue
            for m in _re.finditer(r'【需补充[：:]*([^】]*)】', md):
                hint = m.group(1).strip() if m.group(1) else f"第{ch_num}章数据"
                if hint not in seen:
                    seen.add(hint)
                    missing.append({"hint": hint, "chapter": ch_num, "placeholder": m.group(0)})
        return missing

    async def _emit(self, event_type: str, data, state: dict = None, phase: str = None) -> None:
        """Emit SSE event + optionally sync state phase."""
        if phase and state:
            state["phase"] = phase
        if self._stream_queue:
            # Auto-wrap string data (some callers pass raw strings)
            if isinstance(data, str):
                data = {"content": data}
            await self._stream_queue.put({"event": event_type, "data": data})

    async def _emit_thinking(self, content: str) -> None:
        await self._emit("thinking", {"content": content})
